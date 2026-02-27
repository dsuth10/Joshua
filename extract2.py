import sys
from docx import Document

def extract_text(file_path):
    out = []
    out.append(f"--- Extracting {file_path} ---")
    try:
        doc = Document(file_path)
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                out.append(text)
        out.append("--- Tables ---")
        for table in doc.tables:
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip().replace('\n', ' ')
                    if cell_text and cell_text not in row_data:
                        row_data.append(cell_text)
                if row_data:
                    out.append(" | ".join(row_data))
            out.append("---")
    except Exception as e:
        out.append(f"Error reading {file_path}: {e}")
    return "\n".join(out)

text1 = extract_text(r"c:\Users\dsuth\Documents\Joshua\Year34_English_Sequence\K-6 3-4 English Unit 1 Plan 26 .docx")
text2 = extract_text(r"c:\Users\dsuth\Documents\Joshua\Year34_English_Sequence\p-6cpm_eng_v9_y04_u1_at1-1.docx")

with open(r"c:\Users\dsuth\Documents\Joshua\extracted2_utf8.txt", "w", encoding="utf-8") as f:
    f.write(text1 + "\n\n" + text2)
