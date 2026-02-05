import os
import sys

# Add the docx skill path if needed or use the tool directly via npx/python if available.
# Since I have the docx skill, I will use the python library directly in a script.

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_lesson_plan(output_path):
    doc = Document()
    
    # Title
    title = doc.add_heading('Maths Lesson: Rounding Numbers', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Lesson Details
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = 'Year Level: Year 5'
    table.cell(0, 1).text = 'Subject: Mathematics'
    table.cell(1, 0).text = 'Topic: Rounding (Number & Place Value)'
    table.cell(1, 1).text = 'Duration: 60 Minutes'
    
    doc.add_paragraph() # Spacer
    
    # Learning Intention & Success Criteria
    doc.add_heading('Learning Intention', level=1)
    doc.add_paragraph('We are learning to round numbers to the nearest ten, hundred, and thousand to help us estimate and check the reasonableness of our answers.')
    
    doc.add_heading('Success Criteria', level=1)
    doc.add_paragraph('I can:', style='List Bullet')
    doc.add_paragraph('Identify the rounding digit and the digit to its right.', style='List Bullet')
    doc.add_paragraph('Apply the "5 or more, round up; 4 or less, let it rest" rule.', style='List Bullet')
    doc.add_paragraph('Round numbers up to tens of thousands to the nearest 10, 100, or 1000.', style='List Bullet')
    
    # Introduction (10 mins)
    doc.add_heading('Introduction (10 mins)', level=1)
    doc.add_paragraph('Start with a real-world scenario: "If 3864 people attended a concert, would we say there were about 3000 or about 4000 people?" Discuss why we round (estimation, simplicity).')
    doc.add_paragraph('Review the core rule: Look at the digit to the right of the rounding place. If it is 5, 6, 7, 8, or 9, round up. If it is 0, 1, 2, 3, or 4, round down (stay the same).')
    
    # Guided Discovery & Examples (15 mins)
    doc.add_heading('Guided Discovery (15 mins)', level=1)
    doc.add_paragraph('Work through examples on the board/PowerPoint:')
    doc.add_paragraph('Rounding to the nearest 10: 96 becomes 100.', style='List Bullet')
    doc.add_paragraph('Rounding to the nearest 100: 527 becomes 500.', style='List Bullet')
    doc.add_paragraph('Rounding to the nearest 1000: 4510 becomes 5000.', style='List Bullet')
    
    # Independent Practice (30 mins)
    doc.add_heading('Independent Practice (30 mins)', level=1)
    doc.add_paragraph('Students complete Australian Signpost Maths 5, Page 30 (Rounding).')
    doc.add_paragraph('Tasks overview:')
    doc.add_paragraph('Q1: Rounding to the nearest ten (e.g., 37, 82, 1894).', style='List Bullet')
    doc.add_paragraph('Q2: Rounding to the nearest hundred.', style='List Bullet')
    doc.add_paragraph('Q3: Rounding to the nearest thousand.', style='List Bullet')
    doc.add_paragraph('Q4-5: Circle numbers that round to a specific value.', style='List Bullet')
    doc.add_paragraph('Q6: True or False statements.', style='List Bullet')
    
    # Conclusion / Plenary (5 mins)
    doc.add_heading('Conclusion (5 mins)', level=1)
    doc.add_paragraph('Review a few tricky ones together (e.g., 1895 to the nearest 10). Discuss why 4500 rounds to 5000 when rounding to the nearest thousand.')
    
    doc.save(output_path)
    print(f"Lesson plan created at {output_path}")

if __name__ == "__main__":
    output_dir = r"c:\Users\dsuth\Documents\Joshua\Signpost Math Lessons"
    os.makedirs(output_dir, exist_ok=True)
    create_lesson_plan(os.path.join(output_dir, "Lesson_Plan_Rounding.docx"))
