from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor


ROOT = Path(r"C:\Users\dsuth\Documents\Joshua")
QA = ROOT / ".codex_tmp" / "flying_fox_article_qa"
OUT = ROOT / "Units" / "English" / "English_Unit_3" / "Lesson_Plans" / "Lesson_15" / "Flying_Fox_Alternative" / "Lesson_15_Flying_Fox_Reading_Article.docx"

INK = "12343D"
DEEP = "063C46"
TEAL = "0B8E9F"
CORAL = "E84E47"
CREAM = "F7E8C6"
MINT = "E8F6F4"
PALE = "F1F7F7"
LINE = "C6D7D9"
MUTED = "5C7277"
WHITE = "FFFFFF"


def set_run_font(run, size=10.3, color=INK, bold=False, italic=False, name="Arial"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def set_paragraph(p, before=0, after=7, line=1.18, align=None, keep=False):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    pf.keep_with_next = keep
    if align is not None:
        p.alignment = align


def add_rich_paragraph(container, pieces, size=10.3, after=7, line=1.18, color=INK, align=None):
    p = container.add_paragraph()
    set_paragraph(p, after=after, line=line, align=align)
    for text, bold, italic in pieces:
        set_run_font(p.add_run(text), size=size, color=color, bold=bold, italic=italic)
    return p


def add_body(container, text, bold_lead=None, after=7, size=10.3, line=1.18):
    p = container.add_paragraph()
    set_paragraph(p, after=after, line=line)
    if bold_lead and text.startswith(bold_lead):
        set_run_font(p.add_run(bold_lead), size=size, bold=True)
        set_run_font(p.add_run(text[len(bold_lead):]), size=size)
    else:
        set_run_font(p.add_run(text), size=size)
    return p


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, **edges):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge_name, opts in edges.items():
        edge = borders.find(qn(f"w:{edge_name}"))
        if edge is None:
            edge = OxmlElement(f"w:{edge_name}")
            borders.append(edge)
        edge.set(qn("w:val"), opts.get("val", "single"))
        edge.set(qn("w:sz"), str(opts.get("sz", 8)))
        edge.set(qn("w:color"), opts.get("color", LINE))


def set_table_geometry(table, widths_inches, indent_dxa=120):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    total = round(sum(widths_inches) * 1440)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for old in list(grid):
        grid.remove(old)
    for width in widths_inches:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(round(width * 1440)))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            dxa = round(widths_inches[idx] * 1440)
            cell.width = Inches(widths_inches[idx])
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(dxa))
            tc_w.set(qn("w:type"), "dxa")


def remove_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{name}")
        el.set(qn("w:val"), "nil")
        borders.append(el)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_image(doc, path, width, alt):
    p = doc.add_paragraph()
    set_paragraph(p, after=7, line=1.0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shape = p.add_run().add_picture(str(path), width=Inches(width))
    doc_pr = shape._inline.docPr
    doc_pr.set("descr", alt)
    doc_pr.set("title", alt)
    return p


def add_numbered_heading(container, number, title, width, after=5):
    table = container.add_table(rows=1, cols=2)
    set_table_geometry(table, [0.34, width - 0.34])
    remove_table_borders(table)
    ncell, tcell = table.rows[0].cells
    ncell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    tcell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    shade_cell(ncell, TEAL)
    set_cell_margins(ncell, top=55, start=120, bottom=55, end=120)
    set_cell_margins(tcell, top=0, start=120, bottom=0, end=120)
    np = ncell.paragraphs[0]
    set_paragraph(np, after=0, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_run_font(np.add_run(str(number)), size=9, color=WHITE, bold=True)
    tp = tcell.paragraphs[0]
    set_paragraph(tp, after=0, line=1.0, keep=True)
    set_run_font(tp.add_run(title), size=14.2, color="066B78", bold=True)
    spacer = container.add_paragraph()
    set_paragraph(spacer, after=after, line=1.0)
    return table


def add_callout(container, title, text, fill=MINT, accent=TEAL, dark=False, width=4.15):
    table = container.add_table(rows=1, cols=1)
    set_table_geometry(table, [width])
    cell = table.cell(0, 0)
    shade_cell(cell, DEEP if dark else fill)
    set_cell_margins(cell, top=150, start=120, bottom=150, end=120)
    set_cell_border(cell, start={"color": accent, "sz": 18})
    p = cell.paragraphs[0]
    set_paragraph(p, after=4, line=1.1)
    set_run_font(p.add_run(title), size=10.7, color=WHITE if dark else DEEP, bold=True)
    p2 = cell.add_paragraph()
    set_paragraph(p2, after=0, line=1.17)
    set_run_font(p2.add_run(text), size=9.4, color=WHITE if dark else INK, bold=dark)
    spacer = container.add_paragraph()
    set_paragraph(spacer, after=5, line=1.0)
    return table


def add_top_rule(cell, color=CORAL, sz=18):
    set_cell_border(cell, top={"color": color, "sz": sz})


def add_column_break(doc):
    p = doc.add_paragraph()
    set_paragraph(p, after=0, line=1.0)
    p.add_run().add_break(WD_BREAK.COLUMN)


def set_columns(section, widths=None, space=360):
    sect_pr = section._sectPr
    cols = sect_pr.find(qn("w:cols"))
    if cols is not None:
        sect_pr.remove(cols)
    cols = OxmlElement("w:cols")
    if widths:
        cols.set(qn("w:num"), str(len(widths)))
        cols.set(qn("w:equalWidth"), "0")
        cols.set(qn("w:space"), str(space))
        for width in widths:
            col = OxmlElement("w:col")
            col.set(qn("w:w"), str(width))
            col.set(qn("w:space"), str(space))
            cols.append(col)
    else:
        cols.set(qn("w:num"), "1")
    sect_pr.append(cols)


def set_paragraph_bottom_border(p, color="9EB7BB", sz=5, space=3):
    p_pr = p._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(sz))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_page_field(paragraph):
    set_run_font(paragraph.add_run("READING "), size=7, color=MUTED, bold=True)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    r = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), "Arial")
    rfonts.set(qn("w:hAnsi"), "Arial")
    rpr.append(rfonts)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), MUTED)
    rpr.append(color)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "14")
    rpr.append(sz)
    r.append(rpr)
    t = OxmlElement("w:t")
    t.text = "1"
    r.append(t)
    fld.append(r)
    paragraph._p.append(fld)


def configure_section(section):
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(13)
    section.bottom_margin = Mm(13)
    section.left_margin = Mm(13)
    section.right_margin = Mm(13)
    section.header_distance = Mm(6)
    section.footer_distance = Mm(6)


def style_doc(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(10.3)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.18
    for style_name, size, color in (("Heading 1", 15, "066B78"), ("Heading 2", 12.5, DEEP), ("Heading 3", 10.5, DEEP)):
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)
        style.paragraph_format.keep_with_next = True


def add_footer(section):
    footer = section.footer
    footer.is_linked_to_previous = False
    p0 = footer.paragraphs[0]
    p0._element.getparent().remove(p0._element)
    table = footer.add_table(rows=1, cols=2, width=Inches(7.24))
    set_table_geometry(table, [5.95, 1.29], indent_dxa=0)
    remove_table_borders(table)
    left, right = table.rows[0].cells
    set_cell_margins(left, top=0, start=0, bottom=0, end=0)
    set_cell_margins(right, top=0, start=0, bottom=0, end=0)
    lp = left.paragraphs[0]
    set_paragraph(lp, after=0, line=1.0)
    set_run_font(lp.add_run("Lesson 15 | Flying-fox evidence reading"), size=7, color=MUTED)
    rp = right.paragraphs[0]
    set_paragraph(rp, after=0, line=1.0, align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_page_field(rp)


doc = Document()
style_doc(doc)
configure_section(doc.sections[0])
set_columns(doc.sections[0])

# Page 1 hero and editable article body.
add_image(doc, QA / "flying_fox_banner.png", 7.24, "Flying-foxes leaving a Queensland tree camp at dusk; article title and reading mission overlaid.")

page1_body = doc.add_section(WD_SECTION.CONTINUOUS)
configure_section(page1_body)
set_columns(page1_body, [6100, 3740], space=380)

add_numbered_heading(doc, 1, "Night travellers and daytime camps", 4.24)
add_body(doc, "Flying-foxes are large fruit- and nectar-eating bats. They are nocturnal: they rest together in trees during the day, then fly out around dusk to feed. The trees used for resting and breeding are called a camp or roost. Food and camp use change with flowering, fruiting and seasons, so numbers at a camp can rise, fall or move.", size=9.7, line=1.16)
add_body(doc, "As flying-foxes travel between feeding areas, they can carry pollen and seeds. This helps native forests reproduce and connect across the landscape. Their long-distance movement makes them important pollinators and seed dispersers. Australia monitors flying-fox populations because grey-headed and spectacled flying-foxes are nationally listed as threatened.", size=9.7, line=1.16)
add_numbered_heading(doc, 2, "Why camps appear near people", 4.24)
add_body(doc, "Flying-foxes need flowering trees, feeding habitat and safe places to rest. Forest loss and fragmentation have reduced some natural habitat. Remaining patches of trees in towns, parks and along waterways can become useful camp sites. This can bring wildlife and people into closer contact.", size=9.7, line=1.16)
add_body(doc, "A large or seasonal camp can create real challenges. Residents may experience noise, strong odour, droppings on cars or washing, damaged vegetation and interrupted sleep. These impacts are not imaginary simply because flying-foxes are ecologically valuable. A fair community message should describe the exact impact instead of labelling the animals.", size=9.7, line=1.16)

add_column_break(doc)
side_table = doc.add_table(rows=1, cols=1)
set_table_geometry(side_table, [2.55])
side_cell = side_table.cell(0, 0)
shade_cell(side_cell, CREAM)
set_cell_margins(side_cell, top=170, start=120, bottom=170, end=120)
add_top_rule(side_cell)
p = side_cell.paragraphs[0]
set_paragraph(p, after=7, line=1.0)
set_run_font(p.add_run("Quick evidence check"), size=13, color=CORAL, bold=True)
check = side_cell.add_table(rows=4, cols=2)
set_table_geometry(check, [1.05, 1.18])
set_repeat_table_header(check.rows[0])
for r_idx, row in enumerate(check.rows):
    for c_idx, cell in enumerate(row.cells):
        set_cell_margins(cell, top=65, start=120, bottom=65, end=120)
        set_cell_border(cell, top={"color": "E1CFA5", "sz": 5}, bottom={"color": "E1CFA5", "sz": 5}, start={"color": "E1CFA5", "sz": 5}, end={"color": "E1CFA5", "sz": 5})
        if c_idx == 0:
            shade_cell(cell, "FFFDFA")
texts = [
    ("What could be checked?", "What detail would make it precise?"),
    ("The number of bats at a camp", "Place, date, counting method and species"),
    ("Residents reported noise", "Who was asked, how many people and when"),
    ("Flying-foxes move seeds", "Which species, plants, distance and research source"),
]
for r_idx, values in enumerate(texts):
    for c_idx, value in enumerate(values):
        cp = check.cell(r_idx, c_idx).paragraphs[0]
        set_paragraph(cp, after=0, line=1.05)
        set_run_font(cp.add_run(value), size=7.4, bold=(c_idx == 0 or r_idx == 0))

alert = side_cell.add_table(rows=1, cols=1)
set_table_geometry(alert, [2.23])
alert_cell = alert.cell(0, 0)
shade_cell(alert_cell, WHITE)
set_cell_margins(alert_cell, top=105, start=120, bottom=105, end=120)
set_cell_border(alert_cell, start={"color": TEAL, "sz": 16})
ap = alert_cell.paragraphs[0]
set_paragraph(ap, after=2, line=1.05)
set_run_font(ap.add_run("Language alert"), size=8.7, bold=True)
ap2 = alert_cell.add_paragraph()
set_paragraph(ap2, after=0, line=1.1)
set_run_font(ap2.add_run('“A camp is present beside the park” is checkable. “A disgusting invasion has stolen our park” adds loaded judgement and metaphor.'), size=7.7)

wp = side_cell.add_paragraph()
set_paragraph(wp, before=6, after=5, line=1.0)
set_run_font(wp.add_run("Word help"), size=12, color=CORAL, bold=True)
for term, definition in [
    ("nocturnal", "active at night"),
    ("camp or roost", "trees where flying-foxes rest and breed"),
    ("pollinator", "an animal that moves pollen between flowers"),
    ("fragmentation", "habitat broken into smaller patches"),
]:
    add_rich_paragraph(side_cell, [(term, True, False), (" - " + definition, False, False)], size=7.7, after=3, line=1.08)
add_callout(side_cell, "Evidence can hold two truths", "Ecological value and community impact can both be true. Strong evidence describes each one precisely.", dark=True, width=2.23)

# Page 2.
page2_head = doc.add_section(WD_SECTION.NEW_PAGE)
configure_section(page2_head)
set_columns(page2_head)
add_image(doc, QA / "page2_header.png", 7.24, "Health, safety and lawful coexistence section banner over a flying-fox camp at dusk.")

page2_body = doc.add_section(WD_SECTION.CONTINUOUS)
configure_section(page2_body)
set_columns(page2_body, [6100, 3740], space=380)
add_numbered_heading(doc, 3, "What the health advice actually says", 4.24)
add_body(doc, "Queensland Government advice says that catching disease directly from flying-foxes is extremely unlikely.", size=10.1, line=1.17)
add_body(doc, "Australian Bat Lyssavirus can be transmitted through an untreated bite or scratch from an infected bat. Living, walking or playing near a camp does not expose a person to the virus. There is no evidence that Hendra virus passes directly from flying-foxes to people.", size=9.8, line=1.16)
add_callout(doc, "No touch means safe distance", "Never touch, feed or try to rescue a flying-fox. Keep people and pets away from an injured or grounded bat and contact a trained, vaccinated wildlife rescuer. If a person is bitten or scratched, seek medical advice immediately.", width=4.24)
add_numbered_heading(doc, 4, "Practical responses", 4.24)
add_body(doc, "People living near a seasonal camp can reduce some impacts by planning ahead. Official suggestions include bringing washing in at night, parking vehicles under cover, covering outdoor items and maintaining buildings or water systems appropriately. Avoiding disturbance can also reduce extra noise and stress.", size=9.7, line=1.16)
add_body(doc, "Flying-foxes and their roosts are protected under Queensland conservation law. Some low-impact activities and council management actions are allowed only when current rules and codes are followed. Management must not harm the animals.", size=9.7, line=1.16)
add_body(doc, "Moving a camp is not a simple “push them away” solution: flying-foxes are mobile, impacts can shift elsewhere, and disturbance can affect animals and nearby residents.", size=9.7, line=1.16)

add_column_break(doc)
panel = doc.add_table(rows=1, cols=1)
set_table_geometry(panel, [2.55])
panel_cell = panel.cell(0, 0)
shade_cell(panel_cell, CREAM)
set_cell_margins(panel_cell, top=170, start=120, bottom=170, end=120)
add_top_rule(panel_cell)
hp = panel_cell.paragraphs[0]
set_paragraph(hp, after=6, line=1.0)
set_run_font(hp.add_run("Practical, lawful action"), size=13, color=CORAL, bold=True)
for title, text in [
    ("Plan around nightly movement", "Bring washing inside, cover outdoor items and park vehicles under cover where possible."),
    ("Leave rescue to trained people", "Keep people and pets back from an injured or grounded bat and contact a vaccinated wildlife rescuer."),
    ("Check the current rules", "Council action must follow Queensland conservation law, current codes and animal-welfare requirements."),
]:
    action = panel_cell.add_table(rows=1, cols=1)
    set_table_geometry(action, [2.23])
    ac = action.cell(0, 0)
    shade_cell(ac, WHITE)
    set_cell_margins(ac, top=115, start=120, bottom=115, end=120)
    ap = ac.paragraphs[0]
    set_paragraph(ap, after=3, line=1.05)
    set_run_font(ap.add_run(title), size=9.4, color=DEEP, bold=True)
    ap2 = ac.add_paragraph()
    set_paragraph(ap2, after=0, line=1.12)
    set_run_font(ap2.add_run(text), size=8.2)
    sp = panel_cell.add_paragraph()
    set_paragraph(sp, after=3, line=1.0)
add_callout(panel_cell, "Management is not simple", "Disturbance can create more noise, stress animals and move impacts to another place. A strong response checks evidence, safety, law and likely consequences.", dark=True, width=2.23)
add_rich_paragraph(panel_cell, [("Evidence habit: ", True, False), ("When a health claim sounds alarming, check the pathway of exposure. “Near a camp” and “bitten or scratched by a bat” are not the same situation.", False, False)], size=8.0, after=0, line=1.12)

# Page 3.
page3 = doc.add_section(WD_SECTION.NEW_PAGE)
configure_section(page3)
set_columns(page3)
add_image(doc, QA / "page3_header.png", 7.24, "Teal section banner: Audit the message before the council publishes it.")
add_numbered_heading(doc, 5, "A fair council message needs three things", 7.24)
add_body(doc, "These messages are fictional and deliberately positioned. Audit the language; neither is the official answer.", size=9.7, line=1.14)

criteria = doc.add_table(rows=1, cols=3)
set_table_geometry(criteria, [2.35, 2.35, 2.36])
criteria_data = [
    ("Accuracy", MINT, TEAL, "Checkable facts and source limits", "Rumour or certainty without evidence"),
    ("Respect", "FFF1EF", CORAL, "Real community impacts and ecological value", "Mocking residents or demonising wildlife"),
    ("Action", "FFF7DD", "D6A11E", "Safe, lawful and practical steps", "Handling bats or unauthorised disturbance"),
]
for idx, (title, fill, accent, include, avoid) in enumerate(criteria_data):
    cell = criteria.cell(0, idx)
    shade_cell(cell, fill)
    set_cell_margins(cell, top=110, start=120, bottom=110, end=120)
    set_cell_border(cell, top={"color": accent, "sz": 14})
    cp = cell.paragraphs[0]
    set_paragraph(cp, after=4, line=1.0)
    set_run_font(cp.add_run(title), size=10.2, color=DEEP, bold=True)
    add_rich_paragraph(cell, [("Include: ", True, False), (include, False, False)], size=7.8, after=3, line=1.08)
    add_rich_paragraph(cell, [("Avoid: ", True, False), (avoid, False, False)], size=7.8, after=0, line=1.08, color="9A3B35")

add_callout(doc, "Remember", "Objective language is not automatically proven. A council should still check its source, date, place, measurement and limits.", fill=PALE, width=7.24)

titlep = doc.add_paragraph()
set_paragraph(titlep, before=2, after=5, line=1.0, keep=True)
set_run_font(titlep.add_run("Rivergum Council: two draft community messages"), size=13.5, color=DEEP, bold=True)

drafts = doc.add_table(rows=2, cols=2)
set_table_geometry(drafts, [3.55, 3.55])
set_repeat_table_header(drafts.rows[0])
headers = [("MESSAGE A | CLEAR THE CAMP NOW", CORAL), ("MESSAGE B | WELCOME OUR NIGHT GARDENERS", TEAL)]
for idx, (text, fill) in enumerate(headers):
    cell = drafts.cell(0, idx)
    shade_cell(cell, fill)
    set_cell_margins(cell, top=90, start=120, bottom=90, end=120)
    p = cell.paragraphs[0]
    set_paragraph(p, after=0, line=1.0)
    set_run_font(p.add_run(text), size=9.3, color=WHITE, bold=True)
draft_texts = [
    ("Resident Action Group draft", "A deafening invasion has taken over Rivergum Park. Families are forced to endure unbearable noise, foul smells and dangerous animals beside their homes. The camp is ruining our neighbourhood, and every responsible council would clear it immediately. Rivergum residents deserve their peaceful park back - no excuses and no delays."),
    ("Wildlife Network draft", "Our wonderful night gardeners are bringing life to Rivergum Park. Flying-foxes move pollen and seeds through native forests, so their camp is a precious gift. Noise and droppings are only a tiny inconvenience, and there is nothing for residents to worry about. Everyone should welcome the camp and leave nature completely alone."),
]
for idx, (label, text) in enumerate(draft_texts):
    cell = drafts.cell(1, idx)
    set_cell_margins(cell, top=105, start=120, bottom=105, end=120)
    set_cell_border(cell, bottom={"color": LINE, "sz": 7}, start={"color": LINE, "sz": 7}, end={"color": LINE, "sz": 7})
    lp = cell.paragraphs[0]
    set_paragraph(lp, after=4, line=1.0)
    set_run_font(lp.add_run(label.upper()), size=7.5, color=MUTED, bold=True)
    tp = cell.add_paragraph()
    set_paragraph(tp, after=0, line=1.1)
    set_run_font(tp.add_run(text), size=8.3)

audit = doc.add_table(rows=1, cols=1)
set_table_geometry(audit, [7.24])
ac = audit.cell(0, 0)
shade_cell(ac, CREAM)
set_cell_margins(ac, top=110, start=120, bottom=110, end=120)
set_cell_border(ac, top={"color": CORAL, "sz": 14})
ah = ac.paragraphs[0]
set_paragraph(ah, after=4, line=1.0)
set_run_font(ah.add_run("Audit before you advise"), size=10.2, color=CORAL, bold=True)
for item in [
    "Underline one checkable claim in each message.",
    "Circle one judgement, emotion or loaded phrase in each message.",
    "Place a star beside one statement that mixes information and positioning.",
    "Write one question the council should answer before publishing either message.",
]:
    p = ac.add_paragraph(style="List Number")
    set_paragraph(p, after=2, line=1.05)
    p.paragraph_format.left_indent = Inches(0.28)
    p.paragraph_format.first_line_indent = Inches(-0.16)
    set_run_font(p.add_run(item), size=8.2)

qb = doc.add_table(rows=1, cols=1)
set_table_geometry(qb, [7.24])
qc = qb.cell(0, 0)
shade_cell(qc, "FBFDFD")
set_cell_margins(qc, top=90, start=120, bottom=90, end=120)
set_cell_border(qc, top={"color": LINE, "sz": 7}, bottom={"color": LINE, "sz": 7}, start={"color": LINE, "sz": 7}, end={"color": LINE, "sz": 7})
qh = qc.paragraphs[0]
set_paragraph(qh, after=3, line=1.0)
set_run_font(qh.add_run("My question for Rivergum Council"), size=9.4, color=DEEP, bold=True)
for _ in range(2):
    line = qc.add_paragraph()
    set_paragraph(line, after=4, line=1.0)
    set_paragraph_bottom_border(line)

source = doc.add_paragraph()
set_paragraph(source, before=4, after=0, line=1.02)
set_run_font(source.add_run("Source trail: "), size=6.8, color=DEEP, bold=True)
set_run_font(source.add_run("Evidence briefing: original summary of Queensland Government Living near flying-foxes; Queensland Government Bats and human health; Queensland Flying-fox roost management guideline; and Australian Government DCCEEW flying-fox monitoring and conservation pages. Full links are in the Teacher Guide. Checked 29 July 2026. Hero image is AI-generated and illustrative, not field evidence."), size=6.8, color=MUTED)

# Consistent linked footer and section geometry.
for idx, section in enumerate(doc.sections):
    configure_section(section)
    if idx == 0:
        add_footer(section)
    else:
        section.footer.is_linked_to_previous = True

# Encourage Word to update page fields on open.
settings = doc.settings._element
update = settings.find(qn("w:updateFields"))
if update is None:
    update = OxmlElement("w:updateFields")
    settings.append(update)
update.set(qn("w:val"), "true")

doc.core_properties.title = "Lesson 15 - Living With Flying-Foxes"
doc.core_properties.subject = "Adapted evidence reading article"
doc.core_properties.author = ""
doc.core_properties.keywords = "flying-foxes, evidence, community, Queensland"
doc.core_properties.comments = "AI-assisted layout; flying-fox hero image is illustrative."

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)
