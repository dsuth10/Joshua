from pathlib import Path

import pytest
from docx import Document

from audiobook_studio.contracts import HeadingRangeSelector, PageRangeSelector
from audiobook_studio.errors import SourceSelectionError
from audiobook_studio.extractors.docx import DocxExtractor


def _document(path: Path) -> None:
    document = Document()
    document.add_heading("Chapter One", level=1)
    paragraph = document.add_paragraph()
    paragraph.add_run("Bold opening").bold = True
    paragraph.add_run(" and ordinary text.")
    document.add_heading("Chapter Two", level=1)
    document.add_paragraph("Outside selection.")
    document.save(path)


def test_docx_heading_structure_and_emphasis(tmp_path: Path) -> None:
    source = tmp_path / "sample.docx"
    _document(source)
    extractor = DocxExtractor()
    assert [heading.text for heading in extractor.inspect(source).headings] == [
        "Chapter One",
        "Chapter Two",
    ]
    result = extractor.extract(
        source,
        HeadingRangeSelector(
            type="heading_range",
            start_heading="Chapter One",
            end_before_heading="Chapter Two",
        ),
    )
    assert result.narration_text == "Bold opening and ordinary text.\n"
    assert "**Bold opening**" in result.original_selection


def test_docx_page_range_fails_with_guidance(tmp_path: Path) -> None:
    source = tmp_path / "sample.docx"
    _document(source)
    with pytest.raises(SourceSelectionError, match="LibreOffice"):
        DocxExtractor().extract(
            source, PageRangeSelector(type="page_range", start_page=1, end_page=1)
        )


def test_docx_nonempty_table_is_rejected(tmp_path: Path) -> None:
    source = tmp_path / "table.docx"
    document = Document()
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Narrative in a table"
    document.save(source)
    with pytest.raises(SourceSelectionError, match="non-empty table"):
        DocxExtractor().inspect(source)
