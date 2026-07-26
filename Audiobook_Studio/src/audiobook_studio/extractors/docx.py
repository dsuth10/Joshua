"""Structural DOCX extraction using paragraph styles and runs."""

from pathlib import Path

from audiobook_studio.contracts import (
    DocumentIndex,
    DocumentParagraph,
    ExtractedSelection,
    Heading,
    PageRangeSelector,
    Selector,
)
from audiobook_studio.errors import SourceSelectionError
from audiobook_studio.extractors.markdown import count_words
from audiobook_studio.extractors.structured import StructuralHeading, select_paragraphs


class DocxExtractor:
    def _load(self, source: Path) -> tuple[list[str], list[str], list[StructuralHeading]]:
        try:
            from docx import Document
        except ImportError as exc:
            raise SourceSelectionError("DOCX support requires python-docx") from exc
        document = Document(str(source))
        table_cells = (
            cell for table in document.tables for row in table.rows for cell in row.cells
        )
        if any(cell.text.strip() for cell in table_cells):
            raise SourceSelectionError(
                "DOCX contains a non-empty table; narrative table linearisation is not configured"
            )
        plain: list[str] = []
        marked: list[str] = []
        headings: list[StructuralHeading] = []
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            paragraph_index = len(plain)
            style_name = paragraph.style.name if paragraph.style else ""
            level = None
            if style_name.lower().startswith("heading"):
                suffix = style_name.split()[-1]
                level = int(suffix) if suffix.isdigit() else 1
                headings.append(
                    StructuralHeading(
                        text=text,
                        level=level,
                        paragraph_index=paragraph_index,
                    )
                )
            runs: list[str] = []
            for run in paragraph.runs:
                value = run.text
                if not value:
                    continue
                if run.bold:
                    value = f"**{value}**"
                elif run.italic:
                    value = f"*{value}*"
                runs.append(value)
            plain.append(text)
            marked.append("".join(runs) or text)
        return plain, marked, headings

    def inspect(self, source: Path) -> DocumentIndex:
        plain, _, headings = self._load(source)
        return DocumentIndex(
            source_path=str(source),
            source_format="docx",
            headings=[
                Heading(
                    level=heading.level,
                    text=heading.text,
                    line_number=heading.paragraph_index + 1,
                )
                for heading in headings
            ],
            paragraphs=[
                DocumentParagraph(
                    paragraph_id=f"p{index:04d}",
                    text=text,
                    location=f"paragraph:{index}",
                    heading_level=next(
                        (
                            heading.level
                            for heading in headings
                            if heading.paragraph_index == index - 1
                        ),
                        None,
                    ),
                )
                for index, text in enumerate(plain, start=1)
            ],
            warnings=["DOCX physical pages are renderer-dependent; prefer heading selectors."],
        )

    def extract(self, source: Path, selector: Selector) -> ExtractedSelection:
        if isinstance(selector, PageRangeSelector):
            raise SourceSelectionError(
                "DOCX page selection requires deterministic LibreOffice rendering; "
                "use embedded page-range headings for Berani"
            )
        plain, marked, headings = self._load(source)
        selected, heading, start, end = select_paragraphs(plain, selector, headings=headings)
        original = "\n\n".join(marked[start - 1 : end]).strip() + "\n"
        narration = "\n\n".join(selected).strip() + "\n"
        return ExtractedSelection(
            source_path=str(source),
            selector=selector,
            heading=heading,
            start_line=start,
            end_line=end,
            original_selection=original,
            narration_text=narration,
            word_count=count_words(narration),
            paragraph_count=len(selected),
        )
