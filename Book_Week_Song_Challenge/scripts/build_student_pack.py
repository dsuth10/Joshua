from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / 'Book_Week_Song_Challenge_Student_Pack.docx'

NAVY = '152A4A'; ORANGE = 'EF8B3C'; GOLD = 'FFC84D'; TEAL = '1D8B91'; PALE = 'FFF5DE'; LINE = 'C8D2DE'; WHITE = 'FFFFFF'; INK = '152235'

def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr(); shd = tc_pr.find(qn('w:shd'))
    if shd is None: shd = OxmlElement('w:shd'); tc_pr.append(shd)
    shd.set(qn('w:fill'), fill)

def set_cell_border(cell, color=LINE, size='8'):
    tc_pr = cell._tc.get_or_add_tcPr(); borders = tc_pr.first_child_found_in('w:tcBorders')
    if borders is None: borders = OxmlElement('w:tcBorders'); tc_pr.append(borders)
    for edge in ('top','left','bottom','right'):
        tag = 'w:' + edge; elem = borders.find(qn(tag))
        if elem is None: elem = OxmlElement(tag); borders.append(elem)
        elem.set(qn('w:val'),'single'); elem.set(qn('w:sz'),size); elem.set(qn('w:color'),color)

def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr(); mar = tcPr.first_child_found_in('w:tcMar')
    if mar is None: mar = OxmlElement('w:tcMar'); tcPr.append(mar)
    for side, value in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        node = mar.find(qn('w:'+side))
        if node is None: node=OxmlElement('w:'+side); mar.append(node)
        node.set(qn('w:w'),str(value)); node.set(qn('w:type'),'dxa')

def fixed_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT; table.autofit = False
    tblPr = table._tbl.tblPr; tblW = tblPr.first_child_found_in('w:tblW')
    if tblW is None: tblW = OxmlElement('w:tblW'); tblPr.append(tblW)
    tblW.set(qn('w:w'),'9360'); tblW.set(qn('w:type'),'dxa')
    tblInd = tblPr.first_child_found_in('w:tblInd')
    if tblInd is None: tblInd=OxmlElement('w:tblInd'); tblPr.append(tblInd)
    tblInd.set(qn('w:w'),'120'); tblInd.set(qn('w:type'),'dxa')
    grid=table._tbl.tblGrid
    for col,width in zip(grid.gridCol_lst,widths): col.set(qn('w:w'),str(width))
    for row in table.rows:
        for cell,width in zip(row.cells,widths):
            cell.width=Inches(width/1440); cell._tc.tcPr.tcW.set(qn('w:w'),str(width)); cell._tc.tcPr.tcW.set(qn('w:type'),'dxa')
            set_cell_margins(cell); set_cell_border(cell)

def set_run(run, size=11, bold=False, colour=INK):
    run.font.name='Arial'; run._element.rPr.rFonts.set(qn('w:ascii'),'Arial'); run._element.rPr.rFonts.set(qn('w:hAnsi'),'Arial'); run.font.size=Pt(size); run.bold=bold; run.font.color.rgb=RGBColor.from_string(colour)

def p_text(p, text='', size=11, bold=False, colour=INK, align=None, before=0, after=4):
    p.paragraph_format.space_before=Pt(before); p.paragraph_format.space_after=Pt(after); p.paragraph_format.line_spacing=1.15
    if align is not None: p.alignment=align
    r=p.add_run(text); set_run(r,size,bold,colour); return p

def clear_cell(cell):
    p=cell.paragraphs[0]; p.clear(); return p

def label(cell, title, subtitle='', fill=PALE):
    set_cell_shading(cell,fill); p=clear_cell(cell); p_text(p,title,12,True,NAVY,after=2); 
    if subtitle: p_text(cell.add_paragraph(),subtitle,9,False,INK,after=0)

def blank_lines(cell, count=2, prefix=''):
    for i in range(count):
        p=cell.add_paragraph(); p.paragraph_format.space_after=Pt(7); r=p.add_run(prefix + '________________________________________________________________'); set_run(r,10,False,'667487')

def setup_doc():
    doc=Document(); sec=doc.sections[0]; sec.top_margin=Inches(.55); sec.bottom_margin=Inches(.55); sec.left_margin=Inches(.62); sec.right_margin=Inches(.62); sec.header_distance=Inches(.3); sec.footer_distance=Inches(.3)
    styles=doc.styles
    normal=styles['Normal']; normal.font.name='Arial'; normal._element.rPr.rFonts.set(qn('w:ascii'),'Arial'); normal._element.rPr.rFonts.set(qn('w:hAnsi'),'Arial'); normal.font.size=Pt(10.5); normal.paragraph_format.space_after=Pt(4); normal.paragraph_format.line_spacing=1.15
    for style,size,colour,before,after in [('Heading 1',18,NAVY,10,5),('Heading 2',14,NAVY,8,4),('Heading 3',12,ORANGE,6,3)]:
        s=styles[style]; s.font.name='Arial'; s._element.rPr.rFonts.set(qn('w:ascii'),'Arial'); s._element.rPr.rFonts.set(qn('w:hAnsi'),'Arial'); s.font.size=Pt(size); s.font.color.rgb=RGBColor.from_string(colour); s.font.bold=True; s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after)
    footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.RIGHT; p_text(footer,'Book Week Song Challenge | Dingo bus mascot',8,False,'667487',after=0)
    return doc

def title(doc, level, subtitle):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(2); p.alignment=WD_ALIGN_PARAGRAPH.LEFT; r=p.add_run('BOOK WEEK SONG CHALLENGE'); set_run(r,10,True,ORANGE)
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(4); r=p.add_run(level); set_run(r,24,True,NAVY)
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(9); r=p.add_run(subtitle); set_run(r,11,False,INK)

def banner(doc, text):
    table=doc.add_table(rows=1,cols=1); fixed_width(table,[9360]); cell=table.cell(0,0); set_cell_shading(cell,NAVY); p=clear_cell(cell); p_text(p,text,12,True,WHITE,align=WD_ALIGN_PARAGRAPH.CENTER,after=0)

def add_upper_lyric_page(doc):
    title(doc,'Years 5-6: Songwriting Studio','Write for a real disco audience: clear image, strong rhythm, memorable repeat.')
    banner(doc,'CRAFT GOAL: Original words, purposeful repetition, one dingo shout-out, then one visible revision.')
    t=doc.add_table(rows=2,cols=3); fixed_width(t,[3120,3120,3120])
    for idx,(h,s) in enumerate([('IMAGE BANK','nouns, places, colours, sounds'),('ACTION BANK','verbs, problems, journeys'),('HOOK BANK','repeats, rhyme, contrast')]): label(t.cell(0,idx),h,s,[PALE,'E4F5F3','EEE8FA'][idx]); blank_lines(t.cell(1,idx),2)
    p=doc.add_paragraph(); p_text(p,'Draft the lyric',12,True,NAVY,after=4)
    t=doc.add_table(rows=4,cols=2); fixed_width(t,[1800,7560])
    rows=[('VERSE 1','Place the listener inside the story world.'),('CHORUS','A short hook. DINGO must be prominent.'),('VERSE 2','Action/change. Earn the final chorus.'),('CHORUS','Repeat the strongest idea.')]
    for r,(h,helptext) in enumerate(rows):
        label(t.cell(r,0),h,'',[PALE,GOLD,'E4F5F3',GOLD][r]); p=clear_cell(t.cell(r,1)); p_text(p,helptext,9,False,'667487',after=4); blank_lines(t.cell(r,1),3)
    p=doc.add_paragraph(); p_text(p,'Revision: Our line was ______________________. We changed it to ______________________ because ______________________.',10.5,True,NAVY,after=0)

def add_handover_page(doc):
    doc.add_page_break(); title(doc,'Music Producer Hand-in','Complete this page clearly. Photograph or deliver it by 2:30 pm.')
    banner(doc,'FINAL CHECK: original class words | chorus repeats | DINGO is named | one sound choice in every row')
    t=doc.add_table(rows=3,cols=2); fixed_width(t,[3120,6240])
    fields=[('Class / teacher',''),('Song title',''),('Dingo shout-out is in:','☐ Verse 1     ☐ Chorus     ☐ Verse 2')]
    for r,(h,answer) in enumerate(fields):
        label(t.cell(r,0),h,'',PALE); p=clear_cell(t.cell(r,1)); p_text(p,answer,11,False,INK,after=0); 
        if not answer: blank_lines(t.cell(r,1),1)
    p=doc.add_paragraph(); p_text(p,'Final lyric (write neatly or attach the completed lyric page)',12,True,NAVY,after=4)
    t=doc.add_table(rows=1,cols=1); fixed_width(t,[9360]); blank_lines(t.cell(0,0),10)
    p=doc.add_paragraph(); p_text(p,'Music choices',12,True,NAVY,after=4)
    t=doc.add_table(rows=3,cols=2); fixed_width(t,[2160,7200])
    choices=[('GENRE','☐ Pop   ☐ Rock   ☐ K-Pop   ☐ Hip-Hop   ☐ Dance/EDM   ☐ Musical Theatre   ☐ Country'),('SPEED','☐ Slow   ☐ Medium   ☐ Fast'),('MOOD','☐ Happy   ☐ Funny   ☐ Cool   ☐ Epic   ☐ Magical   ☐ Spooky')]
    for r,(h,opts) in enumerate(choices):
        label(t.cell(r,0),h,'',[PALE,'E4F5F3','EEE8FA'][r]); p=clear_cell(t.cell(r,1)); p_text(p,opts,10,False,INK,after=0)
    p=doc.add_paragraph(); p_text(p,'Teacher confirmation: ☐ Read word-for-word  ☐ Original wording  ☐ Safe and school-appropriate  ☐ Ready for production',10.5,True,NAVY,after=0)

def main():
    doc=setup_doc(); add_upper_lyric_page(doc); add_handover_page(doc); doc.save(OUT); print(OUT)

if __name__=='__main__': main()
