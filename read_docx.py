import sys
from docx import Document

def extract_text(file_path):
    print(f"--- Extracting {file_path} ---")
    try:
        doc = Document(file_path)
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                print(text)
        print("--- Tables ---")
        for table in doc.tables:
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip().replace('\n', ' ')
                    if cell_text and cell_text not in row_data:
                        row_data.append(cell_text)
                if row_data:
                    print(" | ".join(row_data))
            print("---")
    except Exception as e:
        print(f"Error reading {file_path}: {e}")

if __name__ == "__main__":
    for arg in sys.argv[1:]:
        extract_text(arg)
