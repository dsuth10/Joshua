import os
import sys
from docx import Document

students = [
    "Jack", "Jaycee", "Kianni", "Jordan", "Savannah", "Lucas", "Harry", 
    "Patrick", "Daina", "Lakota", "Sommer", "Willow", "Freya", "Carolyn", 
    "Elouise", "Maxwell", "Jacey", "Felicity"
]

def search_file(file_path):
    try:
        doc = Document(file_path)
        text = []
        for p in doc.paragraphs:
            text.append(p.text)
        for t in doc.tables:
            for r in t.rows:
                for c in r.cells:
                    text.append(cell_text := c.text)
        
        full_text = "\n".join(text)
        
        # Check if any student name is in the text
        found_students = [s for s in students if s.lower() in full_text.lower()]
        if found_students:
            # Check if there is a pattern like "Name:" or "Text:" that looks like a draft
            if "draft" in full_text.lower() or "first draft" in full_text.lower() or "report" in full_text.lower():
                print(f"MATCH in {file_path}: found students {found_students}")
                # Print first 200 chars
                print(full_text[:300].strip())
                print("---")
    except Exception as e:
        pass

def main():
    root_dir = r"c:\Users\dsuth\Documents\Joshua"
    for dirpath, dirnames, filenames in os.walk(root_dir):
        if ".git" in dirpath or "node_modules" in dirpath or ".gemini" in dirpath:
            continue
        for f in filenames:
            if f.endswith(".docx"):
                search_file(os.path.join(dirpath, f))

if __name__ == "__main__":
    main()
