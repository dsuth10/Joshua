import os
from docx import Document

def make_handout():
    input_path = "Units/English/English_Unit_2/Lesson_Plans/Lesson_25.2/Lesson_25.2_Worksheet.docx"
    output_path = "Units/English/English_Unit_2/Lesson_Plans/Lesson_25.2/Lesson_25.2_Student_Handout.docx"
    
    doc = Document(input_path)
    
    # 1. Remove all tables (which contain the reading text and boundary/definition tables)
    for table in list(doc.tables):
        tbl_el = table._element
        tbl_el.getparent().remove(tbl_el)
    
    # 2. Remove orphan heading paragraphs that referred to the deleted tables/figures
    paragraphs_to_remove = [
        "Plate Boundary Types (Field Comparison Table)",
        "Figure 1: Labeled Cross-Section of an Earthquake"
    ]
    
    for para in list(doc.paragraphs):
        text = para.text.strip()
        if text in paragraphs_to_remove:
            para_el = para._element
            para_el.getparent().remove(para_el)
            
    # Save the modified document
    doc.save(output_path)
    print(f"Successfully generated student handout at: {os.path.abspath(output_path)}")

if __name__ == "__main__":
    make_handout()
