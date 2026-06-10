import os
from docx import Document

terms = ["melbourne", "1934", "yarra"]

def search_file(file_path):
    try:
        doc = Document(file_path)
        text = []
        for p in doc.paragraphs:
            text.append(p.text)
        for t in doc.tables:
            for r in t.rows:
                for c in r.cells:
                    text.append(c.text)
        
        full_text = "\n".join(text).lower()
        
        for term in terms:
            if term in full_text:
                print(f"MATCH in {file_path}: found '{term}'")
                # find a snippet around the term
                idx = full_text.find(term)
                start = max(0, idx - 100)
                end = min(len(full_text), idx + 200)
                print(f"Snippet: ... {full_text[start:end].strip()} ...")
                print("---")
    except Exception as e:
        print(f"Error reading {file_path}: {e}")

def main():
    root_dir = r"c:\Users\dsuth\Documents\Joshua\Units\English\English_Unit_2"
    for dirpath, dirnames, filenames in os.walk(root_dir):
        for f in filenames:
            if f.endswith(".docx"):
                search_file(os.path.join(dirpath, f))

if __name__ == "__main__":
    main()
