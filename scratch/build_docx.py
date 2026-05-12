import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os

base_dir = r"c:\Users\dsuth\Documents\Joshua\Units\Science\Unit 2 Natural disasters\Lessons_06_08\Lesson_06"

def create_assessment():
    doc = docx.Document()
    
    # Title
    p = doc.add_paragraph()
    r = p.add_run("Lesson 6 Assessment: Introducing Tropical Cyclones")
    r.bold = True
    r.font.size = Pt(16)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    questions = [
        ("1. What is the minimum sea surface temperature required for a tropical cyclone to form?", ["A. 20.5°C", "B. 24.0°C", "C. 26.5°C", "D. 30.0°C"], "C"),
        ("2. At what wind speed is a tropical low classified as a Category 1 tropical cyclone?", ["A. 40 km/h", "B. 62 km/h", "C. 100 km/h", "D. 120 km/h"], "B"),
        ("3. In which hemisphere do tropical cyclones rotate in a clockwise direction?", ["A. Northern Hemisphere", "B. Southern Hemisphere", "C. Both Hemispheres", "D. Neither Hemisphere"], "B"),
        ("4. What is the term used for a tropical cyclone in the Atlantic and Northeast Pacific?", ["A. Typhoon", "B. Tornado", "C. Hurricane", "D. Storm Surge"], "C"),
        ("5. Which of these is a major difference between a cyclone and a tornado?", ["A. Tornadoes form over water; cyclones form over land.", "B. Cyclones are much larger and last longer than tornadoes.", "C. Tornadoes only happen in Australia.", "D. There is no difference."], "B"),
        ("6. What is the name of the calm centre of a tropical cyclone?", ["A. The Eye Wall", "B. The Eye", "C. The Rain Band", "D. The Core"], "B"),
        ("7. Why do tropical cyclones weaken when they move over land?", ["A. Because they hit mountains.", "B. Because they lose their warm ocean water energy source.", "C. Because it is too cold on land.", "D. Because tornadoes destroy them."], "B"),
        ("8. Which part of the cyclone generally has the strongest, most destructive winds?", ["A. The Eye", "B. The outer edges", "C. The Eye Wall", "D. The top of the clouds"], "C"),
        ("9. What is a 'storm surge'?", ["A. A sudden increase in wind speed.", "B. A rise in sea level caused by a cyclone.", "C. A type of electrical storm.", "D. A dry wind from the desert."], "B"),
        ("10. Which government agency is responsible for issuing cyclone warnings in Australia?", ["A. Bureau of Meteorology (BOM)", "B. Department of Education", "C. The Police", "D. NASA"], "A")
    ]
    
    for q_text, options, ans in questions:
        doc.add_paragraph(q_text)
        for opt in options:
            p = doc.add_paragraph(opt)
            p.paragraph_format.left_indent = Inches(0.5)
        p_ans = doc.add_paragraph(f"ANSWER: {ans}")
        p_ans.bold = True
        p_pt = doc.add_paragraph("POINT: 1")
        p_pt.bold = True
        doc.add_paragraph()
        
    doc.save(os.path.join(base_dir, "Lesson_06_Assessment.docx"))

def create_handout():
    doc = docx.Document()
    
    # Title
    p = doc.add_paragraph()
    r = p.add_run("Lesson 6 Handout: Introducing Tropical Cyclones")
    r.bold = True
    r.font.size = Pt(16)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_heading("Section A: Watch & Record", level=2)
    doc.add_paragraph("As you watch the 'Tropical Cyclones' video, record the following information:")
    doc.add_paragraph("1. Minimum sea surface temperature needed: ________________________")
    doc.add_paragraph("2. What happens to warm, moist air? ________________________")
    doc.add_paragraph("3. Wind speed to be classified as a Category 1 cyclone: ________________________")
    doc.add_paragraph("4. One major impact on Earth's surface mentioned: ________________________")
    
    doc.add_heading("Section B: Terminology Comparison Table", level=2)
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Term'
    hdr_cells[1].text = 'Where Used'
    hdr_cells[2].text = 'Hemisphere'
    for i, term in enumerate(["Cyclone", "Hurricane", "Typhoon", "Tornado"], start=1):
        table.rows[i].cells[0].text = term
    doc.add_paragraph()
    
    doc.add_heading("Section C: BTN Key Facts", level=2)
    doc.add_paragraph("1. Fact 1 about how cyclones affect communities: ________________________________________________")
    doc.add_paragraph("2. Fact 2 about how cyclones affect communities: ________________________________________________")
    doc.add_paragraph("3. Fact 3 about how cyclones affect communities: ________________________________________________")
    doc.add_paragraph("4. One question the video raised for you: ________________________________________________")
    
    doc.add_heading("Section D: Eye Spy", level=2)
    doc.add_paragraph("1. What is the job of emergency services during a cyclone?")
    doc.add_paragraph("________________________________________________________________________________________________")
    
    doc.add_heading("Section E: Reading Task", level=2)
    doc.add_paragraph("Circle the text you read:   Cyclone Yasi     /      Cyclone Larry")
    doc.add_paragraph("1. ________________________________________________________________________________")
    doc.add_paragraph("2. ________________________________________________________________________________")
    doc.add_paragraph("3. ________________________________________________________________________________")
    doc.add_paragraph("4. ________________________________________________________________________________")
    p = doc.add_paragraph()
    r = p.add_run("Text Highlighting Task: ")
    r.bold = True
    p.add_run("Highlight 1 cause, 1 effect, 1 statistic, and 1 surprising fact in your reading text.")
    
    doc.add_heading("Section F: Case Study - Cyclone Katryn Warning Message", level=2)
    doc.add_paragraph("1. What is the name of the severe tropical cyclone?")
    doc.add_paragraph("________________________________________________________________________________")
    doc.add_paragraph("2. How fast is the severe tropical cyclone moving?")
    doc.add_paragraph("________________________________________________________________________________")
    doc.add_paragraph("3. Where could very destructive winds be expected?")
    doc.add_paragraph("________________________________________________________________________________")
    doc.add_paragraph("4. How do the people of Marrow Island need to prepare? Explain.")
    doc.add_paragraph("________________________________________________________________________________")
    doc.add_paragraph("5. Where could you access more information to ensure the forecast is accurate?")
    doc.add_paragraph("________________________________________________________________________________")
    
    doc.add_heading("Section G: Larry Wind Gusts", level=2)
    doc.add_paragraph("1. Look at the destruction in the video. What category do you think this cyclone was? Use the category table to justify your answer.")
    doc.add_paragraph("________________________________________________________________________________")
    doc.add_paragraph("________________________________________________________________________________")
    
    doc.add_heading("Wrap-Up: Vocabulary Consolidation", level=2)
    doc.add_paragraph("Write a definition in your own words for:")
    doc.add_paragraph("1. cyclone: ____________________________________________________________________")
    doc.add_paragraph("2. eye: ________________________________________________________________________")
    doc.add_paragraph("3. storm surge: ________________________________________________________________")
    doc.add_paragraph("4. low pressure system: ________________________________________________________")
    doc.add_paragraph("5. category: ___________________________________________________________________")

    doc.save(os.path.join(base_dir, "Lesson_06_Handout.docx"))

def create_handout_support():
    doc = docx.Document()
    
    p = doc.add_paragraph()
    r = p.add_run("Lesson 6 Handout (Support): Introducing Tropical Cyclones")
    r.bold = True
    r.font.size = Pt(16)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_heading("Section A: Watch & Record", level=2)
    p = doc.add_paragraph()
    r = p.add_run("Word Bank: 26.5°C, rises, 62 km/h, flooding")
    r.bold = True
    doc.add_paragraph("1. Minimum sea surface temperature needed: ________________________")
    doc.add_paragraph("2. Warm, moist air ________________________ into the atmosphere.")
    doc.add_paragraph("3. Wind speed to be classified as a Category 1 cyclone: ________________________")
    doc.add_paragraph("4. One major impact on Earth's surface is ________________________.")
    
    doc.add_heading("Section B: Terminology Comparison Table", level=2)
    doc.add_paragraph("Draw a line to match the term to where it is used:")
    doc.add_paragraph("Cyclone              Americas, Caribbean")
    doc.add_paragraph("Hurricane            Australia, South Asia")
    doc.add_paragraph("Typhoon              Over land (anywhere)")
    doc.add_paragraph("Tornado              Western Pacific, Asia")
    
    doc.add_heading("Section C: BTN Key Facts", level=2)
    doc.add_paragraph("Circle 3 facts about how cyclones affect communities:")
    doc.add_paragraph("  - They cause floods\n  - They destroy houses\n  - They bring snow\n  - They cut off electricity")
    doc.add_paragraph("Write one question you have: ________________________________________________")
    
    doc.add_heading("Section D: Eye Spy", level=2)
    doc.add_paragraph("Finish the sentence:")
    doc.add_paragraph("The job of emergency services during a cyclone is to help keep people s _ _ _ and clean up d _ _ _ _ _.")
    
    doc.add_heading("Section E: Reading Task", level=2)
    doc.add_paragraph("Circle the text you read:   Cyclone Yasi     /      Cyclone Larry")
    doc.add_paragraph("Circle the correct answers for your text.")
    doc.add_paragraph("1. Was the cyclone a Category 5? ( Yes / No )")
    doc.add_paragraph("2. Did it cause a lot of damage? ( Yes / No )")
    p = doc.add_paragraph()
    r = p.add_run("Text Highlighting Task: ")
    r.bold = True
    p.add_run("Highlight 1 statistic (a number) in your reading text.")
    
    doc.add_heading("Section F: Case Study - Cyclone Katryn Warning Message", level=2)
    doc.add_paragraph("1. What is the name of the cyclone?   K _ _ _ _ _")
    doc.add_paragraph("2. Where could very destructive winds be expected?  M _ _ _ _ _  Island")
    doc.add_paragraph("3. What alert colour is Marrow Island on? ( Red / Yellow / Blue )")
    
    doc.add_heading("Section G: Larry Wind Gusts", level=2)
    doc.add_paragraph("Finish the sentence:")
    doc.add_paragraph("I think this was a Category 5 cyclone because the winds were strong enough to ________________________________________________.")
    
    doc.add_heading("Wrap-Up: Vocabulary Consolidation", level=2)
    doc.add_paragraph("Draw a line to match the word to its meaning:")
    doc.add_paragraph("Eye                  The way we measure how strong a cyclone is (1 to 5)")
    doc.add_paragraph("Storm surge          A giant spinning storm over the ocean")
    doc.add_paragraph("Category             A rise in the sea level caused by the storm")
    doc.add_paragraph("Cyclone              The calm, clear centre of the storm")

    doc.save(os.path.join(base_dir, "Lesson_06_Handout_Support.docx"))

if __name__ == "__main__":
    create_assessment()
    create_handout()
    create_handout_support()
    print("Successfully generated all DOCX files.")
