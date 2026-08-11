import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'''
        <w:tcMar {nsdecls("w")}>
            <w:top w:w="{top}" w:type="dxa"/>
            <w:bottom w:w="{bottom}" w:type="dxa"/>
            <w:left w:w="{left}" w:type="dxa"/>
            <w:right w:w="{right}" w:type="dxa"/>
        </w:tcMar>
    ''')
    tcPr.append(tcMar)

def build_forms_assessment():
    output_dir = os.path.dirname(os.path.abspath(__file__))
    output_path = os.path.join(output_dir, "Lesson_1_06_1_08_1_09_Forms_Assessment.docx")
    
    doc = docx.Document()
    
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Arial'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0x1A, 0x23, 0x32)
    
    NAVY = RGBColor(0x0B, 0x4F, 0x6C)
    DARK_GRAY = RGBColor(0x33, 0x4E, 0x68)
    GREEN = RGBColor(0x27, 0xAE, 0x60)

    # Title Header
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    t_run = title.add_run("Signpost Mathematics — Lessons 1:06, 1:08 & 1:09 Formative Assessment")
    t_run.font.size = Pt(16)
    t_run.font.bold = True
    t_run.font.color.rgb = NAVY

    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(16)
    s_run = sub.add_run("Mixed Numbers, Percentages, and Conversions (Microsoft Forms Import Format)")
    s_run.font.size = Pt(11)
    s_run.font.color.rgb = DARK_GRAY

    questions = [
        {
            "num": 1,
            "q": "What is 237/100 expressed as a mixed number and a decimal?",
            "options": ["A. 2 37/100 or 2.37", "B. 23 7/10 or 23.7", "C. 2 37/10 or 2.37", "D. 237.0"],
            "ans": "ANS: A"
        },
        {
            "num": 2,
            "q": "Which of the following represents 6 8/100 as a decimal?",
            "options": ["A. 6.8", "B. 6.08", "C. 68.0", "D. 0.68"],
            "ans": "ANS: B"
        },
        {
            "num": 3,
            "q": "A hundred square has 45 squares shaded. What percentage of the square is NOT coloured?",
            "options": ["A. 45%", "B. 50%", "C. 55%", "D. 65%"],
            "ans": "ANS: C"
        },
        {
            "num": 4,
            "q": "Which fraction, decimal, and percentage are all equivalent?",
            "options": ["A. 3/10 = 0.03 = 30%", "B. 75/100 = 0.75 = 75%", "C. 1/4 = 0.40 = 25%", "D. 5/100 = 0.50 = 5%"],
            "ans": "ANS: B"
        },
        {
            "num": 5,
            "q": "What is 1.05 expressed as a mixed number?",
            "options": ["A. 1 5/100", "B. 1 5/10", "C. 10 5/100", "D. 15/100"],
            "ans": "ANS: A"
        },
        {
            "num": 6,
            "q": "If a phone screen displays 64% battery charge, what decimal of the battery charge remains?",
            "options": ["A. 6.4", "B. 0.64", "C. 0.064", "D. 64.0"],
            "ans": "ANS: B"
        },
        {
            "num": 7,
            "q": "What is 5/100 expressed as a decimal?",
            "options": ["A. 0.5", "B. 0.05", "C. 5.0", "D. 0.55"],
            "ans": "ANS: B"
        },
        {
            "num": 8,
            "q": "A student colours 20% of a 100-grid square. How many individual small squares did they colour?",
            "options": ["A. 2 squares", "B. 20 squares", "C. 5 squares", "D. 80 squares"],
            "ans": "ANS: B"
        },
        {
            "num": 9,
            "q": "Convert 0.43 into a fraction over 100 and a percentage.",
            "options": ["A. 43/100 and 43%", "B. 43/10 and 4.3%", "C. 4/3 and 430%", "D. 43/1000 and 43%"],
            "ans": "ANS: A"
        },
        {
            "num": 10,
            "q": "What is 9.42 expressed as a mixed number?",
            "options": ["A. 9 42/100", "B. 9 4/2", "C. 94 2/10", "D. 9 42/10"],
            "ans": "ANS: A"
        },
        {
            "num": 11,
            "q": "Which expression shows how to calculate the decimal value for 37/100 on a calculator?",
            "options": ["A. 37 + 100", "B. 37 × 100", "C. 37 ÷ 100", "D. 100 ÷ 37"],
            "ans": "ANS: C"
        },
        {
            "num": 12,
            "q": "What percentage is equivalent to 3 quarters (3/4)?",
            "options": ["A. 25%", "B. 50%", "C. 75%", "D. 100%"],
            "ans": "ANS: C"
        }
    ]

    for item in questions:
        qp = doc.add_paragraph()
        qp.paragraph_format.space_before = Pt(8)
        qp.paragraph_format.space_after = Pt(2)
        q_run = qp.add_run(f"{item['num']}. {item['q']}")
        q_run.bold = True
        
        for opt in item['options']:
            op = doc.add_paragraph()
            op.paragraph_format.left_indent = Inches(0.25)
            op.paragraph_format.space_after = Pt(1)
            op.add_run(opt)
            
        ap = doc.add_paragraph()
        ap.paragraph_format.left_indent = Inches(0.25)
        ap.paragraph_format.space_after = Pt(6)
        a_run = ap.add_run(item['ans'])
        a_run.bold = True
        a_run.font.color.rgb = GREEN

    doc.save(output_path)
    print(f"Forms Assessment successfully generated at: {output_path}")

if __name__ == "__main__":
    build_forms_assessment()
