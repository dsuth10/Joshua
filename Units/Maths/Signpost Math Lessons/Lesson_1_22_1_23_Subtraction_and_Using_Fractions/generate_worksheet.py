import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_shading(cell, color_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=120, bottom=120, left=180, right=180):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'''
        <w:tcMar {nsdecls("w")}>
            <w:top w:w="{top}" w:type="dxa"/>
            <w:bottom w:w="{bottom}" w:type="dxa"/>
            <w:left w:w="{left}" w:type="dxa"/>
            <w:right w:w="{right}" w:type="dxa"/>
        </w:tcMar>
    ''')
    tcPr.append(tcMar)

def set_table_borders(table, color="CCCCCC", sz="4", val="single"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(f'''
        <w:tblBorders {nsdecls("w")}>
            <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:left w:val="none"/>
            <w:right w:val="none"/>
            <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:insideV w:val="none"/>
        </w:tblBorders>
    ''')
    tblPr.append(borders)

def build_worksheet():
    output_dir = os.path.dirname(os.path.abspath(__file__))
    output_path = os.path.join(output_dir, "Lesson_1_22_1_23_Worksheet.docx")
    
    doc = docx.Document()
    
    # 1-inch margins on all sides
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Arial'
    normal_style.font.size = Pt(10.5)
    normal_style.font.color.rgb = RGBColor(0x1A, 0x23, 0x32)
    
    NAVY = RGBColor(0x0B, 0x4F, 0x6C)
    DARK_TEAL = RGBColor(0x1B, 0x9A, 0xAA)
    DARK_GRAY = RGBColor(0x33, 0x4E, 0x68)
    CORAL = RGBColor(0xF0, 0x71, 0x67)
    
    # Title
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(2)
    t_run = title.add_run("Signpost Mathematics — Lessons 1:22 & 1:23")
    t_run.font.size = Pt(18)
    t_run.font.bold = True
    t_run.font.color.rgb = NAVY

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    sub_run = subtitle.add_run("Subtraction from Whole Numbers & Operations with Related Denominators")
    sub_run.font.size = Pt(13)
    sub_run.font.bold = True
    sub_run.font.color.rgb = DARK_GRAY

    # Student metadata table
    info_tbl = doc.add_table(rows=1, cols=3)
    info_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_tbl.autofit = False
    widths = [3400, 3000, 2960]
    headers = ["Name: _______________________", "Date: _______________", "Class: ____________"]
    for i, cell in enumerate(info_tbl.rows[0].cells):
        cell.width = widths[i]
        p = cell.paragraphs[0]
        p.add_run(headers[i]).font.bold = True
        set_cell_margins(cell, top=80, bottom=80, left=100, right=100)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # Concept Summary Box
    concept_box = doc.add_table(rows=1, cols=1)
    concept_box.alignment = WD_TABLE_ALIGNMENT.CENTER
    concept_box.autofit = False
    concept_box.rows[0].cells[0].width = 9360
    c_cell = concept_box.rows[0].cells[0]
    set_cell_shading(c_cell, "EBF8FF")
    set_cell_margins(c_cell, top=140, bottom=140, left=200, right=200)
    
    cp = c_cell.paragraphs[0]
    c_head = cp.add_run("Key Rules & Concept Summary:\n")
    c_head.bold = True
    c_head.font.size = Pt(11)
    c_head.font.color.rgb = NAVY
    
    r1 = cp.add_run("1. Subtraction from Whole Numbers: ")
    r1.bold = True
    cp.add_run("Rename 1 whole into fractional parts (e.g. 1 = 3/3, 8/8, 10/10). For 3 - 1/4, write 3 as 2 4/4, then subtract 1/4 to get 2 3/4.\n")
    
    r2 = cp.add_run("2. Equivalent Fractions: ")
    r2.bold = True
    cp.add_run("Fractions with different numerators and denominators can have equal value (e.g. 1/2 = 2/4 = 4/8 = 5/10).\n")
    
    r3 = cp.add_run("3. Operations with Related Denominators: ")
    r3.bold = True
    cp.add_run("Convert one fraction so both denominators match before adding or subtracting (e.g. 1/4 + 3/8 = 2/8 + 3/8 = 5/8).")

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Helper function for section headings
    def add_section_header(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text)
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = NAVY
        return p

    # Helper function for 4-column problem tables
    def add_problem_grid_4col(problems):
        table = doc.add_table(rows=(len(problems) + 3) // 4, cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        col_w = 2340
        set_table_borders(table)
        
        for idx, (label, expr) in enumerate(problems):
            r = idx // 4
            c = idx % 4
            cell = table.rows[r].cells[c]
            cell.width = col_w
            set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
            p = cell.paragraphs[0]
            l_run = p.add_run(f"{label}  ")
            l_run.bold = True
            l_run.font.color.rgb = DARK_TEAL
            p.add_run(f"{expr} = ")
            box_run = p.add_run("[         ]")
            box_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            
        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # Section 1: Page 1:22 Q1
    add_section_header("1. Complete, writing answers as whole numbers or mixed numerals:")
    p1_data = [
        ("a", "7/8 + 1/8"), ("b", "4/6 + 2/6"), ("c", "2/3 + 1/3"), ("d", "3/4 + 1/4"),
        ("e", "1 1/6 + 5/6"), ("f", "2 5/8 + 3/8"), ("g", "1 7/10 + 3/10"), ("h", "2 3/5 + 2/5"),
        ("i", "1 3/5 - 3/5"), ("j", "2 7/10 - 4/10"), ("k", "3 3/5 - 1/5"), ("l", "2 7/12 - 3/12")
    ]
    add_problem_grid_4col(p1_data)

    # Section 2: Page 1:22 Q2
    add_section_header("2. Complete (Subtraction from 1 Whole):")
    p2_data = [
        ("a", "1 - 1/6"), ("b", "1 - 1/10"), ("c", "1 - 1/8"), ("d", "1 - 1/12"),
        ("e", "1 - 1/5"), ("f", "1 - 3/4"), ("g", "1 - 7/10"), ("h", "1 - 2/5"),
        ("i", "1 - 2/3"), ("j", "1 - 5/6"), ("k", "1 - 3/8"), ("l", "1 - 5/12")
    ]
    add_problem_grid_4col(p2_data)

    # Section 3: Page 1:22 Q3
    add_section_header("3. Complete (Subtraction from Whole Numbers greater than 1):")
    p3_data = [
        ("a", "3 - 1/2"), ("b", "2 - 1/3"), ("c", "2 - 1/6"),
        ("d", "3 - 1/5"), ("e", "4 - 1/10"), ("f", "4 - 1/12"),
        ("g", "3 - 1/8"), ("h", "4 - 1/4"), ("i", "2 - 3/4")
    ]
    add_problem_grid_4col(p3_data)

    # Section 4: Page 1:22 Q4 & Q5 (Fraction Wall)
    add_section_header("4. Equivalent Fractions & True/False Analysis:")
    
    p_fw = doc.add_paragraph()
    p_fw.add_run("Use the fraction relationships (1 whole = 3/3 = 6/6 = 12/12) to complete:\n").italic = True
    
    # Q4 table
    q4_tbl = doc.add_table(rows=2, cols=3)
    q4_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    q4_tbl.autofit = False
    set_table_borders(q4_tbl)
    q4_data = [
        ("a", "2/12 = [      ]"), ("b", "10/12 = [      ]"), ("c", "4/6 = [      ]"),
        ("d", "2/3 = [      ]"), ("e", "8/12 = [      ]"), ("f", "3/6 = [      ]")
    ]
    for idx, (label, expr) in enumerate(q4_data):
        r = idx // 3
        c = idx % 3
        cell = q4_tbl.rows[r].cells[c]
        cell.width = 3120
        set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
        p = cell.paragraphs[0]
        l_run = p.add_run(f"{label}.  ")
        l_run.bold = True
        l_run.font.color.rgb = DARK_TEAL
        p.add_run(expr)
        
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # Q5 table (True/False)
    p_tf_title = doc.add_paragraph()
    p_tf_title.add_run("State whether each statement is True or False:").bold = True
    
    q5_tbl = doc.add_table(rows=2, cols=2)
    q5_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    q5_tbl.autofit = False
    set_table_borders(q5_tbl)
    q5_data = [
        ("a", "1/3 = 4/12", "[ True / False ]"),
        ("b", "8/12 = 4/6", "[ True / False ]"),
        ("c", "2/3 = 6/12", "[ True / False ]"),
        ("d", "10/12 = 5/6", "[ True / False ]")
    ]
    for idx, (label, expr, ans) in enumerate(q5_data):
        r = idx // 2
        c = idx % 2
        cell = q5_tbl.rows[r].cells[c]
        cell.width = 4680
        set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
        p = cell.paragraphs[0]
        l_run = p.add_run(f"{label}.  ")
        l_run.bold = True
        l_run.font.color.rgb = DARK_TEAL
        p.add_run(f"{expr}   --->   {ans}")

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # Section 5: Page 1:23 Q1 (Ordering)
    add_section_header("5. Order these numbers from smallest to largest:")
    p_ord_intro = doc.add_paragraph()
    p_ord_intro.add_run("Hint: Convert all fractions to a common denominator first!").italic = True
    
    ord_tbl = doc.add_table(rows=3, cols=2)
    ord_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    ord_tbl.autofit = False
    set_table_borders(ord_tbl)
    ord_data = [
        ("a", "3/4,  1/2,  1/4", "___________________________________"),
        ("b", "3/4,  1 1/4,  5/8,  1/2", "___________________________________"),
        ("c", "1/4,  1,  9/8,  1/8", "___________________________________"),
        ("d", "3/8,  1/2,  3/4,  1/8", "___________________________________"),
        ("e", "7/10,  1 1/10,  1/5,  1/2", "___________________________________"),
        ("f", "6/10,  1/5,  3/2,  1/10", "___________________________________")
    ]
    for idx, (label, expr, ans_line) in enumerate(ord_data):
        r = idx // 2
        c = idx % 2
        cell = ord_tbl.rows[r].cells[c]
        cell.width = 4680
        set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
        p = cell.paragraphs[0]
        l_run = p.add_run(f"{label}.  ")
        l_run.bold = True
        l_run.font.color.rgb = DARK_TEAL
        p.add_run(f"{expr}\nAnswer: {ans_line}")

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # Section 6: Page 1:23 Q2 (Addition with Related Denominators)
    add_section_header("6. Make the denominators the same before adding:")
    p2_add_data = [
        ("a", "1/4 + 3/8 = [   ]/8 + 3/8"), ("b", "1/8 + 1/2 = 1/8 + [   ]/8"),
        ("c", "1/2 + 1/4"), ("d", "3/4 + 1/2"),
        ("e", "3/4 + 1/8"), ("f", "1/10 + 1/5"),
        ("g", "1/8 + 1/4"), ("h", "3/5 + 3/10"),
        ("i", "1/4 + 7/8"), ("j", "3/5 + 1/10"),
        ("k", "3/10 + 2/5"), ("l", "4/5 + 7/10")
    ]
    add_problem_grid_4col(p2_add_data)

    # Section 7: Page 1:23 Q3 (Subtraction with Related Denominators)
    add_section_header("7. Make the denominators the same before subtracting:")
    p3_sub_data = [
        ("a", "3/8 - 1/8"), ("b", "5/8 - 1/4"), ("c", "7/8 - 1/2"), ("d", "5/8 - 1/2"),
        ("e", "9/10 - 1/2"), ("f", "7/10 - 2/5"), ("g", "3/8 - 1/4"), ("h", "7/10 - 1/2"),
        ("i", "6/8 - 1/2"), ("j", "9/8 - 3/4"), ("k", "9/10 - 1/5"), ("l", "3/4 - 3/8"),
        ("m", "7/8 - 3/4"), ("n", "4/5 - 6/10"), ("o", "3/5 - 1/10"), ("p", "11/10 - 1/2")
    ]
    add_problem_grid_4col(p3_sub_data)

    # Section 8: Reasoning & Written Communication (Page 1:22 Q6)
    add_section_header("8. Mathematical Communication & Reasoning:")
    p_reason = doc.add_paragraph()
    p_r_run = p_reason.add_run("Explain how converting fractions to equivalent fractions helps when adding or subtracting fractions with different denominators. Give one example with diagrams or calculations:\n")
    p_r_run.italic = True
    
    # Reasoning box
    r_box = doc.add_table(rows=1, cols=1)
    r_box.alignment = WD_TABLE_ALIGNMENT.CENTER
    r_box.autofit = False
    r_box.rows[0].cells[0].width = 9360
    r_cell = r_box.rows[0].cells[0]
    set_cell_margins(r_cell, top=140, bottom=140, left=180, right=180)
    set_table_borders(r_box, color="0B4F6C", sz="8")
    
    rp = r_cell.paragraphs[0]
    rp.add_run("\n\n\n\n\n\n")

    doc.save(output_path)
    print(f"Successfully generated worksheet DOCX at: {output_path}")

if __name__ == "__main__":
    build_worksheet()
