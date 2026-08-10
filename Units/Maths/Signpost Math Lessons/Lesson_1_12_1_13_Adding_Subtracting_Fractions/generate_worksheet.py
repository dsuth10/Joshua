import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

def create_element(name):
    return OxmlElement(name)

def set_cell_shading(cell, color_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=120, bottom=120, left=180, right=180):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'''
        <w:tcMar {nsdecls("w")}>
            <w:top w:w="{top}" w:type="dxa"/>
            <w:bottom w:w="{bottom}" w:type="dxa"/>
            <w:left w:w="{left}" w:type="dxa"/>
            <w:right w:w="{right}" w:type="dxa"/>
        </w:tcMar>
    ''')
    tcPr.append(tcMar)

def set_table_borders(table, color="CCCCCC", sz="4", val="single"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(f'''
        <w:tblBorders {nsdecls("w")}>
            <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:left w:val="none"/>
            <w:right w:val="none"/>
            <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:insideV w:val="none"/>
        </w:tblBorders>
    ''')
    tblPr.append(borders)

def build_worksheet():
    doc = docx.Document()
    
    # Page setup - 1 inch margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # Styles setup
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Arial'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0x1A, 0x23, 0x32)
    
    NAVY = RGBColor(0x0B, 0x4F, 0x6C)
    RED = RGBColor(0xC0, 0x39, 0x2B)
    GREEN = RGBColor(0x27, 0xAE, 0x60)
    DARK_GRAY = RGBColor(0x33, 0x4E, 0x68)
    LIGHT_BG = "F4F7F9"
    BLUE_BG = "EBF8FF"

    # Header Title
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(2)
    run_title = title.add_run("Signpost Mathematics — Lessons 1:12 & 1:13")
    run_title.font.size = Pt(20)
    run_title.font.bold = True
    run_title.font.color.rgb = NAVY

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    run_sub = subtitle.add_run("Addition & Subtraction of Fractions (Same Denominator)")
    run_sub.font.size = Pt(14)
    run_sub.font.bold = True
    run_sub.font.color.rgb = DARK_GRAY

    # Student Info Table
    info_tbl = doc.add_table(rows=1, cols=3)
    info_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_tbl.autofit = False
    info_widths = [3500, 3000, 2860]
    info_headers = ["Name: _______________________", "Date: _______________", "Class: ____________"]
    
    for i, cell in enumerate(info_tbl.rows[0].cells):
        cell.width = info_widths[i]
        cell.paragraphs[0].add_run(info_headers[i]).font.bold = True
        set_cell_margins(cell, top=80, bottom=80, left=100, right=100)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # Concept Box
    concept_box = doc.add_table(rows=1, cols=1)
    concept_box.alignment = WD_TABLE_ALIGNMENT.CENTER
    concept_box.autofit = False
    concept_box.rows[0].cells[0].width = 9360
    set_cell_margins(concept_box.rows[0].cells[0], top=140, bottom=140, left=200, right=200)
    set_cell_shading(concept_box.rows[0].cells[0], LIGHT_BG)
    
    cp = concept_box.rows[0].cells[0].paragraphs[0]
    cp.add_run("💡 GOLDEN RULE: ").font.bold = True
    cp.runs[0].font.color.rgb = NAVY
    cp.add_run("When we add or subtract fractions, the denominators MUST be the same. Add or subtract the numerators only while keeping the denominator constant!\n").font.size = Pt(11)
    
    cp2 = concept_box.rows[0].cells[0].add_paragraph()
    cp2.add_run("Addition Example:  3/8 + 2/8 = 5/8        |        Subtraction Example:  4/5 - 1/5 = 3/5").font.bold = True
    cp2.runs[0].font.color.rgb = NAVY

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # --- SECTION 1: ADDITION OF FRACTIONS ---
    h1 = doc.add_paragraph()
    h1.paragraph_format.space_before = Pt(10)
    h1.paragraph_format.space_after = Pt(4)
    run_h1 = h1.add_run("Section 1: Addition of Fractions (Signpost 1:12)")
    run_h1.font.size = Pt(14)
    run_h1.font.bold = True
    run_h1.font.color.rgb = NAVY

    # Q1 Area Grids Addition
    p_q1 = doc.add_paragraph()
    p_q1.paragraph_format.space_after = Pt(4)
    p_q1.add_run("1. Add these fractions. Colour part of the last grid to match your answer:").font.bold = True

    add_grid_tbl = doc.add_table(rows=2, cols=2)
    add_grid_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_grid_tbl.autofit = False
    set_table_borders(add_grid_tbl)

    grid_widths = [4680, 4680]
    grid_data = [
        ["a)  [Grid: 1/8 shaded]  +  [Grid: 3/8 shaded]  =  [  ____ / 8  ]",
         "b)  [Grid: 2/8 shaded]  +  [Grid: 5/8 shaded]  =  [  ____ / 8  ]"],
        ["c)  [Grid: 2/5 shaded]  +  [Grid: 1/5 shaded]  =  [  ____ / 5  ]",
         "d)  [Grid: 1/5 shaded]  +  [Grid: 3/5 shaded]  =  [  ____ / 5  ]"]
    ]

    for r_idx, row in enumerate(add_grid_tbl.rows):
        for c_idx, cell in enumerate(row.cells):
            cell.width = grid_widths[c_idx]
            set_cell_margins(cell, top=100, bottom=100, left=140, right=140)
            cell.paragraphs[0].add_run(grid_data[r_idx][c_idx])

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # Q2 Card Addition (Tenths & Eighths)
    p_q2 = doc.add_paragraph()
    p_q2.paragraph_format.space_after = Pt(4)
    p_q2.add_run("2. Use the fraction cards to find the answers (Remember addition is commutative: 4/10 + 6/10 = 6/10 + 4/10):").font.bold = True

    card_add_tbl = doc.add_table(rows=3, cols=4)
    card_add_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    card_add_tbl.autofit = False
    set_table_borders(card_add_tbl)

    col4_widths = [2340, 2340, 2340, 2340]
    card_add_data = [
        ["a)  1/10 + 5/10 = ____", "b)  3/10 + 4/10 = ____", "c)  2/10 + 7/10 = ____", "d)  6/10 + 3/10 = ____"],
        ["e)  5/10 + 2/10 = ____", "f)  4/10 + 5/10 = ____", "g)  7/10 + 1/10 = ____", "h)  1/10 + 8/10 = ____"],
        ["i)  2/8 + 6/8 = ____",   "j)  7/8 + 4/8 = ____",   "k)  5/8 + 5/8 = ____",   "l)  3/8 + 6/8 = ____"]
    ]

    for r_idx, row in enumerate(card_add_tbl.rows):
        for c_idx, cell in enumerate(row.cells):
            cell.width = col4_widths[c_idx]
            set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
            cell.paragraphs[0].add_run(card_add_data[r_idx][c_idx])

    # --- SECTION 2: SUBTRACTION OF FRACTIONS ---
    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    h2 = doc.add_paragraph()
    h2.paragraph_format.space_before = Pt(10)
    h2.paragraph_format.space_after = Pt(4)
    run_h2 = h2.add_run("Section 2: Subtraction of Fractions (Signpost 1:13)")
    run_h2.font.size = Pt(14)
    run_h2.font.bold = True
    run_h2.font.color.rgb = NAVY

    # Q3 Subtraction Takeaway Grids
    p_q3 = doc.add_paragraph()
    p_q3.paragraph_format.space_after = Pt(4)
    p_q3.add_run("3. Subtract these fractions. Cross out part of the grid to show takeaway:").font.bold = True

    sub_grid_tbl = doc.add_table(rows=3, cols=2)
    sub_grid_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    sub_grid_tbl.autofit = False
    set_table_borders(sub_grid_tbl)

    sub_grid_data = [
        ["a)  5/8  -  1/8  =  [  ____ / 8  ]", "b)  7/8  -  2/8  =  [  ____ / 8  ]"],
        ["c)  6/10  -  3/10  =  [  ____ / 10  ]", "d)  9/10  -  4/10  =  [  ____ / 10  ]"],
        ["e)  5/6  -  2/6  =  [  ____ / 6  ]", "f)  4/6  -  1/6  =  [  ____ / 6  ]"]
    ]

    for r_idx, row in enumerate(sub_grid_tbl.rows):
        for c_idx, cell in enumerate(row.cells):
            cell.width = grid_widths[c_idx]
            set_cell_margins(cell, top=100, bottom=100, left=140, right=140)
            cell.paragraphs[0].add_run(sub_grid_data[r_idx][c_idx])

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # Q4 Card Subtraction
    p_q4 = doc.add_paragraph()
    p_q4.paragraph_format.space_after = Pt(4)
    p_q4.add_run("4. Use fraction cards to solve these subtractions:").font.bold = True

    card_sub_tbl = doc.add_table(rows=3, cols=4)
    card_sub_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    card_sub_tbl.autofit = False
    set_table_borders(card_sub_tbl)

    card_sub_data = [
        ["a)  3/10 - 1/10 = ____", "b)  5/10 - 3/10 = ____", "c)  7/10 - 2/10 = ____", "d)  9/10 - 5/10 = ____"],
        ["e)  8/10 - 7/10 = ____", "f)  6/10 - 1/10 = ____", "g)  4/10 - 3/10 = ____", "h)  7/10 - 4/10 = ____"],
        ["i)  7/8 - 3/8 = ____",   "j)  5/8 - 3/8 = ____",   "k)  6/8 - 2/8 = ____",   "l)  4/8 - 1/8 = ____"]
    ]

    for r_idx, row in enumerate(card_sub_tbl.rows):
        for c_idx, cell in enumerate(row.cells):
            cell.width = col4_widths[c_idx]
            set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
            cell.paragraphs[0].add_run(card_sub_data[r_idx][c_idx])

    # --- SECTION 3: INEQUALITY COMPARISONS ---
    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    h3 = doc.add_paragraph()
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(4)
    run_h3 = h3.add_run("Section 3: Inequality Comparisons (< and >)")
    run_h3.font.size = Pt(14)
    run_h3.font.bold = True
    run_h3.font.color.rgb = NAVY

    p_ineq_tip = doc.add_paragraph()
    p_ineq_tip.paragraph_format.space_after = Pt(4)
    p_ineq_tip.add_run("Remember: For < and >, the arrow points to the smaller number (< means less than, > means greater than, e.g. 10 < 30 and 100 > 40).").font.italic = True

    p_q5 = doc.add_paragraph()
    p_q5.paragraph_format.space_after = Pt(4)
    p_q5.add_run("5. Answer True or False for each comparison:").font.bold = True

    tf_tbl = doc.add_table(rows=2, cols=4)
    tf_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tf_tbl.autofit = False
    set_table_borders(tf_tbl)

    tf_data = [
        ["a)  6/8 > 3/8  [      ]", "b)  7/8 < 6/8  [      ]", "c)  8/8 < 5/8  [      ]", "d)  5/8 > 9/8  [      ]"],
        ["e)  9/10 > 6/10 [      ]", "f)  5/10 < 9/10 [      ]", "g)  13/10 > 11/10 [   ]", "h)  8/10 > 12/10 [   ]"]
    ]

    for r_idx, row in enumerate(tf_tbl.rows):
        for c_idx, cell in enumerate(row.cells):
            cell.width = col4_widths[c_idx]
            set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
            cell.paragraphs[0].add_run(tf_data[r_idx][c_idx])

    # --- SECTION 4: APPLIED WORD PROBLEMS ---
    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    h4 = doc.add_paragraph()
    h4.paragraph_format.space_before = Pt(10)
    h4.paragraph_format.space_after = Pt(4)
    run_h4 = h4.add_run("Section 4: Applied Real-World Problems")
    run_h4.font.size = Pt(14)
    run_h4.font.bold = True
    run_h4.font.color.rgb = NAVY

    word_tbl = doc.add_table(rows=4, cols=1)
    word_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    word_tbl.autofit = False
    set_table_borders(word_tbl)

    word_data = [
        "a) Sharon ate 3/12 of a block of chocolate and Franco ate 5/12 of the same block.\n    How much of the block did they eat altogether?\n    Equation: _______________________________    Answer: ___________",
        "b) Ron used 4/10 of the paper roll. Eva used 5/10.\n    How much paper is left from 1 whole roll (10/10)?\n    Equation: _______________________________    Answer: ___________",
        "c) Tim ate 1/8 of a block of chocolate. How much of the block was left?\n    Equation: _______________________________    Answer: ___________",
        "d) Mum gave me 3/5 of her money. What fraction of her money did she keep?\n    Equation: _______________________________    Answer: ___________"
    ]

    for r_idx, row in enumerate(word_tbl.rows):
        cell = row.cells[0]
        cell.width = 9360
        set_cell_margins(cell, top=120, bottom=120, left=160, right=160)
        cell.paragraphs[0].add_run(word_data[r_idx])

    # --- SECTION 5: ANSWER KEY ---
    doc.add_page_break()

    h_ans = doc.add_paragraph()
    h_ans.paragraph_format.space_before = Pt(10)
    h_ans.paragraph_format.space_after = Pt(6)
    run_ans = h_ans.add_run("Answer Key — Teacher & Self-Correction Guide")
    run_ans.font.size = Pt(16)
    run_ans.font.bold = True
    run_ans.font.color.rgb = NAVY

    ans_p = doc.add_paragraph()
    ans_p.paragraph_format.space_after = Pt(6)
    ans_p.add_run("Section 1: Addition of Fractions\n").font.bold = True
    ans_p.add_run("1. a) 4/8    b) 7/8    c) 3/5    d) 4/5\n")
    ans_p.add_run("2. a) 6/10   b) 7/10   c) 9/10   d) 9/10   e) 7/10   f) 9/10   g) 8/10   h) 9/10\n   i) 8/8 (or 1)   j) 11/8 (or 1 3/8)   k) 10/8 (or 1 2/8)   l) 9/8 (or 1 1/8)\n\n")

    ans_p.add_run("Section 2: Subtraction of Fractions\n").font.bold = True
    ans_p.add_run("3. a) 4/8    b) 5/8    c) 3/10   d) 5/10   e) 3/6    f) 3/6\n")
    ans_p.add_run("4. a) 2/10   b) 2/10   c) 5/10   d) 4/10   e) 1/10   f) 5/10   g) 1/10   h) 3/10\n   i) 4/8    j) 2/8    k) 4/8    l) 3/8\n\n")

    ans_p.add_run("Section 3: Inequality Comparisons\n").font.bold = True
    ans_p.add_run("5. a) True   b) False   c) False   d) False   e) True   f) True   g) True   h) False\n\n")

    ans_p.add_run("Section 4: Applied Word Problems\n").font.bold = True
    ans_p.add_run("a) 3/12 + 5/12 = 8/12 of the block eaten altogether.\n")
    ans_p.add_run("b) 10/10 - 4/10 - 5/10 = 1/10 of paper left.\n")
    ans_p.add_run("c) 8/8 - 1/8 = 7/8 of chocolate left.\n")
    ans_p.add_run("d) 5/5 - 3/5 = 2/5 of money kept.\n")

    worksheet_path = os.path.join(os.path.dirname(__file__), "Lesson_1_12_1_13_Worksheet.docx")
    doc.save(worksheet_path)
    print(f"Saved worksheet to {worksheet_path}")

def build_forms_assessment():
    doc = docx.Document()
    
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    title = doc.add_paragraph()
    run = title.add_run("Signpost Math 1:12 & 1:13 — Addition & Subtraction Assessment")
    run.font.size = Pt(16)
    run.font.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    questions = [
        ("1. What is 3/8 + 2/8?", ["A. 5/16", "B. 5/8", "C. 1/8", "D. 6/8"], "B"),
        ("2. What is 4/5 - 1/5?", ["A. 3/5", "B. 3/10", "C. 5/5", "D. 4/0"], "A"),
        ("3. When adding or subtracting fractions with the same denominator, what do you do?", ["A. Add both numerators and denominators", "B. Multiply the top and bottom", "C. Add or subtract numerators and keep the denominator the same", "D. Convert all fractions to whole numbers"], "C"),
        ("4. What is 1/10 + 5/10?", ["A. 6/20", "B. 4/10", "C. 6/10", "D. 5/100"], "C"),
        ("5. What is 7/8 - 3/8?", ["A. 10/8", "B. 4/8", "C. 4/0", "D. 4/16"], "B"),
        ("6. Which inequality statement is TRUE?", ["A. 7/8 < 6/8", "B. 6/8 > 3/8", "C. 5/8 > 9/8", "D. 8/8 < 5/8"], "B"),
        ("7. Sharon ate 3/12 of a chocolate bar and Franco ate 5/12. How much did they eat altogether?", ["A. 8/24", "B. 2/12", "C. 8/12", "D. 15/12"], "C"),
        ("8. Ron used 4/10 of the paper and Eva used 5/10. How much paper is left from 1 whole pack (10/10)?", ["A. 9/10", "B. 1/10", "C. 2/10", "D. 0/10"], "B"),
        ("9. Tim ate 1/8 of a block of chocolate. What fraction of the block was left?", ["A. 7/8", "B. 1/8", "C. 8/8", "D. 0/8"], "A"),
        ("10. What is 7/8 + 4/8?", ["A. 11/16", "B. 3/8", "C. 11/8", "D. 28/8"], "C")
    ]

    for q_text, opts, ans in questions:
        p_q = doc.add_paragraph()
        p_q.add_run(q_text).font.bold = True
        for opt in opts:
            p_opt = doc.add_paragraph()
            p_opt.add_run(opt)
        p_ans = doc.add_paragraph()
        p_ans.add_run(f"ANS: {ans}").font.bold = True
        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    forms_path = os.path.join(os.path.dirname(__file__), "Lesson_1_12_1_13_Forms_Assessment.docx")
    doc.save(forms_path)
    print(f"Saved forms assessment to {forms_path}")

if __name__ == "__main__":
    build_worksheet()
    build_forms_assessment()
