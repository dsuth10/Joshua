import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def build_comprehensive_forms_assessment():
    output_dir = os.path.dirname(os.path.abspath(__file__))
    output_path = os.path.join(output_dir, "Comprehensive_Fractions_Forms_Assessment.docx")
    
    doc = docx.Document()
    
    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # Styles setup
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Arial'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0x1A, 0x23, 0x32)
    
    NAVY = RGBColor(0x0B, 0x4F, 0x6C)
    DARK_GRAY = RGBColor(0x33, 0x4E, 0x68)
    GREEN = RGBColor(0x27, 0xAE, 0x60)
    GRAY = RGBColor(0x55, 0x55, 0x55)

    # Document Header
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    t_run = title.add_run("Signpost Mathematics — Comprehensive Fractions Assessment")
    t_run.font.size = Pt(16)
    t_run.font.bold = True
    t_run.font.color.rgb = NAVY

    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(16)
    s_run = sub.add_run("Formative & Summative Evaluation Deck (Microsoft Forms Import Format — 35 Questions)")
    s_run.font.size = Pt(11)
    s_run.font.color.rgb = DARK_GRAY

    # Question Database
    sections_data = [
        {
            "title": "Section 1: Mixed Numbers, Decimals & Percentages (Lessons 1:06, 1:08 & 1:09)",
            "questions": [
                {
                    "num": 1,
                    "q": "What is 237/100 expressed as a mixed number and a decimal?",
                    "options": ["A. 2 37/100 or 2.37", "B. 23 7/10 or 23.7", "C. 2 37/10 or 2.37", "D. 237.0"],
                    "ans": "ANSWER: A",
                    "point": "POINT: 1"
                },
                {
                    "num": 2,
                    "q": "Which of the following represents 6 8/100 as a decimal?",
                    "options": ["A. 6.8", "B. 6.08", "C. 68.0", "D. 0.68"],
                    "ans": "ANSWER: B",
                    "point": "POINT: 1"
                },
                {
                    "num": 3,
                    "q": "A hundred-square grid has 45 squares shaded. What percentage of the square grid is NOT shaded?",
                    "options": ["A. 45%", "B. 50%", "C. 55%", "D. 65%"],
                    "ans": "ANSWER: C",
                    "point": "POINT: 1"
                },
                {
                    "num": 4,
                    "q": "Which fraction, decimal, and percentage set are all equivalent?",
                    "options": ["A. 3/10 = 0.03 = 30%", "B. 75/100 = 0.75 = 75%", "C. 1/4 = 0.40 = 25%", "D. 5/100 = 0.50 = 5%"],
                    "ans": "ANSWER: B",
                    "point": "POINT: 1"
                },
                {
                    "num": 5,
                    "q": "What is 1.05 expressed as a mixed number in fraction form over 100?",
                    "options": ["A. 1 5/100", "B. 1 5/10", "C. 10 5/100", "D. 15/100"],
                    "ans": "ANSWER: A",
                    "point": "POINT: 1"
                },
                {
                    "num": 6,
                    "q": "If a phone screen displays 64% battery charge, what decimal of the total battery charge remains?",
                    "options": ["A. 6.4", "B. 0.64", "C. 0.064", "D. 64.0"],
                    "ans": "ANSWER: B",
                    "point": "POINT: 1"
                },
                {
                    "num": 7,
                    "q": "Which mathematical operation shows how to convert 43/100 into a decimal on a calculator?",
                    "options": ["A. 43 + 100", "B. 43 × 100", "C. 43 ÷ 100", "D. 100 ÷ 43"],
                    "ans": "ANSWER: C",
                    "point": "POINT: 1"
                }
            ]
        },
        {
            "title": "Section 2: Improper Fractions & Mixed Numbers (Lesson 1:11)",
            "questions": [
                {
                    "num": 8,
                    "q": "Which of the following is an improper fraction?",
                    "options": ["A. 3/4", "B. 1 1/2", "C. 7/5", "D. 2/9"],
                    "ans": "ANSWER: C",
                    "point": "POINT: 1"
                },
                {
                    "num": 9,
                    "q": "Convert 3 2/5 into an improper fraction.",
                    "options": ["A. 15/5", "B. 17/5", "C. 11/5", "D. 6/5"],
                    "ans": "ANSWER: B",
                    "point": "POINT: 1"
                },
                {
                    "num": 10,
                    "q": "Convert 14/3 into a mixed number.",
                    "options": ["A. 4 2/3", "B. 3 2/3", "C. 4 1/3", "D. 14 1/3"],
                    "ans": "ANSWER: A",
                    "point": "POINT: 1"
                },
                {
                    "num": 11,
                    "q": "If pizzas are cut into quarters and you have 9 slices in total, how many whole pizzas and extra quarters do you have?",
                    "options": ["A. 2 whole pizzas and 1/4 pizza", "B. 9 whole pizzas", "C. 1 whole pizza and 5/4 pizza", "D. 3 whole pizzas and 1/4 pizza"],
                    "ans": "ANSWER: A",
                    "point": "POINT: 1"
                },
                {
                    "num": 12,
                    "q": "A baking recipe requires 2 3/4 cups of flour. How many quarter-cups of flour is this in total?",
                    "options": ["A. 8 quarter-cups", "B. 11 quarter-cups", "C. 9 quarter-cups", "D. 6 quarter-cups"],
                    "ans": "ANSWER: B",
                    "point": "POINT: 1"
                },
                {
                    "num": 13,
                    "q": "Which mixed number is equal to 22/4?",
                    "options": ["A. 5 2/4", "B. 4 2/4", "C. 5 1/4", "D. 22 1/4"],
                    "ans": "ANSWER: A",
                    "point": "POINT: 1"
                },
                {
                    "num": 14,
                    "q": "A student says that 13/5 is equal to 2 3/5. Is this student correct or incorrect?",
                    "options": [
                        "A. Correct, because 13 divided by 5 is 2 with a remainder of 3",
                        "B. Incorrect, it should be 3 2/5",
                        "C. Incorrect, it should be 2 1/5",
                        "D. Incorrect, improper fractions cannot be written as mixed numbers"
                    ],
                    "ans": "ANSWER: A",
                    "point": "POINT: 1"
                }
            ]
        },
        {
            "title": "Section 3: Adding & Subtracting Fractions with Same Denominators (Lessons 1:12 & 1:13)",
            "questions": [
                {
                    "num": 15,
                    "q": "Calculate 3/8 + 2/8.",
                    "options": ["A. 5/16", "B. 5/8", "C. 1/8", "D. 6/8"],
                    "ans": "ANSWER: B",
                    "point": "POINT: 1"
                },
                {
                    "num": 16,
                    "q": "Calculate 7/10 - 4/10.",
                    "options": ["A. 3/10", "B. 3/0", "C. 11/10", "D. 3/20"],
                    "ans": "ANSWER: A",
                    "point": "POINT: 1"
                },
                {
                    "num": 17,
                    "q": "What is the Golden Rule when adding or subtracting fractions with the same denominator?",
                    "options": [
                        "A. Add or subtract both the top and bottom numbers",
                        "B. Multiply the denominators together",
                        "C. Keep the denominator the same and add or subtract the numerators",
                        "D. Convert all denominators to 100"
                    ],
                    "ans": "ANSWER: C",
                    "point": "POINT: 1"
                },
                {
                    "num": 18,
                    "q": "Which inequality comparison statement is TRUE?",
                    "options": ["A. 2/8 + 3/8 > 6/8", "B. 7/10 - 2/10 > 4/10", "C. 1/5 + 2/5 < 3/5", "D. 9/12 - 5/12 < 2/12"],
                    "ans": "ANSWER: B",
                    "point": "POINT: 1"
                },
                {
                    "num": 19,
                    "q": "Sharon ate 3/12 of a chocolate block and Franco ate 5/12 of the same block. What fraction did they eat combined?",
                    "options": ["A. 8/24", "B. 2/12", "C. 8/12", "D. 15/12"],
                    "ans": "ANSWER: C",
                    "point": "POINT: 1"
                },
                {
                    "num": 20,
                    "q": "Ron had 10/10 of a block of poster paper. He used 4/10 for a poster and 3/10 for a card. What fraction of the paper block remains?",
                    "options": ["A. 7/10", "B. 3/10", "C. 1/10", "D. 4/10"],
                    "ans": "ANSWER: B",
                    "point": "POINT: 1"
                },
                {
                    "num": 21,
                    "q": "Evaluate the expression: 9/12 - 2/12 - 3/12.",
                    "options": ["A. 4/12", "B. 4/36", "C. 14/12", "D. 5/12"],
                    "ans": "ANSWER: A",
                    "point": "POINT: 1"
                }
            ]
        },
        {
            "title": "Section 4: Equivalent Fractions (Lesson 1:17 & 1:19)",
            "questions": [
                {
                    "num": 22,
                    "q": "A rectangular strip is split into 3 vertical columns with 2 columns shaded (2/3). If a horizontal cut doubles the number of rows to 2, what fraction of the grid is now shaded?",
                    "options": ["A. 2/6", "B. 4/6", "C. 3/6", "D. 4/3"],
                    "ans": "ANSWER: B",
                    "point": "POINT: 1"
                },
                {
                    "num": 23,
                    "q": "On aligned equal-length number lines, which fraction aligns directly below 1/3?",
                    "options": ["A. 2/6", "B. 3/6", "C. 2/9", "D. 4/9"],
                    "ans": "ANSWER: A",
                    "point": "POINT: 1"
                },
                {
                    "num": 24,
                    "q": "According to the Multiplication Golden Rule, what multiplier must be used on both top and bottom to convert 3/5 into 9/15?",
                    "options": ["A. Multiply by 2", "B. Multiply by 3", "C. Multiply by 5", "D. Multiply by 9"],
                    "ans": "ANSWER: B",
                    "point": "POINT: 1"
                },
                {
                    "num": 25,
                    "q": "Find the missing numerator in this equivalent fraction chain: 1/4 = 2/8 = ?/12 = 4/16.",
                    "options": ["A. 2", "B. 3", "C. 4", "D. 5"],
                    "ans": "ANSWER: B",
                    "point": "POINT: 1"
                },
                {
                    "num": 26,
                    "q": "Which fraction is equivalent to 8/12 when simplified by dividing the numerator and denominator by 4?",
                    "options": ["A. 1/3", "B. 2/3", "C. 3/4", "D. 4/6"],
                    "ans": "ANSWER: B",
                    "point": "POINT: 1"
                },
                {
                    "num": 27,
                    "q": "Sam claims that 1/3 is equivalent to 2/4 because he added 1 to both top and bottom. Why is Sam incorrect?",
                    "options": [
                        "A. Adding the same number to numerator and denominator changes the fraction's value; you must multiply or divide",
                        "B. 1/3 is equal to 3/4",
                        "C. Sam should have subtracted 1 from top and bottom",
                        "D. 2/4 is smaller than 1/3"
                    ],
                    "ans": "ANSWER: A",
                    "point": "POINT: 1"
                },
                {
                    "num": 28,
                    "q": "Which list places the fractions in correct order from smallest to largest?",
                    "options": [
                        "A. 3/4, 1/2, 3/8, 1/8",
                        "B. 1/8, 3/8, 1/2, 3/4",
                        "C. 1/2, 1/8, 3/8, 3/4",
                        "D. 1/8, 1/2, 3/8, 3/4"
                    ],
                    "ans": "ANSWER: B",
                    "point": "POINT: 1"
                }
            ]
        },
        {
            "title": "Section 5: Subtraction from Whole Numbers & Related Denominators (Lessons 1:22 & 1:23)",
            "questions": [
                {
                    "num": 29,
                    "q": "Calculate 1 - 3/8.",
                    "options": ["A. 2/8", "B. 5/8", "C. 3/8", "D. 1 3/8"],
                    "ans": "ANSWER: B",
                    "point": "POINT: 1"
                },
                {
                    "num": 30,
                    "q": "Calculate 4 - 1/5.",
                    "options": ["A. 3 4/5", "B. 3 1/5", "C. 4 4/5", "D. 3/5"],
                    "ans": "ANSWER: A",
                    "point": "POINT: 1"
                },
                {
                    "num": 31,
                    "q": "What is 1 7/10 + 3/10?",
                    "options": ["A. 1 10/10", "B. 2", "C. 1 4/10", "D. Both A and B are correct"],
                    "ans": "ANSWER: D",
                    "point": "POINT: 1"
                },
                {
                    "num": 32,
                    "q": "Calculate 1/4 + 3/8 by first converting 1/4 into eighths.",
                    "options": ["A. 4/12", "B. 5/8", "C. 4/8", "D. 1/2"],
                    "ans": "ANSWER: B",
                    "point": "POINT: 1"
                },
                {
                    "num": 33,
                    "q": "Calculate 5/8 - 1/4 by converting 1/4 into eighths.",
                    "options": ["A. 4/4", "B. 3/8", "C. 4/8", "D. 1/4"],
                    "ans": "ANSWER: B",
                    "point": "POINT: 1"
                },
                {
                    "num": 34,
                    "q": "Calculate 3/5 + 3/10 by converting 3/5 into tenths.",
                    "options": ["A. 6/15", "B. 9/10", "C. 6/10", "D. 3/5"],
                    "ans": "ANSWER: B",
                    "point": "POINT: 1"
                },
                {
                    "num": 35,
                    "q": "A container holds 2 litres of water. If children drink 3/4 of a litre during sports, how much water remains in the container?",
                    "options": ["A. 1 1/4 litres", "B. 1 3/4 litres", "C. 2/4 litres", "D. 1/4 litre"],
                    "ans": "ANSWER: A",
                    "point": "POINT: 1"
                }
            ]
        }
    ]

    for sec in sections_data:
        # Section Heading
        sec_p = doc.add_paragraph()
        sec_p.paragraph_format.space_before = Pt(16)
        sec_p.paragraph_format.space_after = Pt(8)
        sec_run = sec_p.add_run(sec["title"])
        sec_run.font.size = Pt(13)
        sec_run.font.bold = True
        sec_run.font.color.rgb = NAVY

        for item in sec["questions"]:
            # Question line
            qp = doc.add_paragraph()
            qp.paragraph_format.space_before = Pt(6)
            qp.paragraph_format.space_after = Pt(3)
            q_run = qp.add_run(f"{item['num']}. {item['q']}")
            q_run.font.bold = True
            q_run.font.size = Pt(11)
            q_run.font.color.rgb = NAVY
            
            # Options A, B, C, D
            for opt in item["options"]:
                op = doc.add_paragraph()
                op.paragraph_format.left_indent = Inches(0.25)
                op.paragraph_format.space_after = Pt(1.5)
                op.add_run(opt)
                
            # Answer line
            ap = doc.add_paragraph()
            ap.paragraph_format.left_indent = Inches(0.25)
            ap.paragraph_format.space_after = Pt(1.5)
            ans_run = ap.add_run(item["ans"])
            ans_run.font.bold = True
            ans_run.font.color.rgb = GREEN

            # Point line
            pp = doc.add_paragraph()
            pp.paragraph_format.left_indent = Inches(0.25)
            pp.paragraph_format.space_after = Pt(8)
            pt_run = pp.add_run(item["point"])
            pt_run.font.bold = True
            pt_run.font.color.rgb = GRAY

    doc.save(output_path)
    print(f"Successfully generated 35-Question MS Forms Assessment DOCX at: {output_path}")

if __name__ == "__main__":
    build_comprehensive_forms_assessment()
