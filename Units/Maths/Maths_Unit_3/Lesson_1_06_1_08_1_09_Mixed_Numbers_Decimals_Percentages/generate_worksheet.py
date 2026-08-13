import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_shading(cell, color_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=120, bottom=120, left=180, right=180):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'''
        <w:tcMar {nsdecls("w")}>
            <w:top w:w="{top}" w:type="dxa"/>
            <w:bottom w:w="{bottom}" w:type="dxa"/>
            <w:left w:w="{left}" w:type="dxa"/>
            <w:right w:w="{right}" w:type="dxa"/>
        </w:tcMar>
    ''')
    tcPr.append(tcMar)

def set_table_borders(table, color="CCCCCC", sz="4", val="single"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(f'''
        <w:tblBorders {nsdecls("w")}>
            <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:left w:val="none"/>
            <w:right w:val="none"/>
            <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:insideV w:val="none"/>
        </w:tblBorders>
    ''')
    tblPr.append(borders)

def build_worksheet():
    output_dir = os.path.dirname(os.path.abspath(__file__))
    output_path = os.path.join(output_dir, "Lesson_1_06_1_08_1_09_Worksheet.docx")
    
    doc = docx.Document()
    
    # 1-inch margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Arial'
    normal_style.font.size = Pt(10)
    normal_style.font.color.rgb = RGBColor(0x1A, 0x23, 0x32)
    
    NAVY = RGBColor(0x0B, 0x4F, 0x6C)
    DARK_TEAL = RGBColor(0x1B, 0x9A, 0xAA)
    DARK_GRAY = RGBColor(0x33, 0x4E, 0x68)
    CORAL = RGBColor(0xF0, 0x71, 0x67)
    
    # Header Title
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(2)
    t_run = title.add_run("Signpost Mathematics — Lessons 1:06, 1:08 & 1:09")
    t_run.font.size = Pt(18)
    t_run.font.bold = True
    t_run.font.color.rgb = NAVY

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(10)
    sub_run = subtitle.add_run("Mixed Numbers, Percentages, and Conversions")
    sub_run.font.size = Pt(12)
    sub_run.font.bold = True
    sub_run.font.color.rgb = DARK_GRAY

    # Student metadata table
    info_tbl = doc.add_table(rows=1, cols=3)
    info_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_tbl.autofit = False
    widths = [3400, 3000, 2960]
    headers = ["Name: _______________________", "Date: _______________", "Class: ____________"]
    for i, cell in enumerate(info_tbl.rows[0].cells):
        cell.width = widths[i]
        p = cell.paragraphs[0]
        p.add_run(headers[i]).font.bold = True
        set_cell_margins(cell, top=60, bottom=60, left=100, right=100)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # Concept Summary Box
    concept_box = doc.add_table(rows=1, cols=1)
    concept_box.alignment = WD_TABLE_ALIGNMENT.CENTER
    concept_box.autofit = False
    concept_box.rows[0].cells[0].width = 9360
    c_cell = concept_box.rows[0].cells[0]
    set_cell_shading(c_cell, "EBF8FF")
    set_cell_margins(c_cell, top=120, bottom=120, left=180, right=180)
    
    cp = c_cell.paragraphs[0]
    c_head = cp.add_run("Key Rules & Concept Summary:\n")
    c_head.bold = True
    c_head.font.size = Pt(10.5)
    c_head.font.color.rgb = NAVY
    
    r1 = cp.add_run("1. Mixed Numbers (1:06): ")
    r1.bold = True
    cp.add_run("100/100 = 1 whole. 237 hundredths = 2 37/100 or 2.37 (2 ones, 3 tenths, 7 hundredths). Use zero as place holder for single-digit hundredths: 6 8/100 = 6.08.\n")
    
    r2 = cp.add_run("2. Percentages (1:08): ")
    r2.bold = True
    cp.add_run("Per cent means 'out of 100'. 100% = 1 whole (1.00), 50% = 1/2 (0.50), 25% = 1/4 (0.25). Uncoloured % = 100% - Coloured %.\n")
    
    r3 = cp.add_run("3. Conversions (1:09): ")
    r3.bold = True
    cp.add_run("Fraction to Decimal: Divide numerator by denominator (e.g., 35/100 = 35 ÷ 100 = 0.35). 3-Way Equivalence: 43/100 = 0.43 = 43%.")

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # Section Helper
    def add_section_header(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        run.font.size = Pt(11.5)
        run.font.bold = True
        run.font.color.rgb = NAVY
        return p

    # SECTION 1: SIGNPOST 1:06
    add_section_header("Section 1: Signpost 1:06 — Mixed Numbers & Decimals")

    # Q1
    p_q1 = doc.add_paragraph()
    p_q1.add_run("1. Write the fraction and the decimal shown for each hundred square:").bold = True
    
    q1_tbl = doc.add_table(rows=1, cols=4)
    q1_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    q1_tbl.autofit = False
    set_table_borders(q1_tbl)
    q1_items = [
        ("a", "23 shaded squares", "[ 23/100 ] or [ 0.23 ]"),
        ("b", "35 shaded squares", "[ 35/100 ] or [ 0.35 ]"),
        ("c", "42 shaded squares", "[ 42/100 ] or [ 0.42 ]"),
        ("d", "17 shaded squares", "[ 17/100 ] or [ 0.17 ]")
    ]
    for idx, (label, desc, ans_template) in enumerate(q1_items):
        cell = q1_tbl.rows[0].cells[idx]
        cell.width = 2340
        set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
        p = cell.paragraphs[0]
        p.add_run(f"{label}. ").bold = True
        p.add_run(f"Grid ({desc}):\n")
        p.add_run("_____ or _____").font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # Q2
    p_q2 = doc.add_paragraph()
    p_q2.add_run("2. Write the mixed number and the decimal for each part (hundred squares > 1 whole):").bold = True
    
    q2_tbl = doc.add_table(rows=1, cols=4)
    q2_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    q2_tbl.autofit = False
    set_table_borders(q2_tbl)
    q2_items = [
        ("a", "2 wholes + 45 hundredths", "Mixed: _______  Dec: _______"),
        ("b", "2 wholes + 12 hundredths", "Mixed: _______  Dec: _______"),
        ("c", "2 wholes + 64 hundredths", "Mixed: _______  Dec: _______"),
        ("d", "2 wholes + 8 hundredths", "Mixed: _______  Dec: _______")
    ]
    for idx, (label, desc, ans) in enumerate(q2_items):
        cell = q2_tbl.rows[0].cells[idx]
        cell.width = 2340
        set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
        p = cell.paragraphs[0]
        p.add_run(f"{label}. ").bold = True
        p.add_run(f"({desc})\n")
        p.add_run(ans).font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # Q3 & Q4 
    p_q34 = doc.add_paragraph()
    p_q34.add_run("3. Write each mixed number as a decimal (Remember: Use zero as a place holder e.g. 6 8/100 = 6.08):").bold = True

    q3_tbl = doc.add_table(rows=2, cols=3)
    q3_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    q3_tbl.autofit = False
    set_table_borders(q3_tbl)
    q3_items = [
        ("a", "2 37/100", "= _______"), ("b", "1 76/100", "= _______"), ("c", "6 8/100", "= _______"),
        ("d", "9 95/100", "= _______"), ("e", "7 81/100", "= _______"), ("f", "5 3/100", "= _______")
    ]
    for idx, (label, expr, ans) in enumerate(q3_items):
        r, c = idx // 3, idx % 3
        cell = q3_tbl.rows[r].cells[c]
        cell.width = 3120
        set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
        p = cell.paragraphs[0]
        p.add_run(f"{label}.  {expr}  {ans}")

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    p_q4 = doc.add_paragraph()
    p_q4.add_run("4. Write each decimal as a mixed number:").bold = True
    
    q4_tbl = doc.add_table(rows=1, cols=3)
    q4_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    q4_tbl.autofit = False
    set_table_borders(q4_tbl)
    q4_items = [("a", "6.25", "= _______"), ("b", "3.04", "= _______"), ("c", "9.42", "= _______")]
    for idx, (label, expr, ans) in enumerate(q4_items):
        cell = q4_tbl.rows[0].cells[idx]
        cell.width = 3120
        set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
        p = cell.paragraphs[0]
        p.add_run(f"{label}.  {expr}  {ans}")

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    p_q56 = doc.add_paragraph()
    p_q56.add_run("5 & 6. Grid Shading Practice:").bold = True
    p_q56_sub = doc.add_paragraph()
    p_q56_sub.add_run("Q5. Shading prompt: Colour 2.75 of 3 hundred squares. Write mixed number: _______ (2 75/100)\n"
                      "Q6. Shading prompt: Colour 1.05 of 2 hundred squares. Write mixed number: _______ (1 5/100)").italic = True

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # SECTION 2: SIGNPOST 1:08
    add_section_header("Section 2: Signpost 1:08 — Percentages")

    p_q1_p8 = doc.add_paragraph()
    p_q1_p8.add_run("1. What percentage of each hundred square is coloured?").bold = True

    p8_q1_tbl = doc.add_table(rows=2, cols=4)
    p8_q1_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    p8_q1_tbl.autofit = False
    set_table_borders(p8_q1_tbl)
    p8_q1_items = [
        ("a", "50 squares shaded", "______ %"), ("b", "45 squares shaded", "______ %"),
        ("c", "72 squares shaded", "______ %"), ("d", "60 squares shaded", "______ %"),
        ("e", "30 squares shaded", "______ %"), ("f", "48 squares shaded", "______ %"),
        ("g", "85 squares shaded", "______ %"), ("h", "95 squares shaded", "______ %")
    ]
    for idx, (label, desc, ans) in enumerate(p8_q1_items):
        r, c = idx // 4, idx % 4
        cell = p8_q1_tbl.rows[r].cells[c]
        cell.width = 2340
        set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
        p = cell.paragraphs[0]
        p.add_run(f"{label}. ").bold = True
        p.add_run(f"{desc}: {ans}")

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    p_q2_p8 = doc.add_paragraph()
    p_q2_p8.add_run("2. What percentage of each square in Question 1 is NOT coloured?").bold = True

    p8_q2_tbl = doc.add_table(rows=2, cols=4)
    p8_q2_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    p8_q2_tbl.autofit = False
    set_table_borders(p8_q2_tbl)
    p8_q2_items = ["a", "b", "c", "d", "e", "f", "g", "h"]
    for idx, label in enumerate(p8_q2_items):
        r, c = idx // 4, idx % 4
        cell = p8_q2_tbl.rows[r].cells[c]
        cell.width = 2340
        set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
        p = cell.paragraphs[0]
        p.add_run(f"{label}.  [ _____ % ]")

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    p_q3_p8 = doc.add_paragraph()
    p_q3_p8.add_run("3. Complete these equivalents (Decimal | Fraction over 100 | Percentage):").bold = True

    p8_q3_tbl = doc.add_table(rows=4, cols=3)
    p8_q3_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    p8_q3_tbl.autofit = False
    set_table_borders(p8_q3_tbl)
    
    # Header row for table
    for i, title_text in enumerate(["Decimal", "Fraction (/100)", "Percentage (%)"]):
        cell = p8_q3_tbl.rows[0].cells[i]
        set_cell_shading(cell, "E6EEF2")
        cell.paragraphs[0].add_run(title_text).bold = True

    eq_rows = [
        ("a) 0.25", "______ / 100", "______ %"),
        ("b) 0.65", "______ / 100", "______ %"),
        ("c) 0.45", "______ / 100", "______ %"),
        ("d) 0.80", "______ / 100", "______ %"),
        ("e) 0.50", "______ / 100", "______ %"),
        ("f) 0.20", "______ / 100", "______ %"),
        ("g) 0.60", "______ / 100", "______ %"),
        ("h) 0.35", "______ / 100", "______ %"),
        ("i) 0.75", "______ / 100", "______ %"),
        ("j) 0.55", "______ / 100", "______ %"),
        ("k) 0.95", "______ / 100", "______ %")
    ]
    
    # We will expand table rows dynamically
    for item in eq_rows:
        row_cells = p8_q3_tbl.add_row().cells
        row_cells[0].paragraphs[0].add_run(item[0])
        row_cells[1].paragraphs[0].add_run(item[1])
        row_cells[2].paragraphs[0].add_run(item[2])
        for c in row_cells:
            set_cell_margins(c, top=60, bottom=60, left=100, right=100)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # SECTION 3: SIGNPOST 1:09
    add_section_header("Section 3: Signpost 1:09 — Using Percentages & Converting")

    p_q12_p9 = doc.add_paragraph()
    p_q12_p9.add_run("1 & 2. Percentage Shading & Complementary (Uncoloured) Calculation:").bold = True

    p9_q12_tbl = doc.add_table(rows=2, cols=4)
    p9_q12_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    p9_q12_tbl.autofit = False
    set_table_borders(p9_q12_tbl)
    p9_items = [
        ("a", "10%", "Squares to shade: 10", "Uncoloured %: _____"),
        ("b", "25%", "Squares to shade: 25", "Uncoloured %: _____"),
        ("c", "20%", "Squares to shade: 20", "Uncoloured %: _____"),
        ("d", "75%", "Squares to shade: 75", "Uncoloured %: _____"),
        ("e", "50%", "Squares to shade: 50", "Uncoloured %: _____"),
        ("f", "15%", "Squares to shade: 15", "Uncoloured %: _____"),
        ("g", "60%", "Squares to shade: 60", "Uncoloured %: _____"),
        ("h", "90%", "Squares to shade: 90", "Uncoloured %: _____")
    ]
    for idx, (label, pct, shade, uncol) in enumerate(p9_items):
        r, c = idx // 4, idx % 4
        cell = p9_q12_tbl.rows[r].cells[c]
        cell.width = 2340
        set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
        p = cell.paragraphs[0]
        p.add_run(f"{label}. Target: {pct}\n").bold = True
        p.add_run(f"{uncol}")

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    p_q3_p9 = doc.add_paragraph()
    p_q3_p9.add_run("3. Complete the Conversion Table (Fraction | Decimal | Percentage):").bold = True

    p9_q3_tbl = doc.add_table(rows=1, cols=3)
    p9_q3_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    p9_q3_tbl.autofit = False
    set_table_borders(p9_q3_tbl)
    for i, title_text in enumerate(["Fraction", "Decimal", "Percentage"]):
        cell = p9_q3_tbl.rows[0].cells[i]
        set_cell_shading(cell, "E6EEF2")
        cell.paragraphs[0].add_run(title_text).bold = True

    conv_rows = [
        ("a) 17/100", "0.______", "______ %"),
        ("b) 76/100", "0.______", "______ %"),
        ("c) 27/100", "0.______", "______ %"),
        ("d) 49/100", "0.______", "______ %"),
        ("e) 98/100", "0.______", "______ %"),
        ("f) 81/100", "0.______", "______ %"),
        ("g) 31/100", "0.______", "______ %"),
        ("h) 12/100", "0.______", "______ %"),
        ("i) 34/100", "0.______", "______ %"),
        ("j) 28/100", "0.______", "______ %"),
        ("k) 63/100", "0.______", "______ %"),
        ("l) 94/100", "0.______", "______ %")
    ]
    for item in conv_rows:
        r_cells = p9_q3_tbl.add_row().cells
        r_cells[0].paragraphs[0].add_run(item[0])
        r_cells[1].paragraphs[0].add_run(item[1])
        r_cells[2].paragraphs[0].add_run(item[2])
        for c in r_cells:
            set_cell_margins(c, top=60, bottom=60, left=100, right=100)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    p_q4_p9 = doc.add_paragraph()
    p_q4_p9.add_run("4. Convert fractions to decimals (Divide numerator by denominator):").bold = True

    p9_q4_tbl = doc.add_table(rows=3, cols=5)
    p9_q4_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    p9_q4_tbl.autofit = False
    set_table_borders(p9_q4_tbl)
    
    div_probs = [
        ("a", "65/100"), ("b", "15/100"), ("c", "95/100"), ("d", "45/100"), ("e", "75/100"),
        ("f", "25/100"), ("g", "5/100"),  ("h", "60/100"), ("i", "80/100"), ("j", "40/100"),
        ("k", "10/100"), ("l", "37/100"), ("m", "91/100"), ("n", "20/100"), ("o", "100/100")
    ]
    for idx, (label, frac) in enumerate(div_probs):
        r, c = idx // 5, idx % 5
        cell = p9_q4_tbl.rows[r].cells[c]
        cell.width = 1872
        set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
        p = cell.paragraphs[0]
        p.add_run(f"{label}. {frac}\n= _____").font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    doc.save(output_path)
    print(f"Worksheet successfully generated at: {output_path}")

if __name__ == "__main__":
    build_worksheet()
