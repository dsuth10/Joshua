import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

def create_element(name):
    return OxmlElement(name)

def set_cell_shading(cell, color_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=120, bottom=120, left=180, right=180):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'''
        <w:tcMar {nsdecls("w")}>
            <w:top w:w="{top}" w:type="dxa"/>
            <w:bottom w:w="{bottom}" w:type="dxa"/>
            <w:left w:w="{left}" w:type="dxa"/>
            <w:right w:w="{right}" w:type="dxa"/>
        </w:tcMar>
    ''')
    tcPr.append(tcMar)

def set_table_borders(table, color="CCCCCC", sz="4", val="single"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(f'''
        <w:tblBorders {nsdecls("w")}>
            <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:left w:val="none"/>
            <w:right w:val="none"/>
            <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:insideV w:val="none"/>
        </w:tblBorders>
    ''')
    tblPr.append(borders)

def build_worksheet():
    doc = docx.Document()
    
    # Page setup - 1 inch margins all around
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # Standard styles setup
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Arial'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0x1A, 0x23, 0x32) # Dark Navy body text
    
    # Colors
    NAVY = RGBColor(0x0B, 0x4F, 0x6C)
    ORANGE = RGBColor(0xF0, 0x71, 0x67)
    GREEN = RGBColor(0x2A, 0x9D, 0x8F)
    GREY_BG = "F7F9FB"
    NAVY_HEX = "0B4F6C"
    BLUE_BG = "EEF5F8"

    # --- DOCUMENT HEADER ---
    header_tbl = doc.add_table(rows=1, cols=2)
    header_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_tbl.autofit = False
    
    # Total printable width = 6.5 in = 9360 DXA
    col_widths = [6200, 3160]
    for row in header_tbl.rows:
        for idx, width in enumerate(col_widths):
            row.cells[idx].width = width
            set_cell_margins(row.cells[idx], top=60, bottom=60, left=60, right=60)
            
    cell_l = header_tbl.cell(0, 0)
    p_title = cell_l.paragraphs[0]
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(2)
    run_t = p_title.add_run("Equivalent Fractions")
    run_t.font.name = 'Arial'
    run_t.font.size = Pt(22)
    run_t.font.bold = True
    run_t.font.color.rgb = NAVY

    p_sub = cell_l.add_paragraph()
    p_sub.paragraph_format.space_before = Pt(0)
    p_sub.paragraph_format.space_after = Pt(0)
    run_s = p_sub.add_run("Year 4/5 Mathematics \u2022 Signpost Math 1:17 & 1:19")
    run_s.font.size = Pt(11)
    run_s.font.color.rgb = ORANGE
    run_s.font.bold = True

    cell_r = header_tbl.cell(0, 1)
    p_meta = cell_r.paragraphs[0]
    p_meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_meta.paragraph_format.space_after = Pt(2)
    p_meta.add_run("Name: ______________________\nDate: ______________________")

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # --- SECTION 1: ALIGNED NUMBER LINES MATCHING ---
    h1 = doc.add_paragraph()
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(4)
    run_h1 = h1.add_run("Section 1: Aligned Number Lines & Equivalent Fractions")
    run_h1.font.size = Pt(14)
    run_h1.font.bold = True
    run_h1.font.color.rgb = NAVY

    p_desc1 = doc.add_paragraph()
    p_desc1.paragraph_format.space_after = Pt(6)
    p_desc1.add_run("Use the aligned equal-length number lines below to find equivalent fractions and answer true or false.")

    # Visual Table Box representing 3 Stacked Number Lines (Thirds, Sixths, Ninths)
    nl_box = doc.add_table(rows=4, cols=1)
    nl_box.alignment = WD_TABLE_ALIGNMENT.CENTER
    nl_box.autofit = False
    for r in nl_box.rows:
        r.cells[0].width = 9360
        set_cell_margins(r.cells[0], top=100, bottom=100, left=180, right=180)
        set_cell_shading(r.cells[0], BLUE_BG)

    nl_box.cell(0, 0).paragraphs[0].add_run("Line 1 (Thirds):    |--- 0 ---|--- 1/3 ---|--- 2/3 ---|--- 1 ---|").font.bold = True
    nl_box.cell(1, 0).paragraphs[0].add_run("Line 2 (Sixths):    |--- 0 ---|--- 1/6 ---|--- 2/6 ---|--- 3/6 ---|--- 4/6 ---|--- 5/6 ---|--- 1 ---|").font.bold = True
    nl_box.cell(2, 0).paragraphs[0].add_run("Line 3 (Ninths):    |--- 0 ---|--- 1/9 ---|--- 2/9 ---|--- 3/9 ---|--- 4/9 ---|--- 5/9 ---|--- 6/9 ---|--- 7/9 ---|--- 8/9 ---|--- 1 ---|").font.bold = True
    nl_box.cell(3, 0).paragraphs[0].add_run("Rule: Equivalent fractions sit at the EXACT SAME VERTICAL POINT on aligned number lines!").font.italic = True

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # Q1 Grid Table (4 columns, 2 rows)
    p_q1 = doc.add_paragraph()
    p_q1.paragraph_format.space_after = Pt(4)
    p_q1.add_run("1. Use the number lines above to write an equivalent fraction for each:").font.bold = True

    q1_tbl = doc.add_table(rows=2, cols=4)
    q1_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    q1_tbl.autofit = False
    set_table_borders(q1_tbl)
    
    q1_widths = [2340, 2340, 2340, 2340]
    q1_data = [
        ["a)  1/3  =  [      ]", "b)  9/9  =  [      ]", "c)  3/9  =  [      ]", "d)  2/3  =  [      ]"],
        ["e)  2/6  =  [      ]", "f)  4/6  =  [      ]", "g)  6/6  =  [      ]", "h)  6/9  =  [      ]"]
    ]

    for r_idx, row in enumerate(q1_tbl.rows):
        for c_idx, cell in enumerate(row.cells):
            cell.width = q1_widths[c_idx]
            set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
            cell.paragraphs[0].add_run(q1_data[r_idx][c_idx])

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # Q2 True or False Table
    p_q2 = doc.add_paragraph()
    p_q2.paragraph_format.space_after = Pt(4)
    p_q2.add_run("2. Use the number lines above to answer true or false:").font.bold = True

    q2_tbl = doc.add_table(rows=2, cols=4)
    q2_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    q2_tbl.autofit = False
    set_table_borders(q2_tbl)

    q2_data = [
        ["a)  1/6 = 2/9  [      ]", "b)  2/3 = 6/9  [      ]", "c)  2/3 = 4/6  [      ]", "d)  6/9 = 4/6  [      ]"],
        ["e)  3/3 = 9/9  [      ]", "f)  3/6 = 5/9  [      ]", "g)  1/3 = 3/9  [      ]", "h)  5/6 = 2/3  [      ]"]
    ]

    for r_idx, row in enumerate(q2_tbl.rows):
        for c_idx, cell in enumerate(row.cells):
            cell.width = q1_widths[c_idx]
            set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
            cell.paragraphs[0].add_run(q2_data[r_idx][c_idx])

    # --- SECTION 2: SUBDIVIDED AREA BLOCKS ---
    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    h2 = doc.add_paragraph()
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(4)
    run_h2 = h2.add_run("Section 2: Subdivided Area Blocks (Cut & Color)")
    run_h2.font.size = Pt(14)
    run_h2.font.bold = True
    run_h2.font.color.rgb = NAVY

    p_desc2 = doc.add_paragraph()
    p_desc2.paragraph_format.space_after = Pt(6)
    p_desc2.add_run("When a block divided into vertical columns has a horizontal line cut through it, the number of parts doubles or triples! Fill in the missing equivalent fractions below:")

    block_tbl = doc.add_table(rows=3, cols=2)
    block_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    block_tbl.autofit = False
    set_table_borders(block_tbl)

    block_widths = [4680, 4680]
    block_data = [
        ["a) A rectangle is cut into 3 columns.\n    2 columns are shaded (2/3).\n    Draw 1 line horizontally across the middle.\n    New fraction shaded:  ____ / 6", 
         "b) A rectangle is cut into 4 columns.\n    3 columns are shaded (3/4).\n    Draw 1 line horizontally across the middle.\n    New fraction shaded:  ____ / 8"],
        ["c) A rectangle is cut into 5 columns.\n    2 columns are shaded (2/5).\n    Draw 1 line horizontally across the middle.\n    New fraction shaded:  ____ / 10", 
         "d) A rectangle is cut into 3 columns.\n    1 column is shaded (1/3).\n    Draw 2 horizontal lines across (3 rows).\n    New fraction shaded:  ____ / 9"],
        ["e) Complete the area rule:\n    Cutting 2/3 horizontally in half creates 4/6.\n    Does the total shaded area change? [ YES / NO ]",
         "f) Complete the area equivalence:\n    3/5 cut horizontally into 2 rows =  [  ____ / ____  ]"]
    ]

    for r_idx, row in enumerate(block_tbl.rows):
        for c_idx, cell in enumerate(row.cells):
            cell.width = block_widths[c_idx]
            set_cell_margins(cell, top=100, bottom=100, left=140, right=140)
            cell.paragraphs[0].add_run(block_data[r_idx][c_idx])

    # --- SECTION 3: THE MULTIPLICATION RULE ---
    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    h3 = doc.add_paragraph()
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(4)
    run_h3 = h3.add_run("Section 3: The Multiplication Rule (Signpost Math 1:19)")
    run_h3.font.size = Pt(14)
    run_h3.font.bold = True
    run_h3.font.color.rgb = NAVY

    p_desc3 = doc.add_paragraph()
    p_desc3.paragraph_format.space_after = Pt(6)
    p_desc3.add_run("Golden Rule: Multiply both the numerator (top) and denominator (bottom) by the SAME number to find an equivalent fraction.")

    mult_tbl = doc.add_table(rows=3, cols=3)
    mult_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    mult_tbl.autofit = False
    set_table_borders(mult_tbl)

    m_widths = [3120, 3120, 3120]
    m_data = [
        ["a) Multiply top & bottom by 2:\n    1/3 (x2)/(x2) =  [  ____ / ____  ]",
         "b) Multiply top & bottom by 3:\n    1/4 (x3)/(x3) =  [  ____ / ____  ]",
         "c) Multiply top & bottom by 5:\n    1/2 (x5)/(x5) =  [  ____ / ____  ]"],
        ["d) Multiply top & bottom by 2:\n    2/5 (x2)/(x2) =  [  ____ / ____  ]",
         "e) Multiply top & bottom by 4:\n    2/3 (x4)/(x4) =  [  ____ / ____  ]",
         "f) Multiply top & bottom by 3:\n    3/4 (x3)/(x3) =  [  ____ / ____  ]"],
        ["g) What multiplier was used?\n    1/4  -->  4/16   ( Multiplier = ____ )",
         "h) What multiplier was used?\n    2/5  -->  6/15   ( Multiplier = ____ )",
         "i) What multiplier was used?\n    3/5  -->  6/10   ( Multiplier = ____ )"]
    ]

    for r_idx, row in enumerate(mult_tbl.rows):
        for c_idx, cell in enumerate(row.cells):
            cell.width = m_widths[c_idx]
            set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
            cell.paragraphs[0].add_run(m_data[r_idx][c_idx])

    # --- SECTION 4: EQUIVALENT FRACTION CHAINS & CHALLENGE ---
    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    h4 = doc.add_paragraph()
    h4.paragraph_format.space_before = Pt(12)
    h4.paragraph_format.space_after = Pt(4)
    run_h4 = h4.add_run("Section 4: Equivalent Fraction Chains & Extension Challenge")
    run_h4.font.size = Pt(14)
    run_h4.font.bold = True
    run_h4.font.color.rgb = NAVY

    chain_tbl = doc.add_table(rows=2, cols=1)
    chain_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    chain_tbl.autofit = False
    for r in chain_tbl.rows:
        r.cells[0].width = 9360
        set_cell_margins(r.cells[0], top=120, bottom=120, left=180, right=180)
        set_cell_shading(r.cells[0], GREY_BG)

    chain_tbl.cell(0, 0).paragraphs[0].add_run("1. Complete these equivalent fraction chains:\n   a)  1/2  =  [  ]/4  =  [  ]/6  =  [  ]/8  =  [  ]/10  =  [  ]/12\n   b)  1/3  =  [  ]/6  =  [  ]/9  =  [  ]/12 =  [  ]/15  =  [  ]/18").font.bold = True

    chain_tbl.cell(1, 0).paragraphs[0].add_run("2. Extension Word Problem:\n   Lucas ate 2/4 of a pizza. Mia ate 4/8 of a pizza of the exact same size. Mia claims she ate more pizza than Lucas because 4 and 8 are bigger numbers. Is Mia correct? Explain using equivalent fractions and draw a quick diagram to prove your answer.\n   ____________________________________________________________________________________________________\n   ____________________________________________________________________________________________________").font.size = Pt(10.5)

    # Save document
    out_dir = r"c:\Users\dsuth\Documents\Joshua\Units\Maths\Signpost Math Lessons\Lesson_1_17_Equivalent_Fractions"
    os.makedirs(out_dir, exist_ok=True)
    out_file = os.path.join(out_dir, "Lesson_1_17_Equivalent_Fractions_Worksheet.docx")
    doc.save(out_file)
    print(f"SUCCESS: Worksheet created successfully at: {out_file}")

if __name__ == "__main__":
    build_worksheet()
