import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def create_element(name):
    return OxmlElement(name)

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)

def add_header(doc, title, subtitle, target_tier, roster_str):
    header_table = doc.add_table(rows=1, cols=1)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = header_table.cell(0, 0)
    set_cell_background(cell, "1E3A8A") # Navy
    set_cell_margins(cell, top=180, bottom=180, left=240, right=240)
    
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(title.upper())
    run.font.name = "Arial"
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)
    
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run2 = p2.add_run(f"{subtitle} | {target_tier}")
    run2.font.name = "Arial"
    run2.font.size = Pt(11)
    run2.font.italic = True
    run2.font.color.rgb = RGBColor(224, 231, 255)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # Student metadata box
    meta_table = doc.add_table(rows=1, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_table.autofit = False
    
    col_widths = [Inches(4.0), Inches(3.0)]
    for row in meta_table.rows:
        for idx, width in enumerate(col_widths):
            row.cells[idx].width = width

    c0 = meta_table.cell(0, 0)
    c1 = meta_table.cell(0, 1)
    set_cell_background(c0, "F1F5F9")
    set_cell_background(c1, "F1F5F9")
    set_cell_margins(c0, top=100, bottom=100, left=150, right=150)
    set_cell_margins(c1, top=100, bottom=100, left=150, right=150)
    
    p_meta0 = c0.paragraphs[0]
    p_meta0.add_run("Student Name: ").bold = True
    p_meta0.add_run("_______________________")
    
    p_meta1 = c1.paragraphs[0]
    p_meta1.add_run("Target Roster: ").bold = True
    p_meta1.add_run(roster_str)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

def add_heading_1(doc, text):
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(14)
    h.paragraph_format.space_after = Pt(6)
    h.paragraph_format.keep_with_next = True
    run = h.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(30, 58, 138) # Navy
    return h

def add_heading_2(doc, text):
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(10)
    h.paragraph_format.space_after = Pt(4)
    h.paragraph_format.keep_with_next = True
    run = h.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(37, 99, 235) # Blue
    return h

def add_instruction_box(doc, text):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = tbl.cell(0, 0)
    set_cell_background(c, "FEF3C7") # Warm yellow
    set_cell_margins(c, top=120, bottom=120, left=180, right=180)
    p = c.paragraphs[0]
    r = p.add_run("💡 KEY STRATEGY: ")
    r.bold = True
    r.font.color.rgb = RGBColor(146, 64, 14)
    r2 = p.add_run(text)
    r2.font.color.rgb = RGBColor(120, 53, 15)
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

# -------------------------------------------------------------
# BUILD TIER 3 WORKSHEET (Intensive Support — 4 Pages)
# -------------------------------------------------------------
def build_tier3_worksheet():
    doc = Document()
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    add_header(
        doc, 
        "Fractions & Decimals Misconception Clinic", 
        "Tier 3 Intensive Support & Guided Remediation Pack", 
        "Targeted Small Group",
        "jtayl1104, dmcdo222"
    )

    add_instruction_box(
        doc,
        "Work through this pack with your teacher at the small-group table. Use your tactile fraction towers and place value charts to build every single problem before writing your answer!"
    )

    add_heading_1(doc, "Section 1: Decimal Place Value & Zero Placeholders (Hundreds vs Tenths)")
    p = doc.add_paragraph("Remember: The first number after the decimal point is TENTHS (1/10). The second number is HUNDREDTHS (1/100). If there are NO tenths, put a 0 as a placeholder!")
    p.paragraph_format.space_after = Pt(8)

    # Q Table 1
    tbl1 = doc.add_table(rows=5, cols=4)
    tbl1.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Fraction / Mixed Number", "Ones . Tenths  Hundredths", "Decimal Form", "Teacher Check"]
    for idx, h in enumerate(headers):
        cell = tbl1.cell(0, idx)
        set_cell_background(cell, "1E3A8A")
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)

    data1 = [
        ("6 8/100 (6 wholes and 8 hundredths)", "6 . 0 8", "6.08", "[   ] Correct"),
        ("3 4/100 (3 wholes and 4 hundredths)", "3 . __ __", "________", "[   ] Check"),
        ("10 5/100 (10 wholes and 5 hundredths)", "10 . __ __", "________", "[   ] Check"),
        ("8 5/100 (8 wholes and 5 hundredths)", "8 . __ __", "________", "[   ] Check"),
    ]

    for row_idx, row_data in enumerate(data1, start=1):
        for col_idx, val in enumerate(row_data):
            cell = tbl1.cell(row_idx, col_idx)
            if row_idx % 2 == 1:
                set_cell_background(cell, "F8FAFC")
            cell.paragraphs[0].add_run(val)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    add_heading_1(doc, "Section 2: Renaming Whole Fractions (n/n = 1 Whole)")
    add_instruction_box(
        doc,
        "Whenever your numerator equals your denominator (like 10/10 or 8/8), it equals 1 WHOLE! Swap n/n for 1 and add it to your whole numbers."
    )

    data2 = [
        ("Problem", "Step 1: Add Numerators", "Step 2: Rename n/n = 1", "Final Simplified Answer"),
        ("1 7/10 + 3/10", "1 10/10", "1 + 1", "2"),
        ("4 5/8 + 3/8", "4 8/8", "4 + 1", "______"),
        ("2 4/5 + 1/5", "2 5/5", "2 + __", "______"),
        ("3 6/10 + 4/10", "3 __/10", "__ + __", "______"),
    ]
    tbl2 = doc.add_table(rows=len(data2), cols=4)
    tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r_idx, row in enumerate(data2):
        for c_idx, val in enumerate(row):
            cell = tbl2.cell(r_idx, c_idx)
            if r_idx == 0:
                set_cell_background(cell, "1E3A8A")
                p = cell.paragraphs[0]
                r = p.add_run(val)
                r.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
            else:
                if r_idx % 2 == 1:
                    set_cell_background(cell, "F8FAFC")
                cell.paragraphs[0].add_run(val)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    add_heading_1(doc, "Section 3: The Golden Rule (Adding Related Denominators)")
    p = doc.add_paragraph("Golden Rule: You cannot add different slice sizes! Multiply top and bottom of the smaller denominator to match the bigger denominator BEFORE adding.")
    
    data3 = [
        ("Original Problem", "Step 1: Convert Smaller Denominator", "Step 2: Add Numerators", "Final Answer"),
        ("1/4 + 3/8", "1/4 = (1 x 2)/(4 x 2) = 2/8", "2/8 + 3/8", "5/8"),
        ("3/5 + 3/10", "3/5 = (3 x 2)/(5 x 2) = __/10", "__/10 + 3/10", "______"),
        ("5/8 - 1/4", "1/4 = (1 x 2)/(4 x 2) = 2/8", "5/8 - 2/8", "______"),
        ("1/4 + 5/8", "1/4 = (1 x 2)/(4 x 2) = __/8", "__/8 + 5/8", "______"),
    ]
    tbl3 = doc.add_table(rows=len(data3), cols=4)
    tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r_idx, row in enumerate(data3):
        for c_idx, val in enumerate(row):
            cell = tbl3.cell(r_idx, c_idx)
            if r_idx == 0:
                set_cell_background(cell, "1E3A8A")
                p = cell.paragraphs[0]
                r = p.add_run(val)
                r.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
            else:
                if r_idx % 2 == 1:
                    set_cell_background(cell, "F8FAFC")
                cell.paragraphs[0].add_run(val)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    add_heading_1(doc, "Section 4: Simplifying Fractions to Lowest Terms")
    p = doc.add_paragraph("To simplify a fraction, divide BOTH the numerator and denominator by their Highest Common Factor (HCF).")

    data4 = [
        ("Fraction", "Divide Top and Bottom By", "Working", "Simplest Form"),
        ("8/12", "Divide by 4", "(8 ÷ 4) / (12 ÷ 4)", "2/3"),
        ("4/6", "Divide by 2", "(4 ÷ 2) / (6 ÷ 2)", "______"),
        ("6/10", "Divide by 2", "(6 ÷ 2) / (10 ÷ 2)", "______"),
        ("10/12", "Divide by 2", "(10 ÷ 2) / (12 ÷ 2)", "______"),
    ]
    tbl4 = doc.add_table(rows=len(data4), cols=4)
    tbl4.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r_idx, row in enumerate(data4):
        for c_idx, val in enumerate(row):
            cell = tbl4.cell(r_idx, c_idx)
            if r_idx == 0:
                set_cell_background(cell, "1E3A8A")
                p = cell.paragraphs[0]
                r = p.add_run(val)
                r.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
            else:
                if r_idx % 2 == 1:
                    set_cell_background(cell, "F8FAFC")
                cell.paragraphs[0].add_run(val)

    doc.save(os.path.join(os.path.dirname(__file__), "Tier_3_Intensive_Support_Worksheet.docx"))
    print("Created Tier 3 Worksheet!")

# -------------------------------------------------------------
# BUILD TIER 2 WORKSHEET (Core Consolidation — 3 Pages)
# -------------------------------------------------------------
def build_tier2_worksheet():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    add_header(
        doc, 
        "Fractions & Decimals Misconception Clinic", 
        "Tier 2 Core Consolidation & Practice Pack", 
        "Guided Pairs / Independent",
        "jfull212, mreed71, cpono2, lmcdo381, hherz0, shart259, kfiel89"
    )

    add_instruction_box(
        doc,
        "Complete each section carefully. Pay extra attention to zero placeholders in decimals and applying the Golden Rule before adding fractions with different denominators!"
    )

    add_heading_1(doc, "Section 1: Decimal Place Value & Hundredths Placeholders")
    p = doc.add_paragraph("Write each fraction/mixed number as a decimal, and each decimal as a fraction.")

    data1 = [
        ("Fraction / Mixed Number", "Decimal Form", "Explanation / Strategy"),
        ("6 8/100", "6.08", "8 hundredths requires 0 placeholder in tenths column"),
        ("10 5/100", "________", "5 hundredths requires 0 placeholder"),
        ("1.05", "1 5/100", "0 in tenths means 5 is in hundredths column"),
        ("4.03", "________", "Express as a mixed number over 100"),
        ("64% battery charge remaining", "0.64", "64 hundredths = 0.64"),
    ]
    tbl1 = doc.add_table(rows=len(data1), cols=3)
    tbl1.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r_idx, row in enumerate(data1):
        for c_idx, val in enumerate(row):
            cell = tbl1.cell(r_idx, c_idx)
            if r_idx == 0:
                set_cell_background(cell, "1E3A8A")
                p = cell.paragraphs[0]
                r = p.add_run(val)
                r.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
            else:
                if r_idx % 2 == 1:
                    set_cell_background(cell, "F8FAFC")
                cell.paragraphs[0].add_run(val)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    add_heading_1(doc, "Section 2: Renaming Whole Fractions in Mixed Sums")
    
    data2 = [
        ("Expression", "Intermediate Step", "Renamed Whole", "Final Answer"),
        ("1 7/10 + 3/10", "1 10/10", "1 + 1", "2"),
        ("3 6/10 + 4/10", "3 10/10", "3 + 1", "______"),
        ("4 5/8 + 3/8", "4 8/8", "______", "______"),
        ("2 4/5 + 1/5", "______", "______", "______"),
    ]
    tbl2 = doc.add_table(rows=len(data2), cols=4)
    tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r_idx, row in enumerate(data2):
        for c_idx, val in enumerate(row):
            cell = tbl2.cell(r_idx, c_idx)
            if r_idx == 0:
                set_cell_background(cell, "1E3A8A")
                p = cell.paragraphs[0]
                r = p.add_run(val)
                r.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
            else:
                if r_idx % 2 == 1:
                    set_cell_background(cell, "F8FAFC")
                cell.paragraphs[0].add_run(val)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    add_heading_1(doc, "Section 3: Golden Rule Related Denominators & Simplification")
    
    data3 = [
        ("Problem", "Golden Rule Conversion", "Addition / Subtraction", "Simplest Form"),
        ("1/4 + 3/8", "1/4 = 2/8", "2/8 + 3/8 = 5/8", "5/8"),
        ("3/5 + 3/10", "3/5 = __/10", "__/10 + 3/10 = __/10", "________"),
        ("5/8 - 1/4", "1/4 = __/8", "5/8 - __/8 = __/8", "________"),
        ("Simplify 8/12", "Divide top & bottom by 4", "(8 ÷ 4) / (12 ÷ 4)", "2/3"),
    ]
    tbl3 = doc.add_table(rows=len(data3), cols=4)
    tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r_idx, row in enumerate(data3):
        for c_idx, val in enumerate(row):
            cell = tbl3.cell(r_idx, c_idx)
            if r_idx == 0:
                set_cell_background(cell, "1E3A8A")
                p = cell.paragraphs[0]
                r = p.add_run(val)
                r.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
            else:
                if r_idx % 2 == 1:
                    set_cell_background(cell, "F8FAFC")
                cell.paragraphs[0].add_run(val)

    doc.save(os.path.join(os.path.dirname(__file__), "Tier_2_Core_Consolidation_Worksheet.docx"))
    print("Created Tier 2 Worksheet!")

# -------------------------------------------------------------
# BUILD TIER 1 WORKSHEET (Extension & Peer Mastery — 2 Pages)
# -------------------------------------------------------------
def build_tier1_worksheet():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    add_header(
        doc, 
        "Fractions & Decimals Misconception Clinic", 
        "Tier 1 Extension Challenge & Peer Mastery Pack", 
        "Independent / Peer Coach",
        "epryo13, jbinn27, jbart350, lheck4, fwend2, smorg220, fpick8, wnich33"
    )

    add_instruction_box(
        doc,
        "Part 1 is your rapid 5-minute Diagnostic Warmup to double-check your missed items. Part 2 is an advanced investigation into 3-fraction expressions and real-world capacity problems!"
    )

    add_heading_1(doc, "Part 1: Rapid Diagnostic Error-Check Warmup (5 Mins)")
    p = doc.add_paragraph("Solve your specific target items below:")

    data1 = [
        ("Student / Item Target", "Problem Statement", "Your Correction & Explanation"),
        ("Place Value Check (wnich33, jbart350, lheck4)", "Write 6 8/100 as a decimal.", "6.08 (Needs 0 placeholder in tenths column)"),
        ("Mixed Sum Renaming (fwend2, smorg220)", "Calculate 1 7/10 + 3/10.", "1 10/10 = 1 + 1 = 2"),
        ("Golden Rule Check (fpick8, epryo13)", "Calculate 1/4 + 3/8.", "1/4 = 2/8 -> 2/8 + 3/8 = 5/8"),
        ("Simplification (fpick8, fwend2)", "Simplify 8/12 to lowest terms.", "8/12 = 2/3 (Divide by HCF of 4)"),
    ]
    tbl1 = doc.add_table(rows=len(data1), cols=3)
    tbl1.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r_idx, row in enumerate(data1):
        for c_idx, val in enumerate(row):
            cell = tbl1.cell(r_idx, c_idx)
            if r_idx == 0:
                set_cell_background(cell, "1E3A8A")
                p = cell.paragraphs[0]
                r = p.add_run(val)
                r.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
            else:
                if r_idx % 2 == 1:
                    set_cell_background(cell, "F8FAFC")
                cell.paragraphs[0].add_run(val)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    add_heading_1(doc, "Part 2: Advanced Fraction Investigations & Capacity Challenges")
    
    add_heading_2(doc, "Challenge 1: Multi-Step 3-Fraction Expressions")
    p1 = doc.add_paragraph("Evaluate: 3/4 + 1/8 - 1/2. Convert all fractions to eighths first!")
    p1_box = doc.add_paragraph("Working: 3/4 = 6/8,  1/2 = 4/8  -->  6/8 + 1/8 - 4/8 = 3/8")
    p1_box.paragraph_format.space_after = Pt(8)

    add_heading_2(doc, "Challenge 2: Real-World Sports Container Capacity")
    p2 = doc.add_paragraph("A sports jug holds 4 litres of water. In the first half, players drink 1 3/4 litres. In the second half, players drink 1 5/8 litres. How much water remains in the jug?")
    p2_ans = doc.add_paragraph("Working space: ____________________________________________________________________\nFinal Answer: ________ Litres")
    p2_ans.paragraph_format.space_after = Pt(8)

    add_heading_2(doc, "Challenge 3: Peer Coaching Role")
    p3 = doc.add_paragraph("Paired with a Tier 2 classmate: Guide them through the Golden Rule for 3/5 + 3/10 without giving away the answer. Ask them: 'What do we multiply 5 by to get 10?'")

    doc.save(os.path.join(os.path.dirname(__file__), "Tier_1_Extension_Mastery_Worksheet.docx"))
    print("Created Tier 1 Worksheet!")

if __name__ == "__main__":
    build_tier3_worksheet()
    build_tier2_worksheet()
    build_tier1_worksheet()
