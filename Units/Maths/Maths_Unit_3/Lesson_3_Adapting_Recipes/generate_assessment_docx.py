import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def build_assessment_docx():
    doc = docx.Document()

    # Set page margins (0.75 inch)
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Base styling
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Source Sans 3'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

    # Title
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(4)
    run_title = title_p.add_run("Diagnostic Quiz: Recipe Scaling")
    run_title.font.name = 'Outfit'
    run_title.font.size = Pt(22)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0x0F, 0x4C, 0x81)

    # Subtitle / Metadata
    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_after = Pt(18)
    sub_run = sub_p.add_run("Year 5 Maths | Multipliers & Shopping Units | Microsoft Forms Standard")
    sub_run.font.size = Pt(10)
    sub_run.font.italic = True
    sub_run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    questions = [
        {
            "num": 1,
            "q": "A recipe for Anzac Biscuits makes 12 biscuits. You need to make enough for 24 students. What is the Recipe Multiplier?",
            "opts": ["A. 0.5", "B. 2", "C. 3", "D. 12"],
            "ans": "Answer: B"
        },
        {
            "num": 2,
            "q": "If a recipe that serves 6 people requires 3 avocados, how many avocados are needed to make enough for 24 people?",
            "opts": ["A. 6 avocados", "B. 9 avocados", "C. 12 avocados", "D. 18 avocados"],
            "ans": "Answer: C"
        },
        {
            "num": 3,
            "q": "A recipe requires 1/2 cup of sugar for 6 serves. How much sugar is needed for 18 serves?",
            "opts": ["A. 1 cup", "B. 1 1/2 cups", "C. 2 cups", "D. 2 1/2 cups"],
            "ans": "Answer: B"
        },
        {
            "num": 4,
            "q": "You adapt a recipe by multiplying by 4. The original recipe calls for 400g of flour. What is the total mass of flour needed in kilograms?",
            "opts": ["A. 0.16 kg", "B. 1.6 kg", "C. 16 kg", "D. 1600 kg"],
            "ans": "Answer: B"
        },
        {
            "num": 5,
            "q": "A smoothie recipe calls for 250mL of milk for 4 serves. How much milk is required for 20 serves?",
            "opts": ["A. 1.0 L", "B. 1.25 L", "C. 1.5 L", "D. 2.5 L"],
            "ans": "Answer: B"
        },
        {
            "num": 6,
            "q": "Which formula should you use to find the Recipe Multiplier?",
            "opts": [
                "A. Recipe Multiplier = Recipe Makes + People Needed",
                "B. Recipe Multiplier = People Needed ÷ Recipe Makes",
                "C. Recipe Multiplier = Recipe Makes ÷ People Needed",
                "D. Recipe Multiplier = People Needed × Recipe Makes"
            ],
            "ans": "Answer: B"
        },
        {
            "num": 7,
            "q": "A recipe makes 8 mini pizzas and calls for 150g of diced capsicum. If you double the recipe, how many grams of capsicum will you use?",
            "opts": ["A. 75g", "B. 200g", "C. 300g", "D. 450g"],
            "ans": "Answer: C"
        },
        {
            "num": 8,
            "q": "Why might a baker choose to make 3 full batches of biscuits instead of 2.3 batches?",
            "opts": [
                "A. 2.3 is impossible to multiply.",
                "B. Making 3 whole batches avoids splitting eggs or half-packages of ingredients.",
                "C. 3 batches uses less total flour.",
                "D. 3 batches makes exactly 24 biscuits."
            ],
            "ans": "Answer: B"
        },
        {
            "num": 9,
            "q": "Which measurement change is correct when converting to supermarket units?",
            "opts": [
                "A. 1250 g = 12.5 kg",
                "B. 1500 mL = 1.5 L",
                "C. 2400 g = 0.24 kg",
                "D. 3000 mL = 30 L"
            ],
            "ans": "Answer: B"
        },
        {
            "num": 10,
            "q": "A party guacamole recipe calls for 1/4 tsp of salt for 6 serves. How much salt is needed for 24 serves?",
            "opts": ["A. 1/2 tsp", "B. 3/4 tsp", "C. 1 tsp", "D. 2 tsp"],
            "ans": "Answer: C"
        },
        {
            "num": 11,
            "q": "You are adapting a recipe for a school assembly of 120 people. The original recipe makes 12 serves and calls for 100g of butter. How much butter do you need in total?",
            "opts": ["A. 500g", "B. 1 kg", "C. 1.2 kg", "D. 10 kg"],
            "ans": "Answer: B"
        },
        {
            "num": 12,
            "q": "A recipe calls for 30mL of golden syrup for 12 biscuits. For 120 biscuits, how much golden syrup is needed?",
            "opts": [
                "A. 300mL (0.3L)",
                "B. 3000mL (3.0L)",
                "C. 30mL (0.03L)",
                "D. 300mL (3.0L)"
            ],
            "ans": "Answer: A"
        }
    ]

    for item in questions:
        qp = doc.add_paragraph()
        qp.paragraph_format.space_before = Pt(8)
        qp.paragraph_format.space_after = Pt(3)
        q_run = qp.add_run(f"{item['num']}. {item['q']}")
        q_run.font.bold = True
        q_run.font.size = Pt(11)

        for opt in item['opts']:
            op = doc.add_paragraph()
            op.paragraph_format.space_before = Pt(0)
            op.paragraph_format.space_after = Pt(2)
            op.paragraph_format.left_indent = Inches(0.25)
            opt_run = op.add_run(opt)
            opt_run.font.size = Pt(10.5)

        ap = doc.add_paragraph()
        ap.paragraph_format.space_before = Pt(1)
        ap.paragraph_format.space_after = Pt(8)
        ap.paragraph_format.left_indent = Inches(0.25)
        ans_run = ap.add_run(item['ans'])
        ans_run.font.bold = True
        ans_run.font.color.rgb = RGBColor(0xE8, 0x5D, 0x04)

    output_path = r"c:\Users\dsuth\Documents\Joshua\Units\Maths\Maths_Unit_3\Lesson_3_Adapting_Recipes\Assessment_Forms.docx"
    doc.save(output_path)
    print(f"Successfully generated: {output_path}")

if __name__ == "__main__":
    build_assessment_docx()
