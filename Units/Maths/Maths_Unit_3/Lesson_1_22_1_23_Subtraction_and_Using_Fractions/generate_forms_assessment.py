import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'''
        <w:tcMar {nsdecls("w")}>
            <w:top w:w="{top}" w:type="dxa"/>
            <w:bottom w:w="{bottom}" w:type="dxa"/>
            <w:left w:w="{left}" w:type="dxa"/>
            <w:right w:w="{right}" w:type="dxa"/>
        </w:tcMar>
    ''')
    tcPr.append(tcMar)

def build_forms_assessment():
    output_dir = os.path.dirname(os.path.abspath(__file__))
    output_path = os.path.join(output_dir, "Lesson_1_12_1_13_Forms_Assessment.docx") # matching standard naming or Lesson_1_22_1_23_Forms_Assessment.docx
    output_path_22 = os.path.join(output_dir, "Lesson_1_22_1_23_Forms_Assessment.docx")
    
    doc = docx.Document()
    
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Arial'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0x1A, 0x23, 0x32)
    
    NAVY = RGBColor(0x0B, 0x4F, 0x6C)
    DARK_GRAY = RGBColor(0x33, 0x4E, 0x68)
    GREEN = RGBColor(0x27, 0xAE, 0x60)

    # Title Header
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    t_run = title.add_run("Signpost Mathematics — Lessons 1:22 & 1:23 Formative Assessment")
    t_run.font.size = Pt(16)
    t_run.font.bold = True
    t_run.font.color.rgb = NAVY

    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(16)
    s_run = sub.add_run("Subtraction from Whole Numbers & Operations with Related Denominators (Microsoft Forms Import Format)")
    s_run.font.size = Pt(11)
    s_run.font.color.rgb = DARK_GRAY

    questions = [
        {
            "num": 1,
            "q": "What is 1 - 3/8?",
            "options": ["A. 2/8", "B. 5/8", "C. 3/8", "D. 1 3/8"],
            "ans": "ANS: B"
        },
        {
            "num": 2,
            "q": "Calculate 4 - 1/5.",
            "options": ["A. 3 4/5", "B. 3 1/5", "C. 4 4/5", "D. 3/5"],
            "ans": "ANS: A"
        },
        {
            "num": 3,
            "q": "What is 1 7/10 + 3/10?",
            "options": ["A. 1 10/10", "B. 2", "C. 1 4/10", "D. Both A and B are correct"],
            "ans": "ANS: D"
        },
        {
            "num": 4,
            "q": "Which fraction is equivalent to 8/12?",
            "options": ["A. 1/3", "B. 2/3", "C. 3/4", "D. 5/6"],
            "ans": "ANS: B"
        },
        {
            "num": 5,
            "q": "Which of the following is True?",
            "options": ["A. 2/3 = 6/12", "B. 1/3 = 4/12", "C. 8/12 = 3/4", "D. 10/12 = 2/3"],
            "ans": "ANS: B"
        },
        {
            "num": 6,
            "q": "What is the correct order of 3/8, 1/2, 3/4, 1/8 from smallest to largest?",
            "options": ["A. 1/8, 3/8, 1/2, 3/4", "B. 3/4, 1/2, 3/8, 1/8", "C. 1/2, 1/8, 3/8, 3/4", "D. 1/8, 1/2, 3/8, 3/4"],
            "ans": "ANS: A"
        },
        {
            "num": 7,
            "q": "Calculate 1/4 + 3/8.",
            "options": ["A. 4/12", "B. 5/8", "C. 4/8", "D. 1/2"],
            "ans": "ANS: B"
        },
        {
            "num": 8,
            "q": "Calculate 3/5 + 3/10.",
            "options": ["A. 6/15", "B. 9/10", "C. 6/10", "D. 3/5"],
            "ans": "ANS: B"
        },
        {
            "num": 9,
            "q": "What is 5/8 - 1/4?",
            "options": ["A. 4/4", "B. 3/8", "C. 4/8", "D. 1/4"],
            "ans": "ANS: B"
        },
        {
            "num": 10,
            "q": "Calculate 3/2 - 7/10.",
            "options": ["A. 4/5", "B. 8/10", "C. Both A and B are correct", "D. 4/8"],
            "ans": "ANS: C"
        }
    ]

    for item in questions:
        qp = doc.add_paragraph()
        qp.paragraph_format.space_before = Pt(8)
        qp.paragraph_format.space_after = Pt(4)
        q_run = qp.add_run(f"{item['num']}. {item['q']}")
        q_run.font.bold = True
        q_run.font.size = Pt(11.5)
        q_run.font.color.rgb = NAVY

        for opt in item["options"]:
            op = doc.add_paragraph()
            op.paragraph_format.left_indent = Inches(0.25)
            op.paragraph_format.space_after = Pt(2)
            op.add_run(opt)

        ap = doc.add_paragraph()
        ap.paragraph_format.left_indent = Inches(0.25)
        ap.paragraph_format.space_after = Pt(10)
        ans_run = ap.add_run(item["ans"])
        ans_run.font.bold = True
        ans_run.font.color.rgb = GREEN

    doc.save(output_path_22)
    print(f"Successfully generated MS Forms Assessment DOCX at: {output_path_22}")

if __name__ == "__main__":
    build_forms_assessment()
