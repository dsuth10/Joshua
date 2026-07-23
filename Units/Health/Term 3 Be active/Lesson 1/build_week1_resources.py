from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parent.parent
LESSON_DIR = Path(__file__).resolve().parent

NAVY = "17324D"
TEAL = "007C83"
MINT = "DDF3F1"
OCHRE = "C77D18"
PALE_OCHRE = "FFF1D8"
CORAL = "B64C3B"
PALE_CORAL = "FBE7E3"
INK = "18242F"
MUTED = "52606D"
LINE = "CBD5DC"
WHITE = "FFFFFF"
LIGHT = "F4F7F8"

SOURCES = {
    "pa": "https://www.health.gov.au/topics/physical-activity/about-physical-activity",
    "yulunga": "https://www.ausport.gov.au/clearinghouse/research/archive/asc/yulunga-games",
    "tai_chi": "https://www.nccih.nih.gov/health/tai-chi-what-you-need-to-know",
}


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=LINE, size=6, inside=True):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    names = ["top", "left", "bottom", "right"] + (["insideH", "insideV"] if inside else [])
    for name in names:
        edge = borders.find(qn(f"w:{name}"))
        if edge is None:
            edge = OxmlElement(f"w:{name}")
            borders.append(edge)
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), str(size))
        edge.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa, indent=120):
    total = sum(widths_dxa)
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    flag = OxmlElement("w:tblHeader")
    flag.set(qn("w:val"), "true")
    tr_pr.append(flag)


def keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


def set_run(run, size=10.5, bold=False, color=INK, italic=False, font="Calibri"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)
    return run


def configure_styles(doc):
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(0.78)
    sec.bottom_margin = Inches(0.72)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.header_distance = Inches(0.32)
    sec.footer_distance = Inches(0.35)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    for name, size, color, before, after in (
        ("Title", 27, NAVY, 0, 5),
        ("Subtitle", 13, MUTED, 0, 14),
        ("Heading 1", 17, NAVY, 14, 7),
        ("Heading 2", 13.5, TEAL, 11, 5),
        ("Heading 3", 11.5, NAVY, 8, 4),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = name != "Subtitle"
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    return sec


def set_header_footer(doc, label):
    for sec in doc.sections:
        hp = sec.header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hp.paragraph_format.space_after = Pt(0)
        set_run(hp.add_run(label.upper()), 8, True, TEAL)
        fp = sec.footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.paragraph_format.space_before = Pt(0)
        set_run(fp.add_run("LET'S ALL BE ACTIVE  |  YEAR 6 HEALTH  |  WEEK 1"), 8, True, MUTED)


def add_title_block(doc, kicker, title, subtitle, chips):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    set_run(p.add_run(kicker.upper()), 9, True, OCHRE)
    p = doc.add_paragraph(style="Title")
    p.paragraph_format.space_after = Pt(4)
    set_run(p.add_run(title), 27, True, NAVY)
    p = doc.add_paragraph(style="Subtitle")
    set_run(p.add_run(subtitle), 13, False, MUTED)
    table = doc.add_table(rows=1, cols=len(chips))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [9360 // len(chips)] * len(chips)
    widths[-1] += 9360 - sum(widths)
    set_table_geometry(table, widths, indent=120)
    set_table_borders(table, color=TEAL, size=8)
    for i, (label, value) in enumerate(chips):
        cell = table.cell(0, i)
        set_cell_shading(cell, MINT if i % 2 == 0 else LIGHT)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        set_run(p.add_run(label.upper() + "\n"), 7.5, True, TEAL)
        set_run(p.add_run(value), 10.5, True, NAVY)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_callout(doc, label, text, fill=MINT, accent=TEAL):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360], indent=120)
    set_table_borders(table, color=accent, size=10, inside=False)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    set_run(p.add_run(label.upper() + "  "), 9, True, accent)
    set_run(p.add_run(text), 10.5, False, INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_bullet(doc, text, level=0, color=INK):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.left_indent = Inches(0.5 + level * 0.25)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    set_run(p.add_run(text), 10.5, False, color)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    set_run(p.add_run(text), 10.5)
    return p


def add_link(paragraph, text, url):
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), TEAL)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(color)
    r_pr.append(underline)
    run.append(r_pr)
    node = OxmlElement("w:t")
    node.text = text
    run.append(node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_page_break(doc):
    doc.add_page_break()


def build_plan():
    doc = Document()
    configure_styles(doc)
    set_header_footer(doc, "Teacher plan")
    add_title_block(
        doc,
        "Week 1 | Topic 1: Valuing diversity",
        "My activity, our community",
        "A 60-minute lesson in which students use personal and cultural activity evidence to explain how valuing diversity can strengthen community wellbeing.",
        [("Year", "6"), ("Duration", "60 min"), ("Evidence", "Workbook pp. 2-6"), ("Codes", "ACPPS060 | ACPMP064")],
    )
    add_callout(doc, "Lesson destination", "Students produce a personal activity profile, compare cultural activities, and revise an inclusive community-activity proposal using peer feedback.")
    doc.add_heading("Learning and evidence", level=1)
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [4680, 4680])
    set_table_borders(table)
    headers = [("Learning intention", MINT), ("Success evidence", PALE_OCHRE)]
    for i, (txt, fill) in enumerate(headers):
        set_cell_shading(table.cell(0, i), fill)
        set_run(table.cell(0, i).paragraphs[0].add_run(txt), 10, True, NAVY)
    set_repeat_table_header(table.rows[0])
    left = [
        "Explain how personal and cultural physical-activity choices show diversity.",
        "Explain how valuing that diversity can improve belonging, participation and community wellbeing.",
    ]
    right = [
        "I define physical activity and show what, where and why in my own profile.",
        "I compare two cultural activities using accurate evidence.",
        "I propose an inclusive activity and justify one revision with because.",
    ]
    row = table.add_row()
    for text in left:
        p = row.cells[0].add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        set_run(p.add_run(text), 9.5)
    row.cells[0]._element.remove(row.cells[0].paragraphs[0]._element)
    for text in right:
        p = row.cells[1].add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        set_run(p.add_run(text), 9.5)
    row.cells[1]._element.remove(row.cells[1].paragraphs[0]._element)
    doc.add_heading("Before the lesson", level=2)
    for item in [
        "Print the Student Workbook double-sided (one per student). Keep completed workbooks as Week 1 evidence.",
        "Print and cut one set of Cultural Activity Evidence Cards per five groups. Assign one card to each group.",
        "Open Lesson 1_Presentation.html. It works offline; use N for teacher notes, F for fullscreen and R to reset the current slide.",
        "If running a movement demonstration, prepare a soft ball and four markers, complete school risk-management requirements, and check space, footwear, sun safety and equipment.",
        "For First Nations games, follow local school protocols. Where practical, inform or invite local Elders or First Nations groups rather than presenting one activity as representative of all cultures.",
    ]:
        add_bullet(doc, item)
    add_page_break(doc)
    doc.add_heading("Teaching sequence", level=1)
    phases = [
        ("0-5", "Enter: What counts?", "Project six examples. Students silently choose: definitely physical activity, depends, or not physical activity. Take two contrasting justifications before revealing.", "Listen for the misconception that only organised sport or hard exercise counts."),
        ("5-11", "Define and refine", "Co-construct a class definition, then show the Australian Government wording. Clarify that activity can be incidental, exercise, sport or strengthening, and can occur at light, moderate or vigorous intensity.", "Students complete Workbook p.1 and revise one first response."),
        ("11-23", "My movement profile", "Model one honest entry: activity, place, frequency, reasons and influence. Students complete Workbook p.2. Offer oral rehearsal before writing.", "Conference: ask 'What makes this choice yours?' and 'What might help or limit it?'"),
        ("23-29", "Community alphabet scan", "Pairs add activities to the A-Z grid, then star three that could welcome people with different interests, abilities or cultural experiences.", "Do not reward volume alone; ask students to justify one starred choice."),
        ("29-43", "Cultural evidence jigsaw", "Teach the respect protocol. Groups read one evidence card and prepare a 30-second briefing: origin, movement, one community value, one accurate caution. Students then collect evidence from a second group.", "Check that Edor is linked to Aurukun and that cultures are named specifically where the source permits."),
        ("43-54", "Design for belonging", "Present the one-sport community-day scenario. Students propose two changes, justify them using evidence from their profile or cards, then exchange feedback: 'This may include more people because...' and 'A barrier still might be...'.", "Require an immediate revision in a different-coloured pen or with a star."),
        ("54-60", "Exit evidence", "Students answer the claim-evidence-reasoning exit prompt. Collect workbooks. Sort responses after class into Ready / Developing / Reteach.", "Use the sort to decide whether Week 2 needs a short recap on diversity, participation or wellbeing."),
    ]
    table = doc.add_table(rows=1, cols=4)
    set_table_geometry(table, [720, 1800, 4680, 2160])
    set_table_borders(table)
    for i, txt in enumerate(["Min", "Phase", "Teacher moves", "Evidence to notice"]):
        set_cell_shading(table.cell(0, i), NAVY)
        set_run(table.cell(0, i).paragraphs[0].add_run(txt), 9, True, WHITE)
    set_repeat_table_header(table.rows[0])
    for mins, phase, moves, evidence in phases:
        row = table.add_row()
        for i, txt in enumerate([mins, phase, moves, evidence]):
            p = row.cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            set_run(p.add_run(txt), 8.8, i == 1, NAVY if i == 1 else INK)
        if len(table.rows) % 2 == 0:
            for c in row.cells:
                set_cell_shading(c, LIGHT)

    add_page_break(doc)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(8)
    doc.add_heading("Teacher talk and worked model", level=1)
    doc.add_heading("Student-friendly definition", level=2)
    add_callout(doc, "Say", "Physical activity is anything that gets your body moving and uses energy. It includes everyday movement, active travel, exercise, sport, strengthening and cultural movement practices.")
    p = doc.add_paragraph()
    set_run(p.add_run("Source: "), 9, True, MUTED)
    add_link(p, "Australian Government - About physical activity", SOURCES["pa"])
    doc.add_heading("Model profile entry", level=2)
    table = doc.add_table(rows=2, cols=5)
    set_table_geometry(table, [1650, 1400, 1300, 2510, 2500])
    set_table_borders(table)
    for i, txt in enumerate(["Activity", "Where", "How often", "Why I choose it", "What influences my choice"]):
        set_cell_shading(table.cell(0, i), TEAL)
        set_run(table.cell(0, i).paragraphs[0].add_run(txt), 8.6, True, WHITE)
    set_repeat_table_header(table.rows[0])
    values = ["Walking the dog", "Local streets", "Most afternoons", "Time together; fresh air; helps me reset", "Family routine; free access; weather"]
    for i, txt in enumerate(values):
        set_run(table.cell(1, i).paragraphs[0].add_run(txt), 8.8)
    doc.add_heading("Model community claim", level=2)
    add_callout(doc, "Think aloud", "Valuing different activities can strengthen belonging because people see that movement is not limited to one sport or one type of body. For example, including a gentle tai chi session and a team game gives people different ways to participate. Therefore, more people may feel that the event is for them.", PALE_OCHRE, OCHRE)
    doc.add_heading("Questions that deepen thinking", level=2)
    for item in [
        "Whose interests or needs are visible in this activity? Whose might be missing?",
        "Is a choice personal only, or is it also shaped by family, culture, cost, access, place or ability?",
        "How does your evidence show respect rather than treating a culture as a costume or novelty?",
        "What would participation, belonging or wellbeing look like here?",
    ]:
        add_bullet(doc, item)
    doc.add_heading("Misconceptions and responses", level=2)
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [3900, 5460])
    set_table_borders(table)
    for i, txt in enumerate(["If students say...", "Respond with..."]):
        set_cell_shading(table.cell(0, i), PALE_CORAL)
        set_run(table.cell(0, i).paragraphs[0].add_run(txt), 9, True, CORAL)
    set_repeat_table_header(table.rows[0])
    for misconception, response in [
        ("Only sport counts.", "Test everyday movement, active travel, dance and cultural games against the definition."),
        ("People choose freely, so access does not matter.", "Ask how cost, transport, facilities, safety, family and ability shape available choices."),
        ("One game represents all First Nations peoples.", "Name the specific community or region and emphasise the diversity of Aboriginal and Torres Strait Islander cultures."),
        ("Diversity is just having many options.", "Add the valuing test: are differences respected, included and used to shape decisions?"),
    ]:
        row = table.add_row()
        set_run(row.cells[0].paragraphs[0].add_run(misconception), 9)
        set_run(row.cells[1].paragraphs[0].add_run(response), 9)

    add_page_break(doc)
    doc.add_heading("Differentiation without changing the goal", level=1)
    for title, fill, items in [
        ("Support", MINT, ["Offer the word bank: fun, friendship, fitness, family, culture, calm, challenge, cost, access, safety, belonging.", "Allow the personal profile to be dictated, drawn or completed with an adult scribe.", "Give one evidence card with key sentences highlighted and use the oral briefing before written notes.", "Use the claim frame: Valuing ___ can improve ___ because ___. For example, ___."]),
        ("ICP / accessible pathway", PALE_OCHRE, ["Use photo, symbol, spoken or partner-assisted responses while retaining the same concept: choices, evidence and community inclusion.", "Offer seated, low-impact and non-competitive examples without presenting them as lesser alternatives.", "Pre-teach origin, community, movement and belonging with four icons; let the student select evidence and explain orally.", "Follow documented student adjustments and sensory/movement safety plans."]),
        ("Extend", PALE_CORAL, ["Ask whether adding options is enough if cost, language, transport or confidence still blocks participation.", "Require a counterexample: when might an activity with diverse participants still fail to value diversity?", "Ask students to distinguish cultural appreciation, accurate attribution and tokenism in the proposed event."]),
    ]:
        table = doc.add_table(rows=1, cols=1)
        set_table_geometry(table, [9360])
        set_table_borders(table, color=TEAL if title != "Extend" else CORAL, size=8, inside=False)
        cell = table.cell(0, 0)
        set_cell_shading(cell, fill)
        p = cell.paragraphs[0]
        set_run(p.add_run(title), 11, True, NAVY)
        for item in items:
            p = cell.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(3)
            set_run(p.add_run(item), 9.5)
        doc.add_paragraph().paragraph_format.space_after = Pt(0)
    doc.add_heading("Safety and cultural responsibility", level=1)
    for item in [
        "Use only soft equipment for demonstrations; define boundaries; prohibit body contact; keep bats below waist height and balls below knee height for Dabi.",
        "Do not run Edor without a large clear area and a specific risk check. The core lesson can be completed through evidence cards without game play.",
        "Present the Yulunga activities as researched and adapted school versions, not as a complete representation of a living culture.",
        "Follow the ASC recommendation to inform or invite local Elders or First Nations groups where these games are used in a school program.",
        "Avoid asking students to disclose culture, disability, family finances or health information. Personal-profile responses may remain private.",
    ]:
        add_bullet(doc, item)
    add_page_break(doc)
    doc.add_heading("Formative assessment sort", level=1)
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [3120, 3120, 3120])
    set_table_borders(table)
    for i, (txt, fill) in enumerate([("Ready", MINT), ("Developing", PALE_OCHRE), ("Reteach", PALE_CORAL)]):
        set_cell_shading(table.cell(0, i), fill)
        set_run(table.cell(0, i).paragraphs[0].add_run(txt), 10, True, NAVY)
    set_repeat_table_header(table.rows[0])
    row = table.add_row()
    values = [
        "Defines PA; uses specific evidence; links valuing diversity to participation, belonging or wellbeing.",
        "Gives a relevant example but the because link is general or evidence is incomplete.",
        "Lists different activities only; confuses diversity with wellbeing; makes an inaccurate cultural claim.",
    ]
    for i, txt in enumerate(values):
        set_run(row.cells[i].paragraphs[0].add_run(txt), 9)
    doc.add_heading("Source notes", level=1)
    p = doc.add_paragraph()
    set_run(p.add_run("Australian Sports Commission: "), 9, True, MUTED)
    add_link(p, "Yulunga: Traditional Indigenous Games", SOURCES["yulunga"])
    p = doc.add_paragraph()
    set_run(p.add_run("Australian Government: "), 9, True, MUTED)
    add_link(p, "About physical activity", SOURCES["pa"])
    p = doc.add_paragraph()
    set_run(p.add_run("National Center for Complementary and Integrative Health: "), 9, True, MUTED)
    add_link(p, "Tai Chi: What You Need To Know", SOURCES["tai_chi"])
    doc.save(ROOT / "Lesson 1 - Lesson Plan.docx")


def writing_lines(doc, count=3):
    for _ in range(count):
        p = doc.add_paragraph("________________________________________________________________________________")
        p.paragraph_format.space_after = Pt(5)
        set_run(p.runs[0], 9, False, MUTED)


def page_heading(doc, number, title, prompt):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    set_run(p.add_run(f"{number:02d}"), 20, True, OCHRE)
    set_run(p.add_run(f"  {title}"), 20, True, NAVY)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    set_run(p.add_run(prompt), 11, False, MUTED)


def build_workbook():
    doc = Document()
    configure_styles(doc)
    set_header_footer(doc, "Student workbook")
    add_title_block(
        doc,
        "Week 1 | Student workbook",
        "My activity, our community",
        "Name: ____________________________________    Class: __________________    Date: ________________",
        [("Learn", "notice choices"), ("Use", "cultural evidence"), ("Design", "for belonging")],
    )
    add_callout(doc, "Learning goal", "I can use evidence about personal and cultural physical activities to explain how valuing diversity can strengthen community wellbeing.")
    doc.add_heading("Before the definition: what counts?", level=1)
    p = doc.add_paragraph()
    set_run(p.add_run("For each example, circle "), 10.5)
    set_run(p.add_run("YES"), 10.5, True, TEAL)
    set_run(p.add_run(", "), 10.5)
    set_run(p.add_run("DEPENDS"), 10.5, True, OCHRE)
    set_run(p.add_run(" or "), 10.5)
    set_run(p.add_run("NO"), 10.5, True, CORAL)
    set_run(p.add_run(". Be ready to justify one choice."), 10.5)
    examples = ["Walking to school", "Tai chi", "Competitive gaming", "Gardening", "Wheelchair basketball", "Dancing at home"]
    table = doc.add_table(rows=3, cols=2)
    set_table_geometry(table, [4680, 4680])
    set_table_borders(table)
    for idx, item in enumerate(examples):
        cell = table.cell(idx // 2, idx % 2)
        set_cell_shading(cell, LIGHT if idx % 2 else MINT)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run(p.add_run(item + "\n"), 11, True, NAVY)
        set_run(p.add_run("YES   |   DEPENDS   |   NO"), 8.5, True, MUTED)
    doc.add_heading("Build and refine a definition", level=1)
    set_run(doc.add_paragraph().add_run("My first definition:"), 10, True, TEAL)
    writing_lines(doc, 2)
    set_run(doc.add_paragraph().add_run("After the class discussion, revise it:"), 10, True, TEAL)
    writing_lines(doc, 2)
    add_callout(doc, "Useful idea", "Physical activity includes everyday movement, active travel, exercise, sport, strengthening and cultural movement practices. It is not limited to organised sport.", MINT, TEAL)

    add_page_break(doc)
    page_heading(doc, 2, "My physical activity profile", "Record what you choose, where you move and why. Honest answers are more useful than 'perfect' answers.")
    table = doc.add_table(rows=1, cols=5)
    set_table_geometry(table, [1720, 1380, 1300, 2500, 2460])
    set_table_borders(table)
    headers = ["Activity or movement", "Where?", "How often?", "Why do I choose it?", "What shapes this choice?"]
    for i, txt in enumerate(headers):
        set_cell_shading(table.cell(0, i), NAVY)
        set_run(table.cell(0, i).paragraphs[0].add_run(txt), 8.5, True, WHITE)
    set_repeat_table_header(table.rows[0])
    for _ in range(5):
        row = table.add_row()
        for c in row.cells:
            p = c.paragraphs[0]
            p.paragraph_format.space_after = Pt(36)
    doc.add_paragraph()
    add_callout(doc, "Choice bank", "fun | friendship | fitness | family | culture | calm | challenge | cost | access | safety | confidence | belonging", PALE_OCHRE, OCHRE)
    doc.add_heading("Notice a pattern", level=2)
    prompts = [
        "One activity I value is __________________________ because ________________________________.",
        "A person with different interests, experiences or access might choose __________________________ instead because ____________________________________________.",
        "One thing that could make physical activity easier for more people is ____________________________.",
    ]
    for text in prompts:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(10)
        set_run(p.add_run(text), 10.5)

    add_page_break(doc)
    page_heading(doc, 3, "Community activity alphabet", "Work with a partner. Add examples from everyday life, sport, recreation, dance, cultural games and active travel.")
    letters = list("ABCDEFGHIJKLMNOPQRSTUVWXYZ")
    table = doc.add_table(rows=13, cols=2)
    set_table_geometry(table, [4680, 4680])
    set_table_borders(table)
    for idx, letter in enumerate(letters):
        cell = table.cell(idx // 2, idx % 2)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run(p.add_run(letter + "  "), 12, True, TEAL)
        set_run(p.add_run("________________________________"), 8.5, False, MUTED)
        set_cell_margins(cell, top=35, start=120, bottom=35, end=120)
    doc.add_paragraph()
    add_callout(doc, "Star three", "Choose three activities that could welcome people with different interests, abilities or cultural experiences. Star them, then justify one choice below.", MINT, TEAL)
    writing_lines(doc, 3)

    add_page_break(doc)
    page_heading(doc, 4, "Cultural activity evidence", "Listen for specific evidence. A culture is not a costume, a label or one single activity.")
    add_callout(doc, "Respect protocol", "Name the source community or region when known. Describe the school version accurately. Notice both similarities and differences. Do not claim one game represents every Aboriginal, Torres Strait Islander or Asian culture.", PALE_OCHRE, OCHRE)
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [2100, 3630, 3630])
    set_table_borders(table)
    for i, txt in enumerate(["Evidence", "Activity 1", "Activity 2"]):
        set_cell_shading(table.cell(0, i), TEAL)
        set_run(table.cell(0, i).paragraphs[0].add_run(txt), 9, True, WHITE)
    set_repeat_table_header(table.rows[0])
    prompts = ["Name", "Source community / place", "How people move", "What people need to do together", "Why it may matter to community", "One fact I must say carefully"]
    for prompt in prompts:
        row = table.add_row()
        set_cell_shading(row.cells[0], LIGHT)
        set_run(row.cells[0].paragraphs[0].add_run(prompt), 9, True, NAVY)
        for cell in row.cells[1:]:
            cell.paragraphs[0].paragraph_format.space_after = Pt(21)
    doc.add_heading("Compare", level=2)
    p = doc.add_paragraph()
    set_run(p.add_run("Both activities ______________________________________. However, _____________________________. This shows diversity because ________________________________________________."), 10.5)
    p.paragraph_format.space_after = Pt(15)
    doc.add_heading("Connect to wellbeing", level=2)
    p = doc.add_paragraph()
    set_run(p.add_run("Valuing these activities could strengthen community wellbeing by _____________________________ because ________________________________________________________________________."), 10.5)

    add_page_break(doc)
    page_heading(doc, 5, "Design for belonging", "A community activity afternoon offers only competitive soccer. Some people are excited; others do not see a place for themselves.")
    doc.add_heading("1. Diagnose the barrier", level=2)
    p = doc.add_paragraph()
    set_run(p.add_run("Who might be left out or less likely to participate? Why?"), 10.5, True, TEAL)
    writing_lines(doc, 3)
    doc.add_heading("2. Make two purposeful changes", level=2)
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [2380, 3490, 3490])
    set_table_borders(table)
    for i, txt in enumerate(["Change", "What we would add or adjust", "How this supports participation or belonging"]):
        set_cell_shading(table.cell(0, i), NAVY)
        set_run(table.cell(0, i).paragraphs[0].add_run(txt), 8.7, True, WHITE)
    set_repeat_table_header(table.rows[0])
    for label in ["Choice 1", "Choice 2"]:
        row = table.add_row()
        set_run(row.cells[0].paragraphs[0].add_run(label), 9.5, True, TEAL)
        row.cells[1].paragraphs[0].paragraph_format.space_after = Pt(38)
        row.cells[2].paragraphs[0].paragraph_format.space_after = Pt(38)
    doc.add_heading("3. Use evidence", level=2)
    p = doc.add_paragraph()
    set_run(p.add_run("Our strongest change is __________________________ because ________________________________. Evidence from my profile or a cultural activity card shows ________________________________________."), 10.5)
    p.paragraph_format.space_after = Pt(16)
    doc.add_heading("4. Peer feedback and revision", level=2)
    set_run(doc.add_paragraph().add_run("This may include more people because ________________________________________________."), 10.5)
    set_run(doc.add_paragraph().add_run("A barrier still might be ____________________________________________________________."), 10.5)
    add_callout(doc, "Revise now", "Star or rewrite one part of your design. Explain the revision: I changed __________________ because ________________________________.", PALE_CORAL, CORAL)

    add_page_break(doc)
    page_heading(doc, 6, "Exit evidence", "Show what you now understand. Use a specific example and a clear because link.")
    doc.add_heading("Claim - evidence - reasoning", level=1)
    add_callout(doc, "Prompt", "How can valuing diversity in physical activity strengthen community wellbeing?", MINT, TEAL)
    for label, stem in [
        ("CLAIM", "Valuing diversity in physical activity can..."),
        ("EVIDENCE", "For example, today I learned..."),
        ("REASONING", "This strengthens participation, belonging or wellbeing because..."),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(3)
        set_run(p.add_run(label + "  "), 10, True, OCHRE)
        set_run(p.add_run(stem), 10.5, True, NAVY)
        writing_lines(doc, 3 if label != "REASONING" else 4)
    doc.add_heading("Self-check", level=1)
    for text in [
        "[  ] I used a specific activity or profile example.",
        "[  ] I explained how diversity was valued, not just noticed.",
        "[  ] I linked my example to participation, belonging or wellbeing.",
        "[  ] I used accurate and respectful cultural language.",
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(7)
        set_run(p.add_run(text), 10.5)
    add_callout(doc, "My next step", "I feel: READY / DEVELOPING / UNSURE. One thing I still need help with is _________________________________________________.", PALE_OCHRE, OCHRE)
    doc.save(ROOT / "Lesson 1 - Student Worksheet.docx")


CARDS = [
    {
        "name": "Dabi",
        "tag": "Torres Strait Islands",
        "origin": "Dabi is a school-practice version of kokan, a hockey-type game played on Mabuiag Island in the Torres Strait. The name refers to the rough bamboo bat or club used with the kokan (ball).",
        "movement": "Teams use hockey or unihoc sticks to control and pass a soft ball. A player normally uses one touch to control and a second to pass. The activity is continuous rather than focused on scoring.",
        "together": "Players need accurate passing, awareness of others, self-control and shared safety rules.",
        "care": "Use soft equipment. Keep sticks below waist height and the ball below knee height. Call this an adapted school version and name Mabuiag Island rather than saying it represents all Torres Strait Islander cultures.",
        "prompt": "How could a continuous passing activity include beginners while still challenging experienced players?",
    },
    {
        "name": "Edor",
        "tag": "Aurukun, north Queensland",
        "origin": "Edor is a goal-orientated chasing-and-tagging game from the Aurukun Aboriginal community. The name is used in north Queensland and Torres Strait regions; the Yulunga account recognises Aunty Cathy, an Aurukun Elder and school teacher aide, as an authority for the game.",
        "movement": "Two teams work across a large field. One team's chosen runner tries to reach the opposite goal while teammates protect and support. When the runner is tagged, possession and direction change.",
        "together": "The game depends on teamwork, quick decisions, communication, agility and supporting a runner.",
        "care": "Name the Aurukun source context. The activity needs a large clear space, non-contact tagging rules, active supervision and a school risk check.",
        "prompt": "How can a team protect and support one runner without unsafe contact?",
    },
    {
        "name": "Keentan",
        "tag": "North-west central Queensland",
        "origin": "A catch-ball keep-away game was played by both genders in north-west central Queensland. The Kalkadoon people sometimes described it as 'kangaroo-play' because players jumped to catch. Keentan means 'play' in the Wik-Mungkan language; the name is used because an identifiable Kalkadoon word was unavailable to the source authors.",
        "movement": "Teams pass a ball and try to keep possession. Defenders gain the ball only by catching it while they are in the air. A small jump is enough; no body contact or obstruction is allowed.",
        "together": "Players need timing, passing, spatial awareness, fair play and coordinated team movement.",
        "care": "Explain why the name is used and avoid collapsing Wik-Mungkan language and Kalkadoon cultural history into one group.",
        "prompt": "Which rule could be adjusted so people with different mobility needs can contribute meaningfully?",
    },
    {
        "name": "Wana Wana",
        "tag": "Noongar, south-west Western Australia",
        "origin": "This practice version is based on a game played by young Noongar girls in south-west Western Australia. A wana (or wanna) is a digging stick in some Noongar language. In the historical game, a girl used her wana to stop others hitting a short stick on the ground.",
        "movement": "Four to eight players form a circle. Players take turns to underarm-lob a soft ball to a batter in the centre. The batter gently hits it to the next player to catch. Everyone rotates through the batting role.",
        "together": "The circle structure depends on turn-taking, controlled striking, catching and helping the next person succeed.",
        "care": "Use gentle lobs and strikes, wide spacing and soft equipment. Name the Noongar context and do not present one game as representing every Aboriginal culture.",
        "prompt": "How does rotating every person through the central role support fairness and belonging?",
    },
    {
        "name": "Tai chi",
        "tag": "China",
        "origin": "Tai chi originated as a martial art in China. Today it is also practised as a mind-body activity for health and rehabilitation in many communities around the world.",
        "movement": "A sequence of slow, gentle movements and postures is coordinated with controlled breathing and focused attention. It can be practised alone or in groups and adapted for different bodies.",
        "together": "A group can move in shared rhythm while each person controls their own range and pace. The activity can offer a calm, non-competitive way to participate.",
        "care": "Avoid calling all Asian movement traditions the same. Tai chi is Chinese and has many styles and purposes. Do not make medical promises from a brief classroom experience.",
        "prompt": "Why might a calm, non-competitive option help more people see a place for themselves at an activity event?",
    },
]


def add_card(doc, card):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360], indent=120)
    set_table_borders(table, color=TEAL, size=14, inside=False)
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    set_run(p.add_run(card["name"]), 21, True, NAVY)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    set_run(p.add_run(card["tag"].upper()), 9, True, OCHRE)
    for label, key in [("SOURCE STORY", "origin"), ("HOW PEOPLE MOVE", "movement"), ("WHAT PEOPLE DO TOGETHER", "together"), ("SAY IT CAREFULLY", "care")]:
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(2)
        set_run(p.add_run(label), 9, True, TEAL if label != "SAY IT CAREFULLY" else CORAL)
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        set_run(p.add_run(card[key]), 9.6)
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    set_run(p.add_run("GROUP THINKING PROMPT  "), 9, True, OCHRE)
    set_run(p.add_run(card["prompt"]), 9.6, True, NAVY)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def build_cards():
    doc = Document()
    configure_styles(doc)
    set_header_footer(doc, "Evidence cards")
    add_title_block(
        doc,
        "Week 1 | Group jigsaw",
        "Cultural activity evidence cards",
        "Read closely. Prepare a 30-second briefing: origin, movement, one community value and one fact that must be said carefully.",
        [("Groups", "5"), ("Briefing", "30 sec"), ("Purpose", "compare + connect")],
    )
    add_callout(doc, "Teacher note", "These are concise classroom summaries, not substitutes for local cultural authority. The Yulunga games are researched and adapted school versions. Follow local protocols and, where practical, inform or invite local Elders or First Nations groups.", PALE_OCHRE, OCHRE)
    doc.add_heading("Jigsaw roles", level=1)
    roles = [
        ("Reader", "reads every section aloud"),
        ("Evidence finder", "underlines origin and movement facts"),
        ("Respect checker", "identifies the fact that must be said carefully"),
        ("Connector", "links the activity to participation, belonging or wellbeing"),
        ("Reporter", "delivers the 30-second briefing"),
    ]
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [1900, 7460])
    set_table_borders(table)
    for i, txt in enumerate(["Role", "Job"]):
        set_cell_shading(table.cell(0, i), NAVY)
        set_run(table.cell(0, i).paragraphs[0].add_run(txt), 9, True, WHITE)
    set_repeat_table_header(table.rows[0])
    for role, job in roles:
        row = table.add_row()
        set_run(row.cells[0].paragraphs[0].add_run(role), 9.5, True, TEAL)
        set_run(row.cells[1].paragraphs[0].add_run(job), 9.5)
    for card in CARDS:
        add_page_break(doc)
        add_card(doc, card)
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(5)
        set_run(p.add_run("SOURCE  "), 8, True, MUTED)
        if card["name"] == "Tai chi":
            add_link(p, "NCCIH - Tai Chi: What You Need To Know", SOURCES["tai_chi"])
        else:
            add_link(p, "Australian Sports Commission - Yulunga", SOURCES["yulunga"])
    doc.save(ROOT / "Lesson 1 - Cultural Activity Evidence Cards.docx")


def build_overview():
    content = """# Week 1 - My activity, our community

**Topic:** Valuing diversity influences community wellbeing  
**Year level:** 6  
**Duration:** 60 minutes  
**Curriculum:** ACPPS060, ACPMP064

## Ready-to-teach package

| Resource | Use |
|---|---|
| `Lesson 1 - Lesson Plan.docx` | Full teaching sequence, modelling, differentiation, safety, assessment sort and source notes |
| `Lesson 1 - Student Worksheet.docx` | Self-contained six-page student workbook, including My physical activity, A-Z scan, cultural comparison, inclusive design and exit evidence |
| `Lesson 1 - Cultural Activity Evidence Cards.docx` | Five printable group cards: Dabi, Edor, Keentan, Wana Wana and tai chi |
| `Lesson 1/Lesson 1_Presentation.html` | Offline-safe classroom presentation with notes, timers, navigation, fullscreen and reset |

## Lesson outcome

Students use evidence from their own physical activity choices and from cultural activities to explain how valuing diversity can improve participation, belonging and community wellbeing.

## Print and setup

1. Print one Student Workbook per student (double-sided if preferred).
2. Print and cut one set of Cultural Activity Evidence Cards per five groups.
3. Open the HTML presentation in a browser. Use arrow keys to move, **N** for notes, **F** for fullscreen and **R** to reset the current slide.
4. Movement demonstrations are optional. Complete school risk-management requirements before running them.

## Cultural responsibility

The Yulunga activities are researched and adapted school versions. Name the source community or region when known, avoid presenting any one activity as representative of all Aboriginal or Torres Strait Islander cultures, and follow local protocols. Where practical, inform or invite local Elders or First Nations groups.
"""
    (ROOT / "Lesson 1.md").write_text(content, encoding="utf-8")


if __name__ == "__main__":
    build_plan()
    build_workbook()
    build_cards()
    build_overview()
    print("Built Week 1 Word resources and overview.")
