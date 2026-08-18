"""
Build Assessment Package Resources for Year 6 Health Unit 2: "Let's All Be Active"
Generates:
1. Assessment_Lesson_Plan.docx
2. Student_Game_Design_Booklet.docx
"""

from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

PACKAGE_DIR = Path(__file__).resolve().parent

# Color Palette: Active Motion Theme (Navy, Teal, Mint, Amber, Coral, Ink, Muted, Line, White, Light)
NAVY = "0D1F2D"
TEAL = "00A896"
MINT = "E8F8F5"
AMBER = "F4A261"
PALE_AMBER = "FDF3E7"
CORAL = "E76F51"
PALE_CORAL = "FCEEEB"
INK = "1A202C"
MUTED = "4A5568"
LINE = "CBD5E0"
WHITE = "FFFFFF"
LIGHT = "F7FAFC"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=130, bottom=100, end=130):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=LINE, size=6, inside=True):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    names = ["top", "left", "bottom", "right"] + (["insideH", "insideV"] if inside else [])
    for name in names:
        edge = borders.find(qn(f"w:{name}"))
        if edge is None:
            edge = OxmlElement(f"w:{name}")
            borders.append(edge)
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), str(size))
        edge.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa, indent=120):
    total = sum(widths_dxa)
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    flag = OxmlElement("w:tblHeader")
    flag.set(qn("w:val"), "true")
    tr_pr.append(flag)


def set_run(run, size=10.5, bold=False, color=INK, italic=False, font="Calibri"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)
    return run


def configure_styles(doc):
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(0.75)
    sec.bottom_margin = Inches(0.72)
    sec.left_margin = Inches(0.9)
    sec.right_margin = Inches(0.9)
    sec.header_distance = Inches(0.3)
    sec.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, color, before, after in (
        ("Title", 26, NAVY, 0, 4),
        ("Subtitle", 12.5, MUTED, 0, 12),
        ("Heading 1", 16, NAVY, 13, 6),
        ("Heading 2", 13, TEAL, 10, 4),
        ("Heading 3", 11, NAVY, 7, 3),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = (name != "Subtitle")
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    return sec


def set_header_footer(doc, label):
    for sec in doc.sections:
        hp = sec.header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hp.paragraph_format.space_after = Pt(0)
        set_run(hp.add_run(label.upper()), 8, True, TEAL)
        fp = sec.footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.paragraph_format.space_before = Pt(0)
        set_run(fp.add_run("LET'S ALL BE ACTIVE  |  YEAR 6 HEALTH  |  SUMMATIVE ASSESSMENT"), 8, True, MUTED)


def add_title_block(doc, kicker, title, subtitle, chips):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    set_run(p.add_run(kicker.upper()), 9, True, AMBER)
    p = doc.add_paragraph(style="Title")
    p.paragraph_format.space_after = Pt(3)
    set_run(p.add_run(title), 26, True, NAVY)
    p = doc.add_paragraph(style="Subtitle")
    set_run(p.add_run(subtitle), 12.5, False, MUTED)
    table = doc.add_table(rows=1, cols=len(chips))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [9640 // len(chips)] * len(chips)
    widths[-1] += 9640 - sum(widths)
    set_table_geometry(table, widths, indent=60)
    set_table_borders(table, color=TEAL, size=8)
    for i, (label, value) in enumerate(chips):
        cell = table.cell(0, i)
        set_cell_shading(cell, MINT if i % 2 == 0 else LIGHT)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        set_run(p.add_run(label.upper() + "\n"), 7.5, True, TEAL)
        set_run(p.add_run(value), 10, True, NAVY)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_callout(doc, label, text, fill=MINT, accent=TEAL):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9640], indent=60)
    set_table_borders(table, color=accent, size=10, inside=False)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    set_run(p.add_run(label.upper() + "  "), 9, True, accent)
    set_run(p.add_run(text), 10, False, INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_bullet(doc, text, level=0, color=INK):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.left_indent = Inches(0.4 + level * 0.25)
    p.paragraph_format.first_line_indent = Inches(-0.2)
    p.paragraph_format.space_after = Pt(3.5)
    p.paragraph_format.line_spacing = 1.15
    set_run(p.add_run(text), 10, False, color)
    return p


def writing_lines(doc, count=3):
    for _ in range(count):
        p = doc.add_paragraph("_________________________________________________________________________________")
        p.paragraph_format.space_after = Pt(4)
        set_run(p.runs[0], 9, False, MUTED)


# ==============================================================================
# BUILD LESSON PLAN (2-SESSION TEACHER GUIDE)
# ==============================================================================
def build_lesson_plan():
    doc = Document()
    configure_styles(doc)
    set_header_footer(doc, "Teacher Lesson Plan — 2-Session Sequence")

    add_title_block(
        doc,
        "Summative Assessment | 2-Session Sequence",
        "Let's All Be Active: Inclusive Game Design",
        "A structured two-session sequence guiding Year 6 students to design, analyse, playtest, and evaluate an inclusive 6-player physical activity.",
        [
            ("Year Level", "Year 6"),
            ("Format", "2 × 60 Min"),
            ("Curriculum", "ACPPS054 | ACPPS059 | ACPPS060 | ACPMP064"),
            ("Assessment Mode", "Assignment / Project Portfolio"),
        ],
    )

    add_callout(
        doc,
        "Assessment Destination",
        "Students demonstrate their understanding of health, wellbeing, diversity, and safety by designing a non-eliminating, 100% inclusive physical activity for a group of 6 Year 6 peers, analysing its impact across 4 dimensions of health, playtesting it with peers, and reflecting on design refinements.",
    )

    doc.add_heading("Curriculum Alignment & Achievement Standard Evidence", level=1)
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [4820, 4820], indent=60)
    set_table_borders(table)
    headers = [("Content Descriptors", MINT), ("Achievement Standard Focus", PALE_AMBER)]
    for i, (txt, fill) in enumerate(headers):
        set_cell_shading(table.cell(0, i), fill)
        set_run(table.cell(0, i).paragraphs[0].add_run(txt), 9.5, True, NAVY)
    set_repeat_table_header(table.rows[0])

    row = table.add_row()
    left_items = [
        "ACPPS054: Investigate the role of preventive health in promoting health, safety and wellbeing for individuals and their communities.",
        "ACPPS059: Participate in physical activities that develop health-related and skill-related fitness.",
        "ACPPS060: Examine how physical activity, celebrating diversity and connecting to the environment support community wellbeing and cultural understanding.",
        "ACPMP064: Propose and apply movement concepts and strategies to modify physical activities to enhance participation and enjoyment.",
    ]
    right_items = [
        "Describe the significance of physical activity to health and wellbeing.",
        "Describe their own and others' contributions to health, physical activity, safety and wellbeing.",
        "Examine how physical activity, celebrating diversity and connecting to the environment support community wellbeing and cultural understanding.",
        "Create an inclusive, safe physical activity that matches guidelines and accounts for diverse needs.",
    ]
    for text in left_items:
        p = row.cells[0].add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2.5)
        set_run(p.add_run(text), 9)
    row.cells[0]._element.remove(row.cells[0].paragraphs[0]._element)
    for text in right_items:
        p = row.cells[1].add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2.5)
        set_run(p.add_run(text), 9)
    row.cells[1]._element.remove(row.cells[1].paragraphs[0]._element)

    doc.add_heading("Session 1: Inclusive Game Design & Health Analysis (60 min)", level=1)
    add_callout(doc, "Session 1 Goal", "Students analyse design constraints, invent an inclusive 6-player activity with step-by-step rules and spatial diagrams, and justify benefits across the 4 Dimensions of Health.", PALE_AMBER, AMBER)

    phases_s1 = [
        ("0–10 min", "Launch & Challenge Deconstruction", "Display Design Challenge. Explain non-negotiables: 6 players, 100% inclusion, ZERO elimination, minimal soft equipment, safe bounds. Review prior unit data (Weeks 1–4).", "Look for students who default to standard knockout sports (e.g. dodgeball). Prompt: 'How does everyone stay moving 100% of the time?'"),
        ("10–25 min", "Game Blueprint & Spatial Diagram", "Model worked example (Noodle Ball adapted from Dabi). Students complete Booklet pp. 2–3: Name, safe space (e.g., netball third), equipment, diagram grid, 5 numbered rules.", "Check diagrams: clear starting positions, spatial boundaries, equipment marks, and flow arrows."),
        ("25–45 min", "4 Dimensions of Health & Diversity Matrix", "Project the 4 Dimensions matrix. Model specific explanations with because links. Students complete Booklet pp. 4–5 (Physical, Social, Emotional, Environmental + Accessibility/Culture adaptations).", "Ensure responses describe mechanisms (e.g., 'Non-competitive format removes fear of mistakes, improving emotional confidence') rather than vague statements ('It is fun')."),
        ("45–60 min", "Pair Safety Pitch & Peer Audit", "Pairs pitch games to each other using the Quick Safety Rubric: (1) Is collision risk low? (2) Is equipment soft? (3) Does everyone have an active role? Students revise rules.", "Listen for immediate adjustments made to prevent bottlenecking or standing still."),
    ]
    table = doc.add_table(rows=1, cols=4)
    set_table_geometry(table, [850, 1900, 4400, 2490], indent=60)
    set_table_borders(table)
    for i, txt in enumerate(["Time", "Phase", "Teacher Moves & Prompts", "Evidence to Notice"]):
        set_cell_shading(table.cell(0, i), NAVY)
        set_run(table.cell(0, i).paragraphs[0].add_run(txt), 9, True, WHITE)
    set_repeat_table_header(table.rows[0])
    for time_str, phase, moves, evidence in phases_s1:
        row = table.add_row()
        for i, txt in enumerate([time_str, phase, moves, evidence]):
            p = row.cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            set_run(p.add_run(txt), 8.8, i == 1, NAVY if i == 1 else INK)
        if len(table.rows) % 2 == 0:
            for c in row.cells:
                set_cell_shading(c, LIGHT)

    doc.add_page_break()

    doc.add_heading("Session 2: Peer Playtest Festival, Refinement & Submission (60 min)", level=1)
    add_callout(doc, "Session 2 Goal", "Students run their 6-player games in an outdoor/court festival, collect structured peer feedback, execute a justified design refinement, and complete their GTMJ self-assessment.", MINT, TEAL)

    phases_s2 = [
        ("0–10 min", "Safety Briefing & Station Setup", "Establish 4–6 outdoor game stations (netball court thirds, shaded oval zones). Assign equipment packs (cones, noodles, soft balls). Set rotation rules and whistle cues.", "Ensure all students wear hats, appropriate footwear, and have water bottles."),
        ("10–35 min", "Peer Activity Festival (Live Playtesting)", "Round 1 (12 min): Group A designers teach & referee their games for 5 peers. Group B plays.\nRound 2 (12 min): Group B designers teach & referee. Group A plays.", "Observe student leadership: Are rules explained clearly? Do all 6 players participate actively? Is safety maintained?"),
        ("35–45 min", "Peer Feedback & Iterative Redesign", "Students return to classroom. Peers complete Booklet p. 6 ('Two Stars and a Safety Wish'). Designers document at least one explicit rule/space change made based on testing.", "Check that designers use the formula: 'I changed [Rule/Setup] because [Peer observation].'"),
        ("45–60 min", "Final Reflection & GTMJ Self-Assessment", "Students finalise Booklet pp. 6–7: School community wellbeing statement and A–E GTMJ self-assessment rubric. Collect completed booklets for grading.", "Check student justifications against the Guide to Making Judgments."),
    ]
    table = doc.add_table(rows=1, cols=4)
    set_table_geometry(table, [850, 1900, 4400, 2490], indent=60)
    set_table_borders(table)
    for i, txt in enumerate(["Time", "Phase", "Teacher Moves & Prompts", "Evidence to Notice"]):
        set_cell_shading(table.cell(0, i), NAVY)
        set_run(table.cell(0, i).paragraphs[0].add_run(txt), 9, True, WHITE)
    set_repeat_table_header(table.rows[0])
    for time_str, phase, moves, evidence in phases_s2:
        row = table.add_row()
        for i, txt in enumerate([time_str, phase, moves, evidence]):
            p = row.cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            set_run(p.add_run(txt), 8.8, i == 1, NAVY if i == 1 else INK)
        if len(table.rows) % 2 == 0:
            for c in row.cells:
                set_cell_shading(c, LIGHT)

    doc.add_heading("Teacher Worked Model: Noodle Ball (Adapted from Dabi)", level=1)
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [3000, 6640], indent=60)
    set_table_borders(table)
    headers = [("Design Component", NAVY), ("Exemplar Content", NAVY)]
    for i, (txt, bg) in enumerate(headers):
        set_cell_shading(table.cell(0, i), bg)
        set_run(table.cell(0, i).paragraphs[0].add_run(txt), 9.5, True, WHITE)
    set_repeat_table_header(table.rows[0])

    exemplar_data = [
        ("Game Title", "Noodle Ball (Adapted from Dabi, Mabuiag Island, Torres Strait)"),
        ("Safe Location", "Two-thirds of a netball court (open boundaries, padded goalposts, asphalt/court surface)."),
        ("Equipment", "6 foam pool noodles, 1 small soft foam ball, 8 boundary cones."),
        ("No-Elimination Mechanism", "Players pass in numbered sequence (1 to 6). On 'Change!', all drop noodles and rotate to the next stick. No goals, no scoring, continuous movement."),
        ("Physical Health", "Improves cardiovascular endurance through continuous running and hand-eye coordination by controlling and passing the ball with pool noodles."),
        ("Social Health", "Requires active communication, cooperation, and equal sharing since everyone must touch the ball before calling 'Change!'."),
        ("Emotional Health", "Eliminating scoring and outs creates a low-pressure environment where students feel confident to try without fear of making mistakes."),
        ("Environmental Health", "Takes place outdoors in fresh air; players are responsible for clearing hazards and litter before and after the game."),
        ("Diversity Adaptation", "Longer pool noodles allow seated or wheelchair players to reach the ground easily; boundaries expand or contract based on group mobility."),
    ]
    for comp, content in exemplar_data:
        row = table.add_row()
        set_cell_shading(row.cells[0], LIGHT)
        set_run(row.cells[0].paragraphs[0].add_run(comp), 9, True, TEAL)
        set_run(row.cells[1].paragraphs[0].add_run(content), 9)

    doc.add_page_break()

    doc.add_heading("Differentiation Strategies", level=1)
    diff_data = [
        ("Support Pathway", MINT, [
            "Provide pre-drawn court templates (netball court, oval zone) with boundary lines ready for player placement.",
            "Offer sentence frames for the 4 Dimensions: 'This activity helps [Dimension] because [Movement/Rule], which leads to [Health Benefit].'",
            "Provide a bank of inclusive mechanics: passing circles, cooperative tag without outs, synchronized movement, beanbag relay.",
            "Pair with a peer partner for oral rehearsal before recording rules in the booklet.",
        ]),
        ("Core Expectation", PALE_AMBER, [
            "Complete all 7 booklet sections independently with clear spatial diagrams and 5 numbered rules.",
            "Provide detailed rationales across all 4 dimensions of health with specific physiological/social mechanisms.",
            "Successfully teach and manage the 6-player group during the outdoor festival, executing a justified refinement.",
        ]),
        ("Extension Pathway", PALE_CORAL, [
            "Incorporate a multi-tier scaling rule (e.g. progressive spatial constraints, variable tempo, dual-role challenges).",
            "Explain how the game could be scaled for whole-school lunchtime implementation or adapted for junior primary buddies.",
            "Analyse potential cultural or socioeconomic barriers to community physical activity and how their game circumvents them.",
        ]),
    ]
    for title, fill, items in diff_data:
        table = doc.add_table(rows=1, cols=1)
        set_table_geometry(table, [9640], indent=60)
        set_table_borders(table, color=TEAL if title != "Extension Pathway" else CORAL, size=8, inside=False)
        cell = table.cell(0, 0)
        set_cell_shading(cell, fill)
        p = cell.paragraphs[0]
        set_run(p.add_run(title), 11, True, NAVY)
        for item in items:
            p = cell.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(2.5)
            set_run(p.add_run(item), 9.2)
        doc.add_paragraph().paragraph_format.space_after = Pt(0)

    doc.add_heading("Guide to Making Judgments (GTMJ) Marking Rubric", level=1)
    gtmj_table = doc.add_table(rows=1, cols=3)
    set_table_geometry(gtmj_table, [1200, 4220, 4220], indent=60)
    set_table_borders(gtmj_table)
    headers = [("Grade", NAVY), ("Knowledge & Understanding", NAVY), ("Investigating & Applying", NAVY)]
    for i, (txt, bg) in enumerate(headers):
        set_cell_shading(gtmj_table.cell(0, i), bg)
        set_run(gtmj_table.cell(0, i).paragraphs[0].add_run(txt), 9, True, WHITE)
    set_repeat_table_header(gtmj_table.rows[0])

    rubric_rows = [
        ("A", "Comprehensive description of the significance of physical activity across all 4 dimensions of health. Creates an exceptionally detailed, safe, and 100% inclusive physical activity matching all guidelines.", "Provides an insightful explanation of how the activity accounts for diversity, connects to the environment, and enhances community wellbeing. Executes a purposeful design refinement."),
        ("B", "Detailed description of the significance of physical activity across multiple health dimensions. Creates a well-structured, safe, and inclusive physical activity that clearly matches guidelines.", "Provides a clear explanation of how the activity accounts for diversity and supports the wellbeing of peers. Describes a relevant design refinement."),
        ("C", "Describes the significance of physical activity to health and wellbeing across physical, social, emotional, or environmental dimensions. Creates a safe physical activity matching key guidelines.", "Describes how the activity accounts for student diversity and supports peer participation. Notes an adjustment made during testing."),
        ("D", "Identifies basic benefits of physical activity. Creates a simple physical activity with some guidelines omitted or safety/inclusion gaps.", "States basic connections between physical activity and wellbeing. Identifies a simple change without clear justification."),
        ("E", "Lists fragmented facts about physical activity. Outlines an incomplete activity that does not meet safety or inclusion criteria.", "Makes isolated statements about personal physical activity without linking to diversity or community wellbeing."),
    ]
    for grade, know, app in rubric_rows:
        row = gtmj_table.add_row()
        set_cell_shading(row.cells[0], MINT if grade in ["A", "B"] else (PALE_AMBER if grade == "C" else PALE_CORAL))
        p0 = row.cells[0].paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run(p0.add_run(grade), 12, True, NAVY)
        set_run(row.cells[1].paragraphs[0].add_run(know), 8.8)
        set_run(row.cells[2].paragraphs[0].add_run(app), 8.8)

    doc.save(PACKAGE_DIR / "Assessment_Lesson_Plan.docx")
    print("Saved Assessment_Lesson_Plan.docx successfully.")


# ==============================================================================
# BUILD STUDENT BOOKLET (PRINTABLE WORKBOOK)
# ==============================================================================
def build_student_booklet():
    doc = Document()
    configure_styles(doc)
    set_header_footer(doc, "Student Game Design Booklet & Assessment Portfolio")

    # Cover Page Block
    add_title_block(
        doc,
        "Year 6 Health Assessment | Unit 2: Let's All Be Active",
        "Student Game Design Booklet",
        "Name: ____________________________________    Class: ___________    Date: ____________",
        [
            ("Mission", "Invent a 6-player inclusive game"),
            ("Test", "Peer Activity Festival"),
            ("Analyse", "4 Dimensions of Health"),
        ],
    )

    add_callout(
        doc,
        "Your Design Challenge",
        "Create an inclusive, safe, and engaging physical activity for a group of six Year 6 students. Your game must ensure that everyone can participate continuously (no elimination) and that diverse abilities and cultural backgrounds are welcomed.",
        MINT,
        TEAL,
    )

    doc.add_heading("Design Brief Non-Negotiables", level=1)
    table = doc.add_table(rows=4, cols=2)
    set_table_geometry(table, [2400, 7240], indent=60)
    set_table_borders(table)
    headers = [("Non-Negotiable Rule", NAVY), ("How My Game Meets This Rule", NAVY)]
    for i, (txt, bg) in enumerate(headers):
        set_cell_shading(table.cell(0, i), bg)
        set_run(table.cell(0, i).paragraphs[0].add_run(txt), 9, True, WHITE)
    set_repeat_table_header(table.rows[0])

    criteria = [
        ("1. 100% Inclusion & No Elimination", "No player is ever 'out' or sitting on the sidelines. Movement is continuous for all 6 players."),
        ("2. Safety by Design", "Clear boundaries, soft equipment (pool noodles, soft balls, cones), non-contact rules, and safe play surface."),
        ("3. Cultural / Activity Inspiration", "Adapted from Traditional Indigenous Games (e.g. Dabi, Edor, Keentan, Wana Wana), world cultures, or fitness games."),
    ]
    for idx, (crit, desc) in enumerate(criteria):
        row = table.rows[idx + 1]
        set_cell_shading(row.cells[0], LIGHT)
        set_run(row.cells[0].paragraphs[0].add_run(crit), 9, True, TEAL)
        set_run(row.cells[1].paragraphs[0].add_run(desc), 9)

    doc.add_heading("Section 1: School Cohort Movement Context", level=1)
    p = doc.add_paragraph()
    set_run(p.add_run("1. What have you discovered about Year 6 students' physical activity preferences at school?"), 10, True, NAVY)
    writing_lines(doc, 2)
    p = doc.add_paragraph()
    set_run(p.add_run("2. How does participating in physical activity support the overall wellbeing of our school community?"), 10, True, NAVY)
    writing_lines(doc, 2)

    doc.add_page_break()

    # Page 2: Game Blueprint
    doc.add_heading("Section 2: Game Blueprint & Setup", level=1)
    add_callout(doc, "Game Specifications", "Fill in your game's identity, location, and minimal equipment requirements.", PALE_AMBER, AMBER)

    spec_table = doc.add_table(rows=4, cols=2)
    set_table_geometry(spec_table, [2600, 7040], indent=60)
    set_table_borders(spec_table)
    for idx, (label, prompt) in enumerate([
        ("Game Title", "Write an engaging, descriptive name for your game:"),
        ("Safe Location", "Where in the school will this be played? (e.g., Netball court third, shaded grass, undercover area)"),
        ("When It Can Be Played", "[   ] Lunch / Recess Breaks     [   ] PE Lesson Warm-Up     [   ] Class Movement Break"),
        ("Equipment List", "List minimal, soft equipment needed (e.g., 6 pool noodles, 1 soft foam ball, 8 cones):"),
    ]):
        row = spec_table.rows[idx]
        set_cell_shading(row.cells[0], MINT if idx % 2 == 0 else LIGHT)
        set_run(row.cells[0].paragraphs[0].add_run(label), 9.5, True, NAVY)
        p = row.cells[1].paragraphs[0]
        set_run(p.add_run(prompt + "\n"), 9, False, MUTED)
        p.paragraph_format.space_after = Pt(24 if idx != 2 else 6)

    doc.add_heading("Court / Spatial Layout Diagram", level=2)
    p = doc.add_paragraph()
    set_run(p.add_run("Draw the spatial layout for your 6-player game. Label player starting positions (P1–P6), boundaries, equipment, and movement flow:"), 9.5, False, MUTED)

    # Large Diagram Box
    diag_table = doc.add_table(rows=1, cols=1)
    set_table_geometry(diag_table, [9640], indent=60)
    set_table_borders(diag_table, color=TEAL, size=10)
    diag_cell = diag_table.cell(0, 0)
    set_cell_shading(diag_cell, LIGHT)
    dp = diag_cell.paragraphs[0]
    dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(dp.add_run("\n\n[ SPATIAL COURT & MOVEMENT DIAGRAM BOX ]\n(Show 6 player positions, boundary markers, and equipment zones)\n\n\n\n\n\n\n\n\n"), 9, True, MUTED)

    doc.add_page_break()

    # Page 3: Step-by-Step Rules & Safety
    doc.add_heading("Section 3: Step-by-Step Game Instructions", level=1)
    p = doc.add_paragraph()
    set_run(p.add_run("Write 5 clear, numbered rules so any Year 6 group can pick up your game and play without confusion:"), 10, False, MUTED)

    for i in range(1, 6):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(2)
        set_run(p.add_run(f"Rule {i}: "), 10, True, TEAL)
        writing_lines(doc, 2)

    doc.add_heading("Inclusion & Safety Mechanism", level=2)
    p = doc.add_paragraph()
    set_run(p.add_run("1. Exactly how does your game ensure NO ONE is eliminated or benched during play?"), 10, True, NAVY)
    writing_lines(doc, 2)

    p = doc.add_paragraph()
    set_run(p.add_run("2. Explain why your chosen school space and movement rules are safe for participants:"), 10, True, NAVY)
    writing_lines(doc, 2)

    doc.add_page_break()

    # Page 4: 4 Dimensions of Health
    doc.add_heading("Section 4: The 4 Dimensions of Health Analysis", level=1)
    add_callout(
        doc,
        "Health & Wellbeing Matrix",
        "Explain specifically how playing your physical activity provides benefits across all four dimensions of health. Use because links to explain the physiological, social, and emotional benefits.",
        PALE_AMBER,
        AMBER,
    )

    dim_table = doc.add_table(rows=5, cols=2)
    set_table_geometry(dim_table, [2400, 7240], indent=60)
    set_table_borders(dim_table)
    headers = [("Dimension of Health", NAVY), ("How Are Health & Wellbeing Benefits Obtained?", NAVY)]
    for i, (txt, bg) in enumerate(headers):
        set_cell_shading(dim_table.cell(0, i), bg)
        set_run(dim_table.cell(0, i).paragraphs[0].add_run(txt), 9.5, True, WHITE)
    set_repeat_table_header(dim_table.rows[0])

    dim_prompts = [
        ("Physical Health\n(Body systems, cardio, motor skills, coordination)", "Describe which body systems (heart, lungs, muscles) are activated and what movement skills (agility, catching, passing) are developed:"),
        ("Social Health\n(Teamwork, communication, sharing, fair play)", "Explain how players cooperate, communicate, share turns, and include every player in the group:"),
        ("Emotional Health\n(Confidence, fun, stress relief, self-esteem)", "Explain how the non-competitive / non-eliminating structure makes players feel confident, happy, and free from fear of mistakes:"),
        ("Environmental Health\n(Connecting to nature, school grounds, safety)", "Explain how the activity connects students to outdoor spaces and encourages caring for the school environment and equipment:"),
    ]
    for idx, (dim_title, prompt) in enumerate(dim_prompts):
        row = dim_table.rows[idx + 1]
        set_cell_shading(row.cells[0], MINT if idx % 2 == 0 else PALE_AMBER)
        set_run(row.cells[0].paragraphs[0].add_run(dim_title), 9.5, True, NAVY)
        p = row.cells[1].paragraphs[0]
        set_run(p.add_run(prompt + "\n"), 8.8, False, MUTED)
        p.paragraph_format.space_after = Pt(40)

    doc.add_page_break()

    # Page 5: Diversity & Adaptations
    doc.add_heading("Section 5: Diversity & Accessibility Adaptations", level=1)
    add_callout(
        doc,
        "Valuing Diversity",
        "Show how your activity is welcoming to all students by adapting for varying mobility, skill levels, and cultural inspirations.",
        MINT,
        TEAL,
    )

    doc.add_heading("1. Cultural Connection / Game Origin", level=2)
    p = doc.add_paragraph()
    set_run(p.add_run("Which traditional, cultural, or fitness activity inspired your design (e.g., Dabi, Edor, Keentan, Wana Wana, Tai chi)? How did you adapt it respectfully?"), 10, True, NAVY)
    writing_lines(doc, 3)

    doc.add_heading("2. Mobility & Wheelchair Adaptation", level=2)
    p = doc.add_paragraph()
    set_run(p.add_run("How can your game be modified for a student using a wheelchair or having differing physical mobility?"), 10, True, NAVY)
    writing_lines(doc, 3)

    doc.add_heading("3. Skill Range & Confidence Adaptation", level=2)
    p = doc.add_paragraph()
    set_run(p.add_run("How can rules or equipment be adjusted so both beginner and experienced players feel challenged and successful?"), 10, True, NAVY)
    writing_lines(doc, 3)

    doc.add_page_break()

    # Page 6: Peer Playtest & Feedback
    doc.add_heading("Section 6: Peer Playtest & Redesign Protocol", level=1)
    add_callout(
        doc,
        "Peer Activity Festival (Session 2)",
        "Teach and referee your game for a group of 6 peers outdoors. Collect their feedback and execute at least one purposeful design refinement.",
        PALE_AMBER,
        AMBER,
    )

    p = doc.add_paragraph()
    set_run(p.add_run("Playtester Peer Names: ____________________________________________________________________"), 9.5, True, NAVY)

    doc.add_heading("Peer Evaluation Ratings (Filled by Peer Testers)", level=2)
    rate_table = doc.add_table(rows=5, cols=2)
    set_table_geometry(rate_table, [4820, 4820], indent=60)
    set_table_borders(rate_table)
    for idx, (crit, score) in enumerate([
        ("Continuous 100% Participation (No one waiting/out)", "[   ] 1    [   ] 2    [   ] 3    [   ] 4    [   ] 5"),
        ("Safety & Clear Boundaries (No collisions/soft gear)", "[   ] 1    [   ] 2    [   ] 3    [   ] 4    [   ] 5"),
        ("Clarity of Rules & Step-by-Step Instructions", "[   ] 1    [   ] 2    [   ] 3    [   ] 4    [   ] 5"),
        ("Fun & Emotional Enjoyment Factor", "[   ] 1    [   ] 2    [   ] 3    [   ] 4    [   ] 5"),
        ("Overall Inclusivity for Differing Abilities", "[   ] 1    [   ] 2    [   ] 3    [   ] 4    [   ] 5"),
    ]):
        row = rate_table.rows[idx]
        set_cell_shading(row.cells[0], LIGHT)
        set_run(row.cells[0].paragraphs[0].add_run(crit), 9, True, NAVY)
        p = row.cells[1].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run(p.add_run(score), 9.5, True, TEAL)

    doc.add_heading("Two Stars and a Safety Wish", level=2)
    p = doc.add_paragraph()
    set_run(p.add_run("Star 1 (What worked really well): "), 10, True, TEAL)
    writing_lines(doc, 2)
    p = doc.add_paragraph()
    set_run(p.add_run("Star 2 (Great team/fun moment): "), 10, True, TEAL)
    writing_lines(doc, 2)
    p = doc.add_paragraph()
    set_run(p.add_run("Wish (One rule, space, or safety adjustment to improve the game): "), 10, True, CORAL)
    writing_lines(doc, 2)

    doc.add_heading("Designer's Post-Playtest Refinement", level=2)
    add_callout(
        doc,
        "Iterative Refinement",
        "Describe one specific change you made to your rules, space, or equipment after playtesting, and explain WHY this improved the game.",
        PALE_CORAL,
        CORAL,
    )
    p = doc.add_paragraph()
    set_run(p.add_run("Based on peer testing, I changed ________________________________________________________________\nbecause __________________________________________________________________________________________.\nThis improved participation, safety, or enjoyment by _______________________________________________."), 10, False, INK)

    doc.add_page_break()

    # Page 7: Final Reflection & Self-Assessment
    doc.add_heading("Section 7: Final Reflection & GTMJ Self-Assessment", level=1)

    doc.add_heading("Final Claim & Evidence Statement", level=2)
    p = doc.add_paragraph()
    set_run(p.add_run("How does participation in inclusive physical activities like your game strengthen the health and wellbeing of our school community?"), 10, True, NAVY)
    writing_lines(doc, 3)

    doc.add_heading("Student Self-Assessment Rubric (Guide to Making Judgments)", level=2)
    rub_table = doc.add_table(rows=6, cols=3)
    set_table_geometry(rub_table, [900, 4370, 4370], indent=60)
    set_table_borders(rub_table)
    headers = [("Grade", NAVY), ("Knowledge & Understanding", NAVY), ("Investigating & Applying", NAVY)]
    for i, (txt, bg) in enumerate(headers):
        set_cell_shading(rub_table.cell(0, i), bg)
        set_run(rub_table.cell(0, i).paragraphs[0].add_run(txt), 9, True, WHITE)
    set_repeat_table_header(rub_table.rows[0])

    s_rubric = [
        ("A", "I created an exceptionally safe, 100% inclusive game matching all guidelines. I comprehensively explained benefits across all 4 health dimensions.", "I provided an insightful explanation of diversity adaptations, environmental connection, and justified an effective redesign revision."),
        ("B", "I created a clear, safe, and inclusive game matching all guidelines. I explained benefits across multiple health dimensions in detail.", "I explained how my activity accounts for diversity and supports peer wellbeing, and noted a clear post-playtest refinement."),
        ("C", "I created a safe physical activity matching the key guidelines and described benefits for health and wellbeing.", "I described how the activity accounts for diversity and supports peer participation, and made a simple testing adjustment."),
        ("D", "I created a basic physical activity, but some safety or inclusion rules were missing.", "I gave simple statements about physical activity and wellbeing with limited detail."),
        ("E", "My game was incomplete and did not meet basic safety or inclusion criteria.", "I gave brief, disconnected comments about personal activity."),
    ]
    for idx, (grd, k_text, a_text) in enumerate(s_rubric):
        row = rub_table.rows[idx + 1]
        set_cell_shading(row.cells[0], MINT if grd in ["A", "B"] else (PALE_AMBER if grd == "C" else PALE_CORAL))
        p0 = row.cells[0].paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run(p0.add_run(grd + "\n[   ]"), 10, True, NAVY)
        set_run(row.cells[1].paragraphs[0].add_run(k_text), 8.5)
        set_run(row.cells[2].paragraphs[0].add_run(a_text), 8.5)

    doc.save(PACKAGE_DIR / "Student_Game_Design_Booklet.docx")
    print("Saved Student_Game_Design_Booklet.docx successfully.")


if __name__ == "__main__":
    build_lesson_plan()
    build_student_booklet()
