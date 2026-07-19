from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE = Path(__file__).resolve().parent
OUT = BASE / "Persuasion Detective - Student Questions.docx"
IMAGE = BASE / "assets" / "storytime-dog-library-crop.png"

# Resolved design system: compact_reference_guide with a named
# "student response booklet" override matching the companion exemplar.
TEAL = "0D5C63"
TEAL_DARK = "173F43"
OCHRE = "D89A2B"
CORAL = "D9654B"
PALE_TEAL = "E7F2F1"
PALE_OCHRE = "FBF3DF"
INK = "203034"
MUTED = "657478"
LINE = "B9C9C9"
WHITE = "FFFFFF"


def set_run_font(run, name="Calibri", size=None, color=INK, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr()
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, name="Calibri"):
    style.font.name = name
    style._element.get_or_add_rPr()
    style._element.rPr.rFonts.set(qn("w:ascii"), name)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), name)


def shade(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def border(paragraph, side="left", color=TEAL, size="18", space="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    edge = OxmlElement(f"w:{side}")
    edge.set(qn("w:val"), "single")
    edge.set(qn("w:sz"), size)
    edge.set(qn("w:space"), space)
    edge.set(qn("w:color"), color)
    p_bdr.append(edge)


def set_image_alt(inline_shape, title, description):
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("title", title)
    doc_pr.set("descr", description)


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    label = paragraph.add_run("PAGE ")
    set_run_font(label, size=8.5, color=MUTED, bold=True)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "1"
    r.append(t)
    fld.append(r)
    paragraph._p.append(fld)


def add_kicker(doc, text):
    p = doc.add_paragraph(style="Eyebrow")
    r = p.add_run(text.upper())
    set_run_font(r, size=9, color=OCHRE, bold=True)
    return p


def add_section(doc, number, title, subtitle):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    num = p.add_run(f"{number}  ")
    set_run_font(num, size=10, color=CORAL, bold=True)
    heading = p.add_run(title)
    set_run_font(heading, size=18, color=TEAL_DARK, bold=True)
    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(8)
    r = sub.add_run(subtitle)
    set_run_font(r, size=10, color=MUTED, italic=True)


def add_question(doc, q_label, prompt, evidence_hint=None, lines=2, stretch=False):
    p = doc.add_paragraph(style="Question")
    p.paragraph_format.keep_with_next = True
    q = p.add_run(f"{q_label}  ")
    set_run_font(q, size=10, color=CORAL if not stretch else OCHRE, bold=True)
    if stretch:
        tag = p.add_run("STRETCH  ")
        set_run_font(tag, size=8.5, color=OCHRE, bold=True)
    r = p.add_run(prompt)
    set_run_font(r, size=10.5, color=INK, bold=True)
    if evidence_hint:
        hint = doc.add_paragraph(style="Hint")
        hint.paragraph_format.keep_with_next = True
        hr = hint.add_run(evidence_hint)
        set_run_font(hr, size=9, color=MUTED, italic=True)
    add_response_lines(doc, lines)


def add_response_lines(doc, count=2):
    for _ in range(count):
        p = doc.add_paragraph(style="Response Line")
        p.paragraph_format.tab_stops.add_tab_stop(
            Inches(7.0), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.LINES
        )
        r = p.add_run("\t")
        set_run_font(r, size=9, color=LINE)


def add_labeled_response(doc, label, lines=1):
    p = doc.add_paragraph(style="Response Label")
    p.paragraph_format.keep_with_next = True
    r = p.add_run(label)
    set_run_font(r, size=9.5, color=TEAL, bold=True)
    add_response_lines(doc, lines)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.62)
section.bottom_margin = Inches(0.58)
section.left_margin = Inches(0.75)
section.right_margin = Inches(0.75)
section.header_distance = Inches(0.28)
section.footer_distance = Inches(0.28)

styles = doc.styles
normal = styles["Normal"]
set_style_font(normal)
normal.font.size = Pt(10.5)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.18

for name, size, color, bold, before, after in [
    ("Heading 1", 18, TEAL_DARK, True, 12, 5),
    ("Heading 2", 14, TEAL, True, 9, 4),
    ("Heading 3", 11.5, TEAL_DARK, True, 7, 3),
]:
    st = styles[name]
    set_style_font(st)
    st.font.size = Pt(size)
    st.font.color.rgb = RGBColor.from_string(color)
    st.font.bold = bold
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

custom_styles = {
    "Eyebrow": (9, OCHRE, True, 3, 1.0),
    "Question": (10.5, INK, True, 3, 1.12),
    "Hint": (9, MUTED, False, 2, 1.05),
    "Response Line": (10, INK, False, 4, 1.0),
    "Response Label": (9.5, TEAL, True, 1, 1.0),
    "Small Note": (8.5, MUTED, False, 3, 1.05),
}
for name, (size, color, bold, after, spacing) in custom_styles.items():
    st = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(st)
    st.font.size = Pt(size)
    st.font.color.rgb = RGBColor.from_string(color)
    st.font.bold = bold
    st.paragraph_format.space_before = Pt(0)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.line_spacing = spacing

footer = section.footer
fp = footer.paragraphs[0]
fp.paragraph_format.space_before = Pt(0)
add_page_field(fp)

# Page 1: workshop-agenda-inspired title stack and first evidence hunt.
add_kicker(doc, "Student response booklet • Years 5–6")

title = doc.add_paragraph()
title.paragraph_format.space_before = Pt(0)
title.paragraph_format.space_after = Pt(2)
r = title.add_run("Persuasion Detective")
set_run_font(r, size=27, color=TEAL_DARK, bold=True)

subtitle = doc.add_paragraph()
subtitle.paragraph_format.space_after = Pt(8)
r = subtitle.add_run("Investigating Let Every Reader Find Their Voice")
set_run_font(r, size=13, color=CORAL, bold=True)

identity = doc.add_paragraph()
identity.paragraph_format.space_after = Pt(8)
r = identity.add_run("Name: ____________________________    Class: ______________    Date: ______________")
set_run_font(r, size=9.5, color=MUTED, bold=True)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(5)
shape = p.add_run().add_picture(str(IMAGE), width=Inches(7.0))
set_image_alt(shape, "Students with a Storytime Dog", "Students choose books in a school library beside a trained dog and its adult handler.")

goal = doc.add_paragraph()
goal.paragraph_format.left_indent = Inches(0.16)
goal.paragraph_format.right_indent = Inches(0.12)
goal.paragraph_format.space_before = Pt(4)
goal.paragraph_format.space_after = Pt(9)
shade(goal, PALE_TEAL)
border(goal, side="left", color=TEAL, size="24", space="9")
r = goal.add_run("LEARNING GOAL  ")
set_run_font(r, size=9, color=TEAL, bold=True)
r = goal.add_run("Identify persuasive choices, explain how they influence the reader, and evaluate which choices are most effective.")
set_run_font(r, size=10.2, color=TEAL_DARK, bold=True)

guide = doc.add_paragraph()
guide.paragraph_format.space_after = Pt(10)
r = guide.add_run("Use this pattern in every answer:  ")
set_run_font(r, size=9.5, color=INK, bold=True)
for label, text, color in [
    ("FIND", "quote or name the feature", CORAL),
    ("EXPLAIN", "describe its effect", TEAL),
    ("CONNECT", "link the effect to the audience or purpose", OCHRE),
]:
    r = guide.add_run(f"   {label} — ")
    set_run_font(r, size=9.3, color=color, bold=True)
    r = guide.add_run(text)
    set_run_font(r, size=9.3, color=MUTED)

add_section(doc, "01", "Find the argument", "Start with audience, purpose and the writer's central position.")

add_question(doc, "Q1", "Who is the intended audience? What decision does the writer want this audience to make?", lines=2)
add_question(doc, "Q2", "Copy the thesis statement from the opening paragraph. Then restate it in your own words.", evidence_hint="Look for the sentence that names the proposed action and the three main reasons.", lines=3)
add_question(doc, "Q3", "How does the opening scene — “Picture the library after lunch” — work as a hook? Explain what the reader is encouraged to imagine or feel.", lines=3)

doc.add_page_break()

# Page 2: structure and language choices.
add_section(doc, "02", "Zoom in on language", "Look closely at the words and images chosen to position the reader.")

add_question(doc, "Q4", "The argument is divided by headings. List the four main stages and explain how this order helps the argument build.", evidence_hint="Use the headings beginning with Confidence, Calm minds, Safety and Start small.", lines=4)

add_question(doc, "Q5", "The opening describes the dog as offering “no frown, no interruption, no judgement.” Name the technique and explain the effect of this three-part pattern.", lines=3)

add_question(doc, "Q6", "Explain the figurative comparison: “reading aloud can feel like walking onto a stage without rehearsing.” What does this help the audience understand about a hesitant reader?", lines=3)

add_question(doc, "Q7", "Find two examples of evaluative or emotive vocabulary. For each example, explain whether it makes the proposal sound valuable, safe or necessary.", lines=1)
add_labeled_response(doc, "Example 1 + effect", 2)
add_labeled_response(doc, "Example 2 + effect", 2)

add_question(doc, "Q8", "Find one example of inclusive language, such as “our,” “we” or “let us.” How does it make the School Council and school community feel involved?", lines=3)

doc.add_page_break()

# Page 3: evidence, modality and counterargument.
add_section(doc, "03", "Test the reasoning", "A strong argument does more than make claims: it uses evidence and responds fairly to other views.")

add_question(doc, "Q9", "The writer often uses cautious modal language such as “could” and “may.” Find three examples. Why is cautious language more trustworthy here than claiming the program “will” solve every problem?", lines=4)

add_question(doc, "Q10", "Identify two authoritative sources used in the reading. What claim does each source support?", lines=1)
add_labeled_response(doc, "Source 1 + supported claim", 2)
add_labeled_response(doc, "Source 2 + supported claim", 2)

add_question(doc, "Q11", "Why does the writer admit that the 2022 study was small and that the review found limited evidence? Explain how this concession can strengthen, rather than weaken, the writer's credibility.", lines=4)

add_question(doc, "Q12", "Break down the paragraph beginning “Of course, some families may worry...”", lines=0)
add_labeled_response(doc, "The counterargument or concern", 1)
add_labeled_response(doc, "The evidence used to respond", 2)
add_labeled_response(doc, "The rebuttal — why the proposal can still work", 2)

add_question(doc, "Q13", "What job does each connective perform: “First,” “Just as importantly,” “Of course,” and “Therefore”? Explain how these words guide the reader through the argument.", lines=4)

doc.add_page_break()

# Page 4: sentence craft and evaluation.
add_section(doc, "04", "Judge the persuasive impact", "Now decide which choices make this exemplar successful.")

add_question(doc, "Q14", "Compare the short sentence “That honest caution matters.” with one longer complex sentence in the same paragraph. How does varying sentence length control emphasis and pace?", lines=4)

add_question(doc, "Q15", "The conclusion says: “Let us give them a calm listener, a fair chance and one more reason to turn the page.” Identify two persuasive techniques in this sentence and explain why it is an effective final appeal.", lines=4)

add_question(doc, "Q16", "Which persuasive technique is most effective in this reading? Support your judgement with a quotation and explain its effect on the School Council.", evidence_hint="Use the full pattern: technique → evidence → effect → audience/purpose.", lines=5)

add_question(doc, "Q17", "Rewrite the bare assertion “A Storytime Dog is a good idea” so that it sounds like it belongs in this A-standard exemplar. Include precise vocabulary, a reason and appropriately cautious modality.", lines=4, stretch=True)

self_check = doc.add_paragraph()
self_check.paragraph_format.space_before = Pt(7)
self_check.paragraph_format.space_after = Pt(5)
shade(self_check, PALE_OCHRE)
border(self_check, side="left", color=OCHRE, size="24", space="9")
r = self_check.add_run("SELF-CHECK  ")
set_run_font(r, size=9, color=OCHRE, bold=True)
r = self_check.add_run("□ I used quotations.   □ I named techniques.   □ I explained effects.   □ I linked my ideas to the audience and purpose.")
set_run_font(r, size=9.5, color=INK, bold=True)

note = doc.add_paragraph(style="Small Note")
note.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = note.add_run("Companion activity for the persuasive exemplar Let Every Reader Find Their Voice.")
set_run_font(r, size=8.5, color=MUTED, italic=True)

doc.core_properties.title = "Persuasion Detective - Student Questions"
doc.core_properties.subject = "Student response activity for a persuasive exemplar"
doc.core_properties.author = "Classroom activity"
doc.core_properties.keywords = "persuasive techniques, student questions, Year 5, Year 6, exemplar"

doc.save(OUT)
print(OUT)
