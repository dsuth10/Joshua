from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image


BASE = Path(__file__).resolve().parent
OUT = BASE / "Storytime Dog - A Standard Persuasive Exemplar.docx"
HERO = BASE / "assets" / "storytime-dog-hero.png"
LIBRARY = BASE / "assets" / "storytime-dog-library.png"
HERO_CROP = BASE / "assets" / "storytime-dog-hero-crop.png"
LIBRARY_CROP = BASE / "assets" / "storytime-dog-library-crop.png"

# Resolved design system: narrative_proposal preset with a named
# "student editorial" override for wider content, left-aligned prose,
# larger body type, magazine colours and image-led page furniture.
TEAL = "0D5C63"
TEAL_DARK = "173F43"
OCHRE = "D89A2B"
CORAL = "D9654B"
CREAM = "F7F2E8"
PALE_TEAL = "E7F2F1"
INK = "203034"
MUTED = "657478"
WHITE = "FFFFFF"


def crop_landscape(source, target, target_ratio, focus_x=0.5, focus_y=0.5):
    """Create a wide editorial crop without stretching the generated art."""
    with Image.open(source) as image:
        width, height = image.size
        current = width / height
        if current < target_ratio:
            new_width = width
            new_height = int(width / target_ratio)
        else:
            new_height = height
            new_width = int(height * target_ratio)
        left = int((width - new_width) * focus_x)
        top = int((height - new_height) * focus_y)
        left = max(0, min(left, width - new_width))
        top = max(0, min(top, height - new_height))
        image.crop((left, top, left + new_width, top + new_height)).save(target)


crop_landscape(HERO, HERO_CROP, 2.42, focus_x=0.55, focus_y=0.48)
crop_landscape(LIBRARY, LIBRARY_CROP, 3.0, focus_x=0.5, focus_y=0.52)


def set_cell_or_paragraph_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_paragraph_border(paragraph, side="left", color=OCHRE, size="20", space="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    edge = OxmlElement(f"w:{side}")
    edge.set(qn("w:val"), "single")
    edge.set(qn("w:sz"), size)
    edge.set(qn("w:space"), space)
    edge.set(qn("w:color"), color)
    p_bdr.append(edge)


def set_run_font(run, name="Calibri", size=None, color=INK, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr()
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_repeat_font(style, name="Calibri"):
    style.font.name = name
    style._element.get_or_add_rPr()
    style._element.rPr.rFonts.set(qn("w:ascii"), name)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), name)


def set_image_alt(inline_shape, title, description):
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("title", title)
    doc_pr.set("descr", description)


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("PAGE ")
    set_run_font(run, size=8.5, color=MUTED, bold=True)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "1"
    r.append(t)
    fld.append(r)
    paragraph._p.append(fld)


def add_label(doc, text, color=TEAL):
    p = doc.add_paragraph(style="Eyebrow")
    r = p.add_run(text.upper())
    set_run_font(r, size=9, color=color, bold=True)
    return p


def add_body(doc, text, first=False):
    p = doc.add_paragraph(style="Body Text")
    if first:
        p.paragraph_format.first_line_indent = Inches(0)
    r = p.add_run(text)
    set_run_font(r, size=10.8, color=INK)
    return p


def add_source(doc, number, title, details, url):
    p = doc.add_paragraph(style="Source Note")
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    n = p.add_run(f"[{number}] ")
    set_run_font(n, size=8.5, color=CORAL, bold=True)
    hyperlink = OxmlElement("w:hyperlink")
    relation_id = p.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink.set(qn("r:id"), relation_id)
    link_run = OxmlElement("w:r")
    link_props = OxmlElement("w:rPr")
    link_color = OxmlElement("w:color")
    link_color.set(qn("w:val"), TEAL_DARK)
    link_props.append(link_color)
    link_bold = OxmlElement("w:b")
    link_props.append(link_bold)
    link_run.append(link_props)
    link_text = OxmlElement("w:t")
    link_text.text = title
    link_run.append(link_text)
    hyperlink.append(link_run)
    p._p.append(hyperlink)
    d = p.add_run(f" — {details}")
    set_run_font(d, size=8.5, color=MUTED)
    return p


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.62)
section.bottom_margin = Inches(0.58)
section.left_margin = Inches(0.75)
section.right_margin = Inches(0.75)
section.header_distance = Inches(0.28)
section.footer_distance = Inches(0.28)

styles = doc.styles
normal = styles["Normal"]
set_repeat_font(normal)
normal.font.size = Pt(10.8)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(7)
normal.paragraph_format.line_spacing = 1.18

body = styles["Body Text"]
set_repeat_font(body)
body.font.size = Pt(10.8)
body.font.color.rgb = RGBColor.from_string(INK)
body.paragraph_format.space_before = Pt(0)
body.paragraph_format.space_after = Pt(7)
body.paragraph_format.line_spacing = 1.18
body.paragraph_format.widow_control = True

for style_name, size, color, before, after in [
    ("Heading 1", 16, TEAL_DARK, 12, 4),
    ("Heading 2", 13, TEAL, 9, 3),
    ("Heading 3", 11.5, TEAL_DARK, 7, 3),
]:
    st = styles[style_name]
    set_repeat_font(st)
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(color)
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

for custom_name, size, color, bold, after in [
    ("Eyebrow", 9, TEAL, True, 3),
    ("Standfirst", 12.5, TEAL_DARK, True, 8),
    ("Pull Quote", 15, TEAL_DARK, True, 9),
    ("Caption", 8.5, MUTED, False, 6),
    ("Source Note", 8.5, MUTED, False, 3),
]:
    if custom_name not in styles:
        st = styles.add_style(custom_name, WD_STYLE_TYPE.PARAGRAPH)
    else:
        st = styles[custom_name]
    set_repeat_font(st)
    st.font.size = Pt(size)
    st.font.color.rgb = RGBColor.from_string(color)
    st.font.bold = bold
    st.paragraph_format.space_before = Pt(0)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.line_spacing = 1.12

# Quiet running footer; the body opens with its own clear reading label.
footer = section.footer
fp = footer.paragraphs[0]
fp.paragraph_format.space_before = Pt(0)
add_page_field(fp)

# Page 1: editorial opening.
add_label(doc, "A persuasive reading for Years 5–6", OCHRE)

title = doc.add_paragraph()
title.paragraph_format.space_before = Pt(0)
title.paragraph_format.space_after = Pt(3)
title.paragraph_format.keep_with_next = True
r = title.add_run("Let Every Reader\nFind Their Voice")
set_run_font(r, size=27.5, color=TEAL_DARK, bold=True)

subtitle = doc.add_paragraph()
subtitle.paragraph_format.space_after = Pt(7)
subtitle.paragraph_format.keep_with_next = True
r = subtitle.add_run("Why our school should trial a Storytime Dog program")
set_run_font(r, size=14, color=CORAL, bold=True)

meta = doc.add_paragraph()
meta.paragraph_format.space_after = Pt(8)
r = meta.add_run("A persuasive proposal to the School Council  •  386 words")
set_run_font(r, size=9, color=MUTED, bold=True)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
shape = p.add_run().add_picture(str(HERO_CROP), width=Inches(7.0))
set_image_alt(shape, "Student reading to a Storytime Dog", "A student reads aloud to a calm trained dog in a school library while an adult handler supervises.")

cap = doc.add_paragraph(style="Caption")
cap.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = cap.add_run("A calm listener can make a brave reader.")
set_run_font(r, size=8.5, color=MUTED, italic=True)

stand = doc.add_paragraph(style="Standfirst")
stand.paragraph_format.left_indent = Inches(0.18)
stand.paragraph_format.right_indent = Inches(0.12)
stand.paragraph_format.space_before = Pt(3)
stand.paragraph_format.space_after = Pt(9)
set_paragraph_border(stand, side="left", color=OCHRE, size="24", space="9")
r = stand.add_run("A dog cannot correct a tricky word — but it can help a nervous reader find the courage to try it.")
set_run_font(r, size=12.5, color=TEAL_DARK, bold=True)

add_body(doc, "Picture the library after lunch. A student opens a book, takes a breath and begins to read. Beside them, a trained dog waits — no frown, no interruption, no judgement. Our school should run a one-term Storytime Dog trial because it could strengthen reading confidence, support wellbeing and help more students feel that the library belongs to them.", first=True)

# Continue page 1 with the first developed reason.
add_label(doc, "The case for a one-term trial", OCHRE)
h = doc.add_paragraph("Confidence grows through practice", style="Heading 1")

add_body(doc, "First, reading improves when students are willing to practise. For a hesitant reader, reading aloud can feel like walking onto a stage without rehearsing. A calm dog changes the audience. In a 2022 study of 24 young readers who needed extra support, researchers found evidence of improved reading performance, particularly after children read to a dog.[1] The study was small, so it does not prove that every child will improve; however, it gives our school a sensible reason to test the idea carefully. A weekly session could turn “I can’t” into “I’ll try.”")

doc.add_page_break()

# Page 2: evidence, response to concerns, action and transparent sources.
add_label(doc, "Evidence, empathy and a fair response", OCHRE)
h = doc.add_paragraph("Calm minds are ready to learn", style="Heading 1")
add_body(doc, "Just as importantly, a Storytime Dog could make reading feel safer. A systematic review found promising links between reading to dogs and greater motivation and confidence, as well as reduced anxiety, although the researchers warned that the overall evidence was still limited.[2] That honest caution matters. We should not promise a miracle. We should offer a welcoming, well-supervised space where students can practise, connect and succeed one page at a time.")

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(2)
p.paragraph_format.space_after = Pt(2)
shape = p.add_run().add_picture(str(LIBRARY_CROP), width=Inches(7.0))
set_image_alt(shape, "Students choosing books with a Storytime Dog", "Two students choose books in a school library beside a calm trained dog and its adult handler.")

cap = doc.add_paragraph(style="Caption")
cap.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = cap.add_run("The goal is not a novelty. It is a safe routine that invites students into reading.")
set_run_font(r, size=8.5, color=MUTED, italic=True)

h = doc.add_paragraph("Safety must come first", style="Heading 1")
add_body(doc, "Of course, some families may worry about allergies, fear or distraction. These concerns are valid — and they are exactly why the program must be planned, not improvised. NSW Department of Education guidance recommends notifying families, identifying allergies or fears, providing handwashing, using a handler and completing a risk assessment.[3] Students who do not wish to participate must have an equally supportive reading option. With clear boundaries, a rest space for the dog and short scheduled sessions, safety and inclusion can guide every decision.")

h = doc.add_paragraph("Start small. Measure carefully. Decide together.", style="Heading 1")
add_body(doc, "Therefore, the School Council should approve a one-term trial for a small group of volunteer readers. Teachers could track attendance, reading confidence and student feedback before reviewing the results with families. A Storytime Dog will not replace skilled teaching, patient adults or daily practice. It could, however, open the door for students who are still waiting to see themselves as readers. Let us give them a calm listener, a fair chance and one more reason to turn the page.")

call = doc.add_paragraph()
call.paragraph_format.left_indent = Inches(0.18)
call.paragraph_format.right_indent = Inches(0.18)
call.paragraph_format.space_before = Pt(8)
call.paragraph_format.space_after = Pt(10)
set_cell_or_paragraph_shading(call, PALE_TEAL)
set_paragraph_border(call, side="left", color=TEAL, size="28", space="10")
r = call.add_run("THE PROPOSAL  ")
set_run_font(r, size=9, color=TEAL, bold=True)
r = call.add_run("Approve a one-term, opt-in Storytime Dog trial with a trained handler, a documented risk assessment and a review of student outcomes.")
set_run_font(r, size=10.5, color=TEAL_DARK, bold=True)

doc.add_paragraph("Evidence used in this exemplar", style="Heading 2")
add_source(doc, 1, "Supporting Young Readers", "A 2022 mixed-methods study of 24 children aged 7–8.", "https://pubmed.ncbi.nlm.nih.gov/36312220/")
add_source(doc, 2, "Children Reading to Dogs", "A 2016 systematic review that found promising but low-quality evidence.", "https://pmc.ncbi.nlm.nih.gov/articles/PMC4763282/")
add_source(doc, 3, "Support Dog Guidelines", "NSW Department of Education guidance on risk, hygiene, allergies, supervision and animal welfare.", "https://education.nsw.gov.au/teaching-and-learning/animals-in-schools/animals-in-schools-species/dogs/dogs-introduction/support-dog-guidelines")

note = doc.add_paragraph(style="Source Note")
note.paragraph_format.space_before = Pt(7)
set_paragraph_border(note, side="top", color="D6E4E3", size="6", space="5")
r = note.add_run("Model-text note: ")
set_run_font(r, size=8.3, color=TEAL_DARK, bold=True)
r = note.add_run("This is an original classroom exemplar. The argument uses qualified claims so students can see that trustworthy persuasion is strong, specific and honest about the evidence.")
set_run_font(r, size=8.3, color=MUTED)

# Document metadata and core properties.
doc.core_properties.title = "Let Every Reader Find Their Voice"
doc.core_properties.subject = "A Standard persuasive exemplar for Years 5–6"
doc.core_properties.author = "Classroom exemplar"
doc.core_properties.keywords = "persuasive writing, exemplar, Storytime Dog, Year 5, Year 6"

doc.save(OUT)
print(OUT)
