"""Build the printable resources and original reading article for Lesson 23 Alternative."""

from pathlib import Path
import importlib.util

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
BASE_PATH = ROOT.parents[1] / "Lesson_17" / "Turtle_Safe_Nights_Alternative" / "scripts" / "build_turtle_safe_nights.py"
spec = importlib.util.spec_from_file_location("lesson_pack_base", BASE_PATH)
b = importlib.util.module_from_spec(spec)
spec.loader.exec_module(b)
b.ROOT = ROOT

# compact_reference_guide with a school-library palette; named A4 classroom override.
b.TOKENS.update({
    "override": "borrow_box_a4_classroom",
    "navy": "26364C", "deep_sea": "314B68", "teal": "497C79", "moon": "E5B85B",
    "coral": "B96559", "pale": "EAF4F0", "warm": "FFF2D8", "mist": "EEF1F5",
    "ink": "26313D", "muted": "65717E", "border": "B8C5D0",
})


def add_footer(section, label):
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    b.style_paragraph(p, after=0, line=1.0)
    r = p.add_run("ENGLISH UNIT 3  |  LESSON 23 ALTERNATIVE")
    b.set_run(r, size=7.5, bold=True, color=b.TOKENS["muted"])
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    b.style_paragraph(p, after=0, line=1.0)
    r = p.add_run(label + "  |  The Borrow Box")
    b.set_run(r, size=7.5, color=b.TOKENS["muted"])


b.add_footer = add_footer


def bullet(doc, text, lead=None):
    p = doc.add_paragraph(style="List Bullet")
    b.style_paragraph(p, after=4)
    if lead and text.startswith(lead):
        r = p.add_run(lead)
        b.set_run(r, size=b.TOKENS["body_size"], bold=True, color=b.TOKENS["deep_sea"])
        r = p.add_run(text[len(lead):])
        b.set_run(r, size=b.TOKENS["body_size"], color=b.TOKENS["ink"])
    else:
        r = p.add_run(text)
        b.set_run(r, size=b.TOKENS["body_size"], color=b.TOKENS["ink"])
    return p


def build_reading_pack():
    doc = b.configure_doc("Student Reading Pack")
    b.add_title_block(doc, "Original classroom feature", "The Borrow Box", "A good lunchtime idea still has to answer fair questions.", "READ  |  MARK C / ? / R / X")
    b.add_callout(doc, "Decision:", "Should the fictional Riverbank Primary trial a Borrow Box: shared lunchtime equipment that students can sign out and return?", fill=b.TOKENS["pale"], accent=b.TOKENS["teal"])
    b.add_labeled_para(doc, "Fiction boundary:", "Riverbank Primary, all people, quotations and school details are invented for this lesson. Read the feature as a shared persuasive text, not a report about a real school.", fill=b.TOKENS["warm"], accent=b.TOKENS["coral"])

    b.add_heading(doc, "The empty crate beside the oval", 1)
    b.add_text(doc, "At the first Student Voice meeting of term, Imani placed an empty blue crate in the middle of the table. 'This is the Borrow Box,' she said. 'Not yet, anyway.' Her plan was simple: at the beginning of lunch, students could borrow skipping ropes, chalk, soft balls, card games and drawing boards. At the end, they would return each item to the labelled crate.")
    b.add_text(doc, "Imani had noticed that the basketball court filled quickly while other students waited near the fence or sat on the steps. Some wanted an active game but did not have equipment. Others wanted a quieter choice. She believed a small shared collection could make lunch feel more welcoming without telling anyone what kind of play they had to choose.")
    b.add_callout(doc, "Imani's claim:", "Riverbank Primary should run a four-week Borrow Box trial at lunch and review what it teaches us.", fill=b.TOKENS["warm"], accent=b.TOKENS["moon"])
    b.add_text(doc, "Her first poster ended with: 'Anyone against the Borrow Box wants students to be bored.' Mr Patel, the teacher helping Student Voice, asked her to pause. The sentence made disagreement sound selfish. It did not tell readers what a careful person might genuinely worry about.")

    b.add_heading(doc, "The questions behind the yes", 1)
    b.add_text(doc, "Jayden, who had helped pack away sports gear after a carnival, liked the idea but pointed to the blue crate. 'What happens when a ball disappears, or when one group keeps the same things all lunch?' he asked. A quiet student, Elif, added another question: 'Will the box only have loud games? Sometimes I want something to do with one friend.'")
    b.add_text(doc, "Mr Patel did not call these students negative. Their questions were useful because they named what a proposal would need to solve: shared equipment must be returned, access needs to be fair, and a helpful option should not create a new exclusion. These were counterarguments - fair concerns that an audience could hold even when they liked the purpose of the idea.")
    b.add_callout(doc, "Mark the feature:", "C = claim or action | ? = genuine concern | R = response, reason or safeguard | X = unfair dismissal or assumed motive", fill=b.TOKENS["pale"], accent=b.TOKENS["teal"])

    doc.add_page_break()
    b.add_heading(doc, "A response that fits the worry", 1)
    b.add_text(doc, "Imani rewrote the poster. She began by admitting that shared equipment could be lost and that a busy collection line could waste lunch time. She explained why this mattered: a system that causes frustration or uses money carelessly would not be fair to students or staff. Then she offered a response instead of a slogan.")
    b.add_callout(doc, "Imani's revised response:", "Some people may worry that equipment will be lost or that a queue will take too long. Those are sensible concerns because a shared system must be fair and easy to use. However, a labelled tub, a quick sign-out card and a five-minute collection window would let us test a routine instead of guessing. The Borrow Box should begin as a four-week trial, with an end-of-lunch count and a record of waiting time.", fill=b.TOKENS["pale"], accent=b.TOKENS["teal"])
    b.add_text(doc, "The revised response does not promise that every item will return or that every student will enjoy the same activity. It offers a small action that can be checked. The group also planned to include chalk and drawing boards alongside active equipment, and to ask students which choices were missing. A good rebuttal answers the concern it has acknowledged; it does not pretend the concern disappeared.")

    b.add_heading(doc, "Three responses - only one is a rebuttal", 1)
    b.add_matrix(doc, ["Response", "What it does"], [
        ["'Fun is more important than waiting.'", "Repeats the original opinion. It does not answer how waiting could be reduced."],
        ["'People who complain about queues never support good ideas.'", "Attacks people and assumes motives. This is not a fair counterargument or rebuttal."],
        ["'A five-minute collection window and class roster could limit waiting time; the trial can record whether that works.'", "Answers the queue concern with a relevant, testable safeguard."],
    ], [4800, 5378], font_size=9.2)

    b.add_heading(doc, "The writer's bridge", 1)
    b.add_matrix(doc, ["Move", "Question", "Useful language"], [
        ["Claim", "What do I recommend?", "The school should..."],
        ["Acknowledge", "What could a careful audience worry about?", "Some people may worry... This matters because..."],
        ["Connect", "What reason, detail or safeguard answers that worry?", "However, we could... This would..."],
        ["Keep the claim", "What measured action should happen next?", "Therefore, the school should trial... and review..."],
    ], [1800, 3500, 4878], font_size=8.9)
    b.add_callout(doc, "Letter challenge:", "Write to the fictional principal. Recommend, change or reject the Borrow Box trial. Whatever you decide, represent one fair concern and answer it with a reasoned, workable response.", fill=b.TOKENS["warm"], accent=b.TOKENS["coral"])

    doc.add_page_break()
    b.add_heading(doc, "Counterargument field guide", 1)
    b.add_matrix(doc, ["Term", "What it means", "Quick test"], [
        ["Counterargument", "A fair concern or opposing view that an audience could genuinely hold.", "Would the other person recognise this as their real concern?"],
        ["Straw man", "A weak, exaggerated or invented version of the other view.", "Have I made the concern easier to knock down?"],
        ["Rebuttal", "A reasoned answer that addresses the exact concern.", "Does my response match the concern word for word or idea for idea?"],
        ["Safeguard", "A practical detail that makes an action safer, fairer or more workable.", "Could someone actually do this during the trial?"],
    ], [2100, 4750, 3328], font_size=9.0)
    b.add_heading(doc, "Independent reading check", 1)
    for item in [
        "Copy Imani's claim in your own words.",
        "Record one fair concern from Jayden or Elif.",
        "Underline the exact safeguard that answers the queue concern.",
        "Explain why Imani's first poster sentence is unfair.",
        "Write one question that would make the Borrow Box proposal stronger.",
    ]:
        bullet(doc, item)
    b.add_labeled_para(doc, "Source note:", "This is an original fictional classroom feature. Its purpose is to provide a shared, bounded text for analysing counterarguments and rebuttals. No outside factual claims are required for the lesson.", fill=b.TOKENS["mist"], accent=b.TOKENS["deep_sea"])
    out = ROOT / "Lesson_23_Borrow_Box_Reading_Pack.docx"
    doc.save(out)
    return out


def build_organiser():
    doc = b.configure_doc("ACK Organiser")
    b.add_title_block(doc, "Persuasive writing organiser", "The ACK Bridge", "Hear the real worry. Answer it. Keep your claim.", "PLAN  |  WRITE  |  REVISE")
    b.add_heading(doc, "1. Name the decision and my claim", 1)
    b.add_form_table(doc, ["Decision", "My claim"], [4400, 5778], blank_rows=2, row_prompts=[["What should happen?", "The school should / should not..."], ["Audience:", "Principal / school community"]], blank_after=30)
    b.add_heading(doc, "2. Find a fair concern", 1)
    b.add_matrix(doc, ["Possible concern", "Fair concern?", "Why might it matter?"], [
        ["Equipment could be lost or damaged.", "", ""],
        ["People who disagree do not care about fun.", "", ""],
        ["The queue could take too much lunch time.", "", ""],
        ["Quiet activities could be left out.", "", ""],
    ], [4700, 1800, 3678], font_size=8.9)
    b.add_labeled_para(doc, "My strongest fair concern:", "________________________________________________________________________________", fill=b.TOKENS["pale"], accent=b.TOKENS["teal"])
    b.add_labeled_para(doc, "Why it matters:", "________________________________________________________________________________", fill=b.TOKENS["warm"], accent=b.TOKENS["coral"])

    doc.add_page_break()
    b.add_heading(doc, "3. Build the ACK bridge", 1)
    b.add_form_table(doc, ["Move", "My words", "Precision check"], [1850, 5300, 3028], blank_rows=3, row_prompts=[
        ["A - Acknowledge", "Some people may worry...", "Would they agree this is their concern?"],
        ["C - Connect", "However... / A useful safeguard is...", "Does it answer that exact worry?"],
        ["K - Keep", "Therefore, the school should...", "Is the next action measured and workable?"],
    ], blank_after=30)
    b.add_heading(doc, "4. Evidence / reason / safeguard bank", 1)
    b.add_form_table(doc, ["Text detail or practical action", "Which concern does it answer?", "What could be measured?"], [3800, 3500, 2878], blank_rows=4, blank_after=30)
    b.add_callout(doc, "Match test:", "Put a line from your counterargument to the exact sentence that answers it. If the line does not make sense, revise the rebuttal.")

    doc.add_page_break()
    b.add_heading(doc, "5. Draft the principal letter", 1)
    for prompt in [
        "Dear Principal, I believe ________________________________________________.",
        "Some people may worry that ______________________________________________.",
        "This is a fair concern because ____________________________________________.",
        "However, ________________________________________________________________.",
        "A practical safeguard would be ___________________________________________.",
        "Therefore, the school should ______________________________________________.",
        "The trial should record / review ___________________________________________.",
    ]:
        b.add_labeled_para(doc, "Write:", prompt, fill=b.TOKENS["pale"], accent=b.TOKENS["teal"])

    doc.add_page_break()
    b.add_heading(doc, "6. Stress-test and revise", 1)
    b.add_form_table(doc, ["Partner's fair question", "My revised sentence"], [5089, 5089], blank_rows=2, row_prompts=[
        ["Does your concern sound real to the other side?", "Original sentence:"],
        ["What practical action would answer it more clearly?", "Revised sentence:"],
    ], blank_after=48)
    b.add_heading(doc, "7. Final self-check", 1)
    for item in [
        "[ ] My claim is clear.", "[ ] My counterargument is fair, not an insult.",
        "[ ] I explain why the concern matters.", "[ ] My rebuttal answers the exact concern.",
        "[ ] I name a safeguard, measure or review.", "[ ] I return to my claim.", "[ ] I made one revision visible.",
    ]:
        bullet(doc, item)
    b.add_callout(doc, "Exit:", "A counterargument strengthens my writing when ______________________________________________ because ______________________________________________.")
    out = ROOT / "Lesson_23_Borrow_Box_ACK_Organiser.docx"
    doc.save(out)
    return out


def build_concise_access_pack():
    doc = b.configure_doc("Concise Access Pack")
    normal = doc.styles["Normal"]
    normal.font.size = Pt(13)
    normal.paragraph_format.line_spacing = 1.25
    b.add_title_block(doc, "Concise persuasive pathway", "The Borrow Box", "Same big idea: hear a real worry, then give a helpful answer.", "READ TOGETHER  |  POINT / SAY / WRITE")
    b.add_heading(doc, "The school idea", 1)
    b.add_text(doc, "Imani wants the school to try a Borrow Box at lunch. Students could borrow balls, chalk, skipping ropes, card games and drawing boards. They would return the things when lunch finishes. Imani thinks more students could have something to do.", size=13)
    b.add_heading(doc, "Real worries", 1)
    b.add_matrix(doc, ["A real worry", "A helpful action"], [
        ["Things could be lost.", "Use a labelled tub and count the things at the end."],
        ["The queue could be too long.", "Use a short collection time and record the wait."],
        ["Quiet students could be left out.", "Include drawing boards and card games too."],
    ], [4100, 6078], font_size=12)
    b.add_callout(doc, "Important:", "A real worry is not someone being mean. We can hear it and still keep our idea.")
    b.add_heading(doc, "Say the ACK bridge", 1)
    b.add_form_table(doc, ["Move", "Choose / point / say"], [2500, 7678], blank_rows=3, row_prompts=[
        ["A - Hear the worry", "Some people worry about: lost things / waiting / being left out"],
        ["C - Helpful answer", "We could: count things / use a roster / include quiet choices"],
        ["K - Say the idea again", "So the school should: try it for four weeks and review it"],
    ], blank_after=32, font_size=12)

    doc.add_page_break()
    b.add_heading(doc, "Build my short letter", 1)
    for prompt in [
        "I think the school should ________________________________________________.",
        "Some people worry about __________________________________________________.",
        "That matters because ______________________________________________________.",
        "We could ________________________________________________________________.",
        "So the school should try _________________________________________________.",
        "We could count / check ___________________________________________________.",
    ]:
        b.add_labeled_para(doc, "Say / write:", prompt, fill=b.TOKENS["pale"], accent=b.TOKENS["teal"])
    b.add_heading(doc, "Make one sentence fairer", 1)
    b.add_matrix(doc, ["Not fair", "Fairer"], [["People who disagree are boring.", "Some people worry the equipment could be lost. We could count it."]], [5089, 5089], font_size=12)
    b.add_callout(doc, "You are finished when:", "You have one idea, one real worry, one helpful action and you say your idea again.")
    out = ROOT / "Lesson_23_Borrow_Box_Concise_Access_Pack.docx"
    doc.save(out)
    return out


def build_article_html():
    html = r'''<!doctype html>
<html lang="en-AU"><head><meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1"><title>The Borrow Box - Reading Article</title>
<style>
@page{size:A4;margin:0}*{box-sizing:border-box}body{margin:0;background:#dce2e7;color:#26313d;font-family:Arial,sans-serif}.page{width:210mm;height:297mm;margin:10px auto;background:white;position:relative;overflow:hidden;box-shadow:0 8px 28px #0003}.hero{height:59mm;background:linear-gradient(120deg,#26364c,#497c79 68%,#e5b85b);padding:13mm;color:#fff;position:relative}.hero:after{content:"";position:absolute;right:15mm;top:12mm;width:33mm;height:33mm;border:4mm solid #fff8;box-shadow:inset 0 0 0 4mm #e5b85b88;border-radius:4mm;transform:rotate(9deg)}.kicker{font-size:8pt;letter-spacing:.16em;text-transform:uppercase;font-weight:800;color:#ffe7af}h1{font:800 29pt/1 Georgia,serif;margin:3mm 0 2mm}.stand{width:130mm;font:700 10.5pt/1.25 Arial}.content{padding:8mm 13mm 13mm;display:grid;grid-template-columns:1.52fr .85fr;gap:7mm}h2{font:800 13pt/1.1 Georgia,serif;color:#314b68;margin:0 0 2mm;border-bottom:2px solid #e5b85b;padding-bottom:1mm}p{font-size:9.2pt;line-height:1.32;margin:0 0 2.4mm}.quote,.fact,.task{padding:3mm;border-radius:2mm;margin:3mm 0}.quote{background:#fff2d8;border-left:4px solid #b96559;font-weight:700}.fact{background:#eaf4f0;border-left:4px solid #497c79}.side{background:#eef1f5;padding:5mm;border-top:5px solid #b96559}.side h3{font-size:10pt;margin:0 0 2mm;color:#b96559}.code{display:grid;grid-template-columns:8mm 1fr;gap:1.5mm;margin:1.5mm 0;font-size:8.5pt}.badge{display:grid;place-items:center;background:#314b68;color:#fff;border-radius:2px;font-weight:800}.task{background:#26364c;color:#fff;font-size:9pt}.footer{position:absolute;left:13mm;right:13mm;bottom:5mm;border-top:1px solid #b8c5d0;padding-top:2mm;font-size:6.3pt;color:#65717e}.choices{display:grid;grid-template-columns:1fr 1fr;gap:4mm}.choice{padding:4mm;background:#eef1f5;border-top:4px solid #b96559}.choice.good{border-color:#497c79}.choice h3{margin:0 0 2mm;font-size:10pt}.ack{display:grid;grid-template-columns:repeat(3,1fr);gap:2mm}.ack div{background:#314b68;color:#fff;padding:2.5mm;border-radius:2mm;font-size:8.2pt}.ack b{display:block;color:#ffe7af;font-size:13pt}.field{display:grid;grid-template-columns:30mm 1fr;gap:2mm;border-bottom:1px solid #cbd4dd;padding:1.5mm 0;font-size:8.5pt}@media print{body{background:#fff}.page{margin:0;box-shadow:none;page-break-after:always}.page:last-child{page-break-after:auto}}</style></head><body>
<section class="page"><header class="hero"><div class="kicker">Original classroom feature | English Unit 3</div><h1>The Borrow Box</h1><div class="stand">A good lunchtime idea still has to answer fair questions.</div></header><main class="content"><article><h2>The empty crate beside the oval</h2><p>At the first Student Voice meeting of term, Imani placed an empty blue crate in the middle of the table. "This is the Borrow Box," she said. "Not yet, anyway." Her plan was simple: at the beginning of lunch, students could borrow skipping ropes, chalk, soft balls, card games and drawing boards. At the end, they would return each item to the labelled crate.</p><p>Imani had noticed that the basketball court filled quickly while other students waited near the fence or sat on the steps. Some wanted an active game but did not have equipment. Others wanted a quieter choice. She believed a small shared collection could make lunch feel more welcoming without telling anyone what kind of play they had to choose.</p><div class="quote">Riverbank Primary should run a four-week Borrow Box trial at lunch and review what it teaches us.</div><p>Her first poster ended with: "Anyone against the Borrow Box wants students to be bored." Mr Patel asked her to pause. The sentence made disagreement sound selfish. It did not tell readers what a careful person might genuinely worry about.</p><h2>The questions behind the yes</h2><p>Jayden liked the idea but pointed to the blue crate. "What happens when a ball disappears, or when one group keeps the same things all lunch?" he asked. Elif added, "Will the box only have loud games? Sometimes I want something to do with one friend."</p><p>These were counterarguments: fair concerns an audience could hold even when they liked the purpose of the idea.</p></article><aside class="side"><h3>Mark the feature</h3><div class="code"><span class="badge">C</span><span>claim or action</span></div><div class="code"><span class="badge">?</span><span>genuine concern</span></div><div class="code"><span class="badge">R</span><span>response or safeguard</span></div><div class="code"><span class="badge">X</span><span>unfair dismissal</span></div><div class="task">Find two fair concerns. Then find the sentence that makes disagreement sound selfish.</div><h3>Ask yourself</h3><p>Would the other person recognise this as their real concern?</p><p>Does the response answer the exact concern?</p></aside></main><footer class="footer">The Borrow Box is an original fictional classroom text. It is not a report about a real school.</footer></section>
<section class="page"><header class="hero"><div class="kicker">Original classroom feature | Page 2</div><h1>Answer the concern</h1><div class="stand">A rebuttal does more than repeat an opinion.</div></header><main class="content"><article><h2>A response that fits the worry</h2><p>Imani rewrote the poster. She began by admitting that shared equipment could be lost and that a busy collection line could waste lunch time. She explained why this mattered: a system that causes frustration or uses money carelessly would not be fair to students or staff.</p><div class="fact">Some people may worry that equipment will be lost or that a queue will take too long. Those are sensible concerns because a shared system must be fair and easy to use. However, a labelled tub, a quick sign-out card and a five-minute collection window would let us test a routine instead of guessing. The Borrow Box should begin as a four-week trial, with an end-of-lunch count and a record of waiting time.</div><p>The revised response does not promise that every item will return or that every student will enjoy the same activity. It offers a small action that can be checked. The group also planned to include chalk and drawing boards alongside active equipment, and to ask students which choices were missing.</p><h2>Only one reply works</h2><div class="choices"><div class="choice"><h3>Repeats the claim</h3><p>"Fun is more important than waiting."</p></div><div class="choice"><h3>Attacks people</h3><p>"People who complain about queues never support good ideas."</p></div><div class="choice good"><h3>Answers the concern</h3><p>"A five-minute collection window and class roster could limit waiting time; the trial can record whether that works."</p></div></div></article><aside class="side"><h3>The ACK bridge</h3><div class="ack"><div><b>A</b>Acknowledge the real worry.</div><div><b>C</b>Connect a reason or safeguard.</div><div><b>K</b>Keep the claim with action.</div></div><h3>Writer's test</h3><p>Put a line from the concern to your answer. Does it match?</p><div class="task">Write to the fictional principal. You may support, change or reject the trial. You must represent one fair concern and answer it with a workable response.</div></aside></main><footer class="footer">The Borrow Box | Claim -> fair concern -> answer -> returned claim</footer></section></body></html>'''
    out = ROOT / "Lesson_23_Borrow_Box_Reading_Article.html"
    out.write_text(html, encoding="utf-8")
    return out


if __name__ == "__main__":
    outputs = [build_reading_pack(), build_organiser(), build_concise_access_pack(), build_article_html()]
    for output in outputs:
        print(output)
