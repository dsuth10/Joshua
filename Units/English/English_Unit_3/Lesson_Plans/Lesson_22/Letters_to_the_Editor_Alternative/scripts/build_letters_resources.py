"""Build Lesson 22 alternative reading, organiser, access pack and article."""

from pathlib import Path
import importlib.util
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
BASE_PATH = ROOT.parents[1] / "Lesson_17" / "Turtle_Safe_Nights_Alternative" / "scripts" / "build_turtle_safe_nights.py"
spec = importlib.util.spec_from_file_location("lesson_pack_base", BASE_PATH)
b = importlib.util.module_from_spec(spec)
spec.loader.exec_module(b)
b.ROOT = ROOT


def add_footer(section, label):
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = b.WD_ALIGN_PARAGRAPH.RIGHT
    b.style_paragraph(hp, after=0, line=1.0)
    run = hp.add_run("ENGLISH UNIT 3  |  LESSON 22 ALTERNATIVE")
    b.set_run(run, size=7.5, bold=True, color=b.TOKENS["muted"])

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = b.WD_ALIGN_PARAGRAPH.CENTER
    b.style_paragraph(fp, after=0, line=1.0)
    run = fp.add_run(label + "  |  Letters from Seabreeze Square")
    b.set_run(run, size=7.5, color=b.TOKENS["muted"])


b.add_footer = add_footer

# compact_reference_guide with a restrained newspaper palette; named A4 classroom override.
b.TOKENS.update({
    "override": "seabreeze_letters_a4_classroom",
    "navy": "202F3B", "deep_sea": "334C5C", "teal": "4C7774", "moon": "D5A84B",
    "coral": "A9564D", "pale": "EAF2F0", "warm": "FFF2D6", "mist": "EEF1F3",
    "ink": "263139", "muted": "66737B", "border": "BAC6CC",
})

LETTERS = [
    (
        "A", "Stop the Friday Food-Truck Circus", "Marina Vale, Square resident", "strongly biased against",
        [
            "Seabreeze Square is about to be surrendered to a noisy food-truck circus. Reckless councillors appear more interested in colourful photographs than in the people who actually live here. Every resident knows that eight engines, queues and music will destroy the quiet of Friday evenings.",
            "The council calls this a six-week trial, but we all know what that means: once the trucks invade, they will never leave. Parking will become impossible, rubbish will tumble through the streets and local cafes will be abandoned. Nobody asked ordinary residents whether we wanted our square turned into a dining yard.",
            "Council must reject this foolish scheme before permanent damage is done. Either it protects the square now or admits that visitors matter more than residents."
        ],
    ),
    (
        "B", "Let the Town Wake Up", "Theo Grant, hospitality student", "strongly biased for",
        [
            "At last, somebody has offered to rescue our lifeless town centre. Food trucks will fill Seabreeze Square with colour, flavour and laughter. Only the usual whining few could oppose an idea this brilliant.",
            "Everyone knows the square is completely empty on Friday nights. The trial will bring families together, save struggling businesses and give young people a reason to stay in town. Critics complain about parking because they fear any change they did not invent themselves.",
            "There are only two choices: support the trucks or accept a dead town centre. Council should ignore the selfish complainers and approve the full trial immediately."
        ],
    ),
    (
        "C", "Parking First, Please", "Ravi Singh, nearby resident", "calm but slanted against",
        [
            "The proposed food-truck trial should not proceed at Seabreeze Square until council solves the parking problem. On each trial evening, twelve spaces beside the square would be unavailable for ordinary vehicles. Those spaces are used by residents and customers of nearby shops.",
            "Removing them will push cars into surrounding streets and make access difficult. The result will be inconvenience for residents and fewer customers for existing businesses. Council should choose another location before approving the trial.",
            "A lively evening program may be worthwhile, but parking must come first."
        ],
    ),
    (
        "D", "A Chance for Young People", "Lena Cho, youth committee member", "calm but slanted for",
        [
            "Seabreeze needs more places where young people and families can spend time together. A Friday food-truck trial would give the square a welcoming purpose after many shops close for the day. At our youth committee meeting, several students said they would attend with friends or family.",
            "The proposal offers different foods and a reason to meet in a public place. Events like this create connection, support new vendors and help young people feel that the town includes them. Many families will enjoy an affordable evening close to home.",
            "Council should approve the trial and show that young residents' ideas matter."
        ],
    ),
    (
        "E", "Trial It, Measure It, Adjust It", "Mara Ellis, community volunteer", "balanced conditional support",
        [
            "The food-truck proposal could bring new activity to Seabreeze Square, but nearby residents are reasonable to ask about noise, parking and waste. The draft is a six-Friday trial from 5:00 pm to 8:00 pm, not a permanent program, and it caps attendance at eight vendors.",
            "Council should proceed only with the stated accessible path and waste stations in place. During the trial it should record complaints, litter, attendance, parking pressure and feedback from nearby businesses. These measures will not tell us whether every person likes the event, but they will provide better local evidence than confident predictions from either side.",
            "I support the trial because it is limited and reviewable. If the evidence shows unacceptable problems, council should change the hours, reduce vendor numbers or choose another location."
        ],
    ),
    (
        "F", "Protect the Square While Testing the Idea", "Owen Park, access advocate", "balanced cautious position",
        [
            "Supporters are right that a shared evening event may help people gather, while residents are right that a public square must remain accessible and manageable. The current proposal includes an accessible path and a later review, but it does not yet explain how the path will be monitored when queues form.",
            "Council should begin with four trucks on the first two Fridays, mark the accessible route clearly and station a marshal at the busiest crossing. It should then publish what it observed before increasing the number of vendors. This smaller start would not remove every risk, but it would test the idea without pretending that support or concern is unreasonable.",
            "A careful trial can welcome new activity and protect ordinary use of the square."
        ],
    ),
]


def add_letter(doc, code, title, author, diagnosis, paragraphs):
    b.add_heading(doc, f"Letter {code} - {title}", 1)
    b.add_text(doc, f"From: {author}", bold=True, color=b.TOKENS["teal"], after=4)
    for text in paragraphs:
        b.add_text(doc, text)
    b.add_labeled_para(doc, "Reader task:", "Underline loaded or evaluative words; circle sweeping or certainty words; box a claim that goes beyond the editor's briefing; note one relevant omission.", fill=b.TOKENS["mist"], accent=b.TOKENS["deep_sea"])


def build_reading_pack():
    doc = b.configure_doc("Letters Reading Pack")
    b.add_title_block(doc, "Original letters to the editor", "Letters from Seabreeze Square", "Six writers. One proposal. Different ways of pulling the reader.", "READ  |  MARK  |  COMPARE")
    b.add_callout(doc, "Shared decision:", "Should fictional Seabreeze Council run a six-Friday food-truck trial at Seabreeze Square?", fill=b.TOKENS["pale"], accent=b.TOKENS["teal"])
    b.add_labeled_para(doc, "Fiction boundary:", "The council, square, proposal, writers and quotations are invented. The editor's briefing below is the only shared local fact base.", fill=b.TOKENS["warm"], accent=b.TOKENS["coral"])
    b.add_heading(doc, "Editor's verified briefing", 1)
    b.add_matrix(doc, ["What the draft says", "What it does not prove"], [[
        "Six Friday evenings; 5:00-8:00 pm; maximum eight food trucks; waste stations; accessible path; council review after the trial.",
        "Future noise, profit, parking failure, litter, community renewal, permanent continuation or universal support."
    ]], [5089, 5089], font_size=9.3)
    b.add_heading(doc, "Bias leaves five fingerprints", 1)
    b.add_matrix(doc, ["Fingerprint", "What to notice"], [
        ["Loaded evaluation", "words that praise, shame, alarm or minimise"],
        ["Sweeping generalisation", "everyone, nobody, always, every, only"],
        ["Assumed motive", "claims about why another person acts"],
        ["Omission / selection", "important detail foregrounded or left out"],
        ["Unsupported certainty / false choice", "prediction presented as fact; only two options allowed"],
    ], [3400, 6778], font_size=9.2)
    for idx, letter in enumerate(LETTERS):
        doc.add_page_break()
        add_letter(doc, *letter)
        if idx in (1, 3, 5):
            b.add_callout(doc, "Pair comparison:", {
                1: "How do A and B use similar bias fingerprints to pull in opposite directions?",
                3: "How do C and D sound calmer while still selecting one side?",
                5: "How do E and F stay balanced while making clear recommendations?",
            }[idx])
    doc.add_page_break()
    b.add_heading(doc, "Balanced-language field guide", 1)
    b.add_matrix(doc, ["Move", "Useful language", "Trust check"], [
        ["Attribute", "The writer argues / claims / predicts...", "Is a viewpoint clearly reported as a viewpoint?"],
        ["Concede", "Residents may reasonably be concerned...", "Is the other concern represented accurately?"],
        ["Qualify", "may, could, if, during the trial", "Does certainty match available evidence?"],
        ["State a boundary", "The proposal does not yet show...", "Is missing knowledge made visible?"],
        ["Recommend", "Council should trial / cap / measure / review...", "Is the position clear and workable?"],
    ], [2100, 3500, 4578], font_size=9.0)
    b.add_callout(doc, "Final challenge:", "Write an 8-10 sentence Editorial Fairness Briefing. Preserve each writer's real position while removing loaded, sweeping or mind-reading wording.")
    out = ROOT / "Lesson_22_Letters_to_Editor_Reading_Pack.docx"
    doc.save(out)
    return out


def build_organiser():
    doc = b.configure_doc("Bias Spectrum Organiser")
    b.add_title_block(doc, "Language analysis organiser", "The Bias Spectrum", "Find the exact pull. Report the viewpoint fairly.", "ANALYSE  |  PLACE  |  REWRITE")
    b.add_heading(doc, "1. Establish the fact base", 1)
    b.add_form_table(doc, ["Verified proposal detail", "A letter's claim or prediction"], [5089, 5089], blank_rows=4, blank_after=28)
    b.add_heading(doc, "2. Track each letter", 1)
    b.add_form_table(doc, ["Letter", "Position", "Exact wording", "Fingerprint", "Effect / omission"], [950, 1800, 2500, 2100, 2828], blank_rows=6, row_prompts=[[c, "", "", "", ""] for c in "ABCDEF"], font_size=8.5, blank_after=26)
    doc.add_page_break()
    b.add_heading(doc, "3. Place the letters on the spectrum", 1)
    b.add_form_table(doc, ["Most heavily positioned", "Middle / subtly slanted", "Most balanced"], [3392, 3393, 3393], blank_rows=2, blank_after=38)
    b.add_labeled_para(doc, "Justification 1:", "Letter ___ belongs here because the phrase '________________' ________________________________.", fill=b.TOKENS["pale"], accent=b.TOKENS["teal"])
    b.add_labeled_para(doc, "Justification 2:", "Letter ___ is more / less balanced than ___ because __________________________________________.", fill=b.TOKENS["warm"], accent=b.TOKENS["coral"])
    b.add_heading(doc, "4. Rewrite without changing the position", 1)
    b.add_form_table(doc, ["Biased sentence", "Neutral attributed version", "What I removed or qualified"], [3500, 4300, 2378], blank_rows=3, blank_after=30)
    doc.add_page_break()
    b.add_heading(doc, "5. Plan the Editorial Fairness Briefing", 1)
    for prompt in [
        "Shared decision and verified details:", "Position 1 reported neutrally:", "Position 2 reported neutrally:",
        "Quoted wording + reader effect 1:", "Quoted wording + reader effect 2:", "Important omission / unsupported certainty:",
        "Neutral rewrite:", "Most trustworthy letter + reason:",
    ]:
        b.add_labeled_para(doc, prompt, "________________________________________________________________________________")
    doc.add_page_break()
    b.add_heading(doc, "6. Feedback and visible revision", 1)
    b.add_form_table(doc, ["Reviewer diagnosis", "Writer's revision"], [5089, 5089], blank_rows=2, row_prompts=[
        ["One phrase that absorbs a letter's judgement as fact is...", "Original sentence:"],
        ["One explanation that needs a more precise effect or omission is...", "Revised sentence:"],
    ], blank_after=46)
    b.add_heading(doc, "Final self-check", 1)
    for item in ["[ ] I reported two positions fairly.", "[ ] I used three short exact language choices.", "[ ] I explained two reader effects.", "[ ] I identified an omission or unsupported certainty.", "[ ] My rewrite preserves the position without the distortion.", "[ ] My trust judgement uses criteria.", "[ ] I made one revision visible."]:
        p = doc.add_paragraph(style="List Bullet"); b.style_paragraph(p, after=4); r = p.add_run(item); b.set_run(r, size=b.TOKENS["body_size"], color=b.TOKENS["ink"])
    out = ROOT / "Lesson_22_Letters_to_Editor_Bias_Spectrum_Organiser.docx"
    doc.save(out)
    return out


def build_concise_pack():
    doc = b.configure_doc("Concise Access Pack")
    doc.styles["Normal"].font.size = Pt(13)
    doc.styles["Normal"].paragraph_format.line_spacing = 1.22
    b.add_title_block(doc, "Concise language pathway", "Letters from Seabreeze Square", "Same big idea: find the pull and say the view fairly.", "READ TOGETHER  |  POINT / SAY / WRITE")
    b.add_heading(doc, "The plan", 1)
    b.add_text(doc, "A fictional council may test food trucks at the town square for six Friday nights. The trucks could be there from 5:00 pm to 8:00 pm. There may be up to eight trucks. The plan includes bins, a clear path and a review.", size=13)
    b.add_heading(doc, "Letter A - Stop this circus", 1)
    b.add_text(doc, "The council wants to hand our square to a noisy food-truck circus. Every resident knows the trucks will ruin Friday nights. The council must stop this foolish plan.", size=13)
    b.add_heading(doc, "Letter B - Wake up the town", 1)
    b.add_text(doc, "Food trucks will rescue our lifeless town. Only selfish complainers could dislike this brilliant idea. The council must approve it now.", size=13)
    doc.add_page_break()
    b.add_heading(doc, "Letter E - Test and check", 1)
    b.add_text(doc, "Food trucks could bring people to the square. Residents may also worry about noise, parking and rubbish. Council should try the six nights, count problems and change the plan if needed.", size=13)
    b.add_heading(doc, "Letter F - Start small", 1)
    b.add_text(doc, "People may enjoy a shared event, but the square must stay easy to use. Council should start with four trucks, keep the path clear and check what happens before adding more.", size=13)
    b.add_heading(doc, "Find the pull", 1)
    b.add_matrix(doc, ["Words to find", "Plain meaning"], [
        ["circus, foolish, brilliant, selfish", "strong judging words"],
        ["every, only", "words that speak for everyone"],
        ["will ruin, will rescue", "too certain - the trial has not happened"],
        ["may, could, if", "careful words that show a limit"],
    ], [4300, 5878], font_size=11.5)
    doc.add_page_break()
    b.add_heading(doc, "Make one fair report", 1)
    for prompt in [
        "Letter ___ is for / against ______________________________________________.",
        "It uses the words '______________________________________________________'.",
        "These words may make readers think / feel __________________________________.",
        "The letter does not say _________________________________________________ .",
        "A fair way to report the view is __________________________________________.",
    ]:
        b.add_labeled_para(doc, "Say / write:", prompt, fill=b.TOKENS["pale"], accent=b.TOKENS["teal"])
    b.add_callout(doc, "You are finished when:", "You name a real viewpoint, one strong wording choice, its effect, one missing detail and one fair sentence.")
    out = ROOT / "Lesson_22_Letters_to_Editor_Concise_Access_Pack.docx"
    doc.save(out)
    return out


def page_html(page_no, title, subtitle, body):
    return f'''<section class="page"><header><div class="mast">THE SEABREEZE GAZETTE <span>LETTERS DESK</span></div><div class="eyebrow">Original classroom text | Page {page_no}</div><h1>{title}</h1><p class="stand">{subtitle}</p></header><main>{body}</main><footer>Fictional classroom publication | English Unit 3 | Lesson 22 Alternative</footer></section>'''


def letter_html(code, title, author, paragraphs, tone):
    return f'''<article class="letter {tone}"><div class="letter-code">{code}</div><h2>{title}</h2><div class="by">{author}</div>{''.join(f'<p>{p}</p>' for p in paragraphs)}</article>'''


def build_article_html():
    a,b_,c,d,e,f = LETTERS
    briefing = '''<aside class="brief"><h2>Editor's verified briefing</h2><p>Seabreeze Council is considering a <b>six-Friday trial</b> at Seabreeze Square. Up to <b>eight food trucks</b> could operate from <b>5:00 pm to 8:00 pm</b>. The draft includes <b>waste stations</b>, an <b>accessible path</b> and a <b>review after the trial</b>.</p><p>The briefing does not prove future noise, profit, parking failure, litter, community renewal or universal support.</p></aside>'''
    fingerprints = '''<aside class="tool"><h2>Five fingerprints of bias</h2><ol><li><b>Loaded evaluation</b> praises, shames or alarms.</li><li><b>Sweeping generalisation</b> speaks for everyone.</li><li><b>Assumed motive</b> claims to know why people act.</li><li><b>Omission</b> hides relevant detail.</li><li><b>Unsupported certainty</b> treats a prediction as fact.</li></ol></aside>'''
    pages = [
        page_html(1,"One square. Six letters.","Before reading an opinion, establish what is actually proposed.",briefing+fingerprints+letter_html(*a[:3],a[4],"hot")),
        page_html(2,"The loud pull","Opposite positions can use the same biased techniques.",letter_html(*b_[:3],b_[4],"hot")+letter_html(*c[:3],c[4],"quiet")),
        page_html(3,"The quiet pull","Calm wording can still select, omit and predict beyond the evidence.",letter_html(*d[:3],d[4],"quiet")+letter_html(*e[:3],e[4],"balanced")),
        page_html(4,"Balanced and still persuasive","A fair writer can acknowledge limits and recommend a clear action.",letter_html(*f[:3],f[4],"balanced")+'''<aside class="field"><h2>Balanced-language field guide</h2><div><b>Attribute</b> The writer argues...</div><div><b>Concede</b> Residents may reasonably...</div><div><b>Qualify</b> may, could, if...</div><div><b>Set a boundary</b> The proposal does not yet show...</div><div><b>Recommend</b> Council should trial, measure and review...</div><p class="challenge"><b>Editorial challenge:</b> preserve each real position while removing loaded, sweeping and mind-reading language.</p></aside>'''),
    ]
    html = '''<!doctype html><html lang="en-AU"><head><meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1"><title>Letters from Seabreeze Square</title><style>
@page{size:A4;margin:0}*{box-sizing:border-box}body{margin:0;background:#d8dee2;color:#263139;font-family:Arial,sans-serif}.page{width:210mm;height:297mm;margin:10px auto;background:#fff;position:relative;overflow:hidden;box-shadow:0 8px 28px #0003}header{height:58mm;padding:10mm 14mm;background:linear-gradient(115deg,#202f3b,#334c5c 70%,#4c7774);color:#fff}.mast{font:800 8pt/1 Arial;letter-spacing:.16em;border-bottom:1px solid #ffffff66;padding-bottom:2mm}.mast span{float:right;color:#f6d991}.eyebrow{margin-top:6mm;color:#f6d991;text-transform:uppercase;letter-spacing:.13em;font-weight:800;font-size:7.5pt}h1{font:800 27pt/1 Georgia,serif;margin:2mm 0}.stand{font-size:10.5pt;font-weight:700;margin:0}main{padding:8mm 14mm 14mm;display:grid;grid-template-columns:1fr 1fr;gap:6mm;align-content:start}.letter,.brief,.tool,.field{position:relative;border-top:4px solid #334c5c;padding:5mm;background:#f6f8f9}.letter.hot{border-color:#a9564d;background:#fff4ee}.letter.quiet{border-color:#d5a84b;background:#fffaf0}.letter.balanced{border-color:#4c7774;background:#edf5f2}.letter-code{position:absolute;right:4mm;top:3mm;font:800 20pt Georgia,serif;color:#334c5c55}.letter h2,.brief h2,.tool h2,.field h2{font:800 13pt/1.1 Georgia,serif;color:#202f3b;margin:0 8mm 1mm 0}.by{font-size:7.8pt;text-transform:uppercase;letter-spacing:.08em;color:#66737b;font-weight:800;margin-bottom:3mm}.letter p,.brief p,.tool li,.field div,.field p{font-size:8.45pt;line-height:1.28;margin:0 0 2.2mm}.brief,.tool{min-height:55mm}.tool ol{margin:2mm 0 0;padding-left:5mm}.field{grid-column:1/-1;display:grid;grid-template-columns:repeat(2,1fr);gap:2mm 6mm}.field h2,.field .challenge{grid-column:1/-1}.challenge{background:#202f3b;color:#fff;padding:3mm;border-radius:2mm}footer{position:absolute;left:14mm;right:14mm;bottom:5mm;border-top:1px solid #bac6cc;padding-top:2mm;color:#66737b;font-size:6.5pt}@media print{body{background:#fff}.page{margin:0;box-shadow:none;page-break-after:always}.page:last-child{page-break-after:auto}}</style></head><body>''' + ''.join(pages) + '</body></html>'
    out = ROOT / "Lesson_22_Letters_to_Editor_Reading_Article.html"
    out.write_text(html, encoding="utf-8")
    return out


if __name__ == "__main__":
    for output in [build_reading_pack(), build_organiser(), build_concise_pack(), build_article_html()]:
        print(output)
