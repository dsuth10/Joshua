from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]

COLORS = {
    "navy": "18324A",
    "navy_2": "244A66",
    "ink": "243846",
    "muted": "5D7180",
    "paper": "FFFFFF",
    "mist": "EFF4F7",
    "bias": "D65D4A",
    "bias_dark": "9E3D31",
    "bias_pale": "FCEDE8",
    "balance": "16877D",
    "balance_dark": "0B625D",
    "balance_pale": "E8F6F3",
    "gold": "D99B2B",
    "gold_pale": "FFF3D7",
    "line": "CBD7DE",
    "white": "FFFFFF",
}

TOPICS = [
    {
        "number": "01",
        "title": "School uniforms",
        "question": "How can one school policy be framed as either a threat or a trade-off?",
        "biased_title": "Uniforms crush student freedom",
        "biased_format": "Opinion column",
        "biased": [
            "It is an absolute outrage that students are forced into dull, uncomfortable uniforms every day. These outdated outfits crush creativity and turn energetic children into identical robots.",
            "Strict rules suppress personal identity, while greedy suppliers charge ridiculous prices for poor-quality clothing. Every student and parent hates uniforms, so schools must scrap this cruel policy immediately!",
        ],
        "balanced_title": "Examining the school uniform debate",
        "balanced_format": "Education briefing",
        "balanced": [
            "School uniform policies remain a subject of discussion among students, families and educators. Supporters argue that a shared dress code may strengthen school identity, reduce morning decisions and limit pressure to wear expensive brands.",
            "Others point out that compulsory uniforms can restrict personal expression and place pressure on family budgets. Schools could respond with flexible options, durable items and second-hand uniform programs. A balanced decision weighs belonging and practicality against cost and choice.",
        ],
        "concise_biased": [
            "Uniforms are a complete disaster. Schools force students to wear dull clothes that destroy creativity.",
            "Greedy shops charge huge prices. Everyone hates uniforms, so schools must ban them now!",
        ],
        "concise_balanced": [
            "Uniforms may help students feel part of a school. They can also reduce pressure to wear costly brands.",
            "However, some families worry about price, and some students want more choice. Schools could offer flexible items and second-hand clothes.",
        ],
        "prompt": "Underline one absolute claim. Box one word that shows a balanced trade-off.",
        "concise_prompt": "Circle an angry or extreme word. Box the word that introduces the other side.",
    },
    {
        "number": "02",
        "title": "E-scooters on footpaths",
        "question": "Does the writer describe a transport option, a public danger, or both?",
        "biased_title": "Pedestrian nightmare on our streets",
        "biased_format": "Letter to the editor",
        "biased": [
            "Our footpaths have become a terrifying warzone because reckless e-scooter riders race through crowded spaces. These silent machines zoom past helpless pedestrians and create total chaos.",
            "Careless riders dump metal hazards across walkways, blocking older people, prams and wheelchairs. It is pure insanity that councils permit this menace. E-scooters must be banned permanently before disaster strikes!",
        ],
        "balanced_title": "Evaluating shared e-scooter schemes",
        "balanced_format": "Transport briefing",
        "balanced": [
            "Shared e-scooter schemes offer a compact, low-emission option for some short journeys. Supporters say they may reduce brief car trips and help commuters travel between public transport stops and nearby destinations.",
            "However, safety reports and pedestrian groups raise concerns about speed, collisions and devices left across paths. Some councils are testing lower-speed zones, designated parking bays and clearer enforcement. These measures aim to balance convenient travel with safe, accessible footpaths.",
        ],
        "concise_biased": [
            "E-scooters have turned footpaths into a terrifying warzone. Reckless riders race past helpless people and dump scooters everywhere.",
            "Councils must ban these dangerous machines before disaster strikes!",
        ],
        "concise_balanced": [
            "E-scooters can be a quick, low-emission way to travel. They may replace some short car trips.",
            "However, fast riding and poorly parked scooters can make paths unsafe. Councils are testing slow zones and parking bays.",
        ],
        "prompt": "Highlight language that positions riders. Find two different stakeholder perspectives.",
        "concise_prompt": "Circle words that make riders sound dangerous. Put a star beside each side of the issue.",
    },
    {
        "number": "03",
        "title": "Zoos and wildlife parks",
        "question": "Which evidence and perspectives are included - and which are left out?",
        "biased_title": "Cages of shame",
        "biased_format": "Campaign flyer",
        "biased": [
            "Locking magnificent wild animals in miserable enclosures for human entertainment is heartless cruelty. Innocent animals are trapped behind cold barriers and condemned to endless boredom.",
            "Profit-obsessed operators pretend to care about conservation while exploiting helpless creatures. All zoos are cruel animal prisons and must close today!",
        ],
        "balanced_title": "Conservation, education and animal welfare",
        "balanced_format": "Science feature",
        "balanced": [
            "Modern zoos and wildlife parks operate between conservation, education, research and animal welfare. Accredited organisations report that breeding programs can support threatened species and that visitor programs may build awareness of habitat loss.",
            "Animal-welfare advocates argue that enclosed spaces can limit natural behaviour, particularly for animals that travel over large areas. Strong welfare standards, suitable habitats and transparent monitoring are therefore important. Judging a facility fairly requires evidence about both conservation outcomes and individual animal wellbeing.",
        ],
        "concise_biased": [
            "Zoos are cruel prisons. Helpless animals are trapped in miserable cages just to entertain people.",
            "Greedy operators only want money. Every zoo must close today!",
        ],
        "concise_balanced": [
            "Some zoos help breed threatened animals and teach visitors about conservation.",
            "However, living in an enclosure can limit an animal's natural behaviour. A fair judgement checks both conservation work and animal wellbeing.",
        ],
        "prompt": "Bracket the single-sided evidence. Locate the condition that makes the balanced judgement more precise.",
        "concise_prompt": "Draw a bracket around the one-sided claims. Underline what a fair judgement must check.",
    },
    {
        "number": "04",
        "title": "Artificial intelligence in schoolwork",
        "question": "How does certainty change what a reader believes about a new technology?",
        "biased_title": "The death of learning",
        "biased_format": "Technology opinion column",
        "biased": [
            "Artificial intelligence is a catastrophic disaster that is destroying education. Lazy students use AI to produce entire assignments, ruining their ability to think, read and write independently.",
            "Teachers are helpless against this wave of cheating, and schools are falling apart. If classrooms do not ban AI immediately, an entire generation will become unable to learn for itself!",
        ],
        "balanced_title": "Opportunities, risks and responsible use",
        "balanced_format": "Digital learning advisory",
        "balanced": [
            "Generative artificial intelligence presents opportunities and challenges for schools. Used with clear boundaries, AI tools may support brainstorming, feedback and practice, while teaching students to question digital information.",
            "Educators also raise concerns about inaccurate output, privacy, academic integrity and over-reliance. Schools could define when AI is allowed, require students to explain their thinking and assess foundational skills directly. Responsible use depends on purpose, transparency and teacher guidance.",
        ],
        "concise_biased": [
            "AI is destroying education. Lazy students use it to avoid thinking and writing.",
            "Teachers are helpless. Schools must ban AI now or students will forget how to learn!",
        ],
        "concise_balanced": [
            "AI may help students brainstorm, practise and get feedback. It can also produce wrong information or be used to avoid learning.",
            "Schools could set clear rules, ask students to explain their thinking and protect private information.",
        ],
        "prompt": "Circle certainty words. Underline three conditions or qualifiers in the balanced text.",
        "concise_prompt": "Circle a word that sounds completely certain. Underline words that show possibility or a condition.",
    },
]


def rgb(hex_value):
    return RGBColor.from_string(hex_value)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=120, start=140, bottom=120, end=140):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, **edges):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge, spec in edges.items():
        tag = "start" if edge == "left" else "end" if edge == "right" else edge
        node = borders.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            borders.append(node)
        for key, value in spec.items():
            node.set(qn(f"w:{key}"), str(value))


def set_table_geometry(table, widths_dxa, indent=0):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def set_run(run, size=10, bold=False, color=None, italic=False, font="Aptos"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = rgb(color)


def set_paragraph(paragraph, before=0, after=4, line=1.08, keep_next=False, alignment=None):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line
    paragraph.paragraph_format.keep_with_next = keep_next
    if alignment is not None:
        paragraph.alignment = alignment


def add_runs(paragraph, parts, size=10, color=None):
    for text, bold, italic in parts:
        set_run(paragraph.add_run(text), size=size, bold=bold, italic=italic, color=color)


def add_para(container, text="", size=10, bold=False, color=None, italic=False, before=0, after=4,
             line=1.08, keep_next=False, alignment=None):
    p = container.add_paragraph()
    set_paragraph(p, before, after, line, keep_next, alignment)
    set_run(p.add_run(text), size=size, bold=bold, color=color, italic=italic)
    return p


def add_cell_para(cell, text, size=10, bold=False, color=None, italic=False, before=0, after=4, line=1.08,
                  alignment=None):
    p = cell.add_paragraph() if cell.paragraphs[0].text else cell.paragraphs[0]
    set_paragraph(p, before, after, line, alignment=alignment)
    set_run(p.add_run(text), size=size, bold=bold, color=color, italic=italic)
    return p


def add_header_footer(section, route_label):
    header = section.header
    hp = header.paragraphs[0]
    hp.text = ""
    set_paragraph(hp, after=0, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_runs(hp, [("ENGLISH UNIT 3", True, False), ("  |  LESSON 21 ALTERNATIVE", False, False)], 8, COLORS["muted"])
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.text = ""
    set_paragraph(fp, after=0, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    add_runs(fp, [("Biased vs Balanced  |  ", False, False), (route_label, True, False)], 8, COLORS["muted"])


def configure_document(route_label):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.header_distance = Cm(0.7)
    section.footer_distance = Cm(0.7)
    add_header_footer(section, route_label)
    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal.font.size = Pt(10)
    normal.font.color.rgb = rgb(COLORS["ink"])
    return doc


def add_label(container, text, fill, color=COLORS["white"], width_dxa=11280):
    table = container.add_table(rows=1, cols=1)
    set_table_geometry(table, [width_dxa])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, 70, 130, 70, 130)
    set_cell_border(cell, top={"val": "nil"}, bottom={"val": "nil"}, start={"val": "nil"}, end={"val": "nil"})
    p = cell.paragraphs[0]
    set_paragraph(p, after=0)
    set_run(p.add_run(text.upper()), size=8.5, bold=True, color=color)
    return table


def add_cover(doc, concise):
    kicker = "STUDENT READING PACK  |  ENGLISH UNIT 3  |  LESSON 21"
    add_para(doc, kicker, size=9, bold=True, color=COLORS["gold"], after=8)
    title = "Biased vs Balanced"
    add_para(doc, title, size=28 if not concise else 26, bold=True, color=COLORS["navy"], after=4, line=0.96)
    subtitle = "How writers position a reader - and how careful readers test the view"
    add_para(doc, subtitle, size=13, bold=True, color=COLORS["navy_2"], after=12, line=1.05)

    mission = doc.add_table(rows=1, cols=1)
    set_table_geometry(mission, [11280])
    cell = mission.cell(0, 0)
    set_cell_shading(cell, COLORS["gold_pale"])
    set_cell_margins(cell, 150, 180, 150, 180)
    set_cell_border(cell, start={"val": "single", "sz": "20", "color": COLORS["gold"]},
                    top={"val": "nil"}, bottom={"val": "nil"}, end={"val": "nil"})
    add_cell_para(cell, "YOUR READING MISSION", 9, True, COLORS["gold"], after=2)
    mission_text = ("Read each pair. Mark the words and evidence that make one version one-sided and the other more measured. "
                    "Your job is to explain how you know - not simply choose the view you agree with.")
    if concise:
        mission_text = ("Read each pair. Find words that push one side strongly. Then find words that show both sides. "
                        "Explain which words helped you decide.")
    add_cell_para(cell, mission_text, 10.5 if not concise else 11.5, True, COLORS["navy"], after=0, line=1.12)

    add_para(doc, "Reader lens", size=16, bold=True, color=COLORS["navy"], before=12, after=6)
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [5640, 5640])
    left, right = table.rows[0].cells
    for cell_x, fill, accent in ((left, COLORS["bias_pale"], COLORS["bias"]), (right, COLORS["balance_pale"], COLORS["balance"])):
        set_cell_shading(cell_x, fill)
        set_cell_margins(cell_x, 150, 170, 150, 170)
        set_cell_border(cell_x, top={"val": "single", "sz": "18", "color": accent},
                        bottom={"val": "single", "sz": "4", "color": COLORS["line"]},
                        start={"val": "single", "sz": "4", "color": COLORS["line"]},
                        end={"val": "single", "sz": "4", "color": COLORS["line"]})
        cell_x.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    add_cell_para(left, "BIASED / ONE-SIDED", 11, True, COLORS["bias_dark"], after=5)
    add_cell_para(right, "BALANCED / MEASURED", 11, True, COLORS["balance_dark"], after=5)
    bias_items = [
        ("Loaded words", "push a strong feeling"),
        ("Loaded verbs", "make people or actions seem worse or better"),
        ("Absolute claims", "always, never, everyone, completely"),
        ("Selective evidence", "shows only the helpful facts"),
    ]
    balance_items = [
        ("Neutral tone", "describes without attacking"),
        ("More than one view", "represents different concerns fairly"),
        ("Qualifiers", "may, could, suggests, some"),
        ("Limits and conditions", "states what depends on context or evidence"),
    ]
    if concise:
        bias_items = [("Big feeling words", "terrible, perfect, disaster"), ("One side only", "leaves out another view"),
                      ("Certain words", "always, never, everyone")]
        balance_items = [("Calm words", "reports, says, explains"), ("Both sides", "uses however or while"),
                         ("Careful words", "may, could, some")]
    for heading, detail in bias_items:
        p = left.add_paragraph()
        set_paragraph(p, after=5, line=1.08)
        add_runs(p, [(heading + ": ", True, False), (detail, False, False)], 9.5 if not concise else 10.5, COLORS["ink"])
    for heading, detail in balance_items:
        p = right.add_paragraph()
        set_paragraph(p, after=5, line=1.08)
        add_runs(p, [(heading + ": ", True, False), (detail, False, False)], 9.5 if not concise else 10.5, COLORS["ink"])

    add_para(doc, "Use two marks", size=14, bold=True, color=COLORS["navy"], before=12, after=5)
    mark = doc.add_table(rows=1, cols=2)
    set_table_geometry(mark, [5640, 5640])
    for cell_x in mark.rows[0].cells:
        set_cell_shading(cell_x, COLORS["mist"])
        set_cell_margins(cell_x, 110, 150, 110, 150)
        set_cell_border(cell_x, top={"val": "nil"}, bottom={"val": "nil"}, start={"val": "nil"}, end={"val": "nil"})
    add_cell_para(mark.cell(0, 0), "YELLOW: words that push one side", 10, True, COLORS["bias_dark"], after=0)
    add_cell_para(mark.cell(0, 1), "GREEN: words that qualify or include another side", 10, True, COLORS["balance_dark"], after=0)

    add_para(doc, "Before you call a text balanced...", size=13, bold=True, color=COLORS["navy"], before=12, after=4)
    warning = ("A balanced text is not automatically true, complete or trustworthy. It can still omit evidence. Check the source, author, date and evidence separately.")
    if concise:
        warning = ("Calm language does not prove a text is true. You still need to check who wrote it and what evidence they used.")
    add_para(doc, warning, size=10.5 if not concise else 11.5, bold=True, color=COLORS["navy_2"], after=0, line=1.12)


def add_text_card(cell, title, format_label, paragraphs, biased, concise):
    accent = COLORS["bias"] if biased else COLORS["balance"]
    pale = COLORS["bias_pale"] if biased else COLORS["balance_pale"]
    dark = COLORS["bias_dark"] if biased else COLORS["balance_dark"]
    set_cell_shading(cell, pale)
    set_cell_margins(cell, 160, 180, 150, 180)
    set_cell_border(cell, top={"val": "single", "sz": "20", "color": accent},
                    bottom={"val": "single", "sz": "4", "color": COLORS["line"]},
                    start={"val": "single", "sz": "4", "color": COLORS["line"]},
                    end={"val": "single", "sz": "4", "color": COLORS["line"]})
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    add_cell_para(cell, "ONE-SIDED VERSION" if biased else "BALANCED VERSION", 8.5, True, dark, after=2)
    add_cell_para(cell, title, 14 if not concise else 15, True, COLORS["navy"], after=3, line=1.0)
    add_cell_para(cell, format_label + "  |  Original classroom text", 8.5, False, COLORS["muted"], italic=True, after=8)
    for idx, text in enumerate(paragraphs):
        add_cell_para(cell, text, 10.4 if not concise else 11.8, False, COLORS["ink"], after=7 if idx < len(paragraphs)-1 else 0,
                      line=1.13 if not concise else 1.16)


def add_topic_page(doc, topic, concise):
    doc.add_page_break()
    header = doc.add_table(rows=1, cols=2)
    set_table_geometry(header, [1700, 9580])
    ncell, tcell = header.rows[0].cells
    set_cell_shading(ncell, COLORS["navy"])
    set_cell_shading(tcell, COLORS["mist"])
    for cell in (ncell, tcell):
        set_cell_margins(cell, 120, 160, 120, 160)
        set_cell_border(cell, top={"val": "nil"}, bottom={"val": "nil"}, start={"val": "nil"}, end={"val": "nil"})
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    add_cell_para(ncell, topic["number"], 21, True, COLORS["white"], after=0, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_cell_para(tcell, topic["title"].upper(), 9, True, COLORS["gold"], after=1)
    add_cell_para(tcell, topic["question"], 13, True, COLORS["navy"], after=0, line=1.04)

    add_para(doc, "Read both versions before deciding which techniques you notice.", size=9.5 if not concise else 10.5,
             italic=True, color=COLORS["muted"], before=8, after=7)

    comparison = doc.add_table(rows=1, cols=2)
    set_table_geometry(comparison, [5640, 5640])
    prevent_row_split(comparison.rows[0])
    biased_text = topic["concise_biased"] if concise else topic["biased"]
    balanced_text = topic["concise_balanced"] if concise else topic["balanced"]
    add_text_card(comparison.cell(0, 0), topic["biased_title"], topic["biased_format"], biased_text, True, concise)
    add_text_card(comparison.cell(0, 1), topic["balanced_title"], topic["balanced_format"], balanced_text, False, concise)

    prompt_text = topic["concise_prompt"] if concise else topic["prompt"]
    prompt = doc.add_table(rows=1, cols=1)
    set_table_geometry(prompt, [11280])
    cell = prompt.cell(0, 0)
    set_cell_shading(cell, COLORS["gold_pale"])
    set_cell_margins(cell, 130, 170, 130, 170)
    set_cell_border(cell, start={"val": "single", "sz": "18", "color": COLORS["gold"]},
                    top={"val": "nil"}, bottom={"val": "nil"}, end={"val": "nil"})
    p = cell.paragraphs[0]
    set_paragraph(p, after=0, line=1.1)
    add_runs(p, [("NOTICE & MARK  ", True, False), (prompt_text, False, False)], 9.5 if not concise else 10.5, COLORS["navy"])

    add_para(doc, "Evidence notes", size=12, bold=True, color=COLORS["navy"], before=10, after=4)
    notes = doc.add_table(rows=3 if not concise else 2, cols=2)
    set_table_geometry(notes, [5640, 5640])
    labels = [
        ("A word or phrase that creates bias", "A word or structure that creates balance"),
        ("What perspective or evidence is missing?", "What different perspective is included?"),
        ("How is the reader positioned?", "Why is the claim more measured?"),
    ]
    if concise:
        labels = [("Strong word I found", "Careful word I found"), ("One side shown", "Other side included")]
    for ridx, row in enumerate(notes.rows):
        prevent_row_split(row)
        for cidx, cell_x in enumerate(row.cells):
            set_cell_shading(cell_x, COLORS["paper"])
            set_cell_margins(cell_x, 90, 130, 300 if concise else 220, 130)
            set_cell_border(cell_x, top={"val": "single", "sz": "3", "color": COLORS["line"]},
                            bottom={"val": "single", "sz": "3", "color": COLORS["line"]},
                            start={"val": "single", "sz": "3", "color": COLORS["line"]},
                            end={"val": "single", "sz": "3", "color": COLORS["line"]})
            add_cell_para(cell_x, labels[ridx][cidx], 8.5 if not concise else 9.5, True,
                          COLORS["bias_dark"] if cidx == 0 else COLORS["balance_dark"], after=0)

    add_para(doc, "Source note: These are original fictional classroom texts. The publication labels show genre, not real publications or authorities.",
             size=8, italic=True, color=COLORS["muted"], before=7, after=0)


def build_pack(concise=False):
    route_label = "Focused reading route" if concise else "Year 5 reading route"
    doc = configure_document(route_label)
    add_cover(doc, concise)
    for topic in TOPICS:
        add_topic_page(doc, topic, concise)
    filename = ("Lesson_21_Biased_vs_Balanced_Focused_Reading_Pack.docx" if concise
                else "Lesson_21_Biased_vs_Balanced_Year_5_Reading_Pack.docx")
    path = ROOT / filename
    doc.save(path)
    return path


if __name__ == "__main__":
    outputs = [build_pack(False), build_pack(True)]
    print("Built reading packs:")
    for output in outputs:
        print(f" - {output.name}")
