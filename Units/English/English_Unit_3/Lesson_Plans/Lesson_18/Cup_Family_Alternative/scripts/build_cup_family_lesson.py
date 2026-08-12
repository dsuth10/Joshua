from pathlib import Path
import importlib.util

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
BASE_PATH = ROOT.parents[1] / "Lesson_17" / "Turtle_Safe_Nights_Alternative" / "scripts" / "build_turtle_safe_nights.py"
spec = importlib.util.spec_from_file_location("lesson_pack_base", BASE_PATH)
b = importlib.util.module_from_spec(spec)
spec.loader.exec_module(b)
b.ROOT = ROOT

# compact_reference_guide with a warm cafe palette.
b.TOKENS.update({
    "override": "cup_family_a4_classroom",
    "navy": "23332B",
    "deep_sea": "29483A",
    "teal": "4E7D62",
    "moon": "E4B44B",
    "coral": "C9664D",
    "pale": "EDF5EE",
    "warm": "FFF3D9",
    "mist": "F0EEE7",
    "ink": "26332C",
    "muted": "657069",
    "border": "B9C5BA",
})


def add_footer(section, label):
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    b.style_paragraph(hp, after=0, line=1.0)
    r = hp.add_run("ENGLISH UNIT 3  |  LESSON 18 ALTERNATIVE")
    b.set_run(r, size=7.5, bold=True, color=b.TOKENS["muted"])

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    b.style_paragraph(fp, after=0, line=1.0)
    r = fp.add_run(label + "  |  One Cup, Two Loyalties")
    b.set_run(r, size=7.5, color=b.TOKENS["muted"])


b.add_footer = add_footer


def add_bullet(doc, text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    b.style_paragraph(p, after=4)
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        b.set_run(r1, size=b.TOKENS["body_size"], bold=True, color=b.TOKENS["deep_sea"])
        r2 = p.add_run(text[len(bold_lead):])
        b.set_run(r2, size=b.TOKENS["body_size"], color=b.TOKENS["ink"])
    else:
        r = p.add_run(text)
        b.set_run(r, size=b.TOKENS["body_size"], color=b.TOKENS["ink"])
    return p


def build_reading_pack():
    doc = b.configure_doc("Student Reading Pack")
    b.add_title_block(
        doc,
        "Original classroom feature",
        "One Cup, Two Loyalties",
        "Can an environmental argument protect both a principle and a family relationship?",
        "READ  |  MARK V / C / E / T",
    )
    b.add_callout(
        doc,
        "Decision:",
        "Should the fictional Rahman family trial a reusable-cup system at Sunrise Corner Cafe for four weeks?",
    )
    b.add_labeled_para(
        doc,
        "Fiction boundary:",
        "Sunrise Corner Cafe, the Rahman family and every quotation are invented. The evidence panel summarises the official sources listed at the end.",
        fill=b.TOKENS["pale"],
        accent=b.TOKENS["teal"],
    )

    b.add_heading(doc, "The poster on the pantry door", 1)
    b.add_text(doc, "At 6:35 on Friday morning, Leila Rahman taped a poster to the pantry door of her family's cafe. In thick green letters it asked, 'How many throwaway cups does one busy morning need?' Underneath, she had written a proposal: for four weeks, Sunrise Corner would reward customers who brought a clean reusable cup, offer ceramic mugs to people drinking inside and test a small returnable-mug shelf for regular customers.")
    b.add_text(doc, "Leila had spent the school holidays helping at the register. She saw towers of cups arrive in cardboard sleeves and watched the bins fill before lunchtime. To her, the cups represented a choice repeated so often that it became invisible. She wanted her family to show that a small business could lead change instead of waiting for somebody else.")
    b.add_text(doc, "Her first slogan sounded powerful: 'Our cafe is choosing convenience over the planet.' When she read it again, however, she imagined her aunt standing beside the coffee machine. The sentence made the adults sound careless. Leila knew that was not true.")

    b.add_heading(doc, "What the slogan leaves out", 1)
    b.add_text(doc, "Aunt Samira opened the cafe most mornings and managed the busiest hour. She checked food deliveries, trained new staff and made sure wages could be paid. She liked Leila's goal, but she did not like being treated as an obstacle.")
    b.add_callout(doc, "Aunt Samira:", "A good idea still has to work at 7:45, when twenty people are waiting and a customer hands us a cup that may not be clean.", fill=b.TOKENS["warm"], accent=b.TOKENS["coral"])
    b.add_text(doc, "Her concern was not a secret wish to keep wasting cups. She wanted answers. Who would inspect a customer's cup? Where would staff place it so it did not touch the service area? What would happen if a cup was dirty, damaged or too small? Would a mug-return shelf create extra washing during the rush? What would the cafe offer a customer who forgot a cup or could not use one?")
    b.add_text(doc, "Leila felt the argument changing. She could repeat that waste mattered, but that bare assertion did not answer the people affected by her proposal. If she wanted a family decision rather than a family fight, she needed to acknowledge the concern accurately and propose a test that could produce local evidence.")
    b.add_callout(doc, "Mark the feature:", "V = value or responsibility  |  C = genuine concern  |  E = checkable evidence  |  T = persuasive technique or framing", fill=b.TOKENS["warm"], accent=b.TOKENS["moon"])

    doc.add_page_break()
    b.add_heading(doc, "Evidence that changes the conversation", 1)
    b.add_text(doc, "Queensland Government guidance for businesses suggests offering customers an incentive to bring a reusable cup or drink in store. It also suggests a mug-return program, possibly using a deposit. These are options, not guarantees. The guidance does not say that every cafe must use the same system or that one system will suit every queue, customer and workspace.")
    b.add_text(doc, "Food Standards Australia New Zealand explains that the Food Standards Code does not specifically require a business to accept a customer's reusable cup. A business may decide whether to accept one. It should have a policy for customer cups and consider cleaning, sanitising and cross-contamination. The business remains responsible for ensuring food is safe and suitable.")
    b.add_text(doc, "Disposal can also depend on local services. Brisbane City Council currently directs people to participating drop-off points for coffee cups rather than presenting them as an ordinary yellow-bin item. That local example is a warning against a sweeping claim such as 'all paper cups are easily recycled'. A careful writer checks the exact cup and the exact local pathway.")
    b.add_labeled_para(doc, "Evidence boundary:", "Official guidance supports avoidance, reuse and a clear safety policy. It does not prove that a four-week trial will succeed, remove all waste or cause no delay. The trial can collect local evidence about participation, workflow and safety concerns.", fill=b.TOKENS["pale"], accent=b.TOKENS["teal"])

    b.add_heading(doc, "Two drafts for the family meeting", 1)
    b.add_matrix(
        doc,
        ["Draft A - Heat", "Draft B - FAIR bridge"],
        [[
            "Our throwaway habit is choking the planet because the cafe refuses to change. How can money matter more than our future? We must ban disposable cups now.",
            "We all want Sunrise Corner to remain welcoming and responsible. A morning rush makes safety and speed genuine concerns. Official guidance shows that a cafe can set a clear customer-cup policy, so we should test a four-week system and measure what happens."
        ]],
        [5089, 5089],
        font_size=9.4,
    )
    b.add_text(doc, "Draft A uses urgent emotive language, a rhetorical question and strong modality. Those techniques create force, but the draft invents an unfair motive and promises an immediate ban without answering practical concerns. Draft B is less heated but more defensible. It uses inclusive language, a fair concession, evidence and a proportionate call to action.")

    b.add_heading(doc, "The proposal Leila can defend", 1)
    add_bullet(doc, "Frame a shared goal: a welcoming, safe and responsible cafe.")
    add_bullet(doc, "Acknowledge a fair concern: staff must protect food safety and morning workflow.")
    add_bullet(doc, "Introduce relevant evidence: a cafe can choose a policy and test incentives or mug return.")
    add_bullet(doc, "Recommend a workable action: trial, safeguard, measure and review.")
    b.add_callout(doc, "Open-letter challenge:", "Write 9-11 sentences to the Rahman family. Be firm about the ethical goal, fair about the family's responsibilities and specific about what the four-week trial will test.", fill=b.TOKENS["warm"], accent=b.TOKENS["coral"])

    doc.add_page_break()
    b.add_heading(doc, "Persuasive technique field guide", 1)
    b.add_matrix(
        doc,
        ["Technique", "Possible job", "Accuracy check"],
        [
            ["Inclusive language", "Frames shared responsibility: we, our, together", "Does 'we' genuinely include the audience?"],
            ["Modality", "Calibrates force: could, should, must", "Is the strength justified by the evidence?"],
            ["Rhetorical question", "Directs attention or opens a possibility", "Does it invite thought or shame the audience?"],
            ["Emotive vocabulary", "Signals importance and feeling", "Does it describe the issue or attack a person?"],
            ["Concession", "Acknowledges a legitimate concern", "Is it the audience's real concern?"],
            ["Call to action", "Tells the audience what to do next", "Is the action specific, safe and workable?"],
        ],
        [2300, 4000, 3878],
        font_size=9.0,
    )
    b.add_heading(doc, "Source trail", 1)
    b.add_text(doc, "Facts checked 12 August 2026. Classroom feature and dialogue are original and fictional.", italic=True, color=b.TOKENS["muted"])
    for source in [
        "Queensland Government - How to reduce consumption of avoidable and single-use plastics: business guide",
        "Queensland Government - Action on single-use plastic items",
        "Food Standards Australia New Zealand - Customer cups and containers",
        "Food Standards Australia New Zealand - Food packaging",
        "Brisbane City Council - What goes in my bins? (local disposal example only)",
    ]:
        add_bullet(doc, source)
    out = ROOT / "Lesson_18_Cup_Family_Reading_Pack.docx"
    doc.save(out)
    return out


def build_organiser():
    doc = b.configure_doc("FAIR Organiser")
    b.add_title_block(doc, "Persuasive writing organiser", "The FAIR Bridge", "Firm enough to matter. Fair enough to persuade.", "PLAN  |  JUSTIFY  |  REVISE")
    b.add_heading(doc, "1. Two viewpoints, one decision", 1)
    b.add_form_table(doc, ["Leila notices / values", "Aunt Samira notices / values"], [5089, 5089], blank_rows=2, row_prompts=[["Evidence or detail:", "Evidence or detail:"], ["Responsibility:", "Responsibility:"]], blank_after=32)
    b.add_labeled_para(doc, "Shared goal:", "Both viewpoints may value ________________________________________________", fill=b.TOKENS["pale"], accent=b.TOKENS["teal"])
    b.add_heading(doc, "2. Fair concern or straw man?", 1)
    b.add_matrix(doc, ["Statement", "Fair concern / straw man", "Why?"], [
        ["The cafe must keep food safe and the queue moving.", "", ""],
        ["My family only cares about money.", "", ""],
        ["Not every customer will bring a suitable cup.", "", ""],
    ], [4800, 2200, 3178], font_size=9.0)
    b.add_text(doc, "Write the genuine concern you will answer:")
    b.add_form_table(doc, ["A fair concern is..."], [10178], blank_rows=1, blank_after=42)

    doc.add_page_break()
    b.add_heading(doc, "3. Build the FAIR bridge", 1)
    b.add_form_table(doc, ["Move", "My note", "Purpose check"], [1750, 5550, 2878], blank_rows=4, row_prompts=[
        ["F - Frame", "Shared goal:", "Does it include the audience?"],
        ["A - Acknowledge", "Genuine concern:", "Is this their real concern?"],
        ["I - Introduce", "Evidence + reasoning:", "Does it answer the concern?"],
        ["R - Recommend", "Trial + safeguard + measure:", "Is it workable?"],
    ], blank_after=38)
    b.add_heading(doc, "4. Evidence bank", 1)
    b.add_form_table(doc, ["Evidence from the reading", "What exact claim can it support?", "Boundary / limit"], [3650, 3650, 2878], blank_rows=3, blank_after=32)

    doc.add_page_break()
    b.add_heading(doc, "5. Choose techniques for this audience", 1)
    b.add_form_table(doc, ["Technique", "Exact words I might use", "Effect on the family audience"], [2400, 3900, 3878], blank_rows=3, row_prompts=[
        ["inclusive language", "", ""],
        ["modality", "", ""],
        ["question / emotive words / triad", "", ""],
    ], blank_after=34)
    b.add_callout(doc, "Technique test:", "A technique is effective when its wording, audience, evidence and purpose fit together.")
    b.add_heading(doc, "6. Lock the open letter", 1)
    for prompt in [
        "Position:", "Shared goal:", "Fair concern:", "Evidence 1 + reasoning:", "Evidence 2 + reasoning:",
        "Four-week action:", "Safety / access safeguard:", "What the family will measure:", "What the trial cannot prove:"
    ]:
        b.add_labeled_para(doc, prompt, "________________________________________________________________________________")

    doc.add_page_break()
    b.add_heading(doc, "7. Feedback and visible revision", 1)
    b.add_form_table(doc, ["Reviewer feedback", "Writer's immediate revision"], [5089, 5089], blank_rows=2, row_prompts=[
        ["The family concern is represented fairly when...", "Original sentence:"],
        ["The action would be more convincing / workable if...", "Revised sentence:"],
    ], blank_after=58)
    b.add_heading(doc, "8. Final self-check", 1)
    for item in [
        "I used all four FAIR moves.", "I used two relevant evidence points and explained their job.",
        "I used at least two deliberate persuasive techniques.", "I included a trial, safeguard and measure.",
        "I stated a boundary instead of making a guarantee.", "I made one revision visible."
    ]:
        add_bullet(doc, "[ ] " + item)
    b.add_callout(doc, "Exit:", "A fair concession strengthens persuasion when __________________ because __________________; it becomes weak when __________________.")
    out = ROOT / "Lesson_18_Cup_Family_FAIR_Organiser.docx"
    doc.save(out)
    return out


def build_lucas():
    doc = b.configure_doc("Lucas / ICP Pack")
    normal = doc.styles["Normal"]
    normal.font.size = Pt(13)
    normal.paragraph_format.line_spacing = 1.25
    b.add_title_block(doc, "Accessible persuasive pathway", "One Cup, Two Loyalties", "Same big idea: care for the planet and treat people fairly.", "READ TOGETHER  |  POINT / SAY / WRITE")
    b.add_heading(doc, "The cafe choice", 1)
    b.add_text(doc, "Leila wants her family's cafe to use fewer throwaway cups. Aunt Samira wants to help, but the cafe is very busy. Staff must keep drinks safe and serve customers. The family can test a new system for four weeks.", size=13)
    b.add_matrix(doc, ["Leila cares about", "Aunt Samira cares about"], [["less waste", "safe drinks and a working cafe"]], [5089, 5089], font_size=12.5)
    b.add_callout(doc, "Shared goal:", "A safe, welcoming cafe that makes less waste.")
    b.add_heading(doc, "FAIR - choose one idea for each card", 1)
    b.add_form_table(doc, ["Move", "Choose / point / say"], [2500, 7678], blank_rows=4, row_prompts=[
        ["F - We both want...", "a good cafe  /  less waste  /  both"],
        ["A - A fair worry is...", "dirty cups  /  slow service"],
        ["I - The evidence says...", "the cafe can make a cup policy"],
        ["R - The cafe should...", "try for four weeks and count what happens"],
    ], blank_after=34, font_size=12)
    doc.add_page_break()
    b.add_heading(doc, "Build my short letter", 1)
    for prompt in [
        "Dear Rahman family, we all want ______________________________________.",
        "I think the cafe should ______________________________________________.",
        "A fair worry is ______________________________________________________.",
        "The evidence says ____________________________________________________.",
        "So the cafe could ____________________________________________________.",
        "The trial may show ____________________, but it cannot prove ____________________."
    ]:
        b.add_labeled_para(doc, "Say / write:", prompt, fill=b.TOKENS["pale"], accent=b.TOKENS["teal"])
    b.add_heading(doc, "Make one sentence better", 1)
    b.add_form_table(doc, ["Before", "After"], [5089, 5089], blank_rows=1, row_prompts=[["The cafe will fix all waste.", "The cafe may... because..."]], blank_after=65, font_size=12)
    b.add_callout(doc, "You are finished when:", "You have a shared goal, a fair worry, one fact, one trial action and one careful word such as may or can.")
    out = ROOT / "Lesson_18_Cup_Family_Lucas_Pack.docx"
    doc.save(out)
    return out


def reading_article_html():
    html = r'''<!doctype html>
<html lang="en-AU"><head><meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1">
<title>One Cup, Two Loyalties - Reading Article</title>
<style>
@page{size:A4;margin:0}*{box-sizing:border-box}body{margin:0;background:#d9ddd7;color:#26332c;font-family:Arial,sans-serif}.page{width:210mm;height:297mm;margin:10px auto;background:#fff;position:relative;overflow:hidden;box-shadow:0 8px 28px #0003}.hero{height:62mm;background:linear-gradient(120deg,#23332b,#4e7d62 68%,#e4b44b);padding:14mm;color:#fff;position:relative}.hero:after{content:"";position:absolute;right:13mm;top:12mm;width:38mm;height:38mm;border:5mm solid #fff8;border-radius:50%;box-shadow:inset 0 0 0 5mm #e4b44b88}.kicker{font-size:8pt;letter-spacing:.16em;text-transform:uppercase;font-weight:800;color:#ffe2a0}h1{font:800 27pt/1 Georgia,serif;margin:3mm 0 2mm}.stand{width:130mm;font:700 10.5pt/1.25 Arial}.content{padding:8mm 13mm 13mm;display:grid;grid-template-columns:1.5fr .85fr;gap:7mm}h2{font:800 13pt/1.1 Georgia,serif;color:#29483a;margin:0 0 2mm;border-bottom:2px solid #e4b44b;padding-bottom:1mm}p{font-size:9.2pt;line-height:1.32;margin:0 0 2.4mm}.quote,.fact,.task{padding:3mm;border-radius:2mm;margin:3mm 0}.quote{background:#fff3d9;border-left:4px solid #c9664d;font-weight:700}.fact{background:#edf5ee;border-left:4px solid #4e7d62}.side{background:#f4f1e8;padding:5mm;border-top:5px solid #c9664d}.side h3{font-size:10pt;margin:0 0 2mm;color:#c9664d}.code{display:grid;grid-template-columns:8mm 1fr;gap:1.5mm;margin:1.5mm 0;font-size:8.5pt}.badge{display:grid;place-items:center;background:#29483a;color:#fff;border-radius:2px;font-weight:800}.task{background:#23332b;color:#fff;font-size:9pt}.footer{position:absolute;left:13mm;right:13mm;bottom:5mm;border-top:1px solid #b9c5ba;padding-top:2mm;font-size:6.3pt;color:#657069}.drafts{display:grid;grid-template-columns:1fr 1fr;gap:4mm}.draft{padding:4mm;background:#f4f1e8;border-top:4px solid #c9664d}.draft.good{border-color:#4e7d62}.draft h3{margin:0 0 2mm;font-size:10pt}.tech{display:grid;grid-template-columns:30mm 1fr;gap:2mm;border-bottom:1px solid #ccd4cc;padding:1.5mm 0;font-size:8.5pt}.fair{display:grid;grid-template-columns:repeat(4,1fr);gap:2mm}.fair div{background:#29483a;color:#fff;padding:2.5mm;border-radius:2mm;font-size:8.2pt}.fair b{display:block;color:#ffe2a0;font-size:13pt}.sources{font-size:7.2pt;line-height:1.2}@media print{body{background:#fff}.page{margin:0;box-shadow:none;page-break-after:always}.page:last-child{page-break-after:auto}}</style></head><body>
<section class="page"><header class="hero"><div class="kicker">Original classroom feature | English Unit 3</div><h1>One Cup, Two Loyalties</h1><div class="stand">Can an environmental argument protect both a principle and a family relationship?</div></header><main class="content"><div><h2>The poster on the pantry door</h2><p>At 6:35 on Friday morning, Leila Rahman taped a poster to the pantry door of her family's cafe. In thick green letters it asked, “How many throwaway cups does one busy morning need?” Underneath, she proposed a four-week trial: reward customers who brought a clean reusable cup, offer ceramic mugs to people drinking inside and test a small returnable-mug shelf.</p><p>Leila had spent the holidays helping at the register. She saw towers of cups arrive and watched bins fill before lunch. To her, the cups represented a choice repeated so often that it became invisible. She wanted her family to lead change.</p><p>Her first slogan sounded powerful: “Our cafe is choosing convenience over the planet.” Then she imagined her aunt beside the coffee machine. The sentence made the adults sound careless. Leila knew that was not true.</p><h2>What the slogan leaves out</h2><p>Aunt Samira opened most mornings, checked deliveries, trained staff and made sure wages could be paid. She liked Leila's goal, but not being treated as an obstacle.</p><div class="quote">“A good idea still has to work at 7:45, when twenty people are waiting and a customer hands us a cup that may not be clean.”</div><p>Her concern was not a secret wish to keep wasting cups. She wanted answers: Who would inspect a cup? Where would staff place it? What if it was dirty, damaged or the wrong size? Would a return shelf add washing during the rush? What would the cafe offer someone who forgot a cup?</p><p>Leila felt the argument changing. Repeating that waste mattered would not answer the people affected. She needed to represent their concern fairly and propose a test that could produce local evidence.</p></div><aside class="side"><h3>MARK THE FEATURE</h3><div class="code"><span class="badge">V</span><span>value or responsibility</span></div><div class="code"><span class="badge">C</span><span>genuine concern</span></div><div class="code"><span class="badge">E</span><span>checkable evidence</span></div><div class="code"><span class="badge">T</span><span>technique or framing</span></div><div class="task"><b>Family decision</b><br>Should Sunrise Corner test a reusable-cup system for four weeks?</div><h3>FICTION BOUNDARY</h3><p>The cafe, family and dialogue are invented. Official sources support the evidence section.</p><h3>WORD HELP</h3><p><b>livelihood:</b> work and income that support a person or family</p><p><b>concession:</b> fair acknowledgement of a point made by another side</p><p><b>straw man:</b> an unfairly simplified version of another view</p><p><b>modality:</b> the force or certainty carried by words such as could, should and must</p></aside></main><div class="footer">Facts checked 12 August 2026. Original classroom feature; fictional family and dialogue.</div></section>
<section class="page"><header class="hero" style="height:45mm"><div class="kicker">Evidence + language laboratory</div><h1 style="font-size:23pt">Build a bridge, not a battle</h1></header><main style="padding:8mm 13mm 13mm"><h2>Evidence that changes the conversation</h2><p>Queensland Government guidance suggests offering an incentive to customers who bring a reusable cup or drink in store. It also suggests a mug-return program, possibly with a deposit. These are options, not guarantees; one system will not suit every queue or workspace.</p><p>Food Standards Australia New Zealand explains that a business may decide whether to accept customer reusable cups. It should use a policy and consider cleaning, sanitising and cross-contamination. The business remains responsible for safe and suitable food.</p><div class="fact"><b>Evidence boundary:</b> Official guidance supports avoidance, reuse and a clear safety policy. It does not prove that a four-week trial will remove all waste or cause no delay. A trial can collect local evidence.</div><h2>Two drafts for the family meeting</h2><div class="drafts"><div class="draft"><h3>Draft A - Heat</h3><p>Our throwaway habit is choking the planet because the cafe refuses to change. How can money matter more than our future? We must ban disposable cups now.</p></div><div class="draft good"><h3>Draft B - FAIR bridge</h3><p>We all want Sunrise Corner to remain welcoming and responsible. A morning rush makes safety and speed genuine concerns. Official guidance shows that a cafe can set a clear cup policy, so we should test a four-week system and measure what happens.</p></div></div><h2 style="margin-top:6mm">The FAIR bridge</h2><div class="fair"><div><b>F</b>Frame a shared goal</div><div><b>A</b>Acknowledge a fair concern</div><div><b>I</b>Introduce relevant evidence</div><div><b>R</b>Recommend a workable next step</div></div><h2 style="margin-top:6mm">Technique field guide</h2><div class="tech"><b>Inclusive language</b><span>Frames shared responsibility: we, our, together.</span></div><div class="tech"><b>Modality</b><span>Calibrates force: could, should, must.</span></div><div class="tech"><b>Rhetorical question</b><span>Directs attention - or unfairly pressures an audience.</span></div><div class="tech"><b>Emotive words</b><span>Signal importance, but may attack instead of explain.</span></div><div class="tech"><b>Concession</b><span>Shows the writer understands a genuine concern.</span></div><div class="tech"><b>Call to action</b><span>Names a specific, safe and workable next step.</span></div><div class="task" style="margin-top:5mm"><b>Open-letter challenge</b><br>Write 9-11 sentences to the Rahman family. Be firm about the ethical goal, fair about the family's responsibilities and specific about what the trial will test.</div><div class="sources"><b>Source trail:</b> Queensland Government business guide and Action on single-use plastic items; Food Standards Australia New Zealand customer cups and food packaging guidance; Brisbane City Council bin guidance as a local disposal example.</div></main><div class="footer">Do not claim that every cup has the same materials or disposal path. A trial gathers evidence; it does not guarantee a result.</div></section></body></html>'''
    out = ROOT / "Lesson_18_Cup_Family_Reading_Article.html"
    out.write_text(html, encoding="utf-8")
    return out


def _presentation_html_legacy():
    html = r'''<!doctype html><html lang="en-AU"><head><meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1"><title>Lesson 18 Alternative | One Cup, Two Loyalties</title>
<style>
:root{--dark:#19251f;--forest:#29483a;--green:#4e7d62;--gold:#e4b44b;--rust:#c9664d;--cream:#fff7e7;--paper:#f1eee4;--ink:#f7fbf7;--muted:#c9d4cc}*{box-sizing:border-box}html,body{height:100%;margin:0;overflow:hidden;background:var(--dark);color:var(--ink);font-family:Arial,sans-serif}.slide{display:none;position:absolute;inset:0;padding:5vh 6vw 12vh;background:radial-gradient(circle at 82% 16%,#466a55 0,#29483a 38%,#19251f 78%)}.slide.active{display:grid;grid-template-rows:auto 1fr;gap:2.5vh;animation:rise .25s ease}.slide:after{content:"";position:absolute;right:-5vw;bottom:5vh;width:28vw;height:28vw;border:4vw solid #ffffff0b;border-radius:50%;pointer-events:none}@keyframes rise{from{opacity:0;transform:translateY(8px)}}.top{display:flex;justify-content:space-between;align-items:start;gap:2rem}.kicker{color:var(--gold);text-transform:uppercase;letter-spacing:.16em;font-size:clamp(12px,1vw,17px);font-weight:850}h1,h2{font-family:Georgia,serif;margin:.2em 0;line-height:1.02}h1{font-size:clamp(48px,6vw,95px)}h2{font-size:clamp(35px,4.3vw,72px)}h3{font-size:clamp(22px,2vw,34px);margin:0 0 .5rem}.tag{border:1px solid #ffffff44;border-radius:999px;padding:.5rem .8rem;color:var(--muted);font-weight:700}.content{display:grid;align-content:center;gap:2.3vh;position:relative;z-index:2}.two{grid-template-columns:1fr 1fr;gap:2vw}.three{grid-template-columns:repeat(3,1fr);gap:1.4vw}.four{grid-template-columns:repeat(4,1fr);gap:1vw}.card,.prompt,.warning{background:#ffffff10;border:1px solid #ffffff2d;border-radius:20px;padding:clamp(1rem,2vw,2rem);box-shadow:0 15px 35px #0003}.card p,.body{font-size:clamp(18px,1.65vw,30px);line-height:1.35;margin:.4em 0}.prompt{background:var(--cream);color:#26332c;border-left:8px solid var(--gold)}.prompt p{font:700 clamp(22px,2vw,36px)/1.3 Georgia,serif;margin:0}.warning{background:#c9664d22;border-color:var(--rust);font-size:clamp(19px,1.6vw,28px);font-weight:750}.headline{font:800 clamp(25px,3vw,52px)/1.12 Georgia,serif;min-height:24vh;display:flex;align-items:center}.headline.hot{border-color:var(--rust)}.headline.calm{border-color:var(--green)}.codes{display:grid;grid-template-columns:repeat(4,1fr);gap:1vw}.code{padding:1.2rem;background:#ffffff10;border-radius:18px;text-align:center;font-size:clamp(16px,1.3vw,24px)}.code b{display:block;font-size:clamp(30px,3.5vw,58px);color:var(--gold)}.fair{display:grid;grid-template-columns:repeat(4,1fr);gap:1vw}.move{padding:1.2rem;background:#ffffff12;border:1px solid #ffffff33;border-radius:18px}.move b{display:block;color:var(--gold);font-size:clamp(34px,4vw,68px)}.move span{font-size:clamp(16px,1.45vw,26px);font-weight:700}.choicegrid,.quiz{display:grid;grid-template-columns:repeat(3,1fr);gap:1vw}.choice,.answer,.reveal,.primary{font:700 clamp(16px,1.25vw,23px)/1.3 Arial;border:1px solid #ffffff44;background:#ffffff10;color:#fff;border-radius:16px;padding:1rem;cursor:pointer}.choice:hover,.answer:hover,.reveal:hover,.primary:hover,.choice:focus,.answer:focus,.reveal:focus,.primary:focus{outline:3px solid var(--gold);outline-offset:2px}.feedback{min-height:3rem;padding:.7rem 1rem;border-radius:12px;background:#ffffff0d;color:var(--muted);font-weight:700}.feedback.good{background:#4e7d6244;color:#eaffef}.feedback.retry{background:#c9664d44;color:#fff1ed}.tech{font-size:clamp(22px,2.4vw,42px);font-weight:850}.tech small{display:block;color:var(--muted);font-size:.52em;margin-top:.7rem;line-height:1.25}.statement{background:#ffffff10;border:1px solid #ffffff2b;border-radius:16px;padding:1rem}.statement p{font-size:clamp(15px,1.25vw,22px);line-height:1.25;min-height:8vh}.answers{display:flex;gap:.4rem}.answer{padding:.55rem;font-size:clamp(13px,1vw,18px);flex:1}.result{font-size:14px;min-height:2.5rem;padding-top:.5rem}.model-line{opacity:.2;padding:.65rem;border-left:5px solid var(--gold);font-size:clamp(16px,1.28vw,23px);line-height:1.25}.model-line.shown{opacity:1;background:#ffffff0b}.criteria{display:grid;grid-template-columns:1fr 1fr;gap:.7rem}.criterion{padding:.8rem;border-radius:12px;background:#ffffff10;font-size:clamp(16px,1.25vw,22px);font-weight:750}.timer{font:800 clamp(65px,9vw,145px)/1 monospace;color:var(--gold)}.timer-controls{display:flex;gap:.7rem}.controls{position:fixed;left:0;right:0;bottom:0;height:9vh;min-height:62px;background:#111a16f2;border-top:1px solid #ffffff25;display:grid;grid-template-columns:1fr auto 1fr;align-items:center;padding:0 2vw;z-index:20}.nav,.tools{display:flex;gap:.55rem}.tools{justify-content:flex-end}.control{border:1px solid #ffffff33;background:#ffffff08;color:#fff;border-radius:12px;padding:.65rem .9rem;cursor:pointer}.control:disabled{opacity:.3}.count{font-weight:800;color:var(--muted)}.progress{position:fixed;left:0;bottom:9vh;height:5px;background:var(--gold);z-index:21;transition:width .2s}.notes{position:fixed;right:2vw;bottom:10.5vh;width:min(520px,44vw);max-height:58vh;overflow:auto;background:#fff;color:#26332c;border-radius:18px;padding:1.1rem;box-shadow:0 20px 55px #0008;z-index:30;display:none;font-size:16px;line-height:1.4}.notes.open{display:block}.notes:before{content:"TEACHER NOTE";display:block;color:#4e7d62;font-size:12px;letter-spacing:.14em;font-weight:850;margin-bottom:.5rem}@media(max-width:900px){.two,.three,.four,.choicegrid,.quiz,.fair,.codes{grid-template-columns:1fr 1fr}.slide{padding:3vh 4vw 12vh}.content{overflow:auto;align-content:start}.notes{width:84vw}}@media(prefers-reduced-motion:reduce){*{animation:none!important;transition:none!important}}</style></head><body><main>
<section class="slide active" data-notes="Begin with the relationship, not administration. Ask: What two loyalties can you already predict from the title?"><div class="top"><div><div class="kicker">Lesson 18 alternative | persuasive ethics</div><h1>One Cup,<br>Two Loyalties</h1></div><span class="tag">Family cafe decision</span></div><div class="content"><div class="prompt"><p>Can a writer hold a firm ethical position without turning family into the enemy?</p></div></div></section>
<section class="slide" data-notes="Students identify selection and framing. Do not decide which headline is correct. Each foregrounds one value and hides another."><div class="top"><div><div class="kicker">Same family | two headlines</div><h2>What does each version hide?</h2></div><span class="tag">4 minutes</span></div><div class="content two"><div class="card headline hot">Our cafe is choosing convenience over the planet.</div><div class="card headline calm">Our cafe is protecting jobs during the morning rush.</div><div class="prompt" style="grid-column:1/-1"><p>What value is foregrounded? Who is simplified?</p></div></div></section>
<section class="slide" data-notes="Name the family and cafe as fictional. Students persuade family decision-makers, not defeat opponents."><div class="top"><div><div class="kicker">Mission | meeting tomorrow</div><h2>Advise the Rahman family</h2></div><span class="tag">2 minutes</span></div><div class="content three"><div class="card"><h3>READ</h3><p>Understand the values and evidence.</p></div><div class="card"><h3>BRIDGE</h3><p>Answer a genuine concern.</p></div><div class="card"><h3>RECOMMEND</h3><p>Design a four-week trial.</p></div><div class="warning" style="grid-column:1/-1">Finished product: 9-11 sentence open letter with one visible revision.</div></div></section>
<section class="slide" data-notes="Use the original classroom feature. Narrative and family voices are fictional; evidence is sourced. Compare one mark only after independent reading."><div class="top"><div><div class="kicker">Reading pack | evidence adviser</div><h2>Mark what each voice carries</h2></div><span class="tag">9 minutes</span></div><div class="content"><div class="codes"><div class="code"><b>V</b>value or responsibility</div><div class="code"><b>C</b>genuine concern</div><div class="code"><b>E</b>checkable evidence</div><div class="code"><b>T</b>technique or framing</div></div><div class="prompt"><p>Find two values, two concerns, three evidence points and two techniques.</p></div></div></section>
<section class="slide" data-notes="Understanding does not mean agreeing. Students identify one shared goal before debating the action."><div class="top"><div><div class="kicker">Viewpoint | organiser 1</div><h2>Different pressures. A possible shared goal.</h2></div><span class="tag">5 minutes</span></div><div class="content two"><div class="card"><h3>Leila notices</h3><p>repeated waste | invisible habits | leadership | being heard</p></div><div class="card"><h3>Aunt Samira notices</h3><p>safe drinks | queue speed | staff workload | wages | access</p></div><div class="prompt" style="grid-column:1/-1"><p>Both may want Sunrise Corner to be ___ and ___.</p></div></div></section>
<section class="slide" data-notes="Teach the four moves explicitly. Students name each job before the description is discussed."><div class="top"><div><div class="kicker">Explicit model | FAIR bridge</div><h2>A concession is not surrender</h2></div><span class="tag">7 minutes</span></div><div class="content"><div class="fair"><div class="move"><b>F</b><span>Frame a shared goal</span></div><div class="move"><b>A</b><span>Acknowledge a fair concern</span></div><div class="move"><b>I</b><span>Introduce relevant evidence</span></div><div class="move"><b>R</b><span>Recommend a workable action</span></div></div><div class="prompt"><p>Fairness builds credibility only when the next evidence answers the concern.</p></div></div></section>
<section class="slide" data-notes="Students decide which version represents the family accurately, justify, then test. A straw man is easier to attack because it is not the real concern."><div class="top"><div><div class="kicker">Fairness test | straw man</div><h2>Which concern deserves an answer?</h2></div><span class="tag">Organiser 2</span></div><div class="content"><div class="choicegrid" id="fairChoices"><button class="choice" data-kind="retry" data-msg="Straw man: this invents a selfish motive instead of representing the actual responsibilities.">My family only cares about money.</button><button class="choice" data-kind="good" data-msg="Fair: this accurately names livelihood, safety and workflow without deciding that they outweigh the environment.">The family must protect wages, safe service and the morning queue.</button><button class="choice" data-kind="retry" data-msg="Straw man: disagreement with one plan does not prove opposition to all environmental action.">Anyone against my plan wants more pollution.</button></div><div class="feedback" id="fairFeedback">Choose, justify, then test. Is this their real concern?</div></div></section>
<section class="slide" data-notes="A technique is judged by fit. Ask students to name audience and purpose in every justification."><div class="top"><div><div class="kicker">Technique lab | force + fit</div><h2>Which tool serves this family audience?</h2></div><span class="tag">6 minutes</span></div><div class="content four"><div class="card tech">we | our | together<small>inclusive language</small></div><div class="card tech">could | should | must<small>modality</small></div><div class="card tech">Why not test...?<small>rhetorical question</small></div><div class="card tech">choking | careless<small>loaded vocabulary</small></div><div class="prompt" style="grid-column:1/-1"><p>X is effective here because it positions the family to ___ without ___.</p></div></div></section>
<section class="slide" data-notes="Students classify and justify before clicking. The central boundary is that a trial can collect local evidence, not guarantee a permanent result."><div class="top"><div><div class="kicker">Evidence boundary</div><h2>Evidence, inference or overclaim?</h2></div><span class="tag">5 minutes</span></div><div class="content"><div class="quiz" id="boundaryQuiz"><div class="statement" data-correct="e"><p>Queensland guidance suggests incentives or mug-return programs.</p><div class="answers"><button class="answer" data-a="e">Evidence</button><button class="answer" data-a="i">Inference</button><button class="answer" data-a="o">Overclaim</button></div><div class="result"></div></div><div class="statement" data-correct="i"><p>A well-designed trial may reduce cup use without unacceptable delays.</p><div class="answers"><button class="answer" data-a="e">Evidence</button><button class="answer" data-a="i">Inference</button><button class="answer" data-a="o">Overclaim</button></div><div class="result"></div></div><div class="statement" data-correct="o"><p>After the trial, no cafe cup will ever enter landfill.</p><div class="answers"><button class="answer" data-a="e">Evidence</button><button class="answer" data-a="i">Inference</button><button class="answer" data-a="o">Overclaim</button></div><div class="result"></div></div></div></div></section>
<section class="slide" data-notes="Read the model whole first. Reveal one move at a time. Students star one FAIR move and underline two techniques to transfer."><div class="top"><div><div class="kicker">Annotated model</div><h2>Firm, fair and testable</h2></div><span class="tag">5 minutes</span></div><div class="content two"><div class="card" id="model"><div class="model-line"><b>F:</b> We all want Sunrise Corner to remain welcoming and responsible.</div><div class="model-line"><b>Position:</b> We should run a four-week reusable-cup trial.</div><div class="model-line"><b>A:</b> Staff must protect food safety and keep the queue moving.</div><div class="model-line"><b>I:</b> Official guidance says a cafe can set a cup policy and manage contamination risks.</div><div class="model-line"><b>R:</b> Test one handover point, refuse unsuitable cups and record delays.</div><div class="model-line"><b>Boundary:</b> The trial cannot solve all waste, but it can produce local evidence.</div></div><div class="card"><h3>Writer's question</h3><p>What job is each sentence doing?</p><button class="primary" id="revealModel">Reveal next move</button></div></div></section>
<section class="slide" data-notes="Confer first with students whose plan attacks a person or includes facts without explaining how they support the proposal."><div class="top"><div><div class="kicker">Plan | organiser 3-6</div><h2>Lock the bridge before writing</h2></div><span class="tag">5 minutes</span></div><div class="content two"><div class="criteria"><div class="criterion">clear position</div><div class="criterion">shared goal</div><div class="criterion">fair concern</div><div class="criterion">2 evidence + reasoning links</div><div class="criterion">trial action</div><div class="criterion">safeguard + measure</div><div class="criterion">2 techniques</div><div class="criterion">evidence boundary</div></div><div class="prompt"><p>Make the next step workable, not merely passionate.</p></div></div></section>
<section class="slide" data-notes="Independent writing. Students may qualify or reject the proposed trial if their alternative remains fair, evidenced and workable."><div class="top"><div><div class="kicker">Write | open letter</div><h2>9-11 sentences. One FAIR bridge.</h2></div><span class="tag">12 minutes</span></div><div class="content two"><div><div class="timer" id="timer">12:00</div><div class="timer-controls"><button class="primary" id="timerStart">Start / pause</button><button class="primary" id="timerReset">Reset</button></div></div><div class="card"><h3>Useful launches</h3><p>We all want...<br>A fair concern is...<br>Official guidance explains...<br>This matters because...<br>Therefore, the cafe should...<br>The trial cannot prove..., but...</p></div></div></section>
<section class="slide" data-notes="Reviewer diagnoses; writer controls the change. Revision must happen immediately and remain visible."><div class="top"><div><div class="kicker">Feedback | revision</div><h2>Stress-test the bridge</h2></div><span class="tag">4 minutes</span></div><div class="content two"><div class="prompt"><p>Your letter represents the family's concern fairly when...</p></div><div class="prompt" style="border-left-color:var(--rust)"><p>Your recommendation would be more convincing or workable if...</p></div><div class="warning" style="grid-column:1/-1;text-align:center">Revise one sentence now. Make the change visible.</div></div></section>
<section class="slide" data-notes="Optional depth. Keep urgency, remove the personal attack, and add a defensible action."><div class="top"><div><div class="kicker">Depth A | optional</div><h2>Turn down the heat, not the importance</h2></div><span class="tag">7 minutes</span></div><div class="content"><div class="warning">Our throwaway habit is choking the planet because my family refuses to change.</div><div class="prompt"><p>Retain urgency. Remove the attack. Add one workable action.</p></div></div></section>
<section class="slide" data-notes="Optional depth. Measures should include environmental effect and human implementation, not only one preferred outcome."><div class="top"><div><div class="kicker">Depth B | optional</div><h2>What would make the trial worth keeping?</h2></div><span class="tag">8 minutes</span></div><div class="content three"><div class="card"><h3>Use</h3><p>disposable cups | customer participation</p></div><div class="card"><h3>Work</h3><p>service delay | washing | staff concerns</p></div><div class="card"><h3>Safety</h3><p>refused cups | procedure problems</p></div><div class="prompt" style="grid-column:1/-1"><p>Continue, change or end the trial if...</p></div></div></section>
<section class="slide" data-notes="Collect this sentence. Group support by straw-man thinking, evidence relevance or vague action. This is the safe stopping point."><div class="top"><div><div class="kicker">Exit evidence</div><h2>When does a concession strengthen persuasion?</h2></div><span class="tag">1 minute</span></div><div class="content"><div class="prompt"><p>A fair concession strengthens persuasion when ___ because ___; it becomes weak when ___.</p></div><p class="body" style="text-align:center;color:var(--muted)">Submit this with your visible revision.</p></div></section>
</main><div class="progress" id="progress"></div><aside class="notes" id="notes" aria-live="polite"></aside><nav class="controls"><div class="nav"><button class="control" id="prev">← Previous</button><button class="control" id="next">Next →</button></div><div class="count" id="count">1 / 16</div><div class="tools"><button class="control" id="notesBtn">Notes</button><button class="control" id="resetBtn">Reset</button><button class="control" id="fullBtn">Fullscreen</button></div></nav>
<script>
const slides=[...document.querySelectorAll('.slide')];let current=0;const prev=document.getElementById('prev'),next=document.getElementById('next'),count=document.getElementById('count'),progress=document.getElementById('progress'),notes=document.getElementById('notes');function show(i){current=Math.max(0,Math.min(slides.length-1,i));slides.forEach((s,n)=>s.classList.toggle('active',n===current));prev.disabled=current===0;next.disabled=current===slides.length-1;count.textContent=`${current+1} / ${slides.length}`;progress.style.width=`${(current+1)/slides.length*100}%`;notes.classList.remove('open');notes.textContent=slides[current].dataset.notes||''}prev.onclick=()=>show(current-1);next.onclick=()=>show(current+1);document.addEventListener('keydown',e=>{if(['INPUT','TEXTAREA','SELECT'].includes(document.activeElement.tagName))return;if(['ArrowRight','PageDown',' '].includes(e.key)){e.preventDefault();show(current+1)}if(['ArrowLeft','PageUp'].includes(e.key)){e.preventDefault();show(current-1)}if(e.key==='Home')show(0);if(e.key==='End')show(slides.length-1);if(e.key.toLowerCase()==='n')notes.classList.toggle('open');if(e.key.toLowerCase()==='f')toggleFull()});document.getElementById('notesBtn').onclick=()=>notes.classList.toggle('open');function toggleFull(){if(!document.fullscreenElement)document.documentElement.requestFullscreen?.();else document.exitFullscreen?.()}document.getElementById('fullBtn').onclick=toggleFull;
document.querySelectorAll('#fairChoices .choice').forEach(btn=>btn.onclick=()=>{const f=document.getElementById('fairFeedback');f.textContent=btn.dataset.msg;f.className='feedback '+btn.dataset.kind});function resetFair(){const f=document.getElementById('fairFeedback');f.textContent='Choose, justify, then test. Is this their real concern?';f.className='feedback'}
document.querySelectorAll('.statement').forEach(card=>card.querySelectorAll('.answer').forEach(btn=>btn.onclick=()=>{const ok=btn.dataset.a===card.dataset.correct;const r=card.querySelector('.result');r.textContent=ok?'Defensible. Explain the boundary.':'Retry: directly stated, cautiously inferred or guaranteed?';r.style.color=ok?'#bff4cd':'#ffd0c5'}));function resetQuiz(){document.querySelectorAll('.result').forEach(r=>{r.textContent='';r.style.color=''})}
let modelIndex=0;const modelLines=[...document.querySelectorAll('.model-line')];document.getElementById('revealModel').onclick=()=>{if(modelIndex<modelLines.length)modelLines[modelIndex++].classList.add('shown')};function resetModel(){modelIndex=0;modelLines.forEach(x=>x.classList.remove('shown'))}
let seconds=720,timerId=null;const timer=document.getElementById('timer');function draw(){timer.textContent=`${String(Math.floor(seconds/60)).padStart(2,'0')}:${String(seconds%60).padStart(2,'0')}`}document.getElementById('timerStart').onclick=()=>{if(timerId){clearInterval(timerId);timerId=null;return}timerId=setInterval(()=>{if(seconds>0){seconds--;draw()}else{clearInterval(timerId);timerId=null}},1000)};document.getElementById('timerReset').onclick=()=>{clearInterval(timerId);timerId=null;seconds=720;draw()};function resetCurrent(){const s=slides[current];if(s.querySelector('#fairChoices'))resetFair();if(s.querySelector('#boundaryQuiz'))resetQuiz();if(s.querySelector('#model'))resetModel();if(s.querySelector('#timer')){clearInterval(timerId);timerId=null;seconds=720;draw()}}document.getElementById('resetBtn').onclick=resetCurrent;show(0);draw();
</script></body></html>'''
    out = ROOT / "Lesson_18_Cup_Family_Persuasive_Presentation.html"
    out.write_text(html, encoding="utf-8")
    return out


def presentation_html():
    presentation_builder_path = Path(__file__).with_name("build_cup_family_presentation.py")
    presentation_spec = importlib.util.spec_from_file_location("cup_family_presentation", presentation_builder_path)
    presentation_builder = importlib.util.module_from_spec(presentation_spec)
    presentation_spec.loader.exec_module(presentation_builder)
    return presentation_builder.build_presentation(ROOT)


if __name__ == "__main__":
    outputs = [presentation_html(), reading_article_html(), build_reading_pack(), build_organiser(), build_lucas()]
    print("Built Lesson 18 alternative package:")
    for output in outputs:
        print(" -", output.name)
