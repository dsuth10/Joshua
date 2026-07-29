from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE


ROOT = Path(__file__).resolve().parents[1]

# compact_reference_guide + named A4 classroom override
TOKENS = {
    "preset": "compact_reference_guide",
    "override": "turtle_safe_nights_a4_classroom",
    "page_width": 11906,
    "page_height": 16838,
    "margin": 864,
    "header_distance": 360,
    "footer_distance": 360,
    "content_width": 10178,
    "table_indent": 120,
    "cell_top": 90,
    "cell_bottom": 90,
    "cell_start": 120,
    "cell_end": 120,
    "font": "Calibri",
    "body_size": 10.5,
    "body_after": 4,
    "body_line": 1.15,
    "title_size": 23,
    "h1_size": 15,
    "h2_size": 12.5,
    "h3_size": 11.5,
    "navy": "071C2C",
    "deep_sea": "10374F",
    "teal": "168C86",
    "moon": "F2C14E",
    "coral": "EF6F61",
    "pale": "EAF7F4",
    "warm": "FFF4D6",
    "mist": "EEF3F5",
    "ink": "162A35",
    "muted": "526873",
    "white": "FFFFFF",
    "border": "9EB8C1",
}


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=None, start=None, bottom=None, end=None):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    values = {
        "top": top if top is not None else TOKENS["cell_top"],
        "start": start if start is not None else TOKENS["cell_start"],
        "bottom": bottom if bottom is not None else TOKENS["cell_bottom"],
        "end": end if end is not None else TOKENS["cell_end"],
    }
    for m, value in values.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, **edges):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge_name, edge in edges.items():
        tag = f"w:{edge_name}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key in ["val", "sz", "space", "color"]:
            if key in edge:
                element.set(qn(f"w:{key}"), str(edge[key]))


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_table_geometry(table, widths, indent=TOKENS["table_indent"]):
    if sum(widths) != TOKENS["content_width"]:
        raise ValueError(f"Table widths must sum to {TOKENS['content_width']}: {widths}")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(TOKENS["content_width"]))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        prevent_row_split(row)
        for idx, cell in enumerate(row.cells):
            width = widths[min(idx, len(widths) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_run(run, size=None, bold=None, color=None, italic=None, font=None):
    name = font or TOKENS["font"]
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if italic is not None:
        run.italic = italic


def style_paragraph(paragraph, before=0, after=None, line=None, keep_next=False):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(TOKENS["body_after"] if after is None else after)
    pf.line_spacing = TOKENS["body_line"] if line is None else line
    pf.keep_with_next = keep_next


def add_text(doc, text="", size=None, bold=False, color=None, italic=False,
             align=None, before=0, after=None, line=None, style=None, keep_next=False):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    style_paragraph(p, before=before, after=after, line=line, keep_next=keep_next)
    if text:
        r = p.add_run(text)
        set_run(r, size=size or TOKENS["body_size"], bold=bold,
                color=color or TOKENS["ink"], italic=italic)
    return p


def add_labeled_para(doc, label, text, fill=None, accent=None, after=5):
    p = doc.add_paragraph()
    style_paragraph(p, after=after)
    if fill:
        p_pr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill)
        p_pr.append(shd)
        borders = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "22")
        left.set(qn("w:space"), "8")
        left.set(qn("w:color"), accent or TOKENS["teal"])
        borders.append(left)
        p_pr.append(borders)
        spacing = p_pr.find(qn("w:spacing"))
        if spacing is None:
            spacing = OxmlElement("w:spacing")
            p_pr.append(spacing)
        spacing.set(qn("w:before"), "90")
        spacing.set(qn("w:after"), "90")
    r1 = p.add_run(label + " ")
    set_run(r1, size=TOKENS["body_size"], bold=True, color=accent or TOKENS["deep_sea"])
    r2 = p.add_run(text)
    set_run(r2, size=TOKENS["body_size"], color=TOKENS["ink"])
    return p


def add_callout(doc, label, text, fill=None, accent=None, size=None):
    return add_labeled_para(
        doc, label, text,
        fill=fill or TOKENS["warm"],
        accent=accent or TOKENS["coral"],
        after=6,
    )


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    style_paragraph(
        p,
        before={1: 9, 2: 7, 3: 5}[level],
        after={1: 5, 2: 4, 3: 3}[level],
        line=1.0,
        keep_next=True,
    )
    r = p.add_run(text)
    set_run(
        r,
        size={1: TOKENS["h1_size"], 2: TOKENS["h2_size"], 3: TOKENS["h3_size"]}[level],
        bold=True,
        color={1: TOKENS["deep_sea"], 2: TOKENS["teal"], 3: TOKENS["deep_sea"]}[level],
    )
    return p


def add_title_block(doc, kicker, title, subtitle, pack_label):
    p = doc.add_paragraph()
    style_paragraph(p, before=0, after=2, line=1.0)
    r = p.add_run(kicker.upper())
    set_run(r, size=9, bold=True, color=TOKENS["teal"])

    p = doc.add_paragraph()
    style_paragraph(p, before=0, after=3, line=1.0, keep_next=True)
    r = p.add_run(title)
    set_run(r, size=TOKENS["title_size"], bold=True, color=TOKENS["navy"])

    p = doc.add_paragraph()
    style_paragraph(p, before=0, after=8, line=1.05)
    r = p.add_run(subtitle)
    set_run(r, size=11.5, color=TOKENS["muted"], italic=True)

    p_pr = p._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "16")
    bottom.set(qn("w:space"), "8")
    bottom.set(qn("w:color"), TOKENS["moon"])
    borders.append(bottom)
    p_pr.append(borders)

    tag = doc.add_paragraph()
    tag.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    style_paragraph(tag, after=6, line=1.0)
    r = tag.add_run(pack_label)
    set_run(r, size=8.5, bold=True, color=TOKENS["muted"])


def set_cell_text(cell, text, size=9.5, bold=False, color=None, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    style_paragraph(p, after=0, line=1.08)
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_run(r, size=size, bold=bold, color=color or TOKENS["ink"])


def add_matrix(doc, headers, rows, widths, header_fill=None, font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    table.style = "Table Grid"
    mark_header_row(table.rows[0])
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, size=9.2, bold=True,
                      color=TOKENS["white"], align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(table.rows[0].cells[i], header_fill or TOKENS["deep_sea"])
    for ridx, row_data in enumerate(rows):
        row = table.add_row()
        for i, value in enumerate(row_data):
            set_cell_text(row.cells[i], value, size=font_size)
            if ridx % 2 == 1:
                set_cell_shading(row.cells[i], "F7FAFB")
    add_text(doc, "", after=2)
    return table


def add_form_table(doc, headers, widths, blank_rows=1, row_prompts=None, font_size=9.5, blank_after=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    mark_header_row(table.rows[0])
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, size=9.2, bold=True,
                      color=TOKENS["white"], align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(table.rows[0].cells[i], TOKENS["deep_sea"])
    for ridx in range(blank_rows):
        row = table.add_row()
        prompts = row_prompts[ridx] if row_prompts and ridx < len(row_prompts) else [""] * len(headers)
        for i, prompt in enumerate(prompts):
            set_cell_text(row.cells[i], prompt, size=font_size, color=TOKENS["muted"])
            cell_p = row.cells[i].paragraphs[0]
            cell_p.paragraph_format.space_after = Pt(
                blank_after if blank_after is not None else (26 if blank_rows <= 2 else 17)
            )
    add_text(doc, "", after=2)
    return table


def add_footer(section, label):
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    style_paragraph(hp, after=0, line=1.0)
    r = hp.add_run("ENGLISH UNIT 3  |  LESSON 17 ALTERNATIVE")
    set_run(r, size=7.5, bold=True, color=TOKENS["muted"])

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(fp, after=0, line=1.0)
    r = fp.add_run(label + "  |  Turtle-Safe Nights")
    set_run(r, size=7.5, color=TOKENS["muted"])


def configure_doc(label):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(TOKENS["page_width"] / 1440)
    section.page_height = Inches(TOKENS["page_height"] / 1440)
    section.top_margin = Inches(TOKENS["margin"] / 1440)
    section.bottom_margin = Inches(TOKENS["margin"] / 1440)
    section.left_margin = Inches(TOKENS["margin"] / 1440)
    section.right_margin = Inches(TOKENS["margin"] / 1440)
    section.header_distance = Inches(TOKENS["header_distance"] / 1440)
    section.footer_distance = Inches(TOKENS["footer_distance"] / 1440)
    add_footer(section, label)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = TOKENS["font"]
    normal._element.rPr.rFonts.set(qn("w:ascii"), TOKENS["font"])
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), TOKENS["font"])
    normal.font.size = Pt(TOKENS["body_size"])
    normal.font.color.rgb = RGBColor.from_string(TOKENS["ink"])
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(TOKENS["body_after"])
    normal.paragraph_format.line_spacing = TOKENS["body_line"]

    for level in [1, 2, 3]:
        style = styles[f"Heading {level}"]
        style.font.name = TOKENS["font"]
        style._element.rPr.rFonts.set(qn("w:ascii"), TOKENS["font"])
        style._element.rPr.rFonts.set(qn("w:hAnsi"), TOKENS["font"])
        style.font.size = Pt({1: TOKENS["h1_size"], 2: TOKENS["h2_size"], 3: TOKENS["h3_size"]}[level])
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(
            {1: TOKENS["deep_sea"], 2: TOKENS["teal"], 3: TOKENS["deep_sea"]}[level]
        )
        style.paragraph_format.space_before = Pt({1: 9, 2: 7, 3: 5}[level])
        style.paragraph_format.space_after = Pt({1: 5, 2: 4, 3: 3}[level])
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = True

    # Real list styles, resolved to the compact_reference_guide alignment.
    for style_name in ["List Bullet", "List Number"]:
        s = styles[style_name]
        s.font.name = TOKENS["font"]
        s.font.size = Pt(TOKENS["body_size"])
        s.paragraph_format.left_indent = Inches(0.375)
        s.paragraph_format.first_line_indent = Inches(-0.188)
        s.paragraph_format.space_after = Pt(4)
        s.paragraph_format.line_spacing = 1.15
    return doc


def build_reading_pack():
    doc = configure_doc("Student Reading Pack")
    add_title_block(
        doc,
        "Council evidence briefing",
        "Turtle-Safe Nights",
        "How one change to the night horizon can create a chain of consequences",
        "READ • MARK E / A / V / ?",
    )
    add_callout(
        doc,
        "Decision:",
        "Should a coastal council strengthen beachfront-lighting rules during locally verified marine turtle nesting and hatching periods?",
    )
    add_heading(doc, "The night horizon is a map", 1)
    add_text(
        doc,
        "Most marine turtle nesting and hatching activity happens at night. A nesting female approaches the coast from the sea, and newly emerged hatchlings must move from the nest to the water. In a naturally dark landscape, the open ocean horizon is usually the brightest and lowest part of the scene. Hatchlings use this light pattern, together with the beach landscape, to orient toward the sea.",
    )
    add_text(
        doc,
        "Artificial light can change that map. A single unshielded light may be visible from a beach, while many smaller lights can combine as skyglow above a town. National guidance reports that brightly lit beaches can disturb or discourage nesting females. Direct light and skyglow can also disorient hatchlings, drawing them inland, making them circle or causing them to remain near the shore.",
    )
    add_labeled_para(
        doc,
        "Consequence boundary:",
        "Disorientation increases risks such as exhaustion, dehydration, overheating, predation or vehicle strike. It does not mean that every light produces the same response in every turtle.",
        fill=TOKENS["pale"],
        accent=TOKENS["teal"],
    )
    add_heading(doc, "An opportunity before the season begins", 1)
    add_text(
        doc,
        "Queensland maps significant turtle nesting areas to help councils and other plan-makers identify places where lighting requires special care. The Sea Turtle Sensitive Area Model Code gives local governments example planning provisions, including ways to avoid direct artificial light reaching beaches, the ocean and the sky. The model is voluntary and can be refined for local conditions.",
    )
    add_text(
        doc,
        "This creates an opportunity: a council can review its rules before the next local nesting and hatching period. It could keep existing arrangements, strengthen rules in verified turtle-sensitive areas or set a broader standard. A defensible decision needs the local nesting season, an audit of existing lights, safety requirements and a way to monitor results.",
    )
    add_callout(
        doc,
        "Mark the text:",
        "E = turtle or lighting evidence   |   A = council action   |   V = a stakeholder value or concern   |   ? = more evidence or a boundary is needed",
        fill=TOKENS["warm"],
        accent=TOKENS["moon"],
    )

    doc.add_page_break()
    add_heading(doc, "Safe light does not mean light everywhere", 1)
    add_text(
        doc,
        "The Australian Government’s best-practice starting point is natural darkness, with light added only for a specific purpose. The guidance also recognises that artificial light supports human safety, amenity and work. Where those goals compete, councils and designers may need creative solutions that meet both safety and wildlife objectives.",
    )
    add_matrix(
        doc,
        ["Lighting move", "What it is designed to do"],
        [
            ["Use light only where and when a task requires it.", "Avoid unnecessary light and reduce the time it operates."],
            ["Keep lights low, directed and shielded.", "Light the intended path or doorway instead of the beach, ocean or sky."],
            ["Use the lowest suitable intensity and adaptive controls.", "Meet the task while reducing excess brightness and duration."],
            ["Reduce blue, violet and ultraviolet wavelengths.", "Use a spectrum generally less disruptive to turtles, alongside all other design measures."],
            ["Audit and monitor.", "Check what turtles can see and improve the plan if light or behaviour outcomes are not acceptable."],
        ],
        [3300, 6878],
        font_size=9.2,
    )
    add_labeled_para(
        doc,
        "Important:",
        "Changing a bulb to amber is not a complete solution. Need, intensity, direction, shielding, height, timing, reflective surfaces and local monitoring still matter.",
        fill=TOKENS["warm"],
        accent=TOKENS["coral"],
    )
    add_heading(doc, "One coast, several responsibilities", 1)
    add_text(
        doc,
        "A turtle researcher may focus on nesting behaviour and hatchling sea-finding. A resident may focus on safe access to a home. A business operator may need clear retrofit requirements and time to make changes. A council officer must consider safety, planning, cost, enforcement and measurable outcomes. These viewpoints can produce different responses even when people share the goal of a safe, healthy coast.",
    )
    add_text(
        doc,
        "Traditional Owners are not simply another interest group. They hold continuing rights, knowledge and responsibilities for caring for Land and Sea Country. Marine turtles have deep cultural importance for many coastal Aboriginal and Torres Strait Islander communities. A real council should identify and consult the culturally appropriate Traditional Owners for that place, and support ongoing co-management, monitoring and ranger work.",
    )
    add_heading(doc, "What the light pathway can — and cannot — show", 1)
    add_text(
        doc,
        "A light pathway can organise a possible chain: a targeted council rule leads to lighting audits and redesign; less avoidable light reaches the beach; natural horizon cues may become clearer; and the risk of disturbance or disorientation may fall. The pathway is useful because it makes each link testable. It cannot calculate the exact number of turtles protected, represent every species or site, or replace local evidence.",
    )
    add_callout(
        doc,
        "Council writing challenge:",
        "Recommend an action, use three accurate evidence points, represent a concern fairly and state what your consequence pathway cannot prove.",
        fill=TOKENS["pale"],
        accent=TOKENS["teal"],
    )
    add_heading(doc, "Source trail", 2)
    sources = [
        "Australian Government — National Light Pollution Guidelines for Wildlife (2023).",
        "Queensland Government — Cut the Glow to help Turtles Go; marine turtle biology and conservation; sea turtle nesting area mapping.",
        "Queensland planning guidance — Sea Turtle Sensitive Area Model Code.",
        "Australian Government — Recovery Plan for Marine Turtles in Australia 2017–2027.",
        "Great Barrier Reef Marine Park Authority — Traditional Use of Marine Resources Agreements.",
    ]
    for source in sources:
        add_text(doc, source, style="List Number", size=8.5, after=1, line=1.0)
    add_text(
        doc,
        "This article is an original classroom synthesis. Full links and teacher accuracy notes are in the Teacher Guide. Facts checked 29 July 2026.",
        size=8.2,
        italic=True,
        color=TOKENS["muted"],
        after=0,
        line=1.0,
    )
    out = ROOT / "Lesson_17_Turtle_Safe_Nights_Reading_Pack.docx"
    doc.save(out)
    return out


def build_organiser():
    doc = configure_doc("Consequence + Viewpoint Organiser")
    add_title_block(
        doc,
        "Evidence-to-action studio",
        "Turtle-Safe Nights",
        "Trace the pathway. Test the viewpoint. Recommend a workable action.",
        "ORGANISER • NAME: ____________________",
    )
    add_callout(
        doc,
        "Decision question:",
        "Should a coastal council strengthen beachfront-lighting rules during locally verified marine turtle nesting and hatching periods?",
    )
    add_heading(doc, "1. Which horizon wins?", 1)
    add_form_table(
        doc,
        ["My prediction", "Evidence I would need before making a strong claim"],
        [4100, 6078],
        blank_rows=1,
        row_prompts=[["I predict…", "I would need to know…"]],
    )
    add_heading(doc, "2. Evidence ledger", 1)
    add_form_table(
        doc,
        ["E — evidence", "A — action", "V — value/concern", "? — boundary/local need"],
        [2600, 2450, 2600, 2528],
        blank_rows=2,
        row_prompts=[
            ["Turtles/light:", "Council could:", "Stakeholder:", "We still need:"],
            ["Turtles/light:", "Council could:", "Stakeholder:", "We still need:"],
        ],
        font_size=8.8,
    )
    add_heading(doc, "3. Opportunity → action → consequence", 1)
    add_form_table(
        doc,
        ["Opportunity", "Council action", "Immediate consequence", "Possible later consequence"],
        [2200, 2750, 2600, 2628],
        blank_rows=1,
        row_prompts=[["Before the season…", "The council…", "Therefore…", "Which means… may…"]],
        font_size=8.8,
    )
    add_text(
        doc,
        "Causal explanation: Because the council __________________________________, __________________________________ may __________________________________.",
        size=9.6,
        after=6,
    )
    add_heading(doc, "4. Viewpoint test", 1)
    add_form_table(
        doc,
        ["Stakeholder lens", "What this person/group may reasonably prioritise", "Evidence or design response needed"],
        [2200, 4200, 3778],
        blank_rows=2,
        row_prompts=[
            ["Lens 1:", "May prioritise… because…", "The council should check/respond by…"],
            ["Lens 2:", "May prioritise… because…", "The council should check/respond by…"],
        ],
        font_size=8.8,
    )
    add_text(
        doc,
        "A shared value or goal: __________________________________________________________________________________",
        size=9.5,
        after=0,
    )

    doc.add_page_break()
    add_heading(doc, "5. Test the light pathway", 1)
    add_form_table(
        doc,
        ["1. Rule or decision", "2. Lighting change", "3. Beach/horizon condition", "4. Possible turtle response"],
        [2500, 2550, 2550, 2578],
        blank_rows=1,
        row_prompts=[["Council…", "People/assets…", "Less/more…", "Turtles may…"]],
        font_size=8.8,
    )
    add_labeled_para(
        doc,
        "Evidence-supported link:",
        "Step ____ → Step ____ because ________________________________________________________________",
        fill=TOKENS["pale"],
        accent=TOKENS["teal"],
    )
    add_labeled_para(
        doc,
        "Model boundary:",
        "This pathway can help explain __________________________________, but it cannot prove ________________________________.",
        fill=TOKENS["warm"],
        accent=TOKENS["coral"],
    )
    add_heading(doc, "6. Plan the council recommendation", 1)
    add_matrix(
        doc,
        ["Move", "My note"],
        [
            ["Position", "The council should / should not / should partly…"],
            ["Evidence 1", "Official guidance explains…"],
            ["Evidence 2 + consequence", "Because…, this may…"],
            ["Fair concern", "A resident/business/council officer may reasonably…"],
            ["Response or qualification", "However / Although…"],
            ["Concrete action", "Require / consult / audit / monitor…"],
            ["Light-pathway boundary", "The pathway suggests…, but cannot prove…"],
        ],
        [2600, 7578],
        font_size=9.0,
    )
    add_heading(doc, "7. Draft criteria", 1)
    add_matrix(
        doc,
        ["Check", "Ready?"],
        [
            ["8–10 sentences and a clear council audience", "Yes / Not yet"],
            ["At least three accurate evidence points", "Yes / Not yet"],
            ["A fair stakeholder concern and response", "Yes / Not yet"],
            ["A because/therefore consequence chain", "Yes / Not yet"],
            ["One concrete action and one model limitation", "Yes / Not yet"],
        ],
        [8000, 2178],
        font_size=8.8,
    )
    add_heading(doc, "8. Feedback → immediate revision", 1)
    add_form_table(
        doc,
        ["Partner feedback", "The sentence I will revise"],
        [5100, 5078],
        blank_rows=1,
        row_prompts=[["Your clearest link is… / More accurate or workable if…", "Before → After"]],
        font_size=8.8,
    )
    add_callout(
        doc,
        "Exit:",
        "If the council __________________, then __________________ may happen because __________________. This is a defensible claim, not a guarantee, because __________________.",
        fill=TOKENS["warm"],
        accent=TOKENS["moon"],
    )
    out = ROOT / "Lesson_17_Turtle_Safe_Nights_Organiser.docx"
    doc.save(out)
    return out


def set_lucas_defaults(doc):
    normal = doc.styles["Normal"]
    normal.font.size = Pt(13.5)
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.2
    for level in [1, 2, 3]:
        doc.styles[f"Heading {level}"].font.size = Pt({1: 18, 2: 15.5, 3: 14.5}[level])


def build_lucas():
    doc = configure_doc("Large-Print Support Pack")
    set_lucas_defaults(doc)
    add_title_block(
        doc,
        "Large-print reading + writing support",
        "Turtle-Safe Nights",
        "Point, speak, copy, type or ask an adult to scribe.",
        "LUCAS / ICP • NAME: ____________________",
    )
    add_callout(
        doc,
        "Big question:",
        "Should the council make stronger light rules near turtle beaches during turtle season?",
    )
    add_heading(doc, "Read: Turtles need the night map", 1)
    add_text(
        doc,
        "Many sea turtles nest and hatch at night. A baby turtle is called a hatchling.",
        size=13.5,
        after=7,
        line=1.2,
    )
    add_text(
        doc,
        "On a naturally dark beach, the horizon over the ocean is often the brightest, lowest part of the view. Hatchlings use this light pattern to move toward the sea.",
        size=13.5,
        after=7,
        line=1.2,
    )
    add_text(
        doc,
        "Bright lights from homes, roads or businesses can change the night map. Some hatchlings may move the wrong way. They can become tired, hot or easier for predators to catch.",
        size=13.5,
        after=7,
        line=1.2,
    )
    add_labeled_para(
        doc,
        "FACT 1:",
        "Artificial light can confuse hatchlings.",
        fill=TOKENS["pale"],
        accent=TOKENS["teal"],
        after=8,
    )
    add_labeled_para(
        doc,
        "FACT 2:",
        "Not every light has the same effect.",
        fill=TOKENS["warm"],
        accent=TOKENS["moon"],
        after=8,
    )
    add_labeled_para(
        doc,
        "FACT 3:",
        "Safety and turtle protection can be planned together.",
        fill=TOKENS["pale"],
        accent=TOKENS["teal"],
        after=8,
    )
    add_heading(doc, "What can the council do?", 1)
    add_text(
        doc,
        "The council can ask people to turn off lights they do not need. Needed lights can be kept low, pointed away from the beach and covered so they light only the path or doorway.",
        size=13.5,
        after=7,
        line=1.2,
    )
    add_text(
        doc,
        "The council should also listen to the Traditional Owners who care for that Sea Country, turtle experts, residents and businesses. It should check whether the rule is working.",
        size=13.5,
        after=7,
        line=1.2,
    )
    add_callout(
        doc,
        "Remember:",
        "A good rule may reduce risk. It cannot promise that every turtle will be safe.",
        fill=TOKENS["warm"],
        accent=TOKENS["coral"],
    )

    doc.add_page_break()
    add_heading(doc, "1. Choose a council action", 1)
    add_matrix(
        doc,
        ["Choice", "Action"],
        [
            ["A", "Keep every light the same."],
            ["B", "Use turtle-safe design near nesting beaches during the local season."],
            ["C", "Turn off every light in the whole town."],
        ],
        [1500, 8678],
        font_size=12.2,
    )
    add_text(doc, "My choice: ________    I chose it because _______________________________________________", size=12.5)
    add_heading(doc, "2. Build the pathway", 1)
    add_form_table(
        doc,
        ["Council action", "Light change", "Possible turtle consequence"],
        [3150, 3400, 3628],
        blank_rows=1,
        row_prompts=[["The council…", "Less light may…", "Hatchlings may…"]],
        font_size=11.7,
        blank_after=16,
    )
    add_text(
        doc,
        "Because __________________________________, hatchlings may __________________________________.",
        size=12.5,
        after=8,
    )
    add_heading(doc, "3. Be fair to another viewpoint", 1)
    add_form_table(
        doc,
        ["Choose one", "A fair concern", "A helpful response"],
        [2200, 3850, 4128],
        blank_rows=1,
        row_prompts=[["Resident / business / ranger", "May worry about…", "The council can…"]],
        font_size=11.3,
        blank_after=16,
    )
    add_heading(doc, "4. Build your council message", 1)
    prompts = [
        "1. The council should ____________________________________________________________.",
        "2. Official information says ______________________________________________________.",
        "3. Because ______________________________, turtles may _____________________________.",
        "4. A resident or business may worry about ________________________________________.",
        "5. The council can help by ________________________________________________________.",
        "6. This pathway cannot prove ______________________________________________________.",
    ]
    for prompt in prompts:
        add_text(doc, prompt, size=11.6, after=2, line=1.05)
    add_labeled_para(
        doc,
        "Response choices:",
        "You may point, speak, copy, type or ask an adult to scribe. Add two more sentences if you are ready.",
        fill=TOKENS["pale"],
        accent=TOKENS["teal"],
        after=6,
    )
    add_heading(doc, "5. Check and improve", 1)
    add_matrix(
        doc,
        ["I included…", "Check"],
        [
            ["a council action", "Yes / Help"],
            ["one fact and one possible consequence", "Yes / Help"],
            ["a fair concern", "Yes / Help"],
            ["the word may or can instead of a promise", "Yes / Help"],
        ],
        [7900, 2278],
        font_size=10.8,
    )
    add_callout(
        doc,
        "My improved sentence:",
        "________________________________________________________________________________________",
        fill=TOKENS["warm"],
        accent=TOKENS["moon"],
    )
    out = ROOT / "Lesson_17_Turtle_Safe_Nights_Lucas_Pack.docx"
    doc.save(out)
    return out


def build_html():
    html = r'''<!doctype html>
<html lang="en-AU">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width,initial-scale=1">
<title>Lesson 17 Alternative — Turtle-Safe Nights</title>
<style>
:root{--night:#041522;--deep:#082d42;--sea:#0a5b67;--teal:#33c7b7;--moon:#f4ca64;--coral:#ff7b68;--foam:#edfdf8;--mist:#b8d7dc;--ink:#f7fbfc;--shadow:rgba(0,0,0,.34)}
*{box-sizing:border-box}html,body{width:100%;height:100%;margin:0;overflow:hidden;background:var(--night);color:var(--ink);font-family:"Segoe UI",Arial,sans-serif}
body{background:radial-gradient(circle at 78% 18%,rgba(244,202,100,.18),transparent 18%),linear-gradient(160deg,#03111c 0%,#082a3e 58%,#063841 100%)}
button{font:inherit}.deck{width:100%;height:100%;position:relative}.slide{position:absolute;inset:0;display:none;padding:5.8vh 6vw 12vh;overflow:hidden}.slide.active{display:grid;grid-template-rows:auto 1fr;gap:2.4vh;animation:rise .32s ease}
@keyframes rise{from{opacity:.2;transform:translateY(8px)}to{opacity:1;transform:none}}
.slide:before{content:"";position:absolute;left:0;right:0;bottom:0;height:17vh;background:linear-gradient(180deg,transparent,rgba(1,12,18,.72)),repeating-radial-gradient(ellipse at 50% 120%,transparent 0 22px,rgba(51,199,183,.09) 24px 27px);pointer-events:none}
.top{display:flex;align-items:flex-start;justify-content:space-between;gap:2vw;position:relative;z-index:2}.kicker{font-weight:800;letter-spacing:.18em;text-transform:uppercase;color:var(--teal);font-size:clamp(12px,1.15vw,20px);margin-bottom:.65vh}.depth .kicker{color:var(--moon)}
h1,h2,h3,p{margin:0}h1{font-size:clamp(44px,6.3vw,104px);line-height:.91;letter-spacing:-.045em;max-width:12ch}h2{font-size:clamp(30px,4.1vw,68px);line-height:1;letter-spacing:-.035em;max-width:18ch}h3{font-size:clamp(20px,2vw,34px);line-height:1.08}.sub{font-size:clamp(18px,2vw,34px);color:var(--mist);line-height:1.25;max-width:44ch}
.content{position:relative;z-index:2;min-height:0;display:grid;align-items:center;gap:2vw}.two{grid-template-columns:1.05fr .95fr}.three{grid-template-columns:repeat(3,1fr)}.four{grid-template-columns:repeat(4,1fr)}
.card{background:linear-gradient(145deg,rgba(255,255,255,.105),rgba(255,255,255,.045));border:1px solid rgba(184,215,220,.25);border-radius:20px;padding:clamp(15px,1.5vw,26px);box-shadow:0 16px 36px var(--shadow);backdrop-filter:blur(8px)}
.card p,.body{font-size:clamp(17px,1.55vw,28px);line-height:1.32}.small{font-size:clamp(14px,1.15vw,20px);line-height:1.3;color:var(--mist)}.bigline{font-size:clamp(26px,3vw,52px);line-height:1.12;font-weight:750;max-width:24ch}.accent{color:var(--moon)}.teal{color:var(--teal)}.coral{color:var(--coral)}
.tag{display:inline-flex;align-items:center;gap:.45em;padding:.4em .7em;border-radius:999px;background:rgba(51,199,183,.13);border:1px solid rgba(51,199,183,.42);color:var(--foam);font-weight:750;font-size:clamp(13px,1vw,18px)}
.prompt{border-left:6px solid var(--moon);padding:1.3rem 1.5rem;background:rgba(244,202,100,.1);border-radius:0 18px 18px 0}.prompt p{font-size:clamp(20px,2vw,36px);line-height:1.24}
.hero{grid-template-columns:1.05fr .95fr}.hero-copy{display:flex;flex-direction:column;justify-content:center;gap:2.2vh}.night-scene{height:min(66vh,650px);position:relative;border-radius:34px;overflow:hidden;background:linear-gradient(#061522 0 47%,#093c50 48% 64%,#d5b878 65% 83%,#082a37 84%);box-shadow:0 25px 60px rgba(0,0,0,.5);border:1px solid rgba(255,255,255,.14)}
.moon{position:absolute;width:15%;aspect-ratio:1;border-radius:50%;right:12%;top:9%;background:var(--moon);box-shadow:0 0 40px rgba(244,202,100,.6)}.town{position:absolute;left:0;right:0;top:35%;height:20%;background:linear-gradient(90deg,transparent 0 6%,#102b36 6% 19%,transparent 19% 23%,#143542 23% 42%,transparent 42% 48%,#0c2935 48% 72%,transparent 72%)}
.town:after{content:"";position:absolute;inset:18% 4%;background:repeating-linear-gradient(90deg,transparent 0 7%,rgba(255,238,163,.9) 7.5% 8.5%,transparent 9% 14%);filter:drop-shadow(0 0 8px #ffe599)}
.beam{position:absolute;left:18%;top:42%;width:55%;height:42%;background:linear-gradient(100deg,rgba(255,235,170,.36),transparent 68%);clip-path:polygon(0 0,100% 88%,0 44%);mix-blend-mode:screen}.turtle{position:absolute;left:48%;bottom:13%;width:15%;height:9%;background:#183d39;border-radius:50% 50% 46% 46%;transform:rotate(-8deg)}.turtle:before,.turtle:after{content:"";position:absolute;width:48%;height:38%;background:#183d39;border-radius:80%;top:34%}.turtle:before{left:-34%;transform:rotate(-20deg)}.turtle:after{right:-34%;transform:rotate(20deg)}
.horizon-label{position:absolute;bottom:6%;right:5%;font-weight:800;color:var(--foam);font-size:clamp(13px,1.1vw,19px);letter-spacing:.08em;text-transform:uppercase}
.choice-grid{display:grid;grid-template-columns:repeat(3,1fr);gap:1.2vw}.choice,.lens-btn,.answer,.reveal-btn,.primary{border:1px solid rgba(184,215,220,.35);background:rgba(255,255,255,.07);color:var(--ink);border-radius:16px;padding:1rem;text-align:left;cursor:pointer;transition:.18s;min-height:58px}.choice:hover,.choice:focus,.lens-btn:hover,.answer:hover,.reveal-btn:hover,.primary:hover{transform:translateY(-2px);border-color:var(--teal);background:rgba(51,199,183,.13);outline:none}.choice strong{display:block;color:var(--moon);font-size:clamp(17px,1.4vw,24px);margin-bottom:.4em}.feedback{min-height:3.8em;padding:.9rem 1rem;border-radius:14px;background:rgba(3,15,22,.5);font-size:clamp(15px,1.25vw,22px);line-height:1.25}.feedback.good{border-left:5px solid var(--teal)}.feedback.retry{border-left:5px solid var(--coral)}
.chain{display:flex;align-items:stretch;gap:.7vw}.node{flex:1;padding:1.1rem;border-radius:18px;background:rgba(255,255,255,.08);border:1px solid rgba(184,215,220,.3);display:flex;flex-direction:column;justify-content:center;min-height:18vh}.node b{color:var(--moon);font-size:clamp(14px,1vw,18px);letter-spacing:.08em;text-transform:uppercase}.node span{font-size:clamp(17px,1.4vw,25px);line-height:1.22;margin-top:.45em}.arrow{align-self:center;color:var(--teal);font-size:clamp(28px,3.2vw,54px)}
.lens-row{display:flex;flex-wrap:wrap;gap:.7rem}.lens-btn{padding:.75rem 1rem;text-align:center;min-height:0}.lens-btn.active{background:var(--teal);color:#041522;border-color:var(--teal);font-weight:800}.lens-panel{margin-top:1rem;min-height:22vh}.lens-panel h3{color:var(--moon);margin-bottom:.6em}.lens-panel p{font-size:clamp(18px,1.5vw,27px);line-height:1.35}
.path-stage{opacity:.15;transform:translateY(8px);transition:.25s}.path-stage.shown{opacity:1;transform:none}.boundary-grid{display:grid;grid-template-columns:repeat(3,1fr);gap:1.1vw}.statement{background:rgba(255,255,255,.08);border-radius:18px;padding:1.1rem;border:1px solid rgba(184,215,220,.25)}.statement p{font-size:clamp(15px,1.25vw,22px);line-height:1.3;min-height:8.5em}.answer-row{display:flex;gap:.4rem;margin-top:.8rem}.answer{padding:.55rem .7rem;min-height:0;text-align:center;flex:1;font-size:clamp(12px,.9vw,16px)}.result{font-size:clamp(12px,.95vw,17px);line-height:1.2;margin-top:.6rem;min-height:2.4em}.result.ok{color:#8ef1d7}.result.no{color:#ffad9f}
.model{counter-reset:line}.model-line{display:grid;grid-template-columns:2.4rem 1fr;gap:.7rem;padding:.6rem .8rem;border-left:4px solid transparent;opacity:.25;transition:.2s;font-size:clamp(14px,1.02vw,18px);line-height:1.26}.model-line:before{counter-increment:line;content:counter(line);color:var(--moon);font-weight:800}.model-line.shown{opacity:1;border-left-color:var(--teal);background:rgba(51,199,183,.07)}.model-label{color:var(--moon);font-weight:800}
.criteria{display:grid;grid-template-columns:repeat(2,1fr);gap:.75rem}.criterion{display:flex;gap:.7rem;align-items:flex-start;font-size:clamp(16px,1.35vw,24px);line-height:1.25}.criterion:before{content:"";width:18px;height:18px;border:2px solid var(--teal);border-radius:5px;flex:0 0 auto;margin-top:.15em}
.timer{font-size:clamp(68px,10vw,168px);font-weight:800;letter-spacing:-.06em;color:var(--moon);text-align:center}.timer-controls{display:flex;justify-content:center;gap:.8rem}.primary{background:var(--teal);color:#041522;font-weight:800;text-align:center;min-width:9rem}.secondary{background:transparent;color:var(--ink)}
.controls{position:fixed;left:0;right:0;bottom:0;height:8.8vh;min-height:62px;background:rgba(2,13,20,.92);border-top:1px solid rgba(184,215,220,.18);display:grid;grid-template-columns:1fr auto 1fr;align-items:center;padding:0 2.2vw;z-index:20;backdrop-filter:blur(14px)}.nav-group,.tool-group{display:flex;gap:.55rem;align-items:center}.tool-group{justify-content:flex-end}.control{border:1px solid rgba(184,215,220,.25);background:rgba(255,255,255,.06);color:var(--ink);border-radius:12px;padding:.62rem .85rem;cursor:pointer;min-width:48px}.control:hover,.control:focus{border-color:var(--teal);outline:none}.control:disabled{opacity:.28;cursor:not-allowed}.slide-count{font-weight:750;color:var(--mist);font-size:clamp(13px,1vw,17px);min-width:6rem;text-align:center}.progress{position:fixed;left:0;bottom:8.8vh;height:5px;background:var(--teal);z-index:21;transition:width .22s}
.notes{position:fixed;right:2vw;bottom:10.5vh;width:min(500px,42vw);max-height:56vh;overflow:auto;background:#fff;color:#11252d;border-radius:18px;padding:1.1rem 1.2rem;box-shadow:0 20px 55px rgba(0,0,0,.5);z-index:30;display:none;font-size:16px;line-height:1.4}.notes.open{display:block}.notes:before{content:"TEACHER NOTE";display:block;color:#0a736d;font-size:12px;letter-spacing:.16em;font-weight:800;margin-bottom:.5rem}
.skip-depth{position:absolute;right:6vw;top:6vh;color:var(--moon);font-size:clamp(13px,1vw,17px);font-weight:800}.sr{position:absolute;width:1px;height:1px;padding:0;margin:-1px;overflow:hidden;clip:rect(0,0,0,0);white-space:nowrap;border:0}
@media(max-width:900px){.slide{padding:4vh 4vw 12vh}.two,.hero,.three,.four,.boundary-grid{grid-template-columns:1fr}.night-scene{height:36vh}.content{overflow:auto;align-content:start}.choice-grid{grid-template-columns:1fr}.chain{flex-direction:column}.arrow{transform:rotate(90deg)}.statement p{min-height:0}.notes{width:80vw}.control span{display:none}}
@media(prefers-reduced-motion:reduce){*{animation:none!important;transition:none!important}}
</style>
</head>
<body>
<main class="deck" aria-live="polite">
<section class="slide active" data-notes="Invite a silent first impression. Ask: What is competing for attention in this night scene? Do not yet teach the turtle fact.">
  <div class="content hero">
    <div class="hero-copy">
      <div><div class="kicker">Lesson 17 alternative • Persuasive analysis</div><h1>Turtle-Safe Nights</h1></div>
      <p class="sub">One coastline. Several responsibilities. A council decision with consequences.</p>
      <span class="tag">Audience: a coastal council</span>
    </div>
    <div class="night-scene" role="img" aria-label="Stylised night coastline with a moon, town lights, a light beam and a turtle hatchling on the beach">
      <div class="moon"></div><div class="town"></div><div class="beam"></div><div class="turtle"></div><div class="horizon-label">Which horizon wins?</div>
    </div>
  </div>
</section>

<section class="slide" data-notes="Students predict, then name what evidence would be needed. Separate an observation about the visual from a factual claim about turtle behaviour.">
  <div class="top"><div><div class="kicker">Entry • Predict</div><h2>Which horizon wins?</h2></div><span class="tag">4 minutes</span></div>
  <div class="content two">
    <div class="night-scene" style="height:52vh"><div class="moon"></div><div class="town"></div><div class="beam"></div><div class="turtle"></div></div>
    <div class="card">
      <p class="bigline">Predict a pathway.</p>
      <div class="prompt" style="margin-top:1.2rem"><p>What would you need to know before calling your prediction a fact?</p></div>
      <p class="small" style="margin-top:1rem">Organiser 1 • prediction + evidence need</p>
    </div>
  </div>
</section>

<section class="slide" data-notes="Frame the issue as purpose-specific lighting design, not lights versus no lights. Students should be able to name audience, purpose and decision.">
  <div class="top"><div><div class="kicker">Mission</div><h2>The council has a choice before turtle season</h2></div><span class="tag">8–10 sentences</span></div>
  <div class="content two">
    <div class="prompt"><p>Should a coastal council strengthen beachfront-lighting rules during locally verified nesting and hatching periods?</p></div>
    <div class="card">
      <h3>Your recommendation must…</h3>
      <div class="criteria" style="margin-top:1rem">
        <div class="criterion">use accurate evidence</div><div class="criterion">trace a consequence</div>
        <div class="criterion">represent a concern fairly</div><div class="criterion">name a concrete action</div>
      </div>
    </div>
  </div>
</section>

<section class="slide" data-notes="Use the original classroom article. The text synthesises official sources and contains no invented stakeholder quotations. Pairs compare one annotation only after independent reading.">
  <div class="top"><div><div class="kicker">Read • Evidence adviser</div><h2>Build the knowledge before the argument</h2></div><span class="tag">8 minutes</span></div>
  <div class="content four">
    <div class="card"><h3 class="teal">E</h3><p>Turtle or lighting evidence</p></div>
    <div class="card"><h3 class="accent">A</h3><p>A possible council action</p></div>
    <div class="card"><h3 class="coral">V</h3><p>A stakeholder value or concern</p></div>
    <div class="card"><h3>?</h3><p>A boundary or local evidence need</p></div>
  </div>
</section>

<section class="slide" data-notes="Students choose the most defensible chain, then explain why. Retry feedback identifies the missing distinction without placing the answer.">
  <div class="top"><div><div class="kicker">Opportunity • Action • Consequence</div><h2>Which pathway can the evidence defend?</h2></div><span class="tag">6 minutes</span></div>
  <div class="content">
    <div class="choice-grid" id="chainChoices">
      <button class="choice" data-kind="retry" data-msg="This starts with a desired outcome, not a council opportunity or action. Rebuild the first link."><strong>Path A</strong>Save turtles → make a rule → the beach changes</button>
      <button class="choice" data-kind="good" data-msg="Defensible: the opportunity and action are distinct, and the later effect is qualified as a reduced risk rather than a guarantee."><strong>Path B</strong>Review rules → target lighting design → less avoidable beachward light → risk may fall</button>
      <button class="choice" data-kind="retry" data-msg="A stakeholder concern belongs in the analysis, but it does not by itself prove the consequence. Add the lighting and turtle links."><strong>Path C</strong>Businesses have concerns → keep the rules → turtles are safe</button>
    </div>
    <div class="feedback" id="chainFeedback">Choose, justify, then test. What is the causal rule?</div>
    <button class="reveal-btn" data-reset="chain" style="justify-self:start">Reset pathway</button>
  </div>
</section>

<section class="slide" data-notes="Students copy the four-step chain in their own words and add because/therefore language. Emphasise the final modal 'may'.">
  <div class="top"><div><div class="kicker">Cause • Consequence</div><h2>A defensible chain names every link</h2></div><span class="tag">Organiser 3</span></div>
  <div class="content">
    <div class="chain">
      <div class="node"><b>Opportunity</b><span>Review rules before the local season</span></div><div class="arrow">→</div>
      <div class="node"><b>Action</b><span>Target need, timing, intensity, direction and shielding</span></div><div class="arrow">→</div>
      <div class="node"><b>Condition</b><span>Less avoidable light reaches beach, ocean and sky</span></div><div class="arrow">→</div>
      <div class="node"><b>Possible consequence</b><span>Disturbance or disorientation risk may fall</span></div>
    </div>
    <p class="bigline" style="justify-self:center"><span class="accent">Because</span> … <span class="teal">therefore</span> … <span class="coral">which means</span> …</p>
  </div>
</section>

<section class="slide" data-notes="These are analytical lenses, not quotations. Do not ask a student to speak for Traditional Owners. A real council consults the culturally appropriate rights-holders for that Sea Country.">
  <div class="top"><div><div class="kicker">Viewpoint laboratory</div><h2>One coast, several responsibilities</h2></div><span class="tag">6 minutes</span></div>
  <div class="content">
    <div class="card">
      <div class="lens-row" id="lensButtons">
        <button class="lens-btn active" data-lens="ranger">Turtle ranger</button>
        <button class="lens-btn" data-lens="resident">Resident</button>
        <button class="lens-btn" data-lens="business">Business</button>
        <button class="lens-btn" data-lens="council">Council officer</button>
        <button class="lens-btn" data-lens="country">Sea Country</button>
      </div>
      <div class="lens-panel" id="lensPanel"></div>
    </div>
    <div class="prompt"><p>This stakeholder may prioritise ___ because ___. A fair council response would need ___.</p></div>
  </div>
</section>

<section class="slide" data-notes="Reveal one step at a time after students predict. The light pathway is an interpretive organiser, not a biological calculator. Require a limitation statement.">
  <div class="top"><div><div class="kicker">Interpretive model • Light pathway</div><h2>Useful because each link can be tested</h2></div><span class="tag">5 minutes</span></div>
  <div class="content">
    <div class="chain" id="pathway">
      <div class="node path-stage"><b>Rule</b><span>Target verified places and periods</span></div><div class="arrow path-stage">→</div>
      <div class="node path-stage"><b>Design</b><span>Audit, shield, direct and control light</span></div><div class="arrow path-stage">→</div>
      <div class="node path-stage"><b>Horizon</b><span>Reduce avoidable beachward glow</span></div><div class="arrow path-stage">→</div>
      <div class="node path-stage"><b>Response</b><span>Risk may be reduced</span></div>
    </div>
    <div style="display:flex;gap:.7rem">
      <button class="primary" id="revealPath">Reveal next link</button>
      <button class="reveal-btn" data-reset="path">Reset</button>
    </div>
    <div class="prompt"><p>The pathway suggests ___, but it cannot prove ___.</p></div>
  </div>
</section>

<section class="slide" data-notes="Students classify each statement and justify before clicking. The interaction teaches the boundary: evidence, a cautious inference, and a guarantee that exceeds the evidence.">
  <div class="top"><div><div class="kicker">Evidence boundary</div><h2>Evidence, inference or overclaim?</h2></div><span class="tag">5 minutes</span></div>
  <div class="content">
    <div class="boundary-grid" id="boundaryQuiz">
      <div class="statement" data-correct="e"><p>National guidance says artificial light can disrupt turtle nesting, sea-finding and dispersal.</p><div class="answer-row"><button class="answer" data-answer="e">Evidence</button><button class="answer" data-answer="i">Inference</button><button class="answer" data-answer="o">Overclaim</button></div><div class="result"></div></div>
      <div class="statement" data-correct="i"><p>A targeted seasonal rule using shielding and adaptive controls would likely reduce avoidable light reaching a nesting beach.</p><div class="answer-row"><button class="answer" data-answer="e">Evidence</button><button class="answer" data-answer="i">Inference</button><button class="answer" data-answer="o">Overclaim</button></div><div class="result"></div></div>
      <div class="statement" data-correct="o"><p>If every bulb becomes amber, no hatchling will become disoriented.</p><div class="answer-row"><button class="answer" data-answer="e">Evidence</button><button class="answer" data-answer="i">Inference</button><button class="answer" data-answer="o">Overclaim</button></div><div class="result"></div></div>
    </div>
    <button class="reveal-btn" data-reset="quiz" style="justify-self:start">Reset classifications</button>
  </div>
</section>

<section class="slide" data-notes="Read the whole model first. Then reveal sentence jobs. Ask students to star one move to imitate and circle a cautious modal verb.">
  <div class="top"><div><div class="kicker">Annotated model</div><h2>Persuade without pretending certainty</h2></div><span class="tag">5 minutes</span></div>
  <div class="content two">
    <div class="card model" id="modelText">
      <div class="model-line"><span><b class="model-label">Position:</b> The council should strengthen seasonal rules in mapped turtle-sensitive areas while keeping necessary public spaces safely lit.</span></div>
      <div class="model-line"><span><b class="model-label">Evidence:</b> Official guidance explains that night lighting can disturb nesting turtles and disorient hatchlings.</span></div>
      <div class="model-line"><span><b class="model-label">Chain:</b> Therefore, purpose-specific lights should be shielded, directed, low-intensity and adaptively controlled.</span></div>
      <div class="model-line"><span><b class="model-label">Fair concern:</b> A beachfront business may reasonably worry about safe access and evening trade.</span></div>
      <div class="model-line"><span><b class="model-label">Response:</b> However, national guidance calls for solutions that meet safety and wildlife objectives together.</span></div>
      <div class="model-line"><span><b class="model-label">Boundary + action:</b> The pathway cannot prove an exact number protected; consult, monitor and improve the rule.</span></div>
    </div>
    <div class="card">
      <p class="bigline">What is each sentence <span class="accent">doing</span>?</p>
      <button class="primary" id="revealModel" style="margin-top:1.2rem">Reveal next move</button>
      <button class="reveal-btn" data-reset="model" style="margin-top:.7rem">Reset model</button>
    </div>
  </div>
</section>

<section class="slide depth" data-notes="Optional depth. Students must combine three design principles and explain how the compromise serves a real lighting task and turtle protection.">
  <span class="skip-depth">OPTIONAL • SKIP WITH →</span>
  <div class="top"><div><div class="kicker">Depth A • Design challenge</div><h2>Protect the path and the horizon</h2></div><span class="tag">8 minutes</span></div>
  <div class="content two">
    <div class="card"><h3>Choose one place</h3><p class="bigline" style="margin-top:1rem">Beach path<br>Car park<br>Business entrance</p></div>
    <div class="prompt"><p>Combine three: <span class="accent">need • timing • intensity • height • direction • shielding • spectrum</span></p><p class="small" style="margin-top:1rem">Explain how your rule supports both safety and wildlife.</p></div>
  </div>
</section>

<section class="slide depth" data-notes="Optional depth. Students identify a monitoring signal and a review trigger. This strengthens evidence-based policy reasoning.">
  <span class="skip-depth">OPTIONAL • SKIP WITH →</span>
  <div class="top"><div><div class="kicker">Depth B • Stress test</div><h2>What evidence could change your recommendation?</h2></div><span class="tag">8 minutes</span></div>
  <div class="content two">
    <div class="card"><h3>Monitor</h3><p class="body" style="margin-top:1rem">Light visible from the beach or water<br><br>Nesting attempts and disorientation reports<br><br>Safety performance and community implementation</p></div>
    <div class="prompt"><p>The council should review the rule if monitoring shows ___, because ___.</p></div>
  </div>
</section>

<section class="slide" data-notes="Students complete all seven organiser moves before drafting. Confer first with students who have facts but no causal relationship.">
  <div class="top"><div><div class="kicker">Plan • Organiser 6</div><h2>Lock the pathway before writing</h2></div><span class="tag">5 minutes</span></div>
  <div class="content two">
    <div class="criteria">
      <div class="criterion">position</div><div class="criterion">three evidence points</div>
      <div class="criterion">because/therefore chain</div><div class="criterion">fair concern</div>
      <div class="criterion">response or qualification</div><div class="criterion">concrete action</div>
      <div class="criterion">model limitation</div>
    </div>
    <div class="prompt"><p>Make the recommendation <span class="accent">workable</span>, not merely passionate.</p></div>
  </div>
</section>

<section class="slide" data-notes="Independent writing. The timer is optional. Students may qualify or disagree with the recommended position if their evidence and alternative are defensible.">
  <div class="top"><div><div class="kicker">Write • Council recommendation</div><h2>8–10 sentences. One defensible pathway.</h2></div><span class="tag">13 minutes</span></div>
  <div class="content two">
    <div>
      <div class="timer" id="timer">13:00</div>
      <div class="timer-controls"><button class="primary" id="timerStart">Start</button><button class="primary secondary" id="timerReset">Reset</button></div>
    </div>
    <div class="card">
      <h3>Sentence launches</h3>
      <p class="body" style="margin-top:1rem">The council should…<br>Official guidance explains…<br>Because…, this may…<br>A reasonable concern is…<br>However, the council can…<br>The pathway suggests…, although…</p>
    </div>
  </div>
</section>

<section class="slide" data-notes="Partners give one precise comment. Writers immediately revise one sentence. A visible revision is the finished condition.">
  <div class="top"><div><div class="kicker">Feedback • Revision</div><h2>Trace one ripple. Improve one sentence.</h2></div><span class="tag">4 minutes</span></div>
  <div class="content two">
    <div class="prompt"><p>Your clearest evidence-to-consequence link is…</p></div>
    <div class="prompt" style="border-color:var(--coral);background:rgba(255,123,104,.1)"><p>Your recommendation would be more accurate or workable if…</p></div>
    <p class="bigline" style="grid-column:1/-1;justify-self:center">Revise <span class="accent">now</span>. Make the change visible.</p>
  </div>
</section>

<section class="slide" data-notes="Collect this exit sentence. Group next support by causal link, evidence boundary and overclaim. This is the safe stopping point.">
  <div class="top"><div><div class="kicker">Exit evidence</div><h2>One action. One consequence. One boundary.</h2></div><span class="tag">2 minutes</span></div>
  <div class="content">
    <div class="prompt"><p>If the council ___, then ___ <span class="accent">may</span> happen because ___. This is a defensible claim, not a guarantee, because ___.</p></div>
    <p class="sub" style="justify-self:center;text-align:center">Submit the sentence with your visible revision.</p>
  </div>
</section>
</main>

<div class="progress" id="progress"></div>
<aside class="notes" id="notes" aria-live="polite"></aside>
<nav class="controls" aria-label="Presentation controls">
  <div class="nav-group"><button class="control" id="prev" aria-label="Previous slide">← <span>Previous</span></button><button class="control" id="next" aria-label="Next slide"><span>Next</span> →</button></div>
  <div class="slide-count" id="slideCount">1 / 16</div>
  <div class="tool-group"><button class="control" id="notesBtn" aria-label="Toggle teacher notes">Notes</button><button class="control" id="resetBtn" aria-label="Reset current slide">Reset</button><button class="control" id="fullBtn" aria-label="Toggle fullscreen">Fullscreen</button></div>
</nav>
<script>
const slides=[...document.querySelectorAll('.slide')];let current=0;
const prev=document.getElementById('prev'),next=document.getElementById('next'),count=document.getElementById('slideCount'),progress=document.getElementById('progress'),notes=document.getElementById('notes');
function show(i){current=Math.max(0,Math.min(slides.length-1,i));slides.forEach((s,n)=>s.classList.toggle('active',n===current));prev.disabled=current===0;next.disabled=current===slides.length-1;count.textContent=`${current+1} / ${slides.length}`;progress.style.width=`${((current+1)/slides.length)*100}%`;notes.classList.remove('open');notes.textContent=slides[current].dataset.notes||''}
prev.onclick=()=>show(current-1);next.onclick=()=>show(current+1);
document.addEventListener('keydown',e=>{if(['INPUT','TEXTAREA','SELECT'].includes(document.activeElement.tagName))return;if(['ArrowRight','PageDown',' '].includes(e.key)){e.preventDefault();show(current+1)}if(['ArrowLeft','PageUp'].includes(e.key)){e.preventDefault();show(current-1)}if(e.key==='Home')show(0);if(e.key==='End')show(slides.length-1);if(e.key.toLowerCase()==='n')notes.classList.toggle('open');if(e.key.toLowerCase()==='f')toggleFull()});
document.getElementById('notesBtn').onclick=()=>notes.classList.toggle('open');
function toggleFull(){if(!document.fullscreenElement)document.documentElement.requestFullscreen?.();else document.exitFullscreen?.()}
document.getElementById('fullBtn').onclick=toggleFull;

document.querySelectorAll('#chainChoices .choice').forEach(btn=>btn.onclick=()=>{const f=document.getElementById('chainFeedback');f.textContent=btn.dataset.msg;f.className='feedback '+btn.dataset.kind});
function resetChain(){const f=document.getElementById('chainFeedback');f.textContent='Choose, justify, then test. What is the causal rule?';f.className='feedback'}

const lenses={
ranger:['Turtle ranger or researcher','Will less avoidable light reach the nesting beach and nearshore water? Evidence: light audits, nesting attempts and disorientation monitoring.'],
resident:['Beachfront resident','Can I safely reach and use my home while reducing unnecessary spill? Evidence: the exact task, safety requirement and suitable design.'],
business:['Business or accommodation operator','What must change, by when, and how can safe access and evening trade continue? Evidence: audit findings, retrofit options and an implementation timetable.'],
council:['Council safety and infrastructure officer','Is the rule targeted, enforceable, affordable and measurable? Evidence: mapped places, local season, standards and review triggers.'],
country:['Traditional Owner ranger or Sea Country representative','How will the decision respect the people who speak for this place and support continuing care, knowledge and monitoring? A real council must engage directly and appropriately.']
};
function setLens(key){document.querySelectorAll('.lens-btn').forEach(b=>b.classList.toggle('active',b.dataset.lens===key));const [h,p]=lenses[key];document.getElementById('lensPanel').innerHTML=`<h3>${h}</h3><p>${p}</p>`}
document.querySelectorAll('.lens-btn').forEach(b=>b.onclick=()=>setLens(b.dataset.lens));setLens('ranger');

let pathIndex=0;const pathStages=[...document.querySelectorAll('.path-stage')];
function revealPath(){if(pathIndex<pathStages.length){pathStages[pathIndex++].classList.add('shown');if(pathIndex<pathStages.length&&pathStages[pathIndex].classList.contains('arrow'))pathStages[pathIndex++].classList.add('shown')}}
function resetPath(){pathIndex=0;pathStages.forEach(x=>x.classList.remove('shown'))}
document.getElementById('revealPath').onclick=revealPath;

document.querySelectorAll('.statement').forEach(card=>card.querySelectorAll('.answer').forEach(btn=>btn.onclick=()=>{const ok=btn.dataset.answer===card.dataset.correct;const r=card.querySelector('.result');r.className='result '+(ok?'ok':'no');r.textContent=ok?'Defensible. Explain the boundary.':'Retry: is this directly stated, cautiously inferred, or guaranteed beyond the evidence?'}));
function resetQuiz(){document.querySelectorAll('.statement .result').forEach(r=>{r.textContent='';r.className='result'})}

let modelIndex=0;const modelLines=[...document.querySelectorAll('.model-line')];
function revealModel(){if(modelIndex<modelLines.length)modelLines[modelIndex++].classList.add('shown')}
function resetModel(){modelIndex=0;modelLines.forEach(x=>x.classList.remove('shown'))}
document.getElementById('revealModel').onclick=revealModel;

let seconds=780,timerId=null;const timer=document.getElementById('timer');
function drawTimer(){timer.textContent=`${String(Math.floor(seconds/60)).padStart(2,'0')}:${String(seconds%60).padStart(2,'0')}`}
document.getElementById('timerStart').onclick=()=>{if(timerId)return;timerId=setInterval(()=>{if(seconds>0){seconds--;drawTimer()}else{clearInterval(timerId);timerId=null}},1000)};
document.getElementById('timerReset').onclick=()=>{clearInterval(timerId);timerId=null;seconds=780;drawTimer()};

function resetCurrent(){
 const s=slides[current];
 if(s.querySelector('#chainChoices'))resetChain();
 if(s.querySelector('#pathway'))resetPath();
 if(s.querySelector('#boundaryQuiz'))resetQuiz();
 if(s.querySelector('#modelText'))resetModel();
 if(s.querySelector('#lensButtons'))setLens('ranger');
 if(s.querySelector('#timer')){clearInterval(timerId);timerId=null;seconds=780;drawTimer()}
}
document.querySelectorAll('[data-reset]').forEach(b=>b.onclick=()=>{const k=b.dataset.reset;if(k==='chain')resetChain();if(k==='path')resetPath();if(k==='quiz')resetQuiz();if(k==='model')resetModel()});
document.getElementById('resetBtn').onclick=resetCurrent;
show(0);drawTimer();
</script>
</body>
</html>'''
    out = ROOT / "Lesson_17_Turtle_Safe_Nights_Presentation.html"
    out.write_text(html, encoding="utf-8")
    return out


if __name__ == "__main__":
    outputs = [build_html(), build_reading_pack(), build_organiser(), build_lucas()]
    print("Built Turtle-Safe Nights package:")
    for output in outputs:
        print(f" - {output.name}")
