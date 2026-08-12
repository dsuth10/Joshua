from pathlib import Path
import sys
ROOT=Path(__file__).resolve().parents[1]; AS=ROOT/'assets'
def doc(path,title,lines):
 from docx import Document
 from docx.shared import Inches, Pt
 d=Document(); sec=d.sections[0]; sec.top_margin=sec.bottom_margin=Inches(.65); sec.left_margin=sec.right_margin=Inches(.7)
 s=d.styles['Normal']; s.font.name='Arial'; s.font.size=Pt(11)
 d.add_heading(title,0); d.add_paragraph('Lesson 27 - The Evidence Archive')
 for head,text in lines:
  d.add_heading(head,1); d.add_paragraph(text)
 d.save(path)
try:
 doc(ROOT/'Lesson_27_Evidence_Archive_Reading_Pack.docx','The Evidence Archive',[('Original fictional reading','At Harbour Street School, the student council kept a folder called the Evidence Archive while proposing a shaded courtyard. Some notes named the writer, date and purpose. Others simply said “a parent told us” or “the internet says”. Zara argued that a proposal should not borrow the second kind of note as proof. “If the principal asks where it came from,” she said, “we need to be able to show the path.”'),('Notice','A label does not make a claim true. It makes checking possible.'),('Questions','1. What makes one note traceable?  2. Why is an anonymous claim not enough?  3. Which part of a research trail could you check first?')])
 doc(ROOT/'Lesson_27_Evidence_Archive_Organiser.docx','My Evidence Trail',[('Research question','_______________________________________________________________'),('Source 1','Who made it? __________________  Title/link/date: __________________\nAccurate note: __________________________________________________\nNAME check: _____________________________________________________'),('Source 2','Who made it? __________________  Title/link/date: __________________\nAccurate note: __________________________________________________\nNAME check: _____________________________________________________')])
 doc(ROOT/'Lesson_27_Evidence_Archive_Concise_Access_Pack.docx','My Evidence Trail',[('Choose a question','What do I need to find out? ______________________________________'),('One source first','Who made it? __________________\nOne idea it says: _______________________________________________\nWhere I found it: _______________________________________________'),('You may','Point, speak, copy or use a partner scribe. If unsure, write: needs checking.')])
except ModuleNotFoundError: print('DOCX dependency unavailable; presentation still compiled')
sys.path.insert(0,str(ROOT.parents[5]/'.agents'/'skills'/'lesson-creator'/'scripts')); from compile_presentation import compile_presentation
compile_presentation(AS/'evidence_archive_slides.html',ROOT/'Lesson_27_Evidence_Archive_Presentation.html',css_path=AS/'evidence_archive.css',js_path=AS/'evidence_archive.js',title='Lesson 27: The Evidence Archive')
html='<!doctype html><meta charset="utf-8"><title>The Evidence Archive</title><style>body{font:18px Georgia;margin:8%;line-height:1.55;color:#17374a}h1{font-size:42px;color:#185c67}.tag{color:#d26b3d;font:700 14px Arial;letter-spacing:2px}</style><p class="tag">ORIGINAL FICTIONAL READING</p><h1>The Evidence Archive</h1><p>At Harbour Street School, the student council kept a folder called the Evidence Archive while proposing a shaded courtyard. Some notes named the writer, date and purpose. Other notes simply said “a parent told us” or “the internet says”. Zara argued that a proposal should not borrow the second kind of note as proof.</p><p>“If the principal asks where it came from,” she said, “we need to be able to show the path.” The archive helped the group separate a promising question from a usable answer. A label did not make a claim true; it made checking possible.</p><p><b>Reading note:</b> This story is fictional. It models research decisions and supplies no evidence for an assessment.</p>'; (ROOT/'Lesson_27_Evidence_Archive_Reading_Article.html').write_text(html,encoding='utf-8')
