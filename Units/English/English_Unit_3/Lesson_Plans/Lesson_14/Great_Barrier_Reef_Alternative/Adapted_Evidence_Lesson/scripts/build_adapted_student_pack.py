import os
from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
ASSETS = ROOT / "assets"
OUTPUT = Path(
    os.environ.get(
        "ADAPTED_PACK_OUTPUT",
        ROOT / "Lesson_14_GBR_Adapted_Student_Pack.docx",
    )
)
BANNER = ASSETS / "gbr-healthy-banner.png"

# compact_reference_guide with a named A4 classroom override:
# - A4 portrait, 16 mm margins, 7.0 in fixed content/table width
# - Arial 11.5 pt body, 1.08 line spacing, 5 pt after
# - student-facing ocean/coral palette and larger headings
NAVY = "073642"
DEEP = "075663"
OCEAN = "087F8C"
CORAL = "D94F45"
SAND = "F6E7C8"
FOAM = "EDF9F7"
PALE_CORAL = "FCE9E5"
INK = "16333A"
MUTED = "526A70"
WHITE = "FFFFFF"
CONTENT_DXA = 10080
TABLE_INDENT_DXA = 120
CELL_MARGIN_DXA = 120


def set_run_font(run, size=11.5, bold=False, color=INK, italic=False, name="Arial"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_spacing(paragraph, before=0, after=5, line=1.08):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def add_para(
    container,
    text="",
    *,
    size=11.5,
    bold=False,
    color=INK,
    italic=False,
    before=0,
    after=5,
    line=1.08,
    align=WD_ALIGN_PARAGRAPH.LEFT,
    keep_with_next=False,
):
    paragraph = container.add_paragraph()
    paragraph.alignment = align
    set_paragraph_spacing(paragraph, before=before, after=after, line=line)
    paragraph.paragraph_format.keep_with_next = keep_with_next
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color, italic=italic)
    return paragraph


def add_labelled_para(container, label, text, *, after=5, fill=None):
    paragraph = container.add_paragraph()
    set_paragraph_spacing(paragraph, after=after, line=1.08)
    if fill:
        shade_paragraph(paragraph, fill)
        paragraph.paragraph_format.left_indent = Pt(8)
        paragraph.paragraph_format.right_indent = Pt(8)
        paragraph.paragraph_format.space_before = Pt(5)
        paragraph.paragraph_format.space_after = Pt(5)
    label_run = paragraph.add_run(label)
    set_run_font(label_run, size=11.5, bold=True, color=OCEAN)
    text_run = paragraph.add_run(text)
    set_run_font(text_run, size=11.5, color=INK)
    return paragraph


def add_heading(container, text, *, level=1, color=None, before=None, after=None):
    if level == 1:
        size, default_color, default_before, default_after = 18, NAVY, 10, 4
    else:
        size, default_color, default_before, default_after = 14, OCEAN, 7, 3
    paragraph = container.add_paragraph(style=f"Heading {level}")
    set_paragraph_spacing(
        paragraph,
        before=default_before if before is None else before,
        after=default_after if after is None else after,
        line=1.0,
    )
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text)
    if color and color != default_color:
        run.font.color.rgb = RGBColor.from_string(color)
    return paragraph


def shade_paragraph(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shading = p_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        p_pr.append(shading)
    shading.set(qn("w:fill"), fill)
    shading.set(qn("w:val"), "clear")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)
    shading.set(qn("w:val"), "clear")


def set_cell_margins(cell, top=CELL_MARGIN_DXA, start=CELL_MARGIN_DXA, bottom=CELL_MARGIN_DXA, end=CELL_MARGIN_DXA):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, *, color="C6DADB", size=10):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa):
    if sum(widths_dxa) != CONTENT_DXA:
        raise ValueError(f"Table widths must total {CONTENT_DXA} DXA: {widths_dxa}")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_width = tbl_pr.find(qn("w:tblW"))
    if tbl_width is None:
        tbl_width = OxmlElement("w:tblW")
        tbl_pr.append(tbl_width)
    tbl_width.set(qn("w:w"), str(CONTENT_DXA))
    tbl_width.set(qn("w:type"), "dxa")

    tbl_indent = tbl_pr.find(qn("w:tblInd"))
    if tbl_indent is None:
        tbl_indent = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_indent)
    tbl_indent.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_indent.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_width = tc_pr.find(qn("w:tcW"))
            if tc_width is None:
                tc_width = OxmlElement("w:tcW")
                tc_pr.append(tc_width)
            tc_width.set(qn("w:w"), str(width))
            tc_width.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def clear_cell(cell):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    set_paragraph_spacing(paragraph, after=0, line=1.05)
    return paragraph


def cell_para(cell, text, *, size=11, bold=False, color=INK, after=2, align=WD_ALIGN_PARAGRAPH.LEFT):
    paragraph = cell.add_paragraph() if cell.paragraphs[0].text else cell.paragraphs[0]
    paragraph.alignment = align
    set_paragraph_spacing(paragraph, after=after, line=1.05)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return paragraph


def add_fixed_table(container, rows, widths_dxa, fills=None, border_color="C6DADB"):
    table = container.add_table(rows=len(rows), cols=len(widths_dxa))
    set_table_geometry(table, widths_dxa)
    for row_index, row_data in enumerate(rows):
        for col_index, value in enumerate(row_data):
            cell = table.rows[row_index].cells[col_index]
            clear_cell(cell)
            set_cell_border(cell, color=border_color)
            if fills and fills[row_index][col_index]:
                set_cell_shading(cell, fills[row_index][col_index])
            if isinstance(value, list):
                for item_index, item in enumerate(value):
                    cell_para(
                        cell,
                        item.get("text", ""),
                        size=item.get("size", 11),
                        bold=item.get("bold", False),
                        color=item.get("color", INK),
                        after=item.get("after", 2),
                        align=item.get("align", WD_ALIGN_PARAGRAPH.LEFT),
                    )
            else:
                cell_para(cell, value, size=11)
    return table


def add_callout(
    container,
    label,
    text,
    *,
    fill=FOAM,
    accent=OCEAN,
    text_size=11.5,
    cell_margin=CELL_MARGIN_DXA,
):
    table = container.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_DXA])
    cell = table.cell(0, 0)
    clear_cell(cell)
    set_cell_margins(cell, top=cell_margin, start=CELL_MARGIN_DXA, bottom=cell_margin, end=CELL_MARGIN_DXA)
    set_cell_shading(cell, fill)
    set_cell_border(cell, color=accent, size=14)
    cell_para(cell, label, size=10, bold=True, color=accent, after=2)
    cell_para(cell, text, size=text_size, bold=True, color=INK, after=0)
    return table


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr_text, fld_char_end])


def set_picture_alt_text(run, description, title):
    doc_pr = run._r.find(".//" + qn("wp:docPr"))
    if doc_pr is not None:
        doc_pr.set("descr", description)
        doc_pr.set("title", title)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(16)
    section.right_margin = Mm(16)
    section.bottom_margin = Mm(16)
    section.left_margin = Mm(16)
    section.header_distance = Mm(8)
    section.footer_distance = Mm(8)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(11.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08

    for style_name, size, color, before, after in (
        ("Heading 1", 18, NAVY, 10, 4),
        ("Heading 2", 14, OCEAN, 7, 3),
    ):
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
        style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = True

    header = section.header
    header_p = header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(header_p, after=0, line=1.0)
    header_run = header_p.add_run("LESSON 14 | ADAPTED EVIDENCE")
    set_run_font(header_run, size=8, bold=True, color=MUTED)

    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_spacing(footer_p, after=0, line=1.0)
    footer_run = footer_p.add_run("REEF EVIDENCE DETECTIVES  |  ")
    set_run_font(footer_run, size=8, bold=True, color=MUTED)
    add_page_number(footer_p)


def make_banner():
    source = ASSETS / "gbr-healthy-panorama.png"
    with Image.open(source) as image:
        target_ratio = 5.0
        crop_height = int(image.width / target_ratio)
        centre_y = int(image.height * 0.62)
        top = max(0, min(image.height - crop_height, centre_y - crop_height // 2))
        cropped = image.crop((0, top, image.width, top + crop_height))
        cropped.save(BANNER, quality=94)


def add_title_block(doc, kicker, title, subtitle, *, page_break_before=False):
    kicker_p = add_para(doc, kicker.upper(), size=9, bold=True, color=CORAL, after=3, line=1.0)
    kicker_p.paragraph_format.page_break_before = page_break_before
    add_para(doc, title, size=27, bold=True, color=NAVY, after=3, line=0.95, keep_with_next=True)
    add_para(doc, subtitle, size=12.5, bold=True, color=DEEP, after=8, line=1.12)


def build_reading_page(doc):
    add_title_block(
        doc,
        "Adapted reading",
        "Help the Great Barrier Reef",
        "The Reef is alive, but it is under pressure. What facts can help us explain why it needs protection?",
    )
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(paragraph, after=6, line=1.0)
    run = paragraph.add_run()
    run.add_picture(str(BANNER), width=Inches(7.0), height=Inches(1.4))
    set_picture_alt_text(
        run,
        "Illustrative underwater view of living coral and fish on the Great Barrier Reef.",
        "Living coral reef",
    )

    add_callout(
        doc,
        "READING MISSION",
        "Underline two facts. Circle one action. Box the sentence that stops an overclaim.",
        fill=SAND,
        accent=CORAL,
    )

    add_heading(doc, "1  A huge living place", level=2, before=6)
    add_para(
        doc,
        "The Great Barrier Reef is not one wall of coral. It is a huge living system along the Queensland coast. "
        "It has about 3,000 reefs, as well as islands, seagrass, mangroves and deep water. Fish, turtles and many "
        "other animals live there.",
        size=11.2,
        after=4,
        line=1.04,
    )
    add_para(
        doc,
        "The Reef is also Sea Country. Aboriginal and Torres Strait Islander Traditional Owners have cared for "
        "this area for a very long time. Today, Traditional Owners use cultural knowledge and science to help "
        "care for the Reef.",
        size=11.2,
        after=4,
        line=1.04,
    )

    add_heading(doc, "2  What scientists found", level=2, before=4)
    add_para(
        doc,
        "Scientists check the Reef to see how it changes. In 2024-25, they checked 124 reefs. They still found "
        "living coral in the north, centre and south, but coral cover had fallen. Hot water and other pressures "
        "had damaged coral.",
        size=11.2,
        after=4,
        line=1.04,
    )
    add_callout(
        doc,
        "KEEP THE FACT ACCURATE",
        "This is evidence that the Reef is alive and needs help. It is not evidence that the whole Reef is dead.",
        fill=FOAM,
        accent=OCEAN,
    )
    add_para(
        doc,
        "Coral animals live with tiny algae that give them food and colour. Very warm water can make coral push "
        "out the algae and turn white. This is called bleaching.",
        size=11.2,
        after=4,
        line=1.04,
    )
    add_para(
        doc,
        "Bleached coral is stressed, but it is not always dead. Some coral can recover if the water cools. Coral "
        "can die if the heat is very strong, lasts too long or returns before the coral has recovered.",
        size=11.2,
        bold=True,
        after=4,
        line=1.04,
    )

    add_heading(doc, "3  Two ways to help", level=2, before=4)
    add_labelled_para(
        doc,
        "Help near the Reef. ",
        "Cleaner water, less rubbish, careful fishing, crown-of-thorns starfish control and Sea Country management "
        "can reduce pressure. These local actions can give coral a better chance to recover.",
        after=3,
    )
    add_labelled_para(
        doc,
        "Reduce ocean warming. ",
        "Greenhouse gases warm Earth and the ocean. Climate action can reduce this warming. Local projects help, "
        "but they cannot cool the whole ocean by themselves.",
        after=3,
    )
    add_para(
        doc,
        "Source note: Simplified summary of the AIMS 2024/25 coral condition report and bleaching information, the Reef "
        "Snapshot, the 2024 Scientific Consensus Statement and Reef Authority Traditional Owner information. Facts "
        "checked in the parent lesson source trail on 29 July 2026. Image is illustrative, not survey evidence.",
        size=7.4,
        color=MUTED,
        after=0,
        line=1.0,
    )


def build_evidence_page(doc):
    add_title_block(
        doc,
        "Student Pack - page 2",
        "Evidence Detective",
        "A fact becomes useful evidence when it matches your claim and you explain what it shows.",
        page_break_before=True,
    )

    add_heading(doc, "1  Find evidence that matches", level=1, before=3)
    claims = [
        (
            "CLAIM A",
            "The Reef is alive, but it needs help.",
            PALE_CORAL,
            CORAL,
        ),
        (
            "CLAIM B",
            "Hot water is dangerous for coral.",
            FOAM,
            OCEAN,
        ),
        (
            "CLAIM C",
            "The Reef needs more than one kind of help.",
            SAND,
            CORAL,
        ),
    ]
    for label, claim, fill, accent in claims:
        table = doc.add_table(rows=1, cols=2)
        set_table_geometry(table, [2600, 7480])
        for cell in table.rows[0].cells:
            clear_cell(cell)
            set_cell_border(cell, color=accent, size=12)
        set_cell_shading(table.cell(0, 0), fill)
        set_cell_shading(table.cell(0, 1), WHITE)
        cell_para(table.cell(0, 0), label, size=9, bold=True, color=accent, after=2)
        cell_para(table.cell(0, 0), claim, size=12, bold=True, color=NAVY, after=0)
        cell_para(table.cell(0, 1), "Evidence from the article:", size=9, bold=True, color=accent, after=3)
        cell_para(table.cell(0, 1), "____________________________________________________________", size=11, after=4)
        cell_para(table.cell(0, 1), "This evidence shows _________________________________________", size=11, after=3)
        cell_para(table.cell(0, 1), "____________________________________________________________", size=11, after=0)
        add_para(doc, "", size=2, after=4, line=1.0)

    add_heading(doc, "2  Evidence or opinion?", level=1, before=4)
    evidence_rows = [
        [
            [
                {"text": "Statement", "size": 10, "bold": True, "color": WHITE, "after": 0},
            ],
            [
                {"text": "Evidence", "size": 10, "bold": True, "color": WHITE, "after": 0, "align": WD_ALIGN_PARAGRAPH.CENTER},
            ],
            [
                {"text": "Opinion", "size": 10, "bold": True, "color": WHITE, "after": 0, "align": WD_ALIGN_PARAGRAPH.CENTER},
            ],
        ],
        ["Very warm water can bleach coral.", "(   )", "(   )"],
        ["The Reef is the prettiest place in Australia.", "(   )", "(   )"],
        ["Cleaner water can reduce pressure on the Reef.", "(   )", "(   )"],
    ]
    fills = [
        [NAVY, NAVY, NAVY],
        [WHITE, WHITE, WHITE],
        [FOAM, FOAM, FOAM],
        [WHITE, WHITE, WHITE],
    ]
    table = add_fixed_table(doc, evidence_rows, [7300, 1390, 1390], fills=fills, border_color="B9D0D2")
    set_repeat_table_header(table.rows[0])
    for row in table.rows[1:]:
        for cell in row.cells[1:]:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(doc, "3  Stop the overclaim", level=1, before=8)
    add_callout(
        doc,
        "NOT SUPPORTED",
        '"The whole Great Barrier Reef is dead."',
        fill=PALE_CORAL,
        accent=CORAL,
    )
    add_para(
        doc,
        "Why is this claim not supported? ____________________________________________________________",
        size=11.2,
        after=7,
    )
    add_para(
        doc,
        "___________________________________________________________________________________________",
        size=11.2,
        after=0,
    )


def build_writing_page(doc):
    add_title_block(
        doc,
        "Student Pack - page 3",
        "Write With Evidence",
        "Write a five-to-six-sentence paragraph for people who can help protect the Great Barrier Reef.",
        page_break_before=True,
    )

    add_heading(doc, "1  Read the model", level=1, before=3)
    add_callout(
        doc,
        "MODEL PARAGRAPH",
        "Australian decision-makers should protect the Great Barrier Reef. In 2024-25, scientists found living "
        "coral in the north, centre and south, but coral cover had fallen. This shows that the Reef is alive and "
        "needs help. Cleaner water and careful Reef management can reduce pressure on coral. Climate action is "
        "also needed because local projects cannot cool the whole ocean. Please use both kinds of action to give "
        "the Reef a better chance to recover.",
        fill=FOAM,
        accent=OCEAN,
        text_size=10.3,
        cell_margin=80,
    )

    add_heading(doc, "2  Plan the five moves", level=1, before=7)
    plan_rows = [
        [
            [{"text": "Move", "size": 10, "bold": True, "color": WHITE, "after": 0}],
            [{"text": "My plan", "size": 10, "bold": True, "color": WHITE, "after": 0}],
        ],
        ["1  Claim", "We should __________________________________________________________"],
        ["2  Problem fact", "The article says ____________________________________________________"],
        ["3  Explain", "This shows _________________________________________________________"],
        ["4  Action fact", "Another fact is _____________________________________________________"],
        ["5  Request", "Please _____________________________________________________________"],
    ]
    fills = [
        [NAVY, NAVY],
        [PALE_CORAL, WHITE],
        [FOAM, WHITE],
        [SAND, WHITE],
        [FOAM, WHITE],
        [PALE_CORAL, WHITE],
    ]
    table = add_fixed_table(doc, plan_rows, [2240, 7840], fills=fills, border_color="BDD2D3")
    set_repeat_table_header(table.rows[0])
    for row in table.rows[1:]:
        for cell in row.cells:
            set_cell_margins(cell, top=60, start=CELL_MARGIN_DXA, bottom=60, end=CELL_MARGIN_DXA)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10.4)
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string(DEEP)

    add_heading(doc, "3  Write your paragraph", level=1, before=8)
    for _ in range(7):
        add_para(
            doc,
            "___________________________________________________________________________________________",
            size=11,
            after=2,
            line=1.0,
        )

    add_heading(doc, "4  Evidence check and revise", level=1, before=2, after=2)
    checklist = doc.add_table(rows=1, cols=3)
    set_table_geometry(checklist, [3360, 3360, 3360])
    checks = [
        "(   ) I used two facts.",
        "(   ) I explained one fact.",
        "(   ) I revised one sentence.",
    ]
    for index, text in enumerate(checks):
        cell = checklist.cell(0, index)
        clear_cell(cell)
        set_cell_shading(cell, SAND if index == 1 else FOAM)
        set_cell_border(cell, color="B9D0D2")
        cell_para(cell, text, size=10.3, bold=True, color=NAVY, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)

    add_heading(doc, "EXIT  Prove one idea", level=2, color=CORAL, before=5, after=2)
    add_para(
        doc,
        "We should __________________ because the article says _________________________________.",
        size=11,
        bold=True,
        after=5,
    )
    add_para(
        doc,
        "This shows _____________________________________________________________________________.",
        size=11,
        bold=True,
        after=0,
    )


def set_core_properties(doc):
    props = doc.core_properties
    props.title = "Lesson 14 Great Barrier Reef Adapted Evidence Student Pack"
    props.subject = "Adapted reading evidence and persuasive writing"
    props.author = "Lesson 14 teaching team"
    props.keywords = "adapted, evidence, persuasion, Great Barrier Reef"
    props.comments = "Generated from the Lesson 14 source package."


def main():
    ASSETS.mkdir(parents=True, exist_ok=True)
    make_banner()
    doc = Document()
    configure_document(doc)
    set_core_properties(doc)
    build_reading_page(doc)
    build_evidence_page(doc)
    build_writing_page(doc)
    doc.save(OUTPUT)
    print(f"Built {OUTPUT}")


if __name__ == "__main__":
    main()
