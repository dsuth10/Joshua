"""Build the supplemental English Unit 3 Lesson 15 flying-fox package.

Original Lesson 15 files are never read or modified by this generator.

DOCX design basis: compact_reference_guide.
Named classroom override: A4 portrait, 12-15 mm margins, Arial typography,
fixed-width tables, forest/night palette, workbook-ready response space.
"""

from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor


ROOT = Path(__file__).resolve().parent.parent

INK = "19312D"
NIGHT = "132A34"
FOREST = "1E5C4A"
LEAF = "4B8B61"
LIME = "B8D97B"
GOLD = "F0C85A"
CREAM = "FFF8E7"
MIST = "EEF5EE"
PALE_GOLD = "FFF3C9"
PALE_RED = "FBE9E5"
RED = "A33B32"
LINE = "A8B9AE"
MUTED = "53665F"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_table_geometry(table, widths, indent=120):
    total = sum(widths)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[min(idx, len(widths) - 1)])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_font(run, size=10.5, bold=False, colour=INK, italic=False):
    run.font.name = "Arial"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(colour)


def set_run_text(p, text, size=10.5, bold=False, colour=INK, italic=False):
    run = p.add_run(text)
    set_font(run, size, bold, colour, italic)
    return run


def configure_doc(doc, title, subtitle, margin_mm=15, large=False):
    sec = doc.sections[0]
    sec.page_width = Mm(210)
    sec.page_height = Mm(297)
    sec.top_margin = Mm(margin_mm)
    sec.bottom_margin = Mm(margin_mm)
    sec.left_margin = Mm(margin_mm)
    sec.right_margin = Mm(margin_mm)
    sec.header_distance = Mm(7)
    sec.footer_distance = Mm(7)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(14 if large else 10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(5 if large else 4)
    normal.paragraph_format.line_spacing = 1.18 if large else 1.12
    for name, size, colour, before, after in (
        ("Title", 26 if large else 24, NIGHT, 0, 4),
        ("Subtitle", 13 if large else 11.5, MUTED, 0, 10),
        ("Heading 1", 19 if large else 16, FOREST, 12, 6),
        ("Heading 2", 16 if large else 13, FOREST, 9, 4),
        ("Heading 3", 14 if large else 11.5, NIGHT, 6, 3),
    ):
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(colour)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = sec.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.paragraph_format.space_after = Pt(0)
    set_run_text(header, "LIVING WITH FLYING-FOXES  |  LESSON 15 ALTERNATIVE", 8.5, True, MUTED)
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_text(footer, "English Unit 3  •  objective, subjective and mixed language", 8, False, MUTED)

    p = doc.add_paragraph(style="Title")
    set_run_text(p, title, 26 if large else 24, True, NIGHT)
    p.paragraph_format.space_after = Pt(2)
    p = doc.add_paragraph(style="Subtitle")
    set_run_text(p, subtitle, 13 if large else 11.5, False, MUTED)


def add_label(doc, text, fill=FOREST, colour=WHITE):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [10080], 120)
    set_cell_shading(table.cell(0, 0), fill)
    p = table.cell(0, 0).paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    set_run_text(p, text.upper(), 9.5, True, colour)
    return table


def add_callout(doc, title, text, fill=MIST, border=LINE, size=10.5):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [10080], 120)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    set_run_text(p, title + "  ", size, True, FOREST)
    set_run_text(p, text, size, False, INK)
    return table


def add_paragraph(doc, text, size=10.5, bold=False, colour=INK, after=4, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    set_run_text(p, text, size, bold, colour, italic)
    return p


def add_bullets(doc, items, size=10.5):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Cm(0.65)
        p.paragraph_format.first_line_indent = Cm(-0.32)
        p.paragraph_format.space_after = Pt(3)
        set_run_text(p, item, size)


def add_lines(doc, count=3, label=None, size=10.5):
    if label:
        add_paragraph(doc, label, size, True, FOREST, 2)
    for _ in range(count):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        set_run_text(p, "________________________________________________________________________________", size, False, LINE)


def write_cell(cell, text, bold=False, fill=None, size=9.2, colour=INK):
    if fill:
        set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    set_run_text(p, text, size, bold, colour)


def add_page_break(doc):
    doc.add_page_break()


def build_reading_pack():
    doc = Document()
    configure_doc(
        doc,
        "Living With Flying-Foxes",
        "Student evidence briefing and competing community messages",
        15,
    )
    add_callout(
        doc,
        "Your reading mission",
        "Mark C for community impact, E for ecological importance and S for a safe or lawful response. Box one sentence that corrects a misconception.",
        PALE_GOLD,
    )
    doc.add_heading("1  Night travellers and daytime camps", level=1)
    add_paragraph(
        doc,
        "Flying-foxes are large fruit- and nectar-eating bats. They are nocturnal: they rest together in trees during the day, then fly out around dusk to feed. The trees used for resting and breeding are called a camp or roost. Food and camp use change with flowering, fruiting and seasons, so numbers at a camp can rise, fall or move.",
    )
    add_paragraph(
        doc,
        "As flying-foxes travel between feeding areas, they can carry pollen and seeds. This helps native forests reproduce and connect across the landscape. Their long-distance movement makes them important pollinators and seed dispersers. Australia monitors flying-fox populations because grey-headed and spectacled flying-foxes are nationally listed as threatened.",
    )
    doc.add_heading("2  Why camps appear near people", level=1)
    add_paragraph(
        doc,
        "Flying-foxes need flowering trees, feeding habitat and safe places to rest. Forest loss and fragmentation have reduced some natural habitat. Remaining patches of trees in towns, parks and along waterways can become useful camp sites. This can bring wildlife and people into closer contact.",
    )
    add_paragraph(
        doc,
        "A large or seasonal camp can create real challenges. Residents may experience noise, strong odour, droppings on cars or washing, damaged vegetation and interrupted sleep. These impacts are not imaginary simply because flying-foxes are ecologically valuable. A fair community message should describe the exact impact instead of labelling the animals.",
    )
    add_callout(
        doc,
        "Language alert",
        "“A camp is present beside the park” is checkable. “A disgusting invasion has stolen our park” adds loaded judgement and metaphor.",
        MIST,
    )
    doc.add_heading("Quick evidence check", level=2)
    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [3000, 7080], 120)
    rows = [
        ("What could be checked?", "What extra detail would make it precise?"),
        ("The number of bats at a camp", "Place, date, counting method and species"),
        ("Residents reported noise", "Who was asked, how many people and when"),
        ("Flying-foxes move seeds", "Which species, plants, distance and research source"),
    ]
    for r, values in enumerate(rows):
        for c, value in enumerate(values):
            write_cell(table.cell(r, c), value, r == 0, FOREST if r == 0 else None, 9.2, WHITE if r == 0 else INK)
    set_repeat_header(table.rows[0])

    add_page_break(doc)
    add_label(doc, "Health, safety and lawful coexistence")
    doc.add_heading("3  What the health advice actually says", level=1)
    add_paragraph(
        doc,
        "Queensland Government advice says that catching disease directly from flying-foxes is extremely unlikely. Australian Bat Lyssavirus can be transmitted through an untreated bite or scratch from an infected bat. Living, walking or playing near a camp does not expose a person to the virus. There is no evidence that Hendra virus passes directly from flying-foxes to people.",
    )
    add_callout(
        doc,
        "No touch means safe distance",
        "Never touch, feed or try to rescue a flying-fox. Keep people and pets away from an injured or grounded bat and contact a trained, vaccinated wildlife rescuer. If a person is bitten or scratched, seek medical advice immediately.",
        PALE_RED,
    )
    doc.add_heading("4  Practical responses", level=1)
    add_paragraph(
        doc,
        "People living near a seasonal camp can reduce some impacts by planning ahead. Official suggestions include bringing washing in at night, parking vehicles under cover, covering outdoor items and maintaining buildings or water systems appropriately. Avoiding disturbance can also reduce extra noise and stress.",
    )
    add_paragraph(
        doc,
        "Flying-foxes and their roosts are protected under Queensland conservation law. Some low-impact activities and council management actions are allowed only when current rules and codes are followed. Management must not harm the animals. Moving a camp is not a simple “push them away” solution: flying-foxes are mobile, impacts can shift elsewhere, and disturbance can affect animals and nearby residents.",
    )
    doc.add_heading("5  A fair council message needs three things", level=1)
    table = doc.add_table(rows=4, cols=3)
    table.style = "Table Grid"
    set_table_geometry(table, [2300, 3900, 3880], 120)
    data = [
        ("Need", "Include", "Avoid"),
        ("Accuracy", "Checkable facts and source limits", "Rumour, certainty without evidence"),
        ("Respect", "Real community impacts and ecological value", "Mocking residents or demonising wildlife"),
        ("Action", "Safe, lawful and practical steps", "Handling bats or unauthorised disturbance"),
    ]
    for r, values in enumerate(data):
        for c, value in enumerate(values):
            write_cell(table.cell(r, c), value, r == 0, FOREST if r == 0 else None, 9.0, WHITE if r == 0 else INK)
    set_repeat_header(table.rows[0])
    add_callout(
        doc,
        "Remember",
        "Objective language is not automatically proven. A council should still check its source, date, place, measurement and limits.",
        PALE_GOLD,
    )

    add_page_break(doc)
    add_label(doc, "Rivergum Council: two draft community messages")
    add_paragraph(
        doc,
        "These messages are fictional and deliberately positioned. Audit the language; neither is the official answer.",
        9.8,
        False,
        MUTED,
        5,
        True,
    )
    doc.add_heading("Message A  |  CLEAR THE CAMP NOW", level=1)
    add_callout(
        doc,
        "Resident Action Group draft",
        "A deafening invasion has taken over Rivergum Park. Families are forced to endure unbearable noise, foul smells and dangerous animals beside their homes. The camp is ruining our neighbourhood, and every responsible council would clear it immediately. Rivergum residents deserve their peaceful park back - no excuses and no delays.",
        PALE_RED,
        size=11,
    )
    doc.add_heading("Message B  |  WELCOME OUR NIGHT GARDENERS", level=1)
    add_callout(
        doc,
        "Wildlife Network draft",
        "Our wonderful night gardeners are bringing life to Rivergum Park. Flying-foxes move pollen and seeds through native forests, so their camp is a precious gift. Noise and droppings are only a tiny inconvenience, and there is nothing for residents to worry about. Everyone should welcome the camp and leave nature completely alone.",
        MIST,
        size=11,
    )
    doc.add_heading("Audit before you advise", level=2)
    add_bullets(
        doc,
        [
            "Underline one checkable claim in each message.",
            "Circle one judgement, emotion or loaded phrase in each message.",
            "Place a star beside one statement that mixes information and positioning.",
            "Write one question the council should answer before publishing either message.",
        ],
    )
    doc.add_heading("Source trail", level=2)
    add_paragraph(
        doc,
        "Evidence briefing: original summary of Queensland Government Living near flying-foxes; Queensland Government Bats and human health; Queensland Flying-fox roost management guideline; and Australian Government DCCEEW flying-fox monitoring and conservation pages. Full links are in the Teacher Guide. Checked 29 July 2026.",
        8.7,
        False,
        MUTED,
    )
    path = ROOT / "Lesson_15_Flying_Fox_Reading_Pack.docx"
    doc.save(path)
    return path


def build_organiser():
    doc = Document()
    configure_doc(
        doc,
        "Flying-Fox Language Audit",
        "Evidence, classification and council advice organiser",
        12,
    )
    add_callout(
        doc,
        "Learning intention",
        "Distinguish checkable information from subjective positioning, explain the effect and advise a council fairly.",
        PALE_GOLD,
    )
    doc.add_heading("A  Shared evidence bank", level=1)
    table = doc.add_table(rows=5, cols=3)
    table.style = "Table Grid"
    set_table_geometry(table, [2450, 4530, 3100], 120)
    headers = ("Category", "Accurate idea from the reading", "What should be checked?")
    for c, value in enumerate(headers):
        write_cell(table.cell(0, c), value, True, FOREST, 8.8, WHITE)
    for r, label in enumerate(("Behaviour/ecology", "Community impact", "Health/safety", "Lawful response"), start=1):
        write_cell(table.cell(r, 0), label, True, MIST, 8.8)
        write_cell(table.cell(r, 1), "\n\n", False, None, 9.2)
        write_cell(table.cell(r, 2), "\n\n", False, None, 9.2)
    set_repeat_header(table.rows[0])
    doc.add_heading("B  Message evidence", level=1)
    table = doc.add_table(rows=3, cols=3)
    table.style = "Table Grid"
    set_table_geometry(table, [1650, 4230, 4200], 120)
    for c, value in enumerate(("Draft", "Checkable wording", "Judgement or loaded wording")):
        write_cell(table.cell(0, c), value, True, FOREST, 8.8, WHITE)
    for r, label in enumerate(("A", "B"), start=1):
        write_cell(table.cell(r, 0), f"Message {label}", True, MIST, 9)
        write_cell(table.cell(r, 1), "\n\n", False, None, 9)
        write_cell(table.cell(r, 2), "\n\n", False, None, 9)
    set_repeat_header(table.rows[0])
    doc.add_heading("C  Classify the language", level=1)
    statements = [
        "Flying-foxes leave their daytime camps around dusk to search for food.",
        "Their unbearable racket ruins every nearby family’s day.",
        "The camp beside Rivergum Park has become a serious neighbourhood problem.",
        "Flying-foxes can spread pollen and seeds as they travel between feeding areas.",
        "These wonderful night gardeners deserve a warm welcome from everyone.",
        "Living near a camp can involve noise, odour and mess, but disease risk is extremely low when people do not handle bats.",
    ]
    table = doc.add_table(rows=7, cols=3)
    table.style = "Table Grid"
    set_table_geometry(table, [6400, 1500, 2180], 120)
    for c, value in enumerate(("Statement", "O / S / M", "Exact wording that proves it")):
        write_cell(table.cell(0, c), value, True, FOREST, 8.6, WHITE)
    for r, statement in enumerate(statements, start=1):
        write_cell(table.cell(r, 0), statement, False, None, 8.7)
        write_cell(table.cell(r, 1), "", False, None, 9)
        write_cell(table.cell(r, 2), "", False, None, 9)
    set_repeat_header(table.rows[0])

    add_page_break(doc)
    add_label(doc, "Positioning and revision")
    doc.add_heading("D  What job does the wording perform?", level=1)
    table = doc.add_table(rows=3, cols=4)
    table.style = "Table Grid"
    set_table_geometry(table, [1350, 2200, 2400, 4130], 120)
    for c, value in enumerate(("Draft", "Exact words", "Effect verb", "How does this position the reader?")):
        write_cell(table.cell(0, c), value, True, FOREST, 8.7, WHITE)
    for r, label in enumerate(("A", "B"), start=1):
        write_cell(table.cell(r, 0), f"Message {label}", True, MIST, 9)
        for c in range(1, 4):
            write_cell(table.cell(r, c), "\n\n", False, None, 9)
    set_repeat_header(table.rows[0])
    add_callout(
        doc,
        "Useful effect verbs",
        "alarms • reassures • minimises • dignifies • blames • invites • pressures • oversimplifies",
        MIST,
    )
    doc.add_heading("E  Revise one unsupported classification", level=1)
    add_lines(doc, 1, "My difficult statement was:")
    add_paragraph(doc, "Underline the checkable part. Circle the judgement or framing part.", 9.5, False, MUTED, 3, True)
    add_lines(doc, 2, "I changed / kept my label because:")
    doc.add_heading("F  Correct or qualify one claim", level=1)
    table = doc.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [3500, 6580], 120)
    pairs = [
        ("Claim that needs work", ""),
        ("Evidence-based correction", ""),
        ("Why the new wording is fairer", ""),
    ]
    for r, (label, value) in enumerate(pairs):
        write_cell(table.cell(r, 0), label, True, MIST, 9)
        write_cell(table.cell(r, 1), "\n\n", False, None, 9)
    doc.add_heading("G  Plan the council advice brief", level=1)
    table = doc.add_table(rows=6, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [3500, 6580], 120)
    prompts = [
        "Message A: words + effect",
        "Message B: words + effect",
        "Fact that corrects/qualifies",
        "Communication principle 1",
        "Communication principle 2",
        "Safe, lawful council action",
    ]
    for r, prompt in enumerate(prompts):
        write_cell(table.cell(r, 0), prompt, True, MIST, 9)
        write_cell(table.cell(r, 1), "", False, None, 9)

    add_page_break(doc)
    add_label(doc, "Write, test, revise")
    doc.add_heading("Persuasive Language Audit and Advice Brief", level=1)
    add_paragraph(
        doc,
        "Write 7-9 sentences to Rivergum Council in your English workbook. Use this organiser as your evidence base.",
        11,
        True,
        NIGHT,
    )
    table = doc.add_table(rows=8, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [900, 9180], 120)
    checklist = [
        "Name the purpose and likely effect of both messages.",
        "Classify at least four examples accurately.",
        "Use two short, exact language choices.",
        "Explain how wording positions readers.",
        "Correct or qualify one misleading or loaded claim.",
        "Recommend two communication principles.",
        "Recommend one practical, safe and lawful action.",
        "Make one visible revision after feedback.",
    ]
    for r, item in enumerate(checklist):
        write_cell(table.cell(r, 0), "☐", False, PALE_GOLD, 13)
        write_cell(table.cell(r, 1), item, False, None, 10)
    doc.add_heading("Peer fair-message test", level=1)
    add_lines(doc, 2, "Your clearest language-effect link is ___ because:")
    add_lines(doc, 2, "One claim that needs checking, qualifying or softening is ___ because:")
    doc.add_heading("My visible revision", level=1)
    add_lines(doc, 3)
    doc.add_heading("Exit evidence", level=1)
    add_lines(doc, 2, "A council can persuade responsibly by ___ instead of ___ because:")
    add_callout(
        doc,
        "Safety check",
        "My advice does not tell anyone to touch, feed, rescue or disturb a flying-fox.",
        PALE_RED,
    )
    path = ROOT / "Lesson_15_Flying_Fox_Language_Organiser.docx"
    doc.save(path)
    return path


def build_lucas_pack():
    doc = Document()
    configure_doc(
        doc,
        "Living With Flying-Foxes",
        "Large-print language choices and council advice",
        14,
        large=True,
    )
    add_callout(
        doc,
        "Today",
        "I can find a fact, a feeling or judgement, or both. I can help the council use calm, clear words.",
        PALE_GOLD,
        size=13,
    )
    doc.add_heading("Read these important ideas", level=1)
    add_paragraph(doc, "Flying-foxes sleep together in trees during the day. They fly out at night to find food.", 14)
    add_paragraph(doc, "They help forests by moving pollen and seeds.", 14)
    add_paragraph(doc, "A large camp can be noisy, smelly and messy for people who live nearby.", 14)
    add_paragraph(doc, "Do not touch a bat. A trained wildlife rescuer should help an injured bat.", 14)
    add_paragraph(doc, "Flying-foxes and their camps are protected. People must follow the rules and must not harm them.", 14)
    add_callout(
        doc,
        "Three labels",
        "FACT = can be checked  •  JUDGEMENT = tells a feeling or opinion  •  BOTH = has a fact and a judgement",
        MIST,
        size=13,
    )
    doc.add_heading("Choose a label", level=1)
    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [6900, 3180], 120)
    examples = [
        ("Sentence", "FACT / JUDGEMENT / BOTH"),
        ("Flying-foxes fly out around dusk.", ""),
        ("The disgusting animals have stolen our park.", ""),
        ("The camp is beside the park and is a serious problem.", ""),
    ]
    for r, values in enumerate(examples):
        for c, value in enumerate(values):
            write_cell(table.cell(r, c), value, r == 0, FOREST if r == 0 else None, 12.5, WHITE if r == 0 else INK)
    set_repeat_header(table.rows[0])
    doc.add_heading("Find the loaded words", level=1)
    add_paragraph(doc, "Circle the words that tell a strong judgement:", 13, True, FOREST)
    add_callout(doc, "Message", "A deafening invasion has taken over our peaceful park.", PALE_RED, size=14)
    add_lines(doc, 2, "The words make the reader feel:")

    add_page_break(doc)
    add_label(doc, "Help Rivergum Council")
    doc.add_heading("Make the message calmer and clearer", level=1)
    add_callout(
        doc,
        "Loaded message",
        "Dangerous bats are ruining everything.",
        PALE_RED,
        size=14,
    )
    add_paragraph(doc, "Choose or say a clearer message:", 14, True, FOREST)
    table = doc.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [900, 9180], 120)
    options = [
        ("☐", "The camp can create noise, odour and mess for nearby residents."),
        ("☐", "The bats are monsters."),
        ("☐", "Nothing matters except the animals."),
    ]
    for r, (box, text) in enumerate(options):
        write_cell(table.cell(r, 0), box, False, PALE_GOLD, 15)
        write_cell(table.cell(r, 1), text, False, None, 13)
    doc.add_heading("Choose safe council advice", level=1)
    table = doc.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [900, 9180], 120)
    options = [
        ("☐", "Do not touch a bat. Ask a trained wildlife rescuer for help."),
        ("☐", "Pick up an injured bat yourself."),
        ("☐", "Chase the camp away."),
    ]
    for r, (box, text) in enumerate(options):
        write_cell(table.cell(r, 0), box, False, MIST, 15)
        write_cell(table.cell(r, 1), text, False, None, 13)
    doc.add_heading("My council advice", level=1)
    add_lines(doc, 2, "Rivergum Council should use:")
    add_lines(doc, 3, "The council should tell people:")
    add_callout(
        doc,
        "You may point, choose, speak, copy or ask someone to scribe.",
        "The goal is a fair message: listen to people, give safe advice and protect wildlife.",
        PALE_GOLD,
        size=12.5,
    )
    path = ROOT / "Lesson_15_Flying_Fox_Lucas_Pack.docx"
    doc.save(path)
    return path


SLIDES = [
    ("FIELD BRIEFING", "Living With Flying-Foxes", "One camp. Two messages. Your council needs language it can trust.", "hero"),
    ("SAME CAMP • DIFFERENT MESSAGE", "What does each version make you notice?", '<div class="duel"><blockquote>A flying-fox camp has formed beside Rivergum Park.</blockquote><blockquote>A deafening invasion has taken over our peaceful park.</blockquote></div><p class="prompt">What changes: the event, the wording—or the reader?</p>', "standard"),
    ("MISSION", "Audit before the council publishes", '<div class="steps"><b>READ</b><span>Build shared knowledge</span><b>AUDIT</b><span>Separate information from positioning</span><b>ADVISE</b><span>Recommend fair persuasive wording</span></div><div class="product">Finished product: 7–9 sentence Language Audit and Advice Brief</div>', "standard"),
    ("READING PACK • 3 CODES", "Read for common ground", '<div class="codegrid"><article><b>C</b><span>community impact</span></article><article><b>E</b><span>ecological importance</span></article><article><b>S</b><span>safe or lawful response</span></article></div><p class="prompt">Box one sentence that corrects a misconception.</p>', "standard"),
    ("SHARED EVIDENCE BANK", "Calm wording still needs checking", '<div class="revealrow"><button class="reveal" data-answer="CHECK • place, date, species and observation">Flying-foxes leave camps around dusk.</button><button class="reveal" data-answer="CHECK • source, exposure pathway and current health advice">Disease risk is extremely low without handling.</button><button class="reveal" data-answer="CHECK • source, species and ecological study">Flying-foxes move pollen and seeds.</button></div><p class="prompt">Click after students name what would verify each claim.</p>', "standard"),
    ("TWO DRAFTS", "Both messages select and frame", '<div class="drafts"><article><small>MESSAGE A</small><h3>Clear the Camp Now</h3><p>noise • odour • mess • urgency</p></article><article><small>MESSAGE B</small><h3>Welcome Our Night Gardeners</h3><p>pollination • conservation • coexistence</p></article></div><p class="prompt">Underline one checkable claim and circle one loaded choice in each.</p>', "standard"),
    ("THE THREE-WAY TEST", "Judge the language—not the side", '<div class="triptych"><article><b>OBJECTIVE</b><p>mainly checkable information</p></article><article><b>SUBJECTIVE</b><p>judgement, emotion or loaded wording</p></article><article><b>MIXED</b><p>checkable claim + framing</p></article></div><div class="warning">Objective ≠ automatically true. Subjective ≠ automatically false.</div>', "standard"),
    ("CLASSIFICATION LAB", "Place it—then defend the boundary", '<div class="quiz" data-answers="O,S,M,M"><div><p>Flying-foxes leave camps around dusk to find food.</p><button>O</button><button>S</button><button>M</button></div><div><p>Their unbearable racket ruins every family’s day.</p><button>O</button><button>S</button><button>M</button></div><div><p>The camp beside the park is a serious neighbourhood problem.</p><button>O</button><button>S</button><button>M</button></div><div><p>Living near a camp can be noisy, but disease risk is extremely low without handling bats.</p><button>O</button><button>S</button><button>M</button></div></div><div class="action"><button id="checkQuiz">Check reasoning</button><button id="resetQuiz">Reset</button><span id="quizFeedback" aria-live="polite"></span></div>', "standard"),
    ("POSITIONING", "What job does the wording perform?", '<div class="wordarc"><span>“deafening invasion”</span><i>→</i><b>alarms</b><i>→</i><p>positions residents to expect danger and loss of control</p></div><div class="wordarc"><span>“tiny inconvenience”</span><i>→</i><b>minimises</b><i>→</i><p>dismisses people who experience genuine impacts</p></div>', "standard"),
    ("REVISION IS EVIDENCE", "Return to your hardest classification", '<div class="revision"><b>1</b> underline the checkable part <b>2</b> circle the framing <b>3</b> keep or change the label <b>4</b> explain why</div><p class="prompt">I changed / kept my label because…</p>', "standard"),
    ("ANNOTATED MODEL", "Fair does not mean neutral mush", '<div class="model"><p><mark>Both drafts use real concerns, but each pushes readers towards a different judgement.</mark> Message A calls the camp a <u>“deafening invasion”</u>, which exaggerates noise and alarms residents. Message B explains pollination, but <u>“welcome from everyone”</u> dismisses people affected by odour or mess. The council should replace loaded labels with specific descriptions, give no-touch safety advice and explain lawful management.</p></div><div class="revealrow"><button class="reveal" data-answer="JOB • establishes a comparative judgement">claim</button><button class="reveal" data-answer="JOB • connects exact wording to reader effect">effect</button><button class="reveal" data-answer="JOB • turns analysis into responsible action">advice</button></div>', "standard"),
    ("PLAN", "Build the advice chain before writing", '<div class="chain"><span>Message A words + effect</span><i>→</i><span>Message B words + effect</span><i>→</i><span>evidence correction</span><i>→</i><span>2 principles + 1 action</span></div>', "standard"),
    ("WRITE • 10 MINUTES", "Persuasive Language Audit and Advice Brief", '<div class="criteria"><p>✓ both messages</p><p>✓ 4 classifications</p><p>✓ 2 exact word choices</p><p>✓ effects on readers</p><p>✓ correction or qualification</p><p>✓ 2 principles + safe action</p></div><div class="timer"><span id="timerDisplay">10:00</span><button id="timerStart">Start / pause</button><button id="timerReset">Reset</button></div>', "standard"),
    ("FEEDBACK + REVISION", "Run the fair-message test", '<div class="feedbackstems"><p>Your clearest language–effect link is ___ because ___.</p><p>One claim that needs checking, qualifying or softening is ___ because ___.</p></div><div class="warning">Revise one sentence immediately.</div>', "standard"),
    ("DEPTH A • OPTIONAL", "Preserve the fact. Change the frame.", '<div class="audiences"><span>resident group</span><span>wildlife group</span><span>council information page</span></div><p class="prompt">Rewrite one event three ways. Which version carries the least visible judgement?</p>', "standard"),
    ("DEPTH B • OPTIONAL", "Audit an objective-sounding claim", '<div class="sourcecheck"><span>named source</span><span>location + date</span><span>definition or measurement</span><span>corroboration</span><span>limits</span></div><p class="prompt">What would the council need before publishing it?</p>', "standard"),
    ("EXIT EVIDENCE", "Persuasive without distortion", '<div class="exit">A council can persuade responsibly by <u>__________</u> instead of <u>__________</u> because <u>____________________</u>.</div>', "standard"),
]


NOTES = [
    "Open with the visual tension: a thriving native animal can also create a difficult human-wildlife interface.",
    "Do not classify yet. Ask what wording changes the frame.",
    "Name Rivergum Council as fictional. Students are language advisers, not wildlife managers.",
    "Preteach camp/roost, nocturnal, pollination, protected and loaded language as required.",
    "Require what would verify the calm-sounding claim. Objective does not mean already proven.",
    "Both drafts are fictional and deliberately incomplete. Neither is the council answer.",
    "Keep checkability separate from truth. Mixed is a useful category, not a failure to decide.",
    "Students answer independently before clicking. Sentence 4 is mixed because it combines impacts, risk scale and a condition.",
    "Move beyond 'emotive'. Ask what the wording invites the audience to believe or do.",
    "A retained label still needs stronger reasoning. Visible revision is the formative evidence.",
    "Read the model whole first. Then reveal the job of each move.",
    "Conference with students who have labels but no reader-effect explanation.",
    "Start only when students have completed the organiser plan. Warn at five and two minutes.",
    "Reviewer diagnoses; writer chooses the revision. Correct unsafe advice immediately.",
    "Optional. Viewpoint can operate through selection and order even when obvious loaded words disappear.",
    "Optional. Use a calm-sounding claim so students do not rely on tone as a credibility test.",
    "Collect the sentence. Group next teaching by classification, effect or factual accuracy.",
]


def build_html():
    slide_html = []
    for idx, (kicker, title, body, kind) in enumerate(SLIDES):
        active = " active" if idx == 0 else ""
        hero = " hero" if kind == "hero" else ""
        if kind == "hero":
            body = """
            <div class="heroart" aria-label="Silhouette of flying-foxes crossing a moonlit forest">
              <div class="moon"></div><div class="canopy c1"></div><div class="canopy c2"></div>
              <div class="bat b1">⌁</div><div class="bat b2">⌁</div><div class="bat b3">⌁</div>
            </div>
            <div class="herocopy"><span>SUPPLEMENTAL LESSON 15</span><h1>Living With<br><em>Flying-Foxes</em></h1><p>One camp. Two messages. Your council needs language it can trust.</p></div>
            """
        slide_html.append(
            f'<section class="slide{active}{hero}" data-notes="{NOTES[idx].replace(chr(34), "&quot;")}">'
            f'<div class="slideinner"><div class="kicker">{kicker}</div>'
            f'{"<h2>"+title+"</h2>" if kind != "hero" else ""}{body}</div></section>'
        )
    html = f"""<!doctype html>
<html lang="en">
<head>
<meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1">
<title>Lesson 15 Alternative | Living With Flying-Foxes</title>
<style>
:root{{--night:#132a34;--forest:#1e5c4a;--leaf:#4b8b61;--lime:#b8d97b;--gold:#f0c85a;--cream:#fff8e7;--mist:#eef5ee;--ink:#19312d;--red:#a33b32}}
*{{box-sizing:border-box}} body{{margin:0;background:#07171d;color:var(--cream);font-family:Arial,sans-serif;overflow:hidden}}
.slide{{display:none;position:absolute;inset:0;padding:5.2vh 6vw 11vh;background:radial-gradient(circle at 78% 18%,#244e4c 0,#132a34 42%,#091c23 100%)}} .slide.active{{display:block}}
.slideinner{{height:100%;max-width:1500px;margin:auto;display:flex;flex-direction:column;justify-content:center}}
.kicker{{color:var(--gold);font-weight:800;letter-spacing:.16em;font-size:clamp(13px,1.15vw,20px);margin-bottom:1.5vh}}
h2{{font-size:clamp(34px,4.8vw,76px);line-height:1.02;margin:0 0 4vh;max-width:1050px}} h3{{font-size:clamp(21px,2.2vw,36px);margin:.35em 0}}
p,blockquote{{font-size:clamp(20px,2vw,34px);line-height:1.35}} button{{font:inherit}}
.hero{{padding:0}} .hero .slideinner{{max-width:none;display:grid;grid-template-columns:1.05fr .95fr}} .herocopy{{padding:10vh 6vw;display:flex;flex-direction:column;justify-content:center}}
.herocopy span{{color:var(--gold);font-weight:800;letter-spacing:.16em}} .herocopy h1{{font-size:clamp(56px,7.6vw,132px);line-height:.87;margin:2vh 0}} .herocopy em{{color:var(--lime);font-style:normal}} .herocopy p{{max-width:760px;color:#d7e5dd}}
.heroart{{position:relative;overflow:hidden;background:linear-gradient(#173747,#386957)}} .moon{{position:absolute;width:31vw;height:31vw;border-radius:50%;background:#fff2bd;right:8vw;top:8vh;box-shadow:0 0 90px #fff3bc66}}
.canopy{{position:absolute;bottom:-8vh;border-radius:50% 50% 0 0;background:#0a2424}} .c1{{left:-10%;width:80%;height:42%}} .c2{{right:-15%;width:75%;height:34%;background:#123a32}}
.bat{{position:absolute;color:#10252a;font-size:clamp(50px,7vw,120px);font-weight:900;transform:rotate(-5deg)}} .b1{{top:20%;left:15%}} .b2{{top:36%;left:42%;font-size:clamp(34px,5vw,80px)}} .b3{{top:14%;right:8%;font-size:clamp(28px,4vw,66px)}}
.duel,.drafts,.triptych,.codegrid{{display:grid;gap:2vw}} .duel,.drafts{{grid-template-columns:1fr 1fr}} .duel blockquote,.drafts article,.triptych article,.codegrid article{{margin:0;padding:3vh 2.2vw;border:2px solid #6da17c;background:#ffffff0e;border-radius:18px}}
.duel blockquote:last-child{{border-color:var(--gold);background:#a33b3222}} .drafts article:first-child{{border-color:#d47167}} .drafts article:last-child{{border-color:var(--lime)}}
.prompt,.product,.warning{{margin-top:4vh;padding:2vh 2vw;border-left:8px solid var(--gold);background:#fff8e710}} .steps{{display:grid;grid-template-columns:auto 1fr;gap:2vh 2vw;font-size:clamp(22px,2.3vw,38px);align-items:center}} .steps b{{color:var(--lime)}}
.codegrid,.triptych{{grid-template-columns:repeat(3,1fr)}} .codegrid b{{display:block;font-size:clamp(42px,5vw,80px);color:var(--gold)}} .codegrid span{{font-size:clamp(18px,1.7vw,28px)}}
.revealrow{{display:flex;gap:1.5vw;flex-wrap:wrap}} .reveal{{flex:1;min-width:220px;padding:2.4vh 1.4vw;border:2px solid var(--lime);border-radius:14px;background:#fff;color:var(--ink);font-weight:700;cursor:pointer}} .reveal.revealed{{background:var(--gold)}} 
.drafts small{{color:var(--gold);font-weight:800}} .triptych b{{color:var(--gold);font-size:clamp(19px,2vw,32px)}} .triptych p{{font-size:clamp(17px,1.55vw,26px)}}
.quiz{{display:grid;grid-template-columns:1fr 1fr;gap:1.2vw}} .quiz>div{{padding:1.4vh 1vw;background:#ffffff10;border-radius:12px}} .quiz p{{font-size:clamp(15px,1.35vw,22px);margin:.2em 0 .7em}} .quiz button{{padding:.55em 1em;margin-right:.5em;border:2px solid #8fb29e;background:transparent;color:#fff;border-radius:9px;cursor:pointer}} .quiz button.selected{{background:var(--gold);color:var(--night)}} .quiz button.correct{{background:var(--leaf)}} .quiz button.wrong{{background:var(--red)}} .action{{margin-top:2vh;display:flex;gap:1vw;align-items:center}} .action>button,.timer button{{padding:.7em 1.2em;border:0;border-radius:10px;background:var(--gold);color:var(--night);font-weight:800;cursor:pointer}} #quizFeedback{{font-weight:700}}
.wordarc{{display:grid;grid-template-columns:1.5fr auto .7fr auto 2.8fr;gap:1vw;align-items:center;margin:2vh 0;padding:2vh 2vw;background:#fff1;border-radius:16px}} .wordarc span{{font-size:clamp(20px,2vw,34px);color:var(--gold)}} .wordarc p{{font-size:clamp(16px,1.5vw,25px)}} .wordarc i{{color:var(--lime)}}
.revision,.chain{{display:flex;align-items:center;justify-content:space-between;gap:1vw;font-size:clamp(18px,2vw,32px)}} .revision b{{display:grid;place-items:center;width:2em;height:2em;border-radius:50%;background:var(--gold);color:var(--night)}} .model{{padding:2vh 2vw;background:#fff;color:var(--ink);border-radius:18px}} .model p{{font-size:clamp(17px,1.55vw,25px);margin:0}} mark{{background:#ffe38a}} .chain span{{padding:2.2vh 1.4vw;background:#fff1;border:2px solid #6da17c;border-radius:14px;text-align:center}} .chain i{{color:var(--gold)}}
.criteria{{display:grid;grid-template-columns:repeat(3,1fr);gap:1vw}} .criteria p{{font-size:clamp(16px,1.5vw,24px);margin:0;padding:1.5vh 1vw;background:#fff1;border-radius:10px}} .timer{{display:flex;gap:1vw;align-items:center;margin-top:4vh}} #timerDisplay{{font-size:clamp(44px,5vw,80px);font-weight:900;color:var(--gold)}}
.feedbackstems p,.exit{{padding:2.4vh 2vw;background:#fff1;border-radius:14px}} .audiences,.sourcecheck{{display:flex;gap:1vw;flex-wrap:wrap}} .audiences span,.sourcecheck span{{padding:2.3vh 1.5vw;border:2px solid var(--lime);border-radius:999px;font-size:clamp(18px,1.7vw,28px)}} .exit{{font-size:clamp(24px,2.8vw,46px);line-height:1.5}}
#controls{{position:fixed;left:0;right:0;bottom:0;height:9vh;display:flex;align-items:center;gap:1vw;padding:0 2vw;background:#07171df2;border-top:1px solid #54766a;z-index:20}} #controls button{{padding:.65em 1em;border:1px solid #709285;background:#132a34;color:#fff;border-radius:9px;cursor:pointer}} #counter{{font-weight:800}} #progress{{height:6px;flex:1;background:#29433c;border-radius:9px;overflow:hidden}} #bar{{height:100%;width:0;background:var(--gold)}} #notes{{position:fixed;right:2vw;bottom:10vh;max-width:520px;padding:1.2em;background:#fff;color:var(--ink);border-radius:12px;box-shadow:0 10px 30px #0008;display:none;z-index:30}} #notes.open{{display:block}}
button:focus-visible{{outline:4px solid #fff;outline-offset:3px}} @media(max-width:900px){{.duel,.drafts,.triptych,.codegrid,.quiz{{grid-template-columns:1fr 1fr}} h2{{margin-bottom:2vh}} .slide{{padding-top:3vh}}}} @media(prefers-reduced-motion:reduce){{*{{scroll-behavior:auto!important;transition:none!important}}}}
</style>
</head>
<body>
{''.join(slide_html)}
<aside id="notes" aria-live="polite"></aside>
<nav id="controls" aria-label="Presentation controls">
<button id="prev" aria-label="Previous slide">← Prev</button><span id="counter">1 / {len(SLIDES)}</span>
<div id="progress"><div id="bar"></div></div>
<button id="notesBtn">Teacher notes</button><button id="resetSlide">Reset</button><button id="fullscreen">Fullscreen</button><button id="next">Next →</button>
</nav>
<script>
const slides=[...document.querySelectorAll('.slide')]; let current=0;
const counter=document.getElementById('counter'),bar=document.getElementById('bar'),notes=document.getElementById('notes');
function show(n){{current=Math.max(0,Math.min(slides.length-1,n));slides.forEach((s,i)=>s.classList.toggle('active',i===current));counter.textContent=`${{current+1}} / ${{slides.length}}`;bar.style.width=`${{(current+1)/slides.length*100}}%`;notes.textContent=slides[current].dataset.notes||'';notes.classList.remove('open')}}
function resetCurrent(){{slides[current].querySelectorAll('.reveal').forEach(b=>{{b.classList.remove('revealed');b.textContent=b.dataset.original||b.textContent}});slides[current].querySelectorAll('.quiz button').forEach(b=>b.className='');if(document.getElementById('quizFeedback'))document.getElementById('quizFeedback').textContent='';}}
document.getElementById('prev').onclick=()=>show(current-1);document.getElementById('next').onclick=()=>show(current+1);document.getElementById('resetSlide').onclick=resetCurrent;
document.getElementById('notesBtn').onclick=()=>notes.classList.toggle('open');document.getElementById('fullscreen').onclick=()=>document.documentElement.requestFullscreen?.();
document.addEventListener('keydown',e=>{{if(['INPUT','TEXTAREA'].includes(document.activeElement.tagName))return;if(['ArrowRight','PageDown',' '].includes(e.key)){{e.preventDefault();show(current+1)}}if(['ArrowLeft','PageUp'].includes(e.key)){{e.preventDefault();show(current-1)}}if(e.key.toLowerCase()==='n')notes.classList.toggle('open')}});
document.querySelectorAll('.reveal').forEach(b=>{{b.dataset.original=b.textContent;b.onclick=()=>{{b.classList.toggle('revealed');b.textContent=b.classList.contains('revealed')?b.dataset.answer:b.dataset.original}}}});
document.querySelectorAll('.quiz>div').forEach(group=>group.querySelectorAll('button').forEach(b=>b.onclick=()=>{{group.querySelectorAll('button').forEach(x=>x.classList.remove('selected','correct','wrong'));b.classList.add('selected')}}));
const check=document.getElementById('checkQuiz');if(check)check.onclick=()=>{{const quiz=document.querySelector('.quiz'),answers=quiz.dataset.answers.split(',');let chosen=0,correct=0;[...quiz.children].forEach((g,i)=>{{const b=g.querySelector('.selected');if(b){{chosen++;if(b.textContent===answers[i]){{b.classList.add('correct');correct++}}else b.classList.add('wrong')}}}});document.getElementById('quizFeedback').textContent=chosen<answers.length?'Choose one label for every statement.':correct===answers.length?'All four fit the boundary. Now justify the hardest one.':`${{correct}} / ${{answers.length}} fit. Recheck the exact wording, then retry.`}};document.getElementById('resetQuiz')?.addEventListener('click',()=>{{document.querySelectorAll('.quiz button').forEach(b=>b.className='');document.getElementById('quizFeedback').textContent=''}});
let remain=600,timer=null;function paint(){{const m=Math.floor(remain/60),s=remain%60;document.getElementById('timerDisplay').textContent=`${{m}}:${{String(s).padStart(2,'0')}}`}}document.getElementById('timerStart')?.addEventListener('click',()=>{{if(timer){{clearInterval(timer);timer=null;return}}timer=setInterval(()=>{{if(remain>0)remain--;paint();if(remain===0){{clearInterval(timer);timer=null}}}},1000)}});document.getElementById('timerReset')?.addEventListener('click',()=>{{clearInterval(timer);timer=null;remain=600;paint()}});
show(0);
</script>
</body></html>"""
    path = ROOT / "Lesson_15_Flying_Fox_Persuasive_Presentation.html"
    path.write_text(html, encoding="utf-8")
    return path


if __name__ == "__main__":
    outputs = [build_html(), build_reading_pack(), build_organiser(), build_lucas_pack()]
    for output in outputs:
        print(output)
