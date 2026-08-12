from __future__ import annotations

"""Build the Lesson 25 Harbour Light alternative lesson package."""

import importlib.util
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib.colors import HexColor
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer


ROOT = Path(__file__).resolve().parents[1]
WORKSPACE = ROOT.parents[6]
SKILL_ROOT = WORKSPACE / ".agents" / "skills" / "lesson-creator"
ASSETS = ROOT / "assets"

ARTICLE_TITLE = "When the Lantern Went Dark"
ARTICLE_DECK = "A fictional community feature about an emergency, a difficult choice and what courage can persuade us to do next."

ARTICLE_PARAGRAPHS = [
    ("A warning no one could ignore", "On the afternoon of the Bay Festival, the lantern on the old Harbour Light Jetty went dark. The jetty was not open to visitors after the storm, but the unlit lantern mattered: it marked the shallow channel that the community fishing boats used to return home. At first, people gathered on the foreshore and waited for an adult to solve it. Then Nia noticed something else: a small rescue dinghy had become caught in drifting rope near the jetty steps."),
    ("The choice", "Nia had been collecting rubbish for the festival clean-up with her cousin, Ravi. They could see the rope tightening whenever the tide pushed the dinghy sideways. Ravi wanted to run straight onto the jetty. Nia stopped him. The storm had loosened boards and the waves were still high. Instead, she used the public emergency phone beside the kiosk, described the rope, the dark lantern and the position of the dinghy, and asked the festival coordinator to keep people back from the water."),
    ("Courage is not careless", "While they waited, Nia did not pretend she felt calm. She worried that the dinghy would scrape against the pylons and that people would blame them for not doing more. Still, she kept giving clear updates. When the marine rescue crew arrived, they used a long pole and safety line to free the dinghy. A technician restored the lantern before dusk. The important turning point was not a dramatic leap onto the jetty. It was Nia choosing a safe action that made expert help possible."),
    ("What changed afterwards", "The next week, the festival committee proposed a 'Look, Call, Clear' sign near the foreshore. The sign would remind visitors to notice hazards, call for help and keep a safe space clear for trained responders. Some residents said a sign would be unnecessary because adults already know what to do. Others worried that the message might make children feel responsible for emergencies. Nia's account offered a more careful idea: a sign cannot replace adults or trained crews, but it can help everyone recognise when stopping, reporting and making room is a brave kind of action."),
    ("A claim worth sharing", "The Harbour Light story suggests that courage is not always loud. It can mean noticing a real problem, accepting that you cannot solve it alone and taking the next safe step. That idea matters beyond the jetty. A community works better when people do not ignore a risk simply because they feel unsure. The proposed sign should be trialled for the next festival because it turns one frightening moment into a practical invitation: see the problem, seek help and leave space for others to act safely."),
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_margins(cell, top=120, start=140, bottom=120, end=140) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    mar = tc_pr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_run(run, size=11, bold=False, colour="173D45"):
    run.font.name = "Aptos"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(colour)


def setup_doc(doc: Document):
    section = doc.sections[0]
    section.top_margin = Inches(0.62)
    section.bottom_margin = Inches(0.62)
    section.left_margin = Inches(0.68)
    section.right_margin = Inches(0.68)
    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.14


def para(doc, text="", size=11, bold=False, colour="173D45", after=6, before=0, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(before)
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_run(r, size, bold, colour)
    return p


def title_block(doc, kicker, title, subtitle):
    para(doc, kicker.upper(), 9, True, "A15C3B", after=2)
    para(doc, title, 25, True, "123E4A", after=2)
    para(doc, subtitle, 11, False, "476D72", after=14)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, "123E4A")
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        r = cell.paragraphs[0].add_run(header)
        set_run(r, 9.5, True, "FFFFFF")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_margins(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            r = cells[i].paragraphs[0].add_run(value)
            set_run(r, 9.5)
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Inches(width)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def build_reading_pack(path: Path):
    doc = Document(); setup_doc(doc)
    title_block(doc, "English Unit 3 | Lesson 25 alternative", ARTICLE_TITLE, ARTICLE_DECK)
    para(doc, "This is an original fictional text. Treat its events as material for close reading, not as real-world evidence.", 9.5, True, "7C4C17", after=10)
    for heading, body in ARTICLE_PARAGRAPHS:
        para(doc, heading, 14, True, "21626A", after=2, before=6)
        para(doc, body, 10.5, False, "173D45", after=6)
    para(doc, "Reading trail", 14, True, "21626A", after=4, before=6)
    add_table(doc, ["Find", "Copy or paraphrase", "What it helps you claim"], [
        ["Turning point", "Nia calls for trained help instead of entering the jetty.", "Courage can include safe, deliberate action."],
        ["Consequence", "The rescue crew can free the dinghy and restore the light.", "One choice can make another action possible."],
        ["Future action", "The committee proposes a 'Look, Call, Clear' sign.", "A story can become a persuasive proposal."],
    ], [1.15, 3.1, 2.0])
    para(doc, "Talk before you write", 14, True, "21626A", after=4, before=8)
    para(doc, "1. Which moment is the climax: the dark lantern, Ravi's first idea, Nia's call, or the sign proposal? Defend your choice.", 10.5)
    para(doc, "2. How does Nia's choice show care without pretending she can solve the emergency alone?", 10.5)
    para(doc, "3. Which detail could support a persuasive claim about courage? Explain the link.", 10.5)
    doc.save(path)


def build_organiser(path: Path, concise=False):
    doc = Document(); setup_doc(doc)
    title_block(doc, "English Unit 3 | Lesson 25 alternative", "Turning point to persuasive claim", "Plan a short paragraph that explains why a brave choice matters to a reader.")
    if concise:
        para(doc, "You can point, speak, copy, draw an arrow or use a partner scribe.", 11, True, "7C4C17", after=10)
        add_table(doc, ["What happened?", "What did Nia do?", "What changed?"], [["The lantern went dark. The dinghy was stuck.", "She used the emergency phone and asked people to stay back.", "The rescue crew could help safely."]], [2.1, 2.1, 2.1])
        para(doc, "My claim", 14, True, "21626A", before=8)
        para(doc, "Courage can mean ________________________________________________.", 12, False, "173D45", after=12)
        para(doc, "Why it matters", 14, True, "21626A")
        para(doc, "This matters because ______________________________________________.", 12, False, "173D45", after=12)
        para(doc, "Say or write", 14, True, "21626A")
        para(doc, "Nia was brave when _______________________________________________.", 12)
        para(doc, "This shows that courage can _______________________________________.", 12)
        para(doc, "□ point  □ speak  □ copy  □ partner scribe", 10, True, "476D72", after=0)
    else:
        para(doc, "1. Zoom in on the turning point", 14, True, "21626A", before=4)
        add_table(doc, ["Precise event", "Character decision", "Consequence"], [["", "", ""]], [2.1, 2.1, 2.1])
        para(doc, "2. Build a defensible claim", 14, True, "21626A", before=8)
        add_table(doc, ["Evidence detail", "Claim about courage", "Why a reader should care"], [["", "", ""]], [2.1, 2.1, 2.1])
        para(doc, "3. Draft and revise (6-8 sentences)", 14, True, "21626A", before=8)
        para(doc, "Start with the event. Explain the decision and consequence. Make a claim about courage. Use one detail. End with why the claim matters to a reader.", 10.5, False, "173D45", after=8)
        for _ in range(6):
            para(doc, "________________________________________________________________________________", 11, False, "476D72", after=8)
        para(doc, "Revision check: □ precise event  □ consequence  □ evidence  □ claim  □ reader purpose", 10, True, "476D72", after=0)
    doc.save(path)


def build_teacher_guide(path: Path):
    doc = Document(); setup_doc(doc)
    title_block(doc, "Teacher guide | Lesson 25 alternative", "Harbour Light: Climax to Claim", "A 50-minute original-text lesson that bridges narrative meaning to persuasive writing.")
    para(doc, "Teaching focus", 14, True, "21626A")
    para(doc, "Students move from recounting an event to explaining how a character's choice changes what is possible, then use that interpretation as the centre of a short persuasive claim. The fictional reading keeps the focus on reasoning rather than factual research.", 10.5)
    para(doc, "Model answer", 14, True, "21626A", before=6)
    para(doc, "When the Harbour Light Jetty lantern goes dark and a rescue dinghy catches in rope, Nia chooses to call for trained help rather than enter the unsafe jetty. Her decision lets the rescue crew work safely and restore the light before dusk. This suggests that courage can mean taking the responsible next step even when you feel afraid or cannot solve a problem alone. Nia does not ignore the danger, but she also does not act carelessly. Her choice matters because readers may think bravery always means doing the dramatic thing. The proposed 'Look, Call, Clear' sign should be trialled because it turns that lesson into a practical reminder for the whole community.", 10.5)
    para(doc, "Misconceptions to surface", 14, True, "21626A", before=6)
    add_table(doc, ["Likely idea", "Responsive prompt"], [
        ["The climax is simply the most exciting event.", "Which event changes what can happen next?"],
        ["Bravery means acting without fear.", "What does Nia feel, and what does she choose despite that feeling?"],
        ["A claim is a plot summary.", "What broader idea does this event suggest to a reader?"],
        ["The story proves a real fact about emergencies.", "What is fictional here? What evidence would we need for a real proposal?"],
    ], [2.4, 3.9])
    para(doc, "Answer guidance", 14, True, "21626A", before=6)
    para(doc, "The strongest climax choice is Nia's emergency call because it redirects the response from unsafe impulse to expert action. Accept another choice only when a student explains how it changes the direction of events. Strong claims are interpretations, not a single correct theme. They should include a precise event, a consequence, a cautious statement about courage and a reader purpose.", 10.5)
    para(doc, "Access and extension", 14, True, "21626A", before=6)
    para(doc, "Green view and concise organiser: reduce language load, retain the same event-to-claim pathway, allow oral/pointing/scribed responses. Extension: ask students to qualify a claim using 'can' or 'may', and explain why a community sign is a proposal rather than proof that every reader will act safely.", 10.5)
    doc.save(path)


def build_pdf(path: Path):
    styles = getSampleStyleSheet()
    story = []
    doc = SimpleDocTemplate(str(path), pagesize=A4, rightMargin=18*mm, leftMargin=18*mm, topMargin=16*mm, bottomMargin=16*mm)
    story.append(Paragraph("ENGLISH UNIT 3 | LESSON 25 ALTERNATIVE", ParagraphStyle("k", parent=styles["Normal"], fontName="Helvetica-Bold", fontSize=8, textColor=HexColor("#A15C3B"), spaceAfter=4)))
    story.append(Paragraph(ARTICLE_TITLE, ParagraphStyle("t", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=24, leading=28, textColor=HexColor("#123E4A"), spaceAfter=5)))
    story.append(Paragraph(ARTICLE_DECK, ParagraphStyle("d", parent=styles["Normal"], fontName="Helvetica", fontSize=10, leading=14, textColor=HexColor("#476D72"), spaceAfter=10)))
    story.append(Paragraph("This is an original fictional text. Treat its events as material for close reading, not as real-world evidence.", ParagraphStyle("b", parent=styles["Normal"], fontName="Helvetica-Bold", fontSize=9, leading=12, textColor=HexColor("#7C4C17"), spaceAfter=10)))
    for heading, body in ARTICLE_PARAGRAPHS:
        story.append(Paragraph(heading, ParagraphStyle("h"+heading[:3], parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=13, leading=16, textColor=HexColor("#21626A"), spaceBefore=7, spaceAfter=3)))
        story.append(Paragraph(body, ParagraphStyle("p"+heading[:3], parent=styles["BodyText"], fontName="Helvetica", fontSize=10.5, leading=15, textColor=HexColor("#173D45"), spaceAfter=6)))
    story.append(Spacer(1, 5))
    story.append(Paragraph("Talk before you write", ParagraphStyle("talk", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=13, textColor=HexColor("#21626A"), spaceAfter=4)))
    for text in ["Which moment is the climax? Defend your choice.", "How does Nia show care without pretending she can solve the emergency alone?", "Which detail could support a persuasive claim about courage? Explain the link."]:
        story.append(Paragraph("• " + text, ParagraphStyle("q"+str(len(story)), parent=styles["BodyText"], fontName="Helvetica", fontSize=10.5, leading=14, textColor=HexColor("#173D45"), leftIndent=6, spaceAfter=3)))
    doc.build(story)


def compile_deck():
    spec = importlib.util.spec_from_file_location("compiler", SKILL_ROOT / "scripts" / "compile_presentation.py")
    compiler = importlib.util.module_from_spec(spec); sys.modules["compiler"] = compiler; spec.loader.exec_module(compiler)
    compiler.compile_presentation(ASSETS / "harbour_light_slides.html", ROOT / "Lesson_25_Harbour_Light_Presentation.html", css_path=ASSETS / "harbour_light.css", js_path=ASSETS / "harbour_light.js", title="Lesson 25 | Harbour Light: Climax to Claim", language="en-AU")


def main():
    ROOT.mkdir(parents=True, exist_ok=True)
    build_reading_pack(ROOT / "Lesson_25_Harbour_Light_Reading_Pack.docx")
    build_organiser(ROOT / "Lesson_25_Harbour_Light_Organiser.docx")
    build_organiser(ROOT / "Lesson_25_Harbour_Light_Concise_Access_Pack.docx", concise=True)
    build_teacher_guide(ROOT / "Lesson_25_Harbour_Light_Teacher_Guide.docx")
    build_pdf(ROOT / "Lesson_25_Harbour_Light_Reading_Article.pdf")
    compile_deck()


if __name__ == "__main__":
    main()
